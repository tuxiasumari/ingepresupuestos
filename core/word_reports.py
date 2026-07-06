# SPDX-License-Identifier: GPL-3.0-or-later
# Copyright (C) 2026 Marco Sumari / Sumari SAC
# This file is part of IngePresupuestos — https://ingepresupuestos.com
# Licensed under the GNU GPL v3.0 or later. See the LICENSE file.
"""Reportes en formato Word (.docx) — editables.

Reutiliza la lógica de cómputo de `core/pdf_reports.py` pero genera un
documento Word nativo con python-docx. Foco inicial: reportes texto-pesados
donde el cliente final quiere editar (Resumen Ejecutivo, ACUs,
Especificaciones Técnicas).

Para los tabulares puros (Presupuesto, Metrados, Insumos, Cronograma)
seguir usando `core/exporter.py` que ya genera Excel.

Convención cross-platform: Word ≈ LibreOffice Writer. El .docx generado
abre sin problemas en MS Office, LibreOffice y Google Docs.
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.database import (
    get_db, calcular_totales, get_decimales_ppto,
)
from core.pdf_reports import (
    _moneda_simbolo, _fmt, _build_pie_rows, _monto_en_letras,
    _png_donut_bytes, _clean_costo_al, strip_titulo_memoria,
    MO_COLOR, MAT_COLOR, EQ_COLOR, SC_COLOR,
)


# ── Helpers de estilo ───────────────────────────────────────────────────────

_SLATE_900 = RGBColor(0x0F, 0x17, 0x2A)
_SLATE_700 = RGBColor(0x27, 0x34, 0x45)
_SLATE_500 = RGBColor(0x48, 0x5A, 0x6C)
_SLATE_300 = RGBColor(0x66, 0x78, 0x85)
_SLATE_100 = RGBColor(0xCB, 0xD5, 0xE1)
_ORANGE    = RGBColor(0xF3, 0x73, 0x29)


def _accent_od_rgb() -> RGBColor:
    """Color `od` actual de `accent_reportes()` como `RGBColor`.

    Espejo del `color: {od}` que usa el CSS del h2 en `pdf_reports._base_css`.
    En modo sobrio devuelve slate-900 (negro); en modo marca, naranja-700.
    Se llama por cada reporte (no se cachea) para que el toggle aplique sin
    reiniciar.
    """
    from utils.theme import accent_reportes
    od = accent_reportes()[1].lstrip('#')
    return RGBColor(int(od[0:2], 16), int(od[2:4], 16), int(od[4:6], 16))


def _accent_o_hex() -> str:
    """Color `o` (primario del acento) como hex sin `#`. Espejo del `{o}`
    del CSS de PDF (naranja `F37329` en marca, slate-700 `273445` en sobrio).
    Útil para `_set_paragraph_border` / `_cell_border` (que esperan hex)."""
    from utils.theme import accent_reportes
    return accent_reportes()[0].lstrip('#')

_FUENTE = 'Inter'
_IDIOMA = 'es-PE'  # Español (Perú) — para corrector ortográfico de Word/LO


def _aplicar_idioma(doc: Document, lang: str = _IDIOMA):
    """Setea el idioma por defecto del documento (es-PE) para que el
    corrector ortográfico de Word/LibreOffice lo reconozca como español
    y no marque las palabras como mal escritas."""
    styles = doc.styles.element
    defaults = styles.find(qn('w:docDefaults'))
    if defaults is None:
        defaults = OxmlElement('w:docDefaults')
        styles.insert(0, defaults)
    rpr_default = defaults.find(qn('w:rPrDefault'))
    if rpr_default is None:
        rpr_default = OxmlElement('w:rPrDefault')
        defaults.append(rpr_default)
    rpr = rpr_default.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        rpr_default.append(rpr)
    lang_el = rpr.find(qn('w:lang'))
    if lang_el is None:
        lang_el = OxmlElement('w:lang')
        rpr.append(lang_el)
    lang_el.set(qn('w:val'),     lang)
    lang_el.set(qn('w:eastAsia'), lang)
    lang_el.set(qn('w:bidi'),    lang)


def _aplicar_fuente_global(doc: Document, nombre: str = _FUENTE):
    """Setea la fuente por defecto del documento en el estilo Normal y
    en los estilos heredados (tablas/encabezados). Espejo del PDF/Excel
    que también usan Inter (Word cae al fallback del sistema si el
    cliente no la tiene instalada — recomendar instalar Inter desde
    fonts.google.com para máxima fidelidad)."""
    for style_name in ('Normal', 'Table Grid', 'Header', 'Footer'):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        try:
            style.font.name = nombre
        except AttributeError:
            pass
        rpr = style.element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.append(rFonts)
        rFonts.set(qn('w:ascii'),    nombre)
        rFonts.set(qn('w:hAnsi'),    nombre)
        rFonts.set(qn('w:cs'),       nombre)
        rFonts.set(qn('w:eastAsia'), nombre)


def _cell_bg(cell, rgb_hex: str):
    """Aplica background color a una celda (no expuesto en python-docx)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex.replace('#', ''))
    tc_pr.append(shd)


def _cell_border(cell, edge: str, sz: int = 6, color: str = 'CBD5E1'):
    """Borde de celda. edge: 'top','bottom','left','right'. sz en 1/8 pt."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(sz))
    border.set(qn('w:color'), color)
    tc_borders.append(border)


def _set_table_compact(tbl, *, top_dxa: int = 20, bottom_dxa: int = 20):
    """Reduce los márgenes verticales de TODAS las celdas de una tabla.
    Valores en twips (1/20 pt). 40 dxa ≈ 2 pt — espejo del padding 4pt del PDF.
    """
    tbl_pr = tbl._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tbl_pr)
    tcm = tbl_pr.find(qn('w:tblCellMar'))
    if tcm is None:
        tcm = OxmlElement('w:tblCellMar')
        tbl_pr.append(tcm)
    for edge, val in (('top', top_dxa), ('bottom', bottom_dxa)):
        e = tcm.find(qn(f'w:{edge}'))
        if e is None:
            e = OxmlElement(f'w:{edge}')
            tcm.append(e)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')


def _set_table_fixed_layout(tbl, col_widths_cm: list = None):
    """Fuerza `<w:tblLayout w:type="fixed"/>` y opcionalmente reescribe el
    `<w:tblGrid>` con los anchos exactos (cm) para que LibreOffice/TextMaker
    respeten `cell.width`. Sin esto los viewers (excepto OnlyOffice/Word)
    distribuyen las columnas por igual ignorando los anchos seteados.

    `col_widths_cm`: lista de anchos en cm (uno por columna). Si se pasa,
    también reescribe `<w:tblGrid>` (más confiable que solo cell.width).
    """
    from docx.shared import Cm as _Cm
    tbl_pr = tbl._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tbl_pr)
    # 1) Forzar layout fijo
    tblLayout = tbl_pr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tbl_pr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')
    # 2) Reescribir tblGrid si se pasaron anchos
    if col_widths_cm:
        widths_twips = [int(_Cm(w).twips) for w in col_widths_cm]
        tblGrid = tbl._element.find(qn('w:tblGrid'))
        if tblGrid is not None:
            tbl._element.remove(tblGrid)
        tblGrid = OxmlElement('w:tblGrid')
        for tw in widths_twips:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(tw))
            tblGrid.append(gc)
        # tblGrid va inmediatamente después de tblPr
        tbl_pr.addnext(tblGrid)


def _para_bottom_border(para, *, sz: int = 16, color: str = 'F37329'):
    """Borde inferior de párrafo (sz en 1/8 pt). Espejo del border-bottom CSS."""
    p_pr = para._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn('w:pBdr'))
    if p_bdr is None:
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)


def _set_cell_text(cell, texto: str, *, bold: bool = False,
                   size: int = 10, color: RGBColor = _SLATE_700,
                   align: str = 'left', italic: bool = False):
    """Limpia y setea texto formateado en una celda. Sin space_after para
    que las filas de tabla queden compactas (espejo del PDF)."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(texto))
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def _add_heading(doc: Document, texto: str, *, size: int = 14,
                 color: RGBColor = _SLATE_900, bold: bool = True,
                 space_before: int = 14, space_after: int = 6):
    """Encabezado de sección con espaciado consistente."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(texto)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _add_header_marca(doc: Document, proy: dict, titulo: str):
    """Encabezado de página: 3 columnas (empresa | título+proyecto | costo+modalidad).
    Espejo del `_draw_header` del PDF. Ancho ADAPTATIVO al tamaño/orientación
    de la sección — sin esto el header queda fijo en 17cm (útil A4 portrait)
    y se ve angosto en docs landscape (Curva S) o A3.
    """
    section = doc.sections[0]
    header = section.header
    # Limpiar el primer párrafo default
    if header.paragraphs:
        p0 = header.paragraphs[0]
        p0.text = ''
        # Reducir el espacio del párrafo placeholder.
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run('')
        r0.font.size = Pt(1)

    # Ancho útil = page_width - left_margin - right_margin (EMU → cm).
    ancho_util_emu = (section.page_width
                       - section.left_margin - section.right_margin)
    ancho_util_cm = ancho_util_emu / 360000  # 1 cm = 360000 EMU
    # Proporciones 4.5 / 8 / 4.5 = 17 (referencia A4 portrait), escaladas
    # al ancho real. Conservan el balance Marco-validado: centro ancho
    # para el nombre del proyecto, izq/der iguales.
    L_W = ancho_util_cm * 4.5 / 17.0
    C_W = ancho_util_cm * 8.0 / 17.0
    R_W = ancho_util_cm * 4.5 / 17.0

    tbl = header.add_table(rows=1, cols=3, width=Cm(ancho_util_cm))
    tbl.autofit = False
    _set_table_fixed_layout(tbl, [L_W, C_W, R_W])
    cell_l, cell_c, cell_r = tbl.cell(0, 0), tbl.cell(0, 1), tbl.cell(0, 2)
    cell_l.width = Cm(L_W)
    cell_c.width = Cm(C_W)
    cell_r.width = Cm(R_W)

    # Izquierda: empresa + subtítulo — desde la configuración (igual que el PDF).
    # Empresa color: `rep_color_marca_dk` custom, fallback al `od` de
    # `accent_reportes()` (slate-900 en sobrio, naranja-700 en marca).
    from core.pdf_reports import get_formato
    _fmt = get_formato()
    _empresa   = _fmt.get('rep_empresa_nombre')    or 'IngePresupuestos'
    _subtitulo = _fmt.get('rep_empresa_subtitulo') or 'Sistema de Presupuestos de Obra Pública'
    _dk_hex = (_fmt.get('rep_color_marca_dk') or '').lstrip('#')
    if not _dk_hex:
        from utils.theme import accent_reportes
        _dk_hex = accent_reportes()[1].lstrip('#')
    _dk_rgb = RGBColor(int(_dk_hex[0:2], 16), int(_dk_hex[2:4], 16), int(_dk_hex[4:6], 16))

    cell_l.text = ''
    p_emp = cell_l.paragraphs[0]
    p_emp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_emp.paragraph_format.space_after = Pt(0)
    r_emp = p_emp.add_run(_empresa)
    r_emp.font.size = Pt(11)
    r_emp.font.color.rgb = _dk_rgb
    r_emp.bold = True

    p_sub = cell_l.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)
    r_sub = p_sub.add_run(_subtitulo)
    r_sub.font.size = Pt(7)
    r_sub.font.color.rgb = _SLATE_300

    # Centro: título reporte + nombre proyecto
    cell_c.text = ''
    p_tit = cell_c.paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.space_after = Pt(0)
    r_tit = p_tit.add_run(titulo)
    r_tit.font.size = Pt(10)
    r_tit.font.color.rgb = _SLATE_900
    r_tit.bold = True

    nombre = (proy.get('nombre') or '').strip()
    p_proy = cell_c.add_paragraph()
    p_proy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_proy.paragraph_format.space_before = Pt(0)
    p_proy.paragraph_format.space_after = Pt(0)
    r_proy = p_proy.add_run(nombre)
    r_proy.font.size = Pt(7)
    r_proy.font.color.rgb = _SLATE_500
    r_proy.italic = True

    # Derecha: costo al + modalidad
    cell_r.text = ''
    p_costo = cell_r.paragraphs[0]
    p_costo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_costo.paragraph_format.space_after = Pt(0)
    r_costo = p_costo.add_run(f"Costo al: {_clean_costo_al(proy.get('costo_al'))}")
    r_costo.font.size = Pt(8)
    r_costo.font.color.rgb = _SLATE_700
    r_costo.bold = True

    p_mod = cell_r.add_paragraph()
    p_mod.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_mod.paragraph_format.space_before = Pt(0)
    p_mod.paragraph_format.space_after = Pt(0)
    r_mod = p_mod.add_run((proy.get('modalidad') or '').strip())
    r_mod.font.size = Pt(7)
    r_mod.font.color.rgb = _SLATE_300
    r_mod.italic = True

    # Línea separadora muy sutil debajo de la tabla del header
    p_sep = header.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after = Pt(0)
    _para_bottom_border(p_sep, sz=4, color='CBD5E1')


def _add_footer(doc: Document, proy: dict):
    """Pie tripartito espejo del `_draw_footer` del PDF:
    IZQ Cliente · CENTRO fecha del reporte · DER Página N de M.

    Implementado con tabla 1×3 (en lugar de tab stops) — más confiable en
    landscape y otros tamaños de página: los tab stops a posiciones grandes
    no siempre se renderizan correctamente en LibreOffice/Word.
    """
    from datetime import datetime as _dt_f
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    from docx.enum.table import WD_ALIGN_VERTICAL as _VA
    section = doc.sections[0]
    footer = section.footer

    # Limpiar párrafo default
    if footer.paragraphs:
        p0 = footer.paragraphs[0]
        p0.text = ''
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after  = Pt(0)
        r0 = p0.add_run('')
        r0.font.size = Pt(1)

    # Ancho útil dinámico (igual que `_add_header_marca`).
    ancho_util_cm = (section.page_width
                       - section.left_margin
                       - section.right_margin) / 360000
    L_W = C_W = R_W = ancho_util_cm / 3.0

    tbl = footer.add_table(rows=1, cols=3, width=Cm(ancho_util_cm))
    tbl.autofit = False
    _set_table_fixed_layout(tbl, [L_W, C_W, R_W])
    cell_l, cell_c, cell_r = tbl.cell(0, 0), tbl.cell(0, 1), tbl.cell(0, 2)
    cell_l.width = Cm(L_W); cell_c.width = Cm(C_W); cell_r.width = Cm(R_W)

    cliente = (proy.get('cliente') or '').strip()
    fecha   = _dt_f.now().strftime('%d/%m/%Y')

    def _set_cell(cell, text, *, align: str = 'left', run_setup=None):
        cell.text = ''
        cell.vertical_alignment = _VA.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        if align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if run_setup is not None:
            run_setup(p)
        else:
            r = p.add_run(text)
            r.font.size = Pt(7)
            r.font.color.rgb = _SLATE_500
        return p

    # IZQ: Cliente
    _set_cell(cell_l, f"Cliente: {cliente}" if cliente else "", align='left')
    # CENTRO: fecha
    p_c = _set_cell(cell_c, fecha, align='center')
    # DER: Página X de N — campos PAGE / NUMPAGES
    def _campo(run, instr_text: str):
        rEl = run._r
        for kind in ('begin', None, 'end'):
            if kind in ('begin', 'end'):
                fc = _OE('w:fldChar')
                fc.set(_qn('w:fldCharType'), kind)
                rEl.append(fc)
            else:
                it = _OE('w:instrText')
                it.set(_qn('xml:space'), 'preserve')
                it.text = instr_text
                rEl.append(it)
    def _runs_pagina(p_der):
        r_pre = p_der.add_run("Página ")
        r_pre.font.size = Pt(7); r_pre.font.color.rgb = _SLATE_500
        r_pg = p_der.add_run()
        r_pg.font.size = Pt(7); r_pg.font.color.rgb = _SLATE_500
        _campo(r_pg, ' PAGE ')
        r_de = p_der.add_run(' de ')
        r_de.font.size = Pt(7); r_de.font.color.rgb = _SLATE_500
        r_n = p_der.add_run()
        r_n.font.size = Pt(7); r_n.font.color.rgb = _SLATE_500
        _campo(r_n, ' NUMPAGES ')
    _set_cell(cell_r, "", align='right', run_setup=_runs_pagina)

    # Línea horizontal sutil arriba de la tabla del pie (espejo del PDF
    # `drawLine(SLATE_100)`). En Word es un border-top 0.5pt slate-100
    # en el párrafo que precede a la tabla (= p0 default del footer).
    if footer.paragraphs:
        p_top = footer.paragraphs[0]
        pPr = p_top._p.get_or_add_pPr()
        pBdr = _OE('w:pBdr')
        bd_bot = _OE('w:bottom')
        bd_bot.set(_qn('w:val'),   'single')
        bd_bot.set(_qn('w:sz'),    '4')          # 0.5pt
        bd_bot.set(_qn('w:space'), '2')
        bd_bot.set(_qn('w:color'), 'CBD5E1')     # slate-100
        pBdr.append(bd_bot)
        pPr.append(pBdr)


# ── Resumen Ejecutivo ─────────────────────────────────────────────────────

def generar_word_resumen_ejecutivo(pid: int, archivo: str) -> str:
    """Genera el Resumen Ejecutivo en formato Word (.docx).

    Estructura espejo del PDF:
      1. Header de marca
      2. Tabla de datos del proyecto
      3. KPIs (4 cuadros)
      4. Estructura del Presupuesto (títulos nivel 1-2)
      5. Pie de Presupuesto + monto en letras
      6. Distribución del Costo Directo (MO/MAT/EQ/SC)
      7. Top 5 Partidas por Monto

    Retorna la ruta del archivo generado.
    """
    conn = get_db()
    proy = dict(conn.execute(
        "SELECT * FROM proyectos WHERE id=?", (pid,)
    ).fetchone() or {})
    conn.close()
    if not proy:
        raise ValueError(f"Proyecto {pid} no encontrado")
    items, totales = calcular_totales(pid)
    sym = _moneda_simbolo(proy.get('moneda') or 'Soles')
    dec = get_decimales_ppto()

    doc = Document()
    _aplicar_fuente_global(doc)
    _aplicar_idioma(doc)
    # Márgenes ajustados
    for section in doc.sections:
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    _add_header_marca(doc, proy, "Resumen Ejecutivo")
    _add_footer(doc, proy)

    # Título principal — espejo del <h2> del PDF (izquierda, 13pt). El color
    # sigue el toggle modo sobrio: en sobrio = slate-900 (negro), en marca =
    # naranja-700. El border-bottom del CSS lo ignora QTextDocument en el PDF,
    # así que tampoco lo dibujamos en Word.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Resumen Ejecutivo")
    r.font.size = Pt(13)
    r.font.color.rgb = _accent_od_rgb()
    r.bold = True

    # KPIs en 4 columnas (igual que el PDF)
    cd = totales.get('cd', 0) or 0
    total = totales.get('total', 0) or 0
    hojas = [e for e in items if not e['partida'].get('es_titulo')]
    n_partidas = len(hojas)
    n_titulos  = len(items) - n_partidas
    pct_cd = (cd / total * 100) if total else 0
    kpis = [
        ("PRESUPUESTO TOTAL", f"{sym} {_fmt(total, dec)}", "Todos los conceptos"),
        ("COSTO DIRECTO",     f"{sym} {_fmt(cd, dec)}",    f"{pct_cd:.1f}% del total"),
        ("PARTIDAS",          f"{n_partidas}",             f"{n_titulos} títulos"),
        ("PLAZO DE OBRA",     f"{proy.get('plazo') or 0} d",
                              proy.get('modalidad') or 'Contrata'),
    ]
    tbl_kpi = doc.add_table(rows=1, cols=4)
    tbl_kpi.autofit = False
    for col, (label, value, sub) in enumerate(kpis):
        cell = tbl_kpi.cell(0, col)
        cell.width = Cm(4.0)
        cell.text = ''
        p_lbl = cell.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(2)
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(8)
        r_lbl.font.color.rgb = _SLATE_500
        r_lbl.bold = True

        p_val = cell.add_paragraph()
        p_val.paragraph_format.space_after = Pt(2)
        r_val = p_val.add_run(value)
        r_val.font.size = Pt(14)
        r_val.font.color.rgb = _SLATE_900
        r_val.bold = True

        p_sub = cell.add_paragraph()
        r_sub = p_sub.add_run(sub)
        r_sub.font.size = Pt(8)
        r_sub.font.color.rgb = _SLATE_500
        r_sub.italic = True
        # Sin marco/sin barra — puramente tipográficas (espejo del PDF).

    # Estructura del Presupuesto (títulos nivel 1 y 2)
    estructura = [e for e in items
                  if e['partida'].get('es_titulo')
                  and int(e['partida'].get('nivel') or 1) <= 2]
    if estructura:
        _add_heading(doc, "Estructura del Presupuesto",
                     size=11, color=_SLATE_700,
                     space_before=18, space_after=4)
        tbl_e = doc.add_table(rows=len(estructura), cols=3)
        tbl_e.autofit = False
        _set_table_compact(tbl_e)
        # Anchos: Item 1.8cm, Descripción 11.7cm, Monto 3.5cm (A4 portrait útil ≈ 17cm)
        _set_table_fixed_layout(tbl_e, [1.8, 11.7, 3.5])
        for row in tbl_e.rows:
            row.cells[0].width = Cm(1.8)
            row.cells[1].width = Cm(11.7)
            row.cells[2].width = Cm(3.5)
        for i, e in enumerate(estructura):
            p_e = e['partida']
            niv = int(p_e.get('nivel') or 1)
            bold = (niv == 1)
            color = _SLATE_900 if bold else _SLATE_500
            # Tabs por nivel en la DESCRIPCIÓN (no en el item); espejo del PDF.
            indent_prefix = "" if niv == 1 else "    " * (niv - 1)
            row = tbl_e.rows[i]
            _set_cell_text(row.cells[0], p_e.get('item') or '',
                           bold=bold, color=color)
            _set_cell_text(row.cells[1], indent_prefix + (p_e.get('descripcion') or ''),
                           bold=bold, color=color)
            _set_cell_text(row.cells[2],
                           f"{sym} {_fmt(e['total'] or 0, dec)}",
                           bold=bold, color=color, align='right')
            # Gridline muy sutil (espejo del border-bottom 1px SLATE_100 del PDF).
            for c in row.cells:
                _cell_border(c, 'bottom', sz=2, color='F1F5F9')

    # Pie de Presupuesto (sin título, igual que el PDF). Separador 18pt
    # (espejo de margin-top:36pt en el PDF, pero compensando el space_before
    # default de los párrafos).
    pie_rows = _build_pie_rows(pid, cd)
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(18)
    p_sep.paragraph_format.space_after = Pt(0)
    tbl_p = doc.add_table(rows=len(pie_rows), cols=2)
    tbl_p.autofit = False
    _set_table_compact(tbl_p)
    # Anchos espejo de la tabla "Estructura del Presupuesto":
    # label = 1.8 + 11.7 = 13.5cm (Item + Descripción de Estructura)
    # monto = 3.5cm (Monto de Estructura). Total 17cm = ancho útil A4.
    _set_table_fixed_layout(tbl_p, [13.5, 3.5])
    for row in tbl_p.rows:
        row.cells[0].width = Cm(13.5)
        row.cells[1].width = Cm(3.5)
    for i, (label, val, cls) in enumerate(pie_rows):
        row = tbl_p.rows[i]
        if cls == 'gran':
            bold = True; sz = 11; color = _SLATE_900
        elif cls == 'sub':
            bold = True; sz = 10; color = _SLATE_900
        else:
            bold = False; sz = 10; color = _SLATE_700
        _set_cell_text(row.cells[0], label, bold=bold, size=sz, color=color)
        _set_cell_text(row.cells[1], f"{sym} {_fmt(val, dec)}",
                       bold=bold, size=sz, color=color, align='right')
        # Bordes según jerarquía — espejo del PDF
        if cls == 'gran':
            for c in row.cells:
                _cell_border(c, 'top',    sz=16, color='273445')
                _cell_border(c, 'bottom', sz=16, color='273445')
        elif cls == 'sub':
            for c in row.cells:
                _cell_border(c, 'top',    sz=4, color='CBD5E1')
                _cell_border(c, 'bottom', sz=4, color='CBD5E1')
        else:
            for c in row.cells:
                _cell_border(c, 'bottom', sz=2, color='F1F5F9')

    # Monto en letras
    letras = _monto_en_letras(total, proy.get('moneda') or 'Soles')
    p_letras = doc.add_paragraph()
    p_letras.paragraph_format.space_before = Pt(8)
    r_l = p_letras.add_run(f"Son: {letras}.")
    r_l.font.size = Pt(9)
    r_l.font.color.rgb = _SLATE_700
    r_l.italic = True

    # Distribución del Costo Directo — página nueva, donut centrado arriba +
    # leyenda full-width abajo (espejo del PDF).
    doc.add_page_break()
    _add_heading(doc, "Distribución del Costo Directo",
                 size=11, color=_SLATE_700,
                 space_before=0, space_after=4)
    from core.database import get_insumos_proyecto
    from io import BytesIO
    conn = get_db()
    insumos = get_insumos_proyecto(conn, pid)
    conn.close()
    tipo_total = {'MO': 0.0, 'MAT': 0.0, 'EQ': 0.0, 'SC': 0.0}
    for r_ins in insumos:
        t = r_ins.get('tipo') if r_ins.get('tipo') in tipo_total else 'MAT'
        tipo_total[t] += float(r_ins.get('parcial_total') or 0)
    leyenda = [
        ('Mano de Obra',   tipo_total['MO'],  MO_COLOR),
        ('Materiales',     tipo_total['MAT'], MAT_COLOR),
        ('Equipos',        tipo_total['EQ'],  EQ_COLOR),
    ]
    if tipo_total['SC'] > 0:
        leyenda.append(('Sub-contratos', tipo_total['SC'], SC_COLOR))
    tot_dist = sum(v for _, v, _ in leyenda) or 1

    # Donut centrado arriba
    png_bytes = _png_donut_bytes(
        tipo_total['MO'], tipo_total['MAT'],
        tipo_total['EQ'], tipo_total['SC'], size=360,
    )
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(12)
    p_img.add_run().add_picture(BytesIO(png_bytes), width=Inches(3.0))

    # Leyenda full-width: tabla 4 col (color, label, monto, %)
    tbl_leg = doc.add_table(rows=len(leyenda), cols=4)
    tbl_leg.autofit = False
    for i, (label, val, color_hex) in enumerate(leyenda):
        pct = val / tot_dist * 100
        row = tbl_leg.rows[i]
        row.cells[0].width = Cm(0.7)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(6.0)
        row.cells[3].width = Cm(2.0)
        # Cuadrito de color (fondo)
        _cell_bg(row.cells[0], color_hex.lstrip('#'))
        row.cells[0].text = ' '
        # Label, monto, pct
        _set_cell_text(row.cells[1], label, bold=True, size=10)
        _set_cell_text(row.cells[2], f"{sym} {_fmt(val, dec)}",
                       align='right', size=10)
        _set_cell_text(row.cells[3], f"{pct:.1f}%", align='right',
                       size=10, color=_SLATE_500)
        for c in row.cells:
            _cell_border(c, 'bottom', sz=2, color='CBD5E1')

    # Top 5 Partidas
    hojas.sort(key=lambda e: e['total'] or 0, reverse=True)
    top5 = hojas[:5]
    if top5:
        _add_heading(doc, "Top 5 Partidas por Monto",
                     size=11, color=_SLATE_700,
                     space_before=18, space_after=4)
        tbl_t = doc.add_table(rows=len(top5) + 1, cols=5)
        tbl_t.autofit = False
        # Anchos: # 1.0 · Ítem 1.8 · Descripción 8.7 · Parcial 3.5 · % CD 2.0
        # → total 17cm = ancho útil A4 (espejo de Estructura/Pie).
        # Descripción concentra el espacio porque suele ser la columna larga.
        _set_table_fixed_layout(tbl_t, [1.0, 1.8, 8.7, 3.5, 2.0])
        col_widths = [Cm(1.0), Cm(1.8), Cm(8.7), Cm(3.5), Cm(2.0)]
        headers = ['#', 'Ítem', 'Descripción', 'Parcial', '% CD']
        for j, h in enumerate(headers):
            cell = tbl_t.cell(0, j)
            cell.width = col_widths[j]
            _set_cell_text(cell, h, bold=True, size=9,
                           color=_SLATE_500,
                           align=('right' if j in (3, 4) else 'left'))
            _cell_border(cell, 'bottom', sz=8, color='273445')
        for i, e in enumerate(top5):
            p_e = e['partida']
            pct = (e['total'] or 0) / cd * 100 if cd else 0
            row = tbl_t.rows[i + 1]
            for j in range(5):
                row.cells[j].width = col_widths[j]
            _set_cell_text(row.cells[0], f"#{i + 1}", bold=True,
                           color=_SLATE_900, align='center')
            _set_cell_text(row.cells[1], p_e.get('item') or '')
            _set_cell_text(row.cells[2], p_e.get('descripcion') or '')
            _set_cell_text(row.cells[3], f"{sym} {_fmt(e['total'] or 0, dec)}",
                           align='right')
            _set_cell_text(row.cells[4], f"{pct:.1f}%", align='right',
                           color=_SLATE_500)
            for c in row.cells:
                _cell_border(c, 'bottom', sz=2, color='CBD5E1')

    doc.save(archivo)
    return archivo


# ── Dispatcher ─────────────────────────────────────────────────────────────

# ── Especificaciones Técnicas ─────────────────────────────────────────────

class _HtmlToDocx:
    """Parser ligero HTML → python-docx. Convierte el HTML del editor de
    especificaciones (rich text) en párrafos/runs/imágenes equivalentes.

    Soporta:
      - <p>, <div> con align="..." y style="text-align:..."
      - <h1>/<h2>/<h3>/<h4> como párrafos bold de tamaño mayor
      - <b>/<strong>, <i>/<em>, <u>, <br>
      - <img src="data:image/...;base64,..."> con ancho opcional
      - <ul>/<ol>/<li> como list paragraphs
    Tags desconocidos se ignoran pero su contenido textual se preserva."""

    def __init__(self, doc, *, line_spacing: float | None = None):
        from html.parser import HTMLParser
        self.doc = doc
        self.current_para = None
        self.fmt_stack = []  # cada elem: {'bold','italic','underline'}
        self.heading_level = None
        self.line_spacing = line_spacing
        self._parser = type('_P', (HTMLParser,), {
            'handle_starttag': lambda s, t, a: self.on_start(t, a),
            'handle_endtag':   lambda s, t: self.on_end(t),
            'handle_data':     lambda s, d: self.on_data(d),
        })(convert_charrefs=True)

    def feed(self, html: str):
        self._parser.feed(html)
        self._parser.close()

    def _apply_line_spacing(self, para):
        if self.line_spacing is not None and para is not None:
            para.paragraph_format.line_spacing = self.line_spacing

    def _ensure_para(self):
        if self.current_para is None:
            self.current_para = self.doc.add_paragraph()
            self.current_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._apply_line_spacing(self.current_para)

    def on_start(self, tag, attrs):
        attrs_d = dict(attrs)
        tag = tag.lower()
        if tag in ('p', 'div'):
            self.current_para = self.doc.add_paragraph()
            align = (attrs_d.get('align') or '').lower()
            style = (attrs_d.get('style') or '').replace(' ', '')
            if align == 'center' or 'text-align:center' in style:
                self.current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right' or 'text-align:right' in style:
                self.current_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                self.current_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._apply_line_spacing(self.current_para)
        elif tag in ('h1', 'h2', 'h3', 'h4'):
            self.current_para = self.doc.add_paragraph()
            self.current_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.heading_level = int(tag[1])
            self._apply_line_spacing(self.current_para)
        elif tag in ('b', 'strong'):
            base = dict(self.fmt_stack[-1]) if self.fmt_stack else {}
            base['bold'] = True
            self.fmt_stack.append(base)
        elif tag in ('i', 'em'):
            base = dict(self.fmt_stack[-1]) if self.fmt_stack else {}
            base['italic'] = True
            self.fmt_stack.append(base)
        elif tag == 'u':
            base = dict(self.fmt_stack[-1]) if self.fmt_stack else {}
            base['underline'] = True
            self.fmt_stack.append(base)
        elif tag == 'span':
            # <span style="font-weight:700; font-style:italic; ...">
            # Algunos editores (incluido Qt rich text) expresan el formato
            # vía CSS en <span> en vez de <b>/<i>/<u>.
            style = (attrs_d.get('style') or '').lower().replace(' ', '')
            base = dict(self.fmt_stack[-1]) if self.fmt_stack else {}
            import re
            mw = re.search(r'font-weight:([0-9]+|bold|bolder)', style)
            if mw:
                w = mw.group(1)
                if w in ('bold', 'bolder') or (w.isdigit() and int(w) >= 600):
                    base['bold'] = True
            if 'font-style:italic' in style or 'font-style:oblique' in style:
                base['italic'] = True
            if 'text-decoration:underline' in style:
                base['underline'] = True
            self.fmt_stack.append(base)
        elif tag == 'br':
            if self.current_para is not None:
                self.current_para.add_run().add_break()
        elif tag == 'li':
            self.current_para = self.doc.add_paragraph(style='List Bullet')
            self.current_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._apply_line_spacing(self.current_para)
        elif tag == 'img':
            self._insert_image(attrs_d)

    def on_end(self, tag):
        tag = tag.lower()
        if tag in ('p', 'div', 'li'):
            self.current_para = None
        elif tag in ('h1', 'h2', 'h3', 'h4'):
            self.current_para = None
            self.heading_level = None
        elif tag in ('b', 'strong', 'i', 'em', 'u', 'span'):
            if self.fmt_stack:
                self.fmt_stack.pop()

    def on_data(self, text):
        if not text:
            return
        if not text.strip() and self.current_para is None:
            return  # whitespace entre tags, ignorar
        self._ensure_para()
        run = self.current_para.add_run(text)
        if self.heading_level:
            run.bold = True
            run.font.size = Pt({1: 14, 2: 12, 3: 11, 4: 11}.get(self.heading_level, 11))
        fmt = self.fmt_stack[-1] if self.fmt_stack else {}
        if fmt.get('bold'):      run.bold = True
        if fmt.get('italic'):    run.italic = True
        if fmt.get('underline'): run.underline = True

    def _insert_image(self, attrs_d):
        src = attrs_d.get('src', '')
        if not src.startswith('data:'):
            return
        import base64
        from io import BytesIO
        try:
            _, b64 = src.split(',', 1)
            img_bytes = base64.b64decode(b64)
        except Exception:
            return
        self._ensure_para()
        run = self.current_para.add_run()
        # Ancho: max 380px → ~3.95 pulgadas a 96dpi
        try:
            w_px = int(attrs_d.get('width', '0'))
        except ValueError:
            w_px = 0
        w_px = min(w_px or 380, 380)
        try:
            run.add_picture(BytesIO(img_bytes), width=Inches(w_px / 96))
        except Exception:
            pass


def _limpiar_html_spec(html: str) -> str:
    """Prepara el HTML del editor antes del parser docx:
    - Elimina `<head>`, `<style>`, `<script>` (Qt rich text editor inserta
      su CSS interno en `<head><style>` que el parser veía como texto).
    - Extrae solo el contenido de `<body>` si existe.
    - Reemplaza guiones de viñeta `- ` al inicio de párrafo por bullet `•`.
    """
    import re
    s = re.sub(r'<head[\s>][\s\S]*?</head>', '', html, flags=re.IGNORECASE)
    s = re.sub(r'<style[\s>][\s\S]*?</style>', '', s, flags=re.IGNORECASE)
    s = re.sub(r'<script[\s>][\s\S]*?</script>', '', s, flags=re.IGNORECASE)
    m = re.search(r'<body[^>]*>([\s\S]*?)</body>', s, re.IGNORECASE)
    if m:
        s = m.group(1)
    # Viñetas de guion → bullet (mismo criterio que el PDF)
    s = re.sub(r'(<(?:p|div|li)[^>]*>)\s*[-–]\s+', r'\1•  ', s, flags=re.IGNORECASE)
    return s


def generar_word_especificaciones(pid: int, archivo: str) -> str:
    """Genera el reporte de Especificaciones Técnicas en .docx editable."""
    from core.database import get_db, calcular_totales
    conn = get_db()
    proy = dict(conn.execute(
        "SELECT * FROM proyectos WHERE id=?", (pid,)
    ).fetchone() or {})
    conn.close()
    if not proy:
        raise ValueError(f"Proyecto {pid} no encontrado")
    items, _ = calcular_totales(pid)

    doc = Document()
    _aplicar_fuente_global(doc)
    _aplicar_idioma(doc)
    for section in doc.sections:
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    _add_header_marca(doc, proy, "Especificaciones Técnicas")
    _add_footer(doc, proy)

    # Título principal — espejo del <h2> del PDF (izquierda, 13pt, color
    # `od` de accent_reportes: slate-900 en sobrio, naranja-700 en marca).
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Especificaciones Técnicas")
    r.font.size = Pt(13)
    r.font.color.rgb = _accent_od_rgb()
    r.bold = True

    od_rgb = _accent_od_rgb()
    o_hex  = _accent_o_hex()

    cnt = 0
    for entry in items:
        partida = entry['partida']
        spec = (partida.get('especificaciones') or '').strip()
        if not spec:
            continue
        cnt += 1

        # Encabezado de partida — espejo del PDF (_chunk_especificacion_partida):
        #   • N1 título  → SLATE_900, 12pt, border-bottom 2pt SLATE_700
        #   • N2+ título → SLATE_700, 11pt, border-bottom 1pt SLATE_500
        #   • Partida    → accent `od`, 11pt, border-bottom 1.5pt accent `o`
        is_titulo = bool(partida.get('es_titulo'))
        nivel = partida.get('nivel') or 1
        if is_titulo and nivel == 1:
            color, size, b_sz, b_color = _SLATE_900, 12, 16, '273445'   # 2pt SLATE_700
        elif is_titulo:
            color, size, b_sz, b_color = _SLATE_700, 11,  8, '485A6C'   # 1pt SLATE_500
        else:
            color, size, b_sz, b_color = od_rgb,    11, 12, o_hex        # 1.5pt accent
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(20)
        p_head.paragraph_format.space_after = Pt(6)
        run_h = p_head.add_run(f"{partida.get('item') or ''}   {partida.get('descripcion') or ''}")
        run_h.bold = True
        run_h.font.size = Pt(size)
        run_h.font.color.rgb = color
        _set_paragraph_border(p_head, 'bottom', sz=b_sz, color=b_color)

        # Cuerpo: convertir HTML del editor o texto plano. line-height 1.5
        # espejo del CSS del PDF.
        if '<' in spec and '>' in spec:
            try:
                limpio = _limpiar_html_spec(spec)
                parser = _HtmlToDocx(doc, line_spacing=1.5)
                parser.feed(limpio)
            except Exception:
                p_body = doc.add_paragraph(spec)
                p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_body.paragraph_format.line_spacing = 1.5
        else:
            for parrafo in spec.split('\n'):
                if not parrafo.strip():
                    continue
                p_body = doc.add_paragraph(parrafo)
                p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_body.paragraph_format.line_spacing = 1.5

    if cnt == 0:
        p = doc.add_paragraph("No hay especificaciones técnicas registradas en este proyecto.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r_p in p.runs:
            r_p.italic = True
            r_p.font.color.rgb = _SLATE_500

    doc.save(archivo)
    return archivo


def _set_paragraph_border(para, edge: str, sz: int = 6, color: str = 'CBD5E1'):
    """Borde de párrafo. edge: 'top','bottom','left','right'. sz en 1/8 pt."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(sz))
    border.set(qn('w:color'), color)
    border.set(qn('w:space'), '4')
    pBdr.append(border)


def generar_word_memoria_descriptiva(pid: int, archivo: str) -> str:
    """Genera el reporte de Memoria Descriptiva en .docx editable (mismos
    encabezados/pie y tratamiento de viñetas que Especificaciones Técnicas)."""
    from core.database import get_db
    conn = get_db()
    proy = dict(conn.execute(
        "SELECT * FROM proyectos WHERE id=?", (pid,)
    ).fetchone() or {})
    conn.close()
    if not proy:
        raise ValueError(f"Proyecto {pid} no encontrado")
    # El título «Memoria Descriptiva» lo pone el encabezado + el título del
    # cuerpo; quitamos el que la IA incluye al inicio del contenido para que
    # no salga duplicado.
    memo = strip_titulo_memoria((proy.get('memoria_descriptiva') or '').strip())

    doc = Document()
    _aplicar_fuente_global(doc)
    _aplicar_idioma(doc)
    for section in doc.sections:
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    _add_header_marca(doc, proy, "Memoria Descriptiva")
    _add_footer(doc, proy)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("MEMORIA DESCRIPTIVA")
    r.font.size = Pt(18)
    r.font.color.rgb = _accent_od_rgb()
    r.bold = True

    if not memo:
        pe = doc.add_paragraph("Este proyecto aún no tiene memoria descriptiva.")
        pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r_p in pe.runs:
            r_p.italic = True
            r_p.font.color.rgb = _SLATE_500
    elif '<' in memo and '>' in memo:
        try:
            limpio = _limpiar_html_spec(memo)
            parser = _HtmlToDocx(doc, line_spacing=1.5)
            parser.feed(limpio)
        except Exception:
            p_body = doc.add_paragraph(memo)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_body.paragraph_format.line_spacing = 1.5
    else:
        for parrafo in memo.split('\n'):
            if not parrafo.strip():
                continue
            p_body = doc.add_paragraph(parrafo)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_body.paragraph_format.line_spacing = 1.5

    doc.save(archivo)
    return archivo


_GENERADORES = {
    'resumen':              generar_word_resumen_ejecutivo,
    'especificaciones':     generar_word_especificaciones,
    'memoria_descriptiva':  generar_word_memoria_descriptiva,
}


def tipos_soportados() -> set[str]:
    """Tipos de reporte con export Word disponible."""
    return set(_GENERADORES.keys())


def _th(cell, txt, *, align='right'):
    """Celda de cabecera de tabla (fondo acento suave, texto oscuro en negrita)."""
    from utils.theme import accent_reportes
    _set_cell_text(cell, txt, bold=True, size=8.5, color=_SLATE_700, align=align)
    _cell_bg(cell, accent_reportes()[2].lstrip('#'))


def generar_word_almacen(pid: int, archivo: str) -> str:
    """Control de materiales (kárdex) en .docx editable."""
    import core.requerimientos as REQ
    from core.database import get_decimales_metrado
    conn = get_db()
    proy = dict(conn.execute("SELECT * FROM proyectos WHERE id=?",
                             (pid,)).fetchone() or {})
    conn.close()
    if not proy:
        raise ValueError(f"Proyecto {pid} no encontrado")
    filas = REQ.control_almacen(pid)
    dm = get_decimales_metrado()
    doc = Document(); _aplicar_fuente_global(doc); _aplicar_idioma(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.0)
    _add_header_marca(doc, proy, "Control de materiales")
    _add_footer(doc, proy)
    _add_heading(doc, "Control de materiales — Kárdex de obra", size=13,
                 color=_accent_od_rgb())
    cols = ["Insumo", "Und", "Pedido", "Ingresado", "Consumido", "Stock",
            "Por llegar"]

    def _tabla_resumen(grupo):
        tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = 'Table Grid'
        for j, c in enumerate(cols):
            _th(tbl.cell(0, j), c, align='left' if j == 0 else 'center' if j == 1
                else 'right')
        for f in grupo:
            r = tbl.add_row().cells
            neg = f['stock'] < -1e-6
            _set_cell_text(r[0], f['descripcion'] or '', size=8.5)
            _set_cell_text(r[1], f['unidad'] or '', size=8.5, align='center')
            vals = [f['pedido'], f['ingresado'], f['consumido'], f['stock'],
                    f['por_llegar']]
            for k, v in enumerate(vals):
                col = RGBColor(0xB7, 0x1C, 0x1C) if (neg and k == 3) else _SLATE_700
                _set_cell_text(r[2 + k], _fmt(v, dm), size=8.5, align='right',
                               color=col, bold=(k == 3))
        _set_table_fixed_layout(tbl, [5.0, 1.6, 2.5, 2.5, 2.5, 2.5, 2.5])

    con_req = [f for f in filas if (f['pedido'] or 0) > 1e-6]
    sin_req = [f for f in filas if (f['pedido'] or 0) <= 1e-6]
    if con_req:
        _add_heading(doc, "Pedidos en requerimientos", size=11,
                     color=_accent_od_rgb(), space_before=8, space_after=3)
        _tabla_resumen(con_req)
    if sin_req:
        _add_heading(doc, "Otros — sin requerimiento", size=11,
                     color=_accent_od_rgb(), space_before=10, space_after=3)
        _tabla_resumen(sin_req)
    if not con_req and not sin_req:
        _tabla_resumen([])

    # ── Kárdex por material (solo los que tuvieron movimiento) ───────────────
    import core.almacen as _alm
    con_mov = [f for f in filas
               if (f['ingresado'] or 0) > 1e-6 or (f['consumido'] or 0) > 1e-6]
    if con_mov:
        def _dmy(iso):
            p = str(iso or '').split('-')
            return f"{p[2]}/{p[1]}/{p[0]}" if len(p) == 3 else str(iso or '')
        doc.add_page_break()
        _add_heading(doc, "Kárdex de almacén por material", size=13,
                     color=_accent_od_rgb())
        for f in con_mov:
            ph = doc.add_paragraph()
            ph.paragraph_format.space_before = Pt(8)
            ph.paragraph_format.space_after = Pt(2)
            r1 = ph.add_run("MATERIAL:  "); r1.bold = True; r1.font.size = Pt(9.5)
            r1.font.color.rgb = _SLATE_700
            r2 = ph.add_run(f['descripcion'] or ''); r2.font.size = Pt(9.5)
            r2.font.color.rgb = _SLATE_900
            r3 = ph.add_run("      UNIDAD:  "); r3.bold = True; r3.font.size = Pt(9.5)
            r3.font.color.rgb = _SLATE_700
            r4 = ph.add_run(f['unidad'] or ''); r4.font.size = Pt(9.5)
            r4.font.color.rgb = _SLATE_900
            kcols = ["FECHA", "ENTRADAS", "SALIDAS", "SALDOS", "OBSERVACIONES"]
            kt = doc.add_table(rows=1, cols=len(kcols)); kt.style = 'Table Grid'
            for j, c in enumerate(kcols):
                _th(kt.cell(0, j), c, align='center')
            movs = _alm.kardex(pid, f['recurso_id'])
            for mv in movs:
                rc = kt.add_row().cells
                _set_cell_text(rc[0], _dmy(mv['fecha']), size=8.5, align='center')
                _set_cell_text(rc[1], _fmt(mv['entrada'], dm) if mv['entrada'] else '',
                               size=8.5, align='right')
                _set_cell_text(rc[2], _fmt(mv['salida'], dm) if mv['salida'] else '',
                               size=8.5, align='right')
                _set_cell_text(rc[3], _fmt(mv['saldo'], dm), size=8.5, align='right')
                _set_cell_text(rc[4], mv['observacion'] or '', size=8.5)
            for _ in range(max(0, 6 - len(movs))):   # filas en blanco (anotar a mano)
                rc = kt.add_row().cells
                for j in range(len(kcols)):
                    _set_cell_text(rc[j], '', size=8.5)
            _set_table_fixed_layout(kt, [3.2, 2.9, 2.9, 2.9, 5.1])
        # Firmas.
        fp = doc.add_paragraph(); fp.paragraph_format.space_before = Pt(24)
        ft = doc.add_table(rows=2, cols=3)
        for j, lab in enumerate(("RESIDENTE", "SUPERVISOR", "ALMACENERO DE OBRA")):
            _set_cell_text(ft.cell(0, j), "______________________", size=9,
                           align='center', color=RGBColor(0x55, 0x55, 0x55))
            _set_cell_text(ft.cell(1, j), lab, size=9, align='center',
                           color=RGBColor(0x55, 0x55, 0x55))
        _set_table_fixed_layout(ft, [5.6, 5.6, 5.6])

    doc.save(archivo)
    return archivo


def generar_word_curva_s(pid: int, archivo: str, *, base: str = 'mes_cal',
                         vis: dict = None) -> str:
    """Curva S real (gráfico + tabla por período) en .docx editable (landscape).
    `vis` muestra/oculta series y columnas (espejo de los toggles del panel)."""
    import base64
    from io import BytesIO
    from docx.enum.section import WD_ORIENT
    import core.curva_s as CS
    from core.pdf_reports import _png_curva_s_real_b64
    if vis is None:
        vis = {'prog': True, 'reprog': True, 'real': True, 'pct': True}
    conn = get_db()
    proy = dict(conn.execute("SELECT * FROM proyectos WHERE id=?",
                             (pid,)).fetchone() or {})
    conn.close()
    if not proy:
        raise ValueError(f"Proyecto {pid} no encontrado")
    d = CS.curva_s_comparada(pid, base)
    dp = get_decimales_ppto()

    def mo(v):
        return _fmt(v, dp) if v is not None else ""

    def pc(v):
        return f"{v:.1f}%" if v is not None else ""

    doc = Document(); _aplicar_fuente_global(doc); _aplicar_idioma(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Cm(1.5)
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    _add_header_marca(doc, proy, "Curva S — avance real")
    _add_footer(doc, proy)
    # Título centrado (espejo del PDF).
    tp0 = doc.add_paragraph(); tp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr0 = tp0.add_run("Curva S — programado vs reprogramado vs real")
    tr0.bold = True; tr0.font.size = Pt(13); tr0.font.color.rgb = _accent_od_rgb()

    # Gráfico (PNG) — respeta los toggles.
    png = base64.b64decode(_png_curva_s_real_b64(d, vis=vis))
    pimg = doc.add_paragraph(); pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pimg.add_run().add_picture(BytesIO(png), width=Cm(24))

    tp = {'p_mon': 0.0, 'r_mon': 0.0, 'x_mon': 0.0}
    ult = {'p_acu': None, 'r_acu': None, 'x_acu': None}
    # Grupos activos (según toggles); «Desv.» acompaña a «Real».
    GROUPS = [('prog', 'Programado', ('p_mon', 'p_eje', 'p_acu')),
              ('reprog', 'Reprogramado', ('r_mon', 'r_eje', 'r_acu')),
              ('real', 'Real', ('x_mon', 'x_eje', 'x_acu'))]
    active = [g for g in GROUPS if vis.get(g[0], True)]
    show_desv = vis.get('real', True)
    ncols = 1 + 3 * len(active) + (1 if show_desv else 0)

    # ── Cabecera de 2 niveles ────────────────────────────────────────────────
    tbl = doc.add_table(rows=2, cols=ncols); tbl.style = 'Table Grid'
    cper = tbl.cell(0, 0).merge(tbl.cell(1, 0))
    _th(cper, "Período", align='left')
    col = 1
    for _k, glbl, _keys in active:
        g = tbl.cell(0, col).merge(tbl.cell(0, col + 2))
        _th(g, glbl, align='center')
        for i, s in enumerate(('Monto', '% Ejecutado', '% Acumulado')):
            _th(tbl.cell(1, col + i), s, align='right')
        col += 3
    if show_desv:
        _th(tbl.cell(0, col).merge(tbl.cell(1, col)), "Desv.", align='right')

    for f in d.get('filas', []):
        r = tbl.add_row().cells
        for k in tp:
            if f.get(k) is not None:
                tp[k] += f[k]
        for k in ult:
            if f.get(k) is not None:
                ult[k] = f[k]
        _set_cell_text(r[0], f['label'], size=8)
        col = 1
        for _k, _glbl, keys in active:
            km, ke, ka = keys
            for i, txt in enumerate((mo(f[km]), pc(f[ke]), pc(f[ka]))):
                _set_cell_text(r[col + i], txt, size=8, align='right')
            col += 3
        if show_desv:
            dv = f.get('desviacion')
            dvs = f"{dv:+.1f}%" if dv is not None else ""
            dcol = (RGBColor(0xB7, 0x1C, 0x1C) if (dv is not None and dv < -1e-6)
                    else RGBColor(0x16, 0xA3, 0x4A) if dv is not None else _SLATE_700)
            _set_cell_text(r[col], dvs, size=8, align='right', color=dcol, bold=True)
    # Fila TOTAL (negro, sin rojo).
    r = tbl.add_row().cells
    _set_cell_text(r[0], "TOTAL", size=8, bold=True, color=_SLATE_900)
    col = 1
    for _k, _glbl, keys in active:
        km, ke, ka = keys
        tmon = {'p_mon': tp['p_mon'], 'r_mon': tp['r_mon'], 'x_mon': tp['x_mon']}[km]
        tacu = {'p_acu': ult['p_acu'], 'r_acu': ult['r_acu'], 'x_acu': ult['x_acu']}[ka]
        for i, txt in enumerate((mo(tmon), "", pc(tacu))):
            _set_cell_text(r[col + i], txt, size=8, bold=True, color=_SLATE_900)
        col += 3
    if show_desv:
        _set_cell_text(r[col], "", size=8, bold=True)
    for c in r:
        _cell_bg(c, 'F0F1F2')

    wdesc = [3.0]
    for _g in active:
        wdesc += [2.7, 2.0, 2.0]
    if show_desv:
        wdesc += [2.0]
    _set_table_fixed_layout(tbl, wdesc)
    doc.save(archivo)
    return archivo


def generar_word_cuaderno(pid: int, parte_ids: list, archivo: str) -> str:
    """Cuaderno de obra (partes diarios seleccionados) en .docx editable."""
    import core.parte_diario as PD
    from core.database import get_decimales_metrado
    from core.pdf_reports import _fecha_larga_es, _cant, _fmt as _f
    conn = get_db()
    proy = dict(conn.execute("SELECT * FROM proyectos WHERE id=?",
                             (pid,)).fetchone() or {})
    conn.close()
    if not proy:
        raise ValueError(f"Proyecto {pid} no encontrado")
    dm = get_decimales_metrado()
    doc = Document(); _aplicar_fuente_global(doc); _aplicar_idioma(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.0)
    _add_header_marca(doc, proy, "Cuaderno de obra")
    _add_footer(doc, proy)
    # Título centrado (espejo del PDF).
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)
    tr = tp.add_run("Cuaderno de obra"); tr.bold = True; tr.font.size = Pt(15)
    tr.font.color.rgb = _accent_od_rgb()

    SZ = 10   # tamaño único para todo el cuerpo del asiento

    def _linea(etq, txt):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        rb = p.add_run("•  "); rb.bold = True; rb.font.size = Pt(SZ)
        rb.font.color.rgb = _SLATE_700
        r = p.add_run(etq + " "); r.bold = True; r.font.size = Pt(SZ)
        r.font.color.rgb = _SLATE_700
        if txt:
            r2 = p.add_run(txt); r2.font.size = Pt(SZ); r2.font.color.rgb = _SLATE_900
        return p

    for pidx in parte_ids:
        parte = PD.get_parte(pidx)
        if not parte:
            continue
        mo = PD.get_recursos_dia(pidx, 'mo')
        acts = PD.get_actividades_dia(pidx)
        recursos = {c: PD.get_recursos_dia(pidx, c) for c in ('mat', 'eq', 'sc')}
        obs = (parte.get('observaciones') or '').strip()
        # Omitir días SIN nada registrado (no ensuciar con asientos vacíos).
        if not (mo or acts or any(recursos.values()) or obs):
            continue
        _add_heading(doc, f"Asiento del día — {_fecha_larga_es(parte['fecha'])}",
                     size=11, color=_accent_od_rgb(), space_before=18, space_after=4)
        if mo:
            _linea("Mano de obra:", " · ".join(
                f"{_cant(x['cantidad'])} {x['descripcion']}".strip() for x in mo))
        if acts:
            _linea("Actividades realizadas:", "")
            tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Table Grid'
            for j, c in enumerate(["Ítem", "Actividad realizada", "Metrado", "Und"]):
                _th(tbl.cell(0, j), c, align='left' if j in (0, 1) else
                    'right' if j == 2 else 'center')
            for a in acts:
                r = tbl.add_row().cells
                _set_cell_text(r[0], a['item'] or '', size=SZ)
                _set_cell_text(r[1], a['descripcion'] or '', size=SZ)
                _set_cell_text(r[2], _f(a['metrado_dia'], dm), size=SZ, align='right')
                _set_cell_text(r[3], a['unidad'] or '', size=SZ, align='center')
            _set_table_fixed_layout(tbl, [2.5, 10.0, 2.5, 2.0])
        for clase, tit in (('mat', 'Materiales'), ('eq', 'Equipos'),
                           ('sc', 'Servicios')):
            rs = recursos.get(clase)
            if rs:
                _linea(f"{tit}:", " · ".join(
                    " ".join(s for s in (x['descripcion'], _f(x['cantidad'], dm),
                             x['unidad']) if s).strip() for x in rs))
        if obs:
            _linea("Observaciones:", "")
            for ln in obs.split('\n'):
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(0.5)
                rr = p.add_run(ln); rr.font.size = Pt(SZ); rr.font.color.rgb = _SLATE_900
    doc.save(archivo)
    return archivo


def _para_bottom_border(p, color: str = '273445', sz: int = 8):
    """Añade un filete inferior a un párrafo (espejo del `border-bottom` del
    encabezado de sección en el PDF)."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def generar_word_tdr(req_id: int, archivo: str) -> str:
    """Requerimiento + TDR / Especificaciones Técnicas en .docx editable.
    Misma tipografía/formato que las Especificaciones (fuente proporcional,
    justificado, viñetas, encabezados) renderizando desde `parse_tdr_bloques`."""
    import json
    import core.requerimientos as REQ
    from core.pdf_reports import parse_tdr_bloques, _proyecto_info
    q = REQ.get_requerimiento(req_id)
    if not q:
        raise ValueError("Requerimiento no encontrado")
    proy = _proyecto_info(q['proyecto_id'])
    try:
        datos = json.loads(q.get('tdr_datos') or '{}')
    except (ValueError, TypeError):
        datos = {}

    od = _accent_od_rgb()
    from core.pdf_reports import _brand_colors, _es_celda_numerica
    o_hexs, _od_hex, _os = _brand_colors()
    o_hex = o_hexs.lstrip('#')

    doc = Document(); _aplicar_fuente_global(doc); _aplicar_idioma(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2.0)
    # Encabezado de página: solo «Requerimiento» (el número va en el título del
    # cuerpo).
    _add_header_marca(doc, proy, "Requerimiento")
    _add_footer(doc, proy)

    def g(k):
        v = datos.get(k)
        return v.strip() if isinstance(v, str) else (v or '')

    # ── Encabezado del documento (memorando público / membrete privado) ─────
    # El número SIEMPRE es el real del requerimiento (no el de tdr_datos, que
    # podía haber quedado viejo).
    numero = str(q['numero'])
    asunto = g('asunto'); fecha = g('fecha'); lugar = g('lugar')
    fl = f"{lugar}, {fecha}".strip(', ') if (lugar or fecha) else ''

    def _titulo_req():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"REQUERIMIENTO N° {numero}")
        r.bold = True; r.font.size = Pt(15); r.font.color.rgb = od

    def _campo(label, value):
        if not value:
            return
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        rl = p.add_run(f"{label}:  "); rl.bold = True
        rl.font.size = Pt(9.5); rl.font.color.rgb = _SLATE_700
        rv = p.add_run(value); rv.font.size = Pt(9.5); rv.font.color.rgb = _SLATE_900

    if datos:
        if g('es_publico'):
            _titulo_req()
            dest = g('destinatario') + (f" — {g('cargo_destinatario')}"
                                        if g('cargo_destinatario') else '')
            deq = g('solicitante') + (f" — {g('cargo_solicitante')}"
                                      if g('cargo_solicitante') else '')
            ent = g('entidad') + (f" / {g('unidad_organica')}"
                                  if g('unidad_organica') else '')
            _campo('A', dest); _campo('ATENCIÓN', g('atencion'))
            _campo('DE', deq); _campo('ENTIDAD', ent)
            _campo('ASUNTO', asunto); _campo('FECHA', fl)
        else:
            if g('solicitante'):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(g('solicitante')); r.bold = True; r.font.size = Pt(12)
                r.font.color.rgb = _SLATE_900
            sub = '  ·  '.join(x for x in (g('ruc') and f"RUC: {g('ruc')}",
                                           g('direccion'), g('telefono')) if x)
            if sub:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(sub); r.font.size = Pt(9); r.font.color.rgb = _SLATE_700
            _titulo_req()
            _campo('ASUNTO', asunto); _campo('FECHA', fl)
        # Divisor bajo el encabezado.
        pd = doc.add_paragraph(); pd.paragraph_format.space_before = Pt(2)
        pd.paragraph_format.space_after = Pt(6)
        _para_bottom_border(pd, color=o_hex, sz=8)

    # ── Cuerpo del TDR (desde el parser compartido) ─────────────────────────
    def _para(texto, *, justify=True, space_after=4, space_before=0):
        p = doc.add_paragraph()
        p.alignment = (WD_ALIGN_PARAGRAPH.JUSTIFY if justify
                       else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.4
        return p

    for blk in parse_tdr_bloques(q.get('tdr') or ''):
        kind = blk[0]
        if kind == 'blank':
            continue
        if kind == 'table':
            filas = blk[1]
            ncols = max(len(f) for f in filas)
            tbl = doc.add_table(rows=0, cols=ncols); tbl.style = 'Table Grid'
            for ri, cells in enumerate(filas):
                cells = list(cells) + [''] * (ncols - len(cells))
                row = tbl.add_row().cells
                for ci, c in enumerate(cells):
                    if ri == 0:
                        _th(row[ci], c, align='center')
                    else:
                        num = _es_celda_numerica(c)
                        _set_cell_text(row[ci], c, size=8.5,
                                       align='right' if num else 'left')
            # Anchos: ítem angosto, DESCRIPCIÓN (col 1) ancha, resto medios —
            # evita el word-wrap feo de la descripción.
            total_cm = 17.0
            if ncols >= 3:
                w = [1.2] + [0.0] * (ncols - 1)
                for j in range(2, ncols):
                    w[j] = 2.3
                w[1] = max(3.0, total_cm - w[0] - sum(w[2:]))
            else:
                w = [total_cm * 0.18, total_cm * 0.82]
            _set_table_fixed_layout(tbl, w)
        elif kind == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.4
            r = p.add_run(blk[1]); r.font.size = Pt(10); r.font.color.rgb = _SLATE_900
        elif kind == 'heading_num':
            depth = blk[2]
            p = _para(None, justify=False,
                      space_before=(13 if depth == 0 else 7), space_after=3)
            r = p.add_run(blk[1]); r.bold = True
            r.font.size = Pt(11 if depth == 0 else 10)
            r.font.color.rgb = od if depth == 0 else _SLATE_700
        elif kind == 'heading_major':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(blk[1]); r.bold = True; r.font.size = Pt(13)
            r.font.color.rgb = od
            _para_bottom_border(p, color=o_hex, sz=8)
        elif kind == 'sublabel':
            p = _para(None, justify=False, space_before=6, space_after=3)
            r = p.add_run(blk[1]); r.bold = True; r.font.size = Pt(9.5)
            r.font.color.rgb = _SLATE_700
        elif kind == 'label':
            p = _para(None, space_after=3)
            rl = p.add_run(f"{blk[1]}:  "); rl.bold = True
            rl.font.size = Pt(10); rl.font.color.rgb = _SLATE_900
            rv = p.add_run(blk[2]); rv.font.size = Pt(10); rv.font.color.rgb = _SLATE_900
        else:   # 'para'
            p = _para(None, space_after=4)
            r = p.add_run(blk[1]); r.font.size = Pt(10); r.font.color.rgb = _SLATE_900

    doc.save(archivo)
    return archivo


def generar_word(tipo: str, pid: int, archivo: str) -> str:
    """Genera un reporte Word. `tipo` es el código del reporte
    (resumen, presupuesto, acus, …). Por ahora solo `resumen` está
    soportado; los demás se irán agregando."""
    fn = _GENERADORES.get(tipo)
    if fn is None:
        raise NotImplementedError(
            f"Word export para tipo «{tipo}» aún no implementado. "
            "Usa PDF para este reporte mientras tanto."
        )
    return fn(pid, archivo)
