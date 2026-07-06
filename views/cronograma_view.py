# SPDX-License-Identifier: GPL-3.0-or-later
# Copyright (C) 2026 Marco Sumari / Sumari SAC
# This file is part of IngePresupuestos — https://ingepresupuestos.com
# Licensed under the GNU GPL v3.0 or later. See the LICENSE file.
"""Vista de Cronograma — Diagrama Gantt + Valorizado + Curva S + Insumos.

Equivalente a templates/cronograma.html del Flask, adaptado a PySide6.
"""
import json as _json
from datetime import datetime, timedelta

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QFrame, QLabel, QPushButton,
    QScrollArea, QSplitter, QTableWidget, QTableWidgetItem, QHeaderView,
    QAbstractItemView, QComboBox, QSpinBox, QGraphicsView, QGraphicsScene,
    QStackedWidget, QSizePolicy, QGraphicsRectItem, QGraphicsLineItem,
    QGraphicsTextItem, QGraphicsItem, QGraphicsPolygonItem, QGraphicsPathItem,
    QGraphicsEllipseItem, QMenu, QFileDialog, QMessageBox, QDialog, QRadioButton,
    QDialogButtonBox, QButtonGroup, QStyledItemDelegate, QStyle,
    QStyleOptionViewItem, QInputDialog, QSlider,
)
from PySide6.QtCore import Qt, QTimer, QRectF, QPointF, QSize, QMarginsF, Signal
from PySide6.QtGui import (
    QFont, QColor, QBrush, QPen, QPainter, QPolygonF, QPainterPath, QPageLayout,
    QPainterPathStroker, QTransform, QLinearGradient,
)
from PySide6.QtPrintSupport import QPrinter

from core.database import (
    get_db, calcular_totales, parcial_wysiwyg, get_decimales_ppto,
    get_decimales_metrado,
    get_acu_items,
)
from core.cronograma import (
    cpm, get_cronograma_map, distribuir_periodos,
    auto_programar, auto_programar_local, auto_programar_ia,
    calcular_duraciones_desde_metrado, cargar_feriados_ia,
)
from core.pdf_reports import get_formato as _get_formato_reportes
from utils.formatting import fmt


# ── Paleta consistente con proyecto_view.py ────────────────────────────────
SLATE_700 = "#273445"
SLATE_500 = "#485A6C"
SLATE_300 = "#667885"
SLATE_100 = "#95A3AB"
SILVER_100 = "#F8F9FA"
SILVER_300 = "#D4D4D4"
BLUE_500 = "#F37329"
BLUE_700 = "#C0621A"
RED_500 = "#C6262E"
GREEN_500 = "#68B723"

# Colores barras gantt — tarea no crítica en arándano (Blueberry, elementary OS)
GANTT_BAR = "#3689E6"        # Blueberry 500 — relleno
GANTT_BAR_TOP = "#64BAFF"    # Blueberry 300 — degradado superior
GANTT_BAR_BORDER = "#0D52BF" # Blueberry 700 — borde
GANTT_CRIT = "#C6262E"
GANTT_CRIT_TOP = "#ED5353"
GANTT_HITO = "#7A36B1"
GANTT_GRID = "#E8EAED"


class CronogramaView(QWidget):
    """Vista completa de Cronograma — usada como page 2 del root_stack del proyecto."""

    def __init__(self, pid: int, proyecto: dict, on_back, parent=None,
                 on_editar=None):
        super().__init__(parent)
        self.pid       = pid
        self._proy     = proyecto
        self._moneda   = proyecto.get('moneda', 'Soles')
        self._on_back  = on_back
        self._on_editar = on_editar   # abre «Editar proyecto» (para cambiar el plazo)
        self._build_ui()

    # ──────────────────────────────────────────────────────────────────────
    # UI
    # ──────────────────────────────────────────────────────────────────────

    def _build_ui(self):
        vl = QVBoxLayout(self)
        vl.setContentsMargins(0, 0, 0, 0)
        vl.setSpacing(0)

        # ── Barra única: ← Presupuesto · pestañas · plazo · acciones ─────────
        # Antes había una fila aparte (CRONOGRAMA DE OBRA / Plazo / Editar);
        # se fusionó con la fila de pestañas para ahorrar alto vertical — las
        # pestañas de PROYECTOS ya viven arriba, en la vista de proyecto.
        bar_frame = QFrame()
        bar_frame.setFixedHeight(36)
        bar_frame.setStyleSheet(f"background:{SLATE_500}; border:none;")
        bar_hl = QHBoxLayout(bar_frame)
        bar_hl.setContentsMargins(8, 4, 10, 4)
        bar_hl.setSpacing(6)

        # ← Volver al presupuesto
        btn_back = QPushButton("← Presupuesto")
        btn_back.setCursor(Qt.PointingHandCursor)
        btn_back.setStyleSheet(
            f"QPushButton {{ background:rgba(255,255,255,0.12); color:white;"
            f" border:1px solid rgba(255,255,255,0.25); border-radius:6px;"
            f" font-size:11px; padding:3px 10px; }}"
            f"QPushButton:hover {{ background:rgba(255,255,255,0.22); }}"
        )
        btn_back.clicked.connect(self._on_back)
        bar_hl.addWidget(btn_back)
        bar_hl.addSpacing(8)

        self._stack = QStackedWidget()
        self._stack.setStyleSheet(f"background:{SILVER_100};")

        self._tabs = ["Diagrama Gantt", "Cronograma Valorizado",
                       "Adquisición de Insumos", "Curva S"]
        self._tab_btns: list[QPushButton] = []

        def _tab_style(sel: bool) -> str:
            bg = BLUE_500 if sel else "transparent"
            hov = ("" if sel
                    else "QPushButton:hover { background:rgba(255,255,255,0.15); color:white; }")
            return (f"QPushButton {{ background:{bg}; color:white; border:none;"
                    f" border-radius:6px; font-size:11px; font-weight:700;"
                    f" padding:4px 14px; }}" + hov)

        def _select_tab(idx: int):
            self._stack.setCurrentIndex(idx)
            for i, b in enumerate(self._tab_btns):
                b.setStyleSheet(_tab_style(i == idx))

        for i, lbl in enumerate(self._tabs):
            btn = QPushButton(lbl)
            btn.setCursor(Qt.PointingHandCursor)
            btn.setStyleSheet(_tab_style(i == 0))
            btn.clicked.connect(lambda _, ix=i: _select_tab(ix))
            bar_hl.addWidget(btn)
            self._tab_btns.append(btn)

        bar_hl.addStretch()

        # Plazo contractual + enlace para editarlo (antes en la topbar propia).
        plazo = self._proy.get('plazo') or 0
        self._lbl_plazo = QLabel(f"Plazo: {plazo} días")
        self._lbl_plazo.setStyleSheet(
            "color:white; font-size:11px; font-weight:600;"
            " background:transparent; border:none;")
        _tip_plazo = ("Plazo = días contractuales del proyecto.\n"
                      "Programados = días que ocupa la programación del Gantt.\n"
                      "Pueden no coincidir. Para cambiar el plazo, ve a "
                      "«Editar proyecto».")
        self._lbl_plazo.setToolTip(_tip_plazo)
        bar_hl.addWidget(self._lbl_plazo)
        if self._on_editar is not None:
            btn_edit_plazo = QPushButton("✏ Editar plazo")
            btn_edit_plazo.setCursor(Qt.PointingHandCursor)
            btn_edit_plazo.setToolTip(_tip_plazo)
            btn_edit_plazo.setStyleSheet(
                "QPushButton { background:transparent; color:#E8EDF2;"
                " border:none; font-size:11px; text-decoration:underline;"
                " padding:0 2px; }"
                "QPushButton:hover { color:white; }")
            btn_edit_plazo.clicked.connect(lambda: self._on_editar())
            bar_hl.addWidget(btn_edit_plazo)
        bar_hl.addSpacing(10)

        # Botones globales (calcular, guardar)
        btn_style = (
            "QPushButton { background:rgba(255,255,255,0.18); color:white;"
            " border:1px solid rgba(255,255,255,0.35); border-radius:4px;"
            " font-size:11px; padding:3px 10px; min-height:0; }"
            "QPushButton:hover { background:rgba(255,255,255,0.30); }"
            "QPushButton::menu-indicator { width:8px; subcontrol-position:right center;"
            " subcontrol-origin:padding; right:4px; }"
        )
        # "Calcular duración", "Auto-programar" y "Guardar" SOLO aplican al
        # Diagrama Gantt — se ocultan en las otras pestañas (valorizado,
        # adquisiciones, curva S son vistas de solo lectura derivadas del Gantt).
        self._btn_dur = QPushButton("Calcular duración")
        self._btn_dur.setCursor(Qt.PointingHandCursor)
        self._btn_dur.setStyleSheet(btn_style)
        self._btn_dur.clicked.connect(self._calcular_duraciones)
        bar_hl.addWidget(self._btn_dur)

        # Auto-programar — botón con menú (3 modos)
        from PySide6.QtWidgets import QMenu
        self._btn_auto = QPushButton("Auto-programar ▾")
        self._btn_auto.setCursor(Qt.PointingHandCursor)
        self._btn_auto.setStyleSheet(btn_style)
        menu_auto = QMenu(self._btn_auto)
        menu_auto.addAction("Por fases (rápido, local)",
                            self._auto_programar_local)
        menu_auto.addAction("Con IA (más realista)",
                            self._auto_programar_ia)
        menu_auto.addSeparator()
        menu_auto.addAction("Secuencial (modo simple)",
                            self._auto_programar)
        self._btn_auto.setMenu(menu_auto)
        bar_hl.addWidget(self._btn_auto)

        self._btn_save = QPushButton("Guardar")
        self._btn_save.setCursor(Qt.PointingHandCursor)
        self._btn_save.setStyleSheet(btn_style)
        self._btn_save.clicked.connect(self._guardar_todo)
        bar_hl.addWidget(self._btn_save)

        vl.addWidget(bar_frame)

        # ── Stack con las 4 sub-vistas ────────────────────────────────────
        self._gantt_w     = GanttWidget(self)
        self._valorz_w    = ValorizadoWidget(self)
        self._curva_w     = CurvaSWidget(self)
        self._insumos_w   = InsumosWidget(self)

        for w in (self._gantt_w, self._valorz_w, self._insumos_w, self._curva_w):
            self._stack.addWidget(w)

        vl.addWidget(self._stack, stretch=1)

        # Cargar datos al cambiar de tab
        self._stack.currentChanged.connect(self._on_tab_changed)

    # ──────────────────────────────────────────────────────────────────────
    # Carga de datos
    # ──────────────────────────────────────────────────────────────────────

    def mostrar_tab(self, idx: int):
        """Activa la pestaña interna por índice — 0=Gantt, 1=Valorizado,
        2=Adquisiciones, 3=Curva S. Llamable desde fuera (e.g. desde el
        menu "Cronogramas ▾" del topbar del proyecto)."""
        if 0 <= idx < len(self._tab_btns):
            self._tab_btns[idx].click()

    def cargar(self):
        """Carga todos los datos del proyecto y refresca las vistas."""
        conn = get_db()
        # Orden AGRUPADO por subpresupuesto (Principal/NULL primero, luego cada
        # subpresupuesto por `orden`), dentro por item — espejo del árbol del
        # presupuesto y base de la numeración estilo MS Project.
        self._partidas = [dict(r) for r in conn.execute(
            """SELECT p.* FROM partidas p
               LEFT JOIN sub_presupuestos s ON s.id = p.sub_presupuesto_id
               WHERE p.proyecto_id=?
               ORDER BY (CASE WHEN p.sub_presupuesto_id IS NULL THEN 0 ELSE 1 END),
                        COALESCE(s.orden, 0), p.sub_presupuesto_id, p.item""",
            (self.pid,)
        ).fetchall()]
        # Mapa subpresupuesto_id → nombre (None = Principal).
        self._sub_nombres = {None: (self._proy.get('sub_presupuesto') or 'Principal')}
        for r in conn.execute(
            "SELECT id, nombre FROM sub_presupuestos WHERE proyecto_id=? ORDER BY orden, id",
            (self.pid,)
        ):
            self._sub_nombres[r['id']] = r['nombre'] or 'Subpresupuesto'
        self._cron_map = get_cronograma_map(conn, self.pid)
        conn.close()
        self._calcular_cpm()
        self._on_tab_changed(self._stack.currentIndex())

    def _calcular_cpm(self):
        plazo = self._proy.get('plazo') or 0
        non_working = self._non_working_set(plazo + 365)  # margen
        self._tasks = cpm(self._cron_map, self._partidas, plazo, non_working)
        # Actualizar plazo en topbar.
        #   plazo  = días contractuales (definidos en el proyecto)
        #   max_ef = días que ocupa la programación (mayor EF del CPM)
        # Son independientes: la programación puede terminar antes o después
        # del plazo contractual; NO se igualan.
        max_ef = max((t['EF'] for t in self._tasks.values() if t['EF'] > 0),
                      default=0)
        if plazo > 0:
            self._lbl_plazo.setText(
                f"Plazo: {plazo} días  ·  {max_ef} programados")
        else:
            self._lbl_plazo.setText(f"{max_ef} días programados")

    def _non_working_set(self, n_dias: int) -> set:
        """Conjunto de día corridos (1..n_dias) que son domingos o feriados,
        según fecha_inicio y feriados del proyecto. Vacío si la opción
        `salta_no_laborables` está deshabilitada."""
        if int(self._proy.get('salta_no_laborables', 1) or 0) == 0:
            return set()
        from datetime import timedelta
        try:
            fi = self._proy.get('fecha_inicio', '') or self._proy.get('costo_al', '')
            if not fi:
                return set()
            f_ini = datetime.strptime(fi, '%Y-%m-%d')
        except Exception:
            return set()
        # Feriados como set de YYYY-MM-DD
        feriados = set()
        raw = (self._proy.get('feriados') or '').strip()
        for tok in raw.replace(';', ',').split(','):
            tok = tok.strip()
            if tok:
                try:
                    datetime.strptime(tok, '%Y-%m-%d')
                    feriados.add(tok)
                except Exception:
                    pass
        out = set()
        for d in range(1, max(1, n_dias) + 1):
            fecha = f_ini + timedelta(days=d - 1)
            if fecha.weekday() == 6 or fecha.strftime('%Y-%m-%d') in feriados:
                out.add(d)
        return out

    def _proj_end(self) -> int:
        """Día corrido en que termina el proyecto (mayor EF, mínimo el plazo)."""
        plazo = self._proy.get('plazo') or 0
        max_ef = max((t['EF'] for t in self._tasks.values() if t['EF'] > 0),
                      default=plazo)
        return max(max_ef, plazo, 1)

    def filas_con_hitos(self) -> list:
        """Lista de VISUALIZACIÓN del Gantt estilo MS Project:
        resumen del proyecto (#1) + hito «Inicio de Obra» (#2) + cabecera de cada
        subpresupuesto con sus títulos/partidas + hito «Termino de Obra». Las
        filas virtuales (proyecto, inicio, cabeceras de subpresupuesto, fin) se
        DERIVAN del CPM (id centinela negativo, `_virtual`): NO se guardan en
        `partidas` ni entran al motor CPM. Su posición 1-based en esta lista ES
        su "#" (coincide con `numerar_filas`). Devuelve la lista cruda si no hay
        partidas."""
        partidas = self._partidas
        if not partidas:
            return list(partidas)
        from core.cronograma import (filas_slots, PROYECTO_PID, INICIO_PID,
                                       FIN_PID)
        tasks = self._tasks
        sub_nombres = getattr(self, '_sub_nombres', {None: 'Principal'})
        # Terminales: hojas que no son predecesoras de ninguna otra tarea.
        con_sucesor = set()
        for t in tasks.values():
            for pr in t.get('preds', []):
                con_sucesor.add(pr.get('pid'))
        terminales = [p['id'] for p in partidas
                      if not p['es_titulo']
                      and p['id'] not in con_sucesor
                      and tasks.get(p['id'], {}).get('EF', 0) > 0]
        sin_pred = [p['id'] for p in partidas
                    if not p['es_titulo']
                    and not tasks.get(p['id'], {}).get('preds')]
        proj_end = self._proj_end()
        by_id = {p['id']: p for p in partidas}
        # Span (min ES, max EF) de cada grupo de subpresupuesto.
        grupo_span = {}
        for p in partidas:
            if p['es_titulo']:
                continue
            t = tasks.get(p['id'], {})
            es, ef = t.get('ES', 0), t.get('EF', 0)
            if es > 0 and ef > 0:
                g = p.get('sub_presupuesto_id')
                cur = grupo_span.get(g)
                grupo_span[g] = ((es, ef) if cur is None
                                 else (min(cur[0], es), max(cur[1], ef)))
        filas = []
        sub_idx = 0
        for tipo, clave in filas_slots(partidas):
            if tipo == 'proyecto':
                filas.append({'id': PROYECTO_PID, '_virtual': 'proyecto', 'item': '',
                              'descripcion': self._proy.get('nombre', 'Proyecto') or 'Proyecto',
                              'es_titulo': 1, 'nivel': 0, 'es_hito': 0,
                              '_ES': 1, '_EF': proj_end, '_dur': proj_end})
            elif tipo == 'inicio':
                filas.append({'id': INICIO_PID, '_virtual': 'inicio', 'item': '',
                              'descripcion': 'Inicio de Obra', 'es_titulo': 0, 'nivel': 1,
                              'es_hito': 2, '_ES': 1, '_EF': 1, '_dur': 0,
                              '_succ_ids': sin_pred})
            elif tipo == 'fin':
                filas.append({'id': FIN_PID, '_virtual': 'fin', 'item': '',
                              'descripcion': 'Termino de Obra', 'es_titulo': 0, 'nivel': 1,
                              'es_hito': 3, '_ES': proj_end, '_EF': proj_end, '_dur': 0,
                              '_pred_ids': terminales})
            elif tipo == 'sub':
                a, b = grupo_span.get(clave, (1, proj_end))
                sub_idx += 1
                filas.append({'id': -(1000 + sub_idx), '_virtual': 'subppto', 'item': '',
                              'descripcion': sub_nombres.get(clave) or 'Subpresupuesto',
                              'es_titulo': 1, 'nivel': 0, 'es_hito': 0,
                              '_ES': a, '_EF': b, '_dur': b - a + 1})
            else:  # 'partida'
                filas.append(by_id[clave])
        return filas

    def _on_tab_changed(self, idx: int):
        # Mostrar/ocultar los botones globales sólo en la pestaña Gantt (idx==0).
        es_gantt = (idx == 0)
        for b in (self._btn_dur, self._btn_auto, self._btn_save):
            b.setVisible(es_gantt)
        if not hasattr(self, '_partidas'):
            return
        w = self._stack.widget(idx)
        if hasattr(w, 'cargar'):
            w.cargar()

    # ──────────────────────────────────────────────────────────────────────
    # Operaciones globales
    # ──────────────────────────────────────────────────────────────────────

    def _calcular_duraciones(self):
        """Calcula duración estimada desde metrado/rendimiento."""
        if not hasattr(self, '_partidas'):
            return
        durs = calcular_duraciones_desde_metrado(self._partidas)
        for pid, d in durs.items():
            if pid in self._cron_map:
                self._cron_map[pid]['duracion'] = d
            else:
                self._cron_map[pid] = {'duracion': d, 'inicio_dia': 1,
                                         'predecesoras': '', 'es_hito': 0,
                                         'segmentos': '', 'partida_id': pid}
        self._guardar_y_refrescar()

    def _aplicar_nuevo_map(self, new_map: dict):
        """Helper: aplica un mapa de cronograma (resultado de auto-programar)
        al `_cron_map` actual y persiste."""
        for pid, data in new_map.items():
            if pid in self._cron_map:
                self._cron_map[pid].update(data)
            else:
                data['partida_id'] = pid
                self._cron_map[pid] = data
        self._guardar_y_refrescar()

    def _ajustar_duraciones_a_plazo(self) -> tuple[int, int, float]:
        """Escala las duraciones de las partidas iterativamente para que la
        ruta crítica encaje dentro de `proyecto.plazo`. Si plazo<=0 no hace
        nada. Retorna (max_ef_antes, max_ef_despues, factor_aplicado).

        Estrategia: factor = plazo / max_ef. Aplicado a todas las partidas
        (mín 1 día). Itera hasta 6 veces porque los días no laborables
        deforman el resultado (skipping de domingos/feriados). No expande
        si max_ef < plazo — el usuario puede preferir holgura."""
        plazo = self._proy.get('plazo') or 0
        if plazo <= 0:
            return (0, 0, 1.0)

        non_working = self._non_working_set(plazo + 730)

        def _max_ef():
            tasks = cpm(self._cron_map, self._partidas, plazo, non_working)
            return max((t['EF'] for t in tasks.values() if t['EF'] > 0),
                        default=0)

        max_ef_pre = _max_ef()
        if max_ef_pre <= plazo or max_ef_pre <= 0:
            # Cabe; nada que ajustar
            return (max_ef_pre, max_ef_pre, 1.0)

        factor_acum = 1.0
        max_ef_post = max_ef_pre
        for _ in range(6):
            factor = plazo / max_ef_post
            factor_acum *= factor
            for pid, cd in self._cron_map.items():
                old_d = int(cd.get('duracion', 1) or 1)
                if old_d <= 1:
                    continue
                new_d = max(1, round(old_d * factor))
                cd['duracion'] = new_d
            max_ef_post = _max_ef()
            if abs(max_ef_post - plazo) <= 1:
                break

        return (max_ef_pre, max_ef_post, factor_acum)

    def _msg_resultado_autoprogramar(self, titulo: str, extra: str = ''):
        from PySide6.QtWidgets import QMessageBox
        plazo = self._proy.get('plazo') or 0
        pre, post, factor = self._ajustar_duraciones_a_plazo()
        if pre != post:
            # Persistir las duraciones ajustadas
            self._guardar_y_refrescar()
        if plazo <= 0:
            cuerpo = (extra + "\n\n⚠ El proyecto no tiene plazo configurado.\n"
                                "Define el plazo en datos del proyecto y vuelve "
                                "a ejecutar auto-programar para ajustar.")
        elif pre <= plazo:
            cuerpo = (extra + f"\n\nDuración programada: {post} días\n"
                                f"Plazo del proyecto: {plazo} días\n"
                                "✔ Encaja dentro del plazo (no se ajustó).")
        else:
            pct = round((1 - factor) * 100, 1)
            cuerpo = (extra + f"\n\nLa ruta crítica original ocupaba {pre} días, "
                                f"excedía el plazo de {plazo} días.\n"
                                f"Se redujeron las duraciones un {pct}% "
                                "para encajar.\n"
                                f"Duración final: {post} días.")
        QMessageBox.information(self, titulo, cuerpo)

    def _auto_programar(self):
        """Modo simple: predecesoras secuenciales (cada partida usa la anterior)."""
        if not hasattr(self, '_partidas'):
            return
        new_map = auto_programar(self._partidas, self._cron_map)
        self._aplicar_nuevo_map(new_map)
        self._msg_resultado_autoprogramar(
            "Auto-programar",
            "Cada partida usa la anterior como predecesora "
            "(modo secuencial simple).")

    def _auto_programar_local(self):
        """Modo local heurístico: agrupa partidas en fases constructivas
        (preliminares, mov.tierras, estructuras, instalaciones…) y permite
        paralelismo dentro de cada fase. Después ajusta duraciones para
        encajar en el plazo del proyecto."""
        if not hasattr(self, '_partidas'):
            return
        new_map = auto_programar_local(self._partidas, self._cron_map)
        self._aplicar_nuevo_map(new_map)
        self._msg_resultado_autoprogramar(
            "Auto-programar — por fases",
            "Partidas agrupadas por fases constructivas.\n"
            "Las partidas de la misma fase se ejecutan en paralelo; "
            "cada fase espera el final de la anterior.")

    def _auto_programar_ia(self):
        """Modo IA: el LLM analiza las descripciones y devuelve dependencias
        realistas. Si la IA falla, hace fallback al modo local. Después
        ajusta duraciones para encajar en el plazo del proyecto."""
        if not hasattr(self, '_partidas'):
            return
        from PySide6.QtCore import Qt
        from PySide6.QtWidgets import QApplication, QMessageBox

        QApplication.setOverrideCursor(Qt.WaitCursor)
        try:
            new_map, err = auto_programar_ia(self._partidas, self._cron_map)
        finally:
            QApplication.restoreOverrideCursor()

        if new_map is None:
            r = QMessageBox.question(self, "Auto-programar con IA",
                f"No se pudo usar la IA:\n\n{err}\n\n"
                "¿Quieres usar el modo local (por fases) en su lugar?",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            if r == QMessageBox.Yes:
                self._auto_programar_local()
            return
        self._aplicar_nuevo_map(new_map)
        self._msg_resultado_autoprogramar(
            "Auto-programar — con IA",
            "Cronograma generado por IA aplicado.\n"
            "Revisa las dependencias críticas y ajusta si es necesario.")

    def _guardar_todo(self):
        """Guarda el cronograma_map a la DB."""
        # Recoger cambios de la vista activa antes de guardar
        cur = self._stack.currentWidget()
        if hasattr(cur, 'volcar_a_map'):
            cur.volcar_a_map()
        self._guardar_a_db()
        self._calcular_cpm()
        self._on_tab_changed(self._stack.currentIndex())

    def _guardar_y_refrescar(self):
        self._guardar_a_db()
        self._calcular_cpm()
        self._on_tab_changed(self._stack.currentIndex())

    def _guardar_a_db(self):
        conn = get_db()
        for pid, d in self._cron_map.items():
            conn.execute(
                """INSERT INTO cronograma_partidas
                    (partida_id, duracion, inicio_dia, predecesoras, es_hito,
                     segmentos, color)
                   VALUES (?,?,?,?,?,?,?)
                   ON CONFLICT(partida_id) DO UPDATE SET
                    duracion=excluded.duracion,
                    inicio_dia=excluded.inicio_dia,
                    predecesoras=excluded.predecesoras,
                    es_hito=excluded.es_hito,
                    segmentos=excluded.segmentos,
                    color=excluded.color""",
                (pid, int(d.get('duracion', 1) or 1),
                 int(d.get('inicio_dia', 1) or 1),
                 d.get('predecesoras', '') or '',
                 int(d.get('es_hito', 0) or 0),
                 d.get('segmentos', '') or '',
                 d.get('color', '') or '')
            )
        conn.commit()
        conn.close()


# ════════════════════════════════════════════════════════════════════════════
# Diálogo de opciones para exportar Gantt a PDF
# ════════════════════════════════════════════════════════════════════════════

class _DialogExportarGanttPdf(QDialog):
    """Pregunta modo (una hoja / multipágina), orientación, tamaño de papel,
    escala del eje temporal y hojas horizontales. Ofrece Vista previa,
    Exportar PDF y Exportar imagen."""

    preview_solicitado = Signal(dict)
    imagen_solicitada  = Signal(dict)

    # Tamaños de papel disponibles (label → enum QPageSize.PageSizeId)
    PAPEL_OPCIONES = [
        ("A4 (210 × 297 mm)", "A4"),
        ("A3 (297 × 420 mm)", "A3"),
        ("A2 (420 × 594 mm)", "A2"),
        ("A1 (594 × 841 mm)", "A1"),
        ("A0 (841 × 1189 mm)", "A0"),
        ("Carta / Letter (216 × 279 mm)", "Letter"),
        ("Tabloide / Ledger (279 × 432 mm)", "Tabloid"),
    ]

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Exportar Gantt")
        self.setMinimumWidth(560)
        self.setMinimumHeight(720)
        self.resize(580, 740)
        self.setModal(True)
        # Forzar paleta clara (evita que en temas oscuros se herede el fondo
        # gris/negro del sistema en labels/radios/checks).
        from PySide6.QtGui import QPalette
        pal = self.palette()
        pal.setColor(QPalette.Window, QColor("white"))
        pal.setColor(QPalette.Base, QColor("white"))
        pal.setColor(QPalette.WindowText, QColor("#273445"))
        pal.setColor(QPalette.Text, QColor("#273445"))
        self.setPalette(pal)
        self.setAutoFillBackground(True)

        self.setStyleSheet(
            "QDialog { background:white; }"
            "QLabel { color:#273445; font-size:12px; background:transparent; }"
            "QRadioButton { color:#273445; font-size:12px; padding:4px 0;"
            "  background:transparent; spacing:8px; }"
            "QRadioButton::indicator { width:16px; height:16px;"
            "  background:transparent; border:none; }"
            "QRadioButton::indicator:unchecked {"
            "  image: url(resources/icons/radio_orange_off.svg); }"
            "QRadioButton::indicator:checked {"
            "  image: url(resources/icons/radio_orange_on.svg); }"
            "QCheckBox { color:#273445; font-size:12px; padding:2px 0;"
            "  background:transparent; spacing:8px; }"
            "QCheckBox::indicator { width:16px; height:16px;"
            "  background:transparent; border:none; }"
            "QCheckBox::indicator:unchecked {"
            "  image: url(resources/icons/check_orange_off.svg); }"
            "QCheckBox::indicator:checked {"
            "  image: url(resources/icons/check_orange_on.svg); }"
            "QGroupBox { font-weight:bold; color:#485A6C; margin-top:8px;"
            "  background:transparent; }"
        )

        vl = QVBoxLayout(self)
        vl.setContentsMargins(18, 16, 18, 14)
        vl.setSpacing(10)

        ttl = QLabel("Opciones de exportación")
        f = ttl.font(); f.setPointSize(13); f.setBold(True); ttl.setFont(f)
        ttl.setStyleSheet("color:#1E2635;")
        vl.addWidget(ttl)

        # Modo paginación
        vl.addWidget(QLabel("Paginación:"))
        self.rb_fit   = QRadioButton("Una hoja (ajustar todo el cronograma)")
        self.rb_multi = QRadioButton("Múltiples hojas (escala legible, recomendado)")
        self.rb_multi.setChecked(True)
        grp_pg = QButtonGroup(self)
        grp_pg.addButton(self.rb_fit)
        grp_pg.addButton(self.rb_multi)
        vl.addWidget(self.rb_fit)
        vl.addWidget(self.rb_multi)

        nota = QLabel(
            "En ambos casos la tabla de partidas se mantiene a la izquierda "
            "y el gantt a la derecha en cada hoja."
        )
        nota.setStyleSheet("color:#667885; font-size:10px;")
        nota.setWordWrap(True)
        vl.addWidget(nota)

        # Tamaño de papel + Orientación (en una fila)
        vl.addSpacing(4)
        hh_pap = QHBoxLayout()
        hh_pap.addWidget(QLabel("Tamaño de papel:"))
        self.cmb_papel = QComboBox()
        for label, _ in self.PAPEL_OPCIONES:
            self.cmb_papel.addItem(label)
        # Default A3 — el formato más usado por ingenieros para Gantt
        self.cmb_papel.setCurrentIndex(1)
        self.cmb_papel.setStyleSheet(
            "QComboBox { padding:3px 6px; border:1px solid #D4D4D4;"
            " border-radius:4px; background:white; min-width:200px; }"
        )
        hh_pap.addWidget(self.cmb_papel, 1)
        vl.addLayout(hh_pap)

        # Orientación
        vl.addWidget(QLabel("Orientación:"))
        self.rb_land = QRadioButton("Horizontal (recomendado)")
        self.rb_port = QRadioButton("Vertical")
        self.rb_land.setChecked(True)
        grp_or = QButtonGroup(self)
        grp_or.addButton(self.rb_land)
        grp_or.addButton(self.rb_port)
        hh = QHBoxLayout()
        hh.addWidget(self.rb_land)
        hh.addWidget(self.rb_port)
        hh.addStretch()
        vl.addLayout(hh)

        # Escala del eje temporal
        vl.addSpacing(4)
        hh2 = QHBoxLayout()
        hh2.addWidget(QLabel("Escala del eje temporal:"))
        self.cmb_escala = QComboBox()
        self.cmb_escala.addItems([
            "Automática (según ancho)",
            "Días — mostrar día y mes",
            "Semanas — S1, S2…",
            "Meses — sólo mes/año",
        ])
        self.cmb_escala.setCurrentIndex(0)
        self.cmb_escala.setStyleSheet(
            "QComboBox { padding:3px 6px; border:1px solid #D4D4D4;"
            " border-radius:4px; background:white; }"
        )
        hh2.addWidget(self.cmb_escala, 1)
        vl.addLayout(hh2)

        # Hojas en eje tiempo — solo aplica en modo multipágina
        from PySide6.QtWidgets import QSpinBox
        hh3 = QHBoxLayout()
        self.lbl_hojas = QLabel("Hojas en eje tiempo (multipágina):")
        self.sp_hojas = QSpinBox()
        self.sp_hojas.setRange(0, 24)
        self.sp_hojas.setValue(0)
        self.sp_hojas.setSpecialValueText("Automático")
        self.sp_hojas.setStyleSheet(
            "QSpinBox { padding:3px 6px; border:1px solid #D4D4D4;"
            " border-radius:4px; background:white; min-width:80px; }"
        )
        hh3.addWidget(self.lbl_hojas)
        hh3.addWidget(self.sp_hojas)
        hh3.addStretch()
        vl.addLayout(hh3)

        def _toggle_hojas():
            on = self.rb_multi.isChecked()
            self.sp_hojas.setEnabled(on)
            self.lbl_hojas.setEnabled(on)
        self.rb_multi.toggled.connect(_toggle_hojas)
        self.rb_fit.toggled.connect(_toggle_hojas)
        _toggle_hojas()

        # Incluir columna Pred.
        from PySide6.QtWidgets import QCheckBox
        self.chk_pred = QCheckBox("Incluir columna de Predecesoras")
        self.chk_pred.setChecked(True)   # importante en un Gantt → activa por defecto
        self.chk_pred.setStyleSheet("color:#273445; font-size:12px;")
        vl.addWidget(self.chk_pred)

        # Encabezado / Pie de página (compartidos con Centro de Reportes)
        vl.addWidget(QLabel("Elementos del reporte:"))
        self.chk_header = QCheckBox("Incluir encabezado (logo, empresa, datos del proyecto)")
        self.chk_header.setChecked(True)
        self.chk_header.setStyleSheet("color:#273445; font-size:12px;")
        vl.addWidget(self.chk_header)

        self.chk_footer = QCheckBox("Incluir pie de página (textos del formato: izq / centro / der)")
        self.chk_footer.setChecked(True)
        self.chk_footer.setStyleSheet("color:#273445; font-size:12px;")
        vl.addWidget(self.chk_footer)

        self.chk_legend = QCheckBox("Incluir leyenda (Tarea / Crítica / Hito / Dependencia / Hoy / Fin plazo)")
        self.chk_legend.setChecked(True)
        self.chk_legend.setStyleSheet("color:#273445; font-size:12px;")
        vl.addWidget(self.chk_legend)

        self.chk_page = QCheckBox("Incluir número de página (N de N)")
        self.chk_page.setChecked(True)
        self.chk_page.setStyleSheet("color:#273445; font-size:12px;")
        vl.addWidget(self.chk_page)

        # Atajo para editar el formato compartido
        hh_fmt = QHBoxLayout()
        hint = QLabel(
            "Logo, nombre de empresa, colores y textos del pie se guardan en "
            "el formato compartido con el Centro de Reportes."
        )
        hint.setStyleSheet("color:#667885; font-size:10px;")
        hint.setWordWrap(True)
        from PySide6.QtWidgets import QPushButton as _QPB
        btn_fmt = _QPB("🎨  Editar formato…")
        btn_fmt.setToolTip("Abrir el editor de formato (logo, empresa, color, pies)")
        btn_fmt.setStyleSheet(
            "QPushButton { background:white; color:#273445; border:1px solid #D4D4D4;"
            " border-radius:4px; padding:5px 12px; font-size:11px; }"
            "QPushButton:hover { background:#FEF5EB; border-color:#F37329; color:#C0621A; }"
        )
        btn_fmt.clicked.connect(self._abrir_formato)
        hh_fmt.addWidget(hint, 1)
        hh_fmt.addWidget(btn_fmt, 0)
        vl.addLayout(hh_fmt)

        vl.addStretch()

        # Botones — fila propia con buen espaciado para que nada quede tapado
        from PySide6.QtWidgets import QPushButton
        bar = QHBoxLayout()
        bar.setSpacing(8)
        btn_cancel = QPushButton("Cancelar")
        btn_cancel.setMinimumWidth(96); btn_cancel.setMinimumHeight(32)
        btn_cancel.setStyleSheet(
            "QPushButton { background:white; color:#485A6C; border:1px solid #D4D4D4;"
            " border-radius:4px; padding:6px 18px; }"
            "QPushButton:hover { background:#F1F5F9; }"
        )
        btn_cancel.clicked.connect(self.reject)

        btn_img = QPushButton("🖼  Imagen…")
        btn_img.setMinimumWidth(120); btn_img.setMinimumHeight(32)
        btn_img.setToolTip("Exportar como imagen (PNG / JPG)")
        btn_img.setStyleSheet(
            "QPushButton { background:white; color:#273445; border:1px solid #D4D4D4;"
            " border-radius:4px; padding:6px 14px; font-weight:600; }"
            "QPushButton:hover { background:#FEF5EB; border-color:#F37329; color:#C0621A; }"
        )
        btn_img.clicked.connect(self._emitir_imagen)

        btn_prev = QPushButton("👁  Vista previa")
        btn_prev.setMinimumWidth(130); btn_prev.setMinimumHeight(32)
        btn_prev.setStyleSheet(
            "QPushButton { background:white; color:#273445; border:1px solid #D4D4D4;"
            " border-radius:4px; padding:6px 14px; font-weight:600; }"
            "QPushButton:hover { background:#FEF5EB; border-color:#F37329; color:#C0621A; }"
        )
        btn_prev.clicked.connect(self._emitir_preview)

        btn_ok = QPushButton("Exportar PDF…")
        btn_ok.setMinimumWidth(130); btn_ok.setMinimumHeight(32)
        btn_ok.setDefault(True)
        btn_ok.setStyleSheet(
            "QPushButton { background:#F37329; color:white; border:none;"
            " border-radius:4px; padding:6px 18px; font-weight:600; }"
            "QPushButton:hover { background:#C0621A; }"
        )
        btn_ok.clicked.connect(self.accept)

        bar.addWidget(btn_cancel)
        bar.addStretch()
        bar.addWidget(btn_img)
        bar.addWidget(btn_prev)
        bar.addWidget(btn_ok)
        vl.addLayout(bar)

    def _opts(self) -> dict:
        escala_map = ['auto', 'dias', 'semanas', 'meses']
        idx = max(0, self.cmb_papel.currentIndex())
        papel = self.PAPEL_OPCIONES[idx][1]
        return {
            'modo':   'fit' if self.rb_fit.isChecked() else 'multi',
            'orient': 'landscape' if self.rb_land.isChecked() else 'portrait',
            'incluir_pred': self.chk_pred.isChecked(),
            'escala': escala_map[self.cmb_escala.currentIndex()],
            'hojas_x': int(self.sp_hojas.value()),  # 0 = auto
            'papel':  papel,
            'incluir_header': self.chk_header.isChecked(),
            'incluir_footer': self.chk_footer.isChecked(),
            'incluir_legend': self.chk_legend.isChecked(),
            'incluir_page':   self.chk_page.isChecked(),
        }

    def _abrir_formato(self):
        try:
            from views.formato_reporte_dialog import FormatoReporteDialog
            FormatoReporteDialog(self).exec()
        except Exception as e:
            from PySide6.QtWidgets import QMessageBox as _Mb
            _Mb.warning(self, "Formato",
                          f"No se pudo abrir el editor de formato:\n{e}")

    def _emitir_preview(self):
        self.preview_solicitado.emit(self._opts())

    def _emitir_imagen(self):
        self.imagen_solicitada.emit(self._opts())

    @classmethod
    def preguntar(cls, parent=None, on_preview=None, on_imagen=None):
        dlg = cls(parent)
        if on_preview is not None:
            dlg.preview_solicitado.connect(on_preview)
        if on_imagen is not None:
            dlg.imagen_solicitada.connect(on_imagen)
        if dlg.exec() != QDialog.Accepted:
            return None
        return dlg._opts()


# ════════════════════════════════════════════════════════════════════════════
# 1. GANTT
# ════════════════════════════════════════════════════════════════════════════

class _GanttBar(QGraphicsRectItem):
    """Barra Gantt interactiva estilo MS Project.

    Gestos (se deciden según el movimiento, no según dónde agarras — clave
    para que las tareas de 1 día no confundan enlazar con redimensionar):
      - arrastre VERTICAL / hacia otra barra  → crear dependencia (FS por
        defecto; el tipo final lo afina dónde sueltas en la barra destino)
      - arrastre HORIZONTAL desde un borde     → redimensionar
      - arrastre HORIZONTAL desde el centro     → mover (reprogramar)
    """
    EDGE = 6   # px desde el borde para detectar resize
    RADIUS = 3 # radio de redondeo

    def __init__(self, x, y, w, h, day_w, partida_id, on_change,
                 critical=False, segmento_idx=0, on_context=None,
                 color=None, gantt_widget=None):
        super().__init__(x, y, w, h)
        self._day_w        = day_w
        self._pid          = partida_id
        self._on_change    = on_change
        self._on_context   = on_context  # callback(pid, qpoint_global)
        self._segmento_idx = segmento_idx
        self._press_zone   = None      # 'move' | 'resize-l' | 'resize-r' (zona presionada)
        self._linking      = False     # arrastrando una dependencia hacia otra barra
        self._press_scene  = None
        self._orig         = None      # rect original al iniciar drag
        self._critical     = critical
        self._color_custom = (color or '').strip() or None
        self._gantt        = gantt_widget
        self._dragged      = False
        self._highlighted  = False     # relacionada con la seleccionada
        self.setAcceptHoverEvents(True)
        self.setFlag(QGraphicsItem.ItemIsSelectable)
        # No usamos pen/brush nativos — pintamos en paint() con bordes redondeados
        self.setPen(QPen(Qt.NoPen))
        self.setBrush(QBrush(Qt.NoBrush))

    def set_highlighted(self, on: bool):
        if self._highlighted != on:
            self._highlighted = on
            self.update()

    def hoverEnterEvent(self, event):
        super().hoverEnterEvent(event)
        # Re-aplicar cursor (super lo ajusta también)
        self.setCursor(Qt.OpenHandCursor)

    def hoverLeaveEvent(self, event):
        super().hoverLeaveEvent(event)

    def contextMenuEvent(self, event):
        if self._on_context:
            self._on_context(self._pid, event.screenPos())
            event.accept()
        else:
            super().contextMenuEvent(event)

    def _colors(self):
        """Devuelve (fill, border) según prioridad: custom > critical > default."""
        if self._color_custom:
            base = QColor(self._color_custom)
            if not base.isValid():
                base = QColor("#F37329")
            return base, base.darker(140)
        if self._critical:
            return QColor("#C6262E"), QColor("#A10705")
        return QColor(GANTT_BAR), QColor(GANTT_BAR_BORDER)

    def paint(self, painter, option, widget=None):
        from PySide6.QtCore import QRectF as _QRectF
        fill, border = self._colors()
        painter.setRenderHint(QPainter.Antialiasing, True)
        # Pen según estado: seleccionada > relacionada > normal
        if self.isSelected():
            # Naranja marca: resalta la barra activa sobre el relleno arándano.
            pen = QPen(QColor("#F37329"), 2.5)
        elif self._highlighted:
            # Relacionadas (pred/suc): contorno azul oscuro visible sobre el fill.
            pen = QPen(QColor("#002E99"), 2.0)
        else:
            pen = QPen(border, 1)
        painter.setPen(pen)
        painter.setBrush(QBrush(fill))
        r = self.rect()
        # Reducir 0.5 px para que el borde no se corte
        rr = _QRectF(r.x() + 0.5, r.y() + 0.5, r.width() - 1, r.height() - 1)
        painter.drawRoundedRect(rr, self.RADIUS, self.RADIUS)

    def _edge(self):
        """Ancho del área de resize, proporcional al tamaño de la barra.
        Garantiza al menos 2px en cada extremo y una zona central de move
        de al menos 2px cuando la barra es muy angosta."""
        w = self.rect().width()
        if w <= 6:
            # Muy angosta: 40% en cada extremo, 20% para move al centro
            return max(1.0, w * 0.4)
        return min(6.0, max(2.0, w / 3.0))

    def hoverMoveEvent(self, event):
        x = event.pos().x() - self.rect().x()
        w = self.rect().width()
        edge = self._edge()
        if x < edge or x > w - edge:
            self.setCursor(Qt.SizeHorCursor)
        else:
            self.setCursor(Qt.OpenHandCursor)
        super().hoverMoveEvent(event)

    def mousePressEvent(self, event):
        if event.button() != Qt.LeftButton:
            return super().mousePressEvent(event)
        x = event.pos().x() - self.rect().x()
        w = self.rect().width()
        edge = self._edge()
        # Zona presionada — decide el gesto HORIZONTAL dentro de la propia fila
        # (resize en los bordes, move en el centro). El gesto que apunta a OTRA
        # barra siempre enlaza, sin importar dónde se agarró.
        if x < edge:
            self._press_zone = 'resize-l'
        elif x > w - edge:
            self._press_zone = 'resize-r'
        else:
            self._press_zone = 'move'
        self._linking = False
        self._press_scene = event.scenePos()
        r = self.rect()
        self._orig = (r.x(), r.y(), r.width(), r.height())
        self._dragged = False
        event.accept()

    def mouseMoveEvent(self, event):
        if self._press_scene is None:
            return super().mouseMoveEvent(event)
        sp = event.scenePos()
        dx = sp.x() - self._press_scene.x()
        dy = sp.y() - self._press_scene.y()

        # ── ¿El arrastre apunta a ENLAZAR? Se reevalúa en cada movimiento
        #    (estilo MS Project): el cursor está sobre OTRA barra, o el cursor
        #    salió verticalmente de la fila propia (yendo hacia otra tarea).
        #    Solo las partidas reales (pid > 0) pueden ser origen de enlace. ──
        want_link = False
        if self._gantt is not None and self._pid > 0:
            over = self._gantt._bar_at_scene(sp)
            if over is not None and over is not self and over._pid > 0:
                want_link = True
            elif abs(dy) > self.rect().height():
                want_link = True

        if want_link:
            if not self._linking:
                self._linking = True
                # deshacer cualquier preview de mover/redimensionar
                ox, oy, ow, oh = self._orig
                self.setRect(ox, oy, ow, oh)
                origin = QPointF(ox + ow, oy + oh / 2)
                self._gantt._dep_drag_start(self._pid, 'finish', origin)
                self.setCursor(Qt.CrossCursor)
            self._gantt._dep_drag_update(sp)
            event.accept()
            return

        # ── Ya no apunta a enlazar: si veníamos enlazando, quitar la línea ──
        if self._linking:
            self._linking = False
            self._gantt._dep_drag_cancel()
            self.setCursor(Qt.OpenHandCursor)

        # ── Mover / redimensionar (umbral 3px para no robar el clic) ──
        if not self._dragged and abs(dx) < 3:
            event.accept()
            return
        delta_days = round(dx / self._day_w)
        if delta_days != 0:
            self._dragged = True
        snap_dx = delta_days * self._day_w
        ox, oy, ow, oh = self._orig
        if self._press_zone == 'resize-l':
            new_x = max(0, ox + snap_dx)
            new_w = max(self._day_w, ow - (new_x - ox))
            self.setRect(new_x, oy, new_w, oh)
        elif self._press_zone == 'resize-r':
            new_w = max(self._day_w, ow + snap_dx)
            self.setRect(ox, oy, new_w, oh)
        else:  # move
            new_x = max(0, ox + snap_dx)
            self.setRect(new_x, oy, ow, oh)
        event.accept()

    def mouseReleaseEvent(self, event):
        # ── Soltó un enlace en curso → resolver dependencia en la barra destino ──
        if self._linking:
            self._linking = False
            self.setCursor(Qt.OpenHandCursor)
            if self._gantt is not None:
                self._gantt._dep_drag_end(event.scenePos())
            self._press_scene = None
            self._dragged = False
            event.accept()
            return
        if self._press_scene is None:
            return super().mouseReleaseEvent(event)
        if self._dragged:
            r = self.rect()
            new_inicio = max(1, round(r.x() / self._day_w) + 1)
            new_dur    = max(1, round(r.width() / self._day_w))
            if self._on_change:
                self._on_change(self._pid, new_inicio, new_dur,
                                  self._press_zone, self._segmento_idx)
        else:
            # Click puro sobre la barra → seleccionarla (resalta pred/suc)
            from PySide6.QtWidgets import QApplication as _QA
            mods = _QA.keyboardModifiers()
            sc = self.scene()
            if sc is not None and not (mods & Qt.ControlModifier):
                sc.clearSelection()
            if mods & Qt.ControlModifier:
                self.setSelected(not self.isSelected())
            else:
                self.setSelected(True)
            if sc is not None:
                for v in sc.views():
                    v.setFocus()
                    break
        self._press_scene = None
        self._dragged = False
        event.accept()


class _GanttMilestone(QGraphicsPolygonItem):
    """Hito como diamante. es_hito: 1=genérico (morado), 2=inicio (azul), 3=fin (verde)."""
    SIZE = 12

    def __init__(self, cx, cy, day_w, partida_id, on_change, es_hito=1,
                 on_context=None):
        s = self.SIZE
        poly = QPolygonF([
            QPointF(cx, cy - s/2),
            QPointF(cx + s/2, cy),
            QPointF(cx, cy + s/2),
            QPointF(cx - s/2, cy),
        ])
        super().__init__(poly)
        self._day_w     = day_w
        self._pid       = partida_id
        self._on_change = on_change
        self._on_context = on_context
        self._es_hito   = es_hito
        self._cx        = cx
        self._cy        = cy
        self._mode      = None
        self._press_scene = None
        self.setAcceptHoverEvents(True)
        self._apply_style()

    def contextMenuEvent(self, event):
        if self._on_context:
            self._on_context(self._pid, event.screenPos())
            event.accept()
        else:
            super().contextMenuEvent(event)

    def _apply_style(self):
        colors = {1: ("#7A36B1", "#4F1F76"),  # morado
                  2: ("#3689E6", "#1E5DA8"),  # azul (inicio)
                  3: ("#3A9104", "#206700")}  # verde (fin)
        bg, brd = colors.get(self._es_hito, colors[1])
        self.setBrush(QBrush(QColor(bg)))
        self.setPen(QPen(QColor(brd), 1.5))

    def hoverEnterEvent(self, event):
        self.setCursor(Qt.OpenHandCursor)
        super().hoverEnterEvent(event)

    def mousePressEvent(self, event):
        if event.button() != Qt.LeftButton:
            return super().mousePressEvent(event)
        self._mode = 'move'
        self._press_scene = event.scenePos()
        event.accept()

    def mouseMoveEvent(self, event):
        if self._mode != 'move':
            return super().mouseMoveEvent(event)
        delta_x = event.scenePos().x() - self._press_scene.x()
        delta_days = round(delta_x / self._day_w)
        self.setX(delta_days * self._day_w)
        event.accept()

    def mouseReleaseEvent(self, event):
        if self._mode != 'move':
            return super().mouseReleaseEvent(event)
        delta_x = self.x()
        delta_days = round(delta_x / self._day_w)
        if self._on_change and delta_days != 0:
            new_inicio = max(1, round(self._cx / self._day_w) + delta_days + 1 - 1)
            self._on_change(self._pid, new_inicio, 0, 'move', 0)
        self._mode = None
        self._press_scene = None
        event.accept()


class _GanttArrow(QGraphicsPathItem):
    """Flecha de dependencia entre barras. Soporta los 4 tipos MS Project
    (FS/SS/FF/SF) y mensajes contextuales (cambiar tipo, editar lag, eliminar).
    Las coordenadas (x1,y1) y (x2,y2) ya vienen calculadas por el caller
    según el tipo: x1 es Fin de la pred. para FS/FF, o Inicio para SS/SF;
    x2 es Inicio del sucesor para FS/SS, o Fin para FF/SF."""
    def __init__(self, x1, y1, x2, y2, critical=False, tipo='FS', lag=0,
                 pct=0, tgt_pct=0, source_pid=None, target_pid=None,
                 on_context=None, on_drag=None, day_w=16,
                 source_at_finish=True, target_at_finish=False,
                 stub_x_override=None):
        # ── Ruteo ortogonal estilo MS Project ──────────────────────────────
        # exit_dir: la pred sale por la derecha si conecta por su Fin, por la
        #           izquierda si conecta por su Inicio.
        # tip_dir : la punta apunta HACIA la barra destino — a la derecha si
        #           entra por el Inicio (la barra se extiende a la derecha), a
        #           la izquierda si entra por el Fin.
        # stub_x_override: cuando un mismo predecesor tiene ≥2 sucesores el
        #   caller pasa una X de stub compartida para que todas las flechas
        #   del grupo bundleen en un solo tronco vertical (estilo MS Project).
        #   En ese caso forzamos la ruta Z para mantener el tronco coherente.
        HEAD_LEN = 7.0
        HEAD_W   = 6.5
        STUB     = 9.0
        exit_dir = 1 if source_at_finish else -1
        tip_dir  = 1 if not target_at_finish else -1
        base_x   = x2 - tip_dir * HEAD_LEN   # unión trazo↔punta
        if stub_x_override is not None:
            sx_stub = stub_x_override
            force_z = True
            skip_trunk = True  # el tronco lo dibuja un item compartido
        else:
            sx_stub = x1 + exit_dir * STUB
            force_z = False
            skip_trunk = False

        if skip_trunk:
            # Bundled: solo la rama lateral (sx_stub, y2) → (base_x, y2). El
            # tronco (x1→sx_stub→y2_máx) lo pinta un item compartido por
            # grupo, para evitar overlapping con colores mezclados.
            pts = [QPointF(sx_stub, y2), QPointF(base_x, y2)]
        else:
            pts = [QPointF(x1, y1), QPointF(sx_stub, y1)]
            if force_z or (base_x - sx_stub) * tip_dir >= 0:
                # Ruta en «Z»: codo, vertical hasta la fila destino, y horizontal
                # avanzando hacia la punta. (caso normal, sin solape)
                pts.append(QPointF(sx_stub, y2))
                pts.append(QPointF(base_x, y2))
            else:
                # Ruta en «C» (solape / lead): rodea por el punto medio entre
                # filas para no atravesar las barras.
                mid_y  = (y1 + y2) / 2.0
                turn_x = base_x - tip_dir * STUB
                pts.append(QPointF(sx_stub, mid_y))
                pts.append(QPointF(turn_x, mid_y))
                pts.append(QPointF(turn_x, y2))
                pts.append(QPointF(base_x, y2))

        tip = QPointF(x2, y2)
        self._elbow_pts = pts
        self._head_poly = QPolygonF([
            tip,
            QPointF(base_x, y2 - HEAD_W / 2.0),
            QPointF(base_x, y2 + HEAD_W / 2.0),
        ])
        self._head_pad = HEAD_W

        path = QPainterPath()
        path.moveTo(pts[0])
        for pt in pts[1:]:
            path.lineTo(pt)
        path.lineTo(tip)   # incluir la punta en bounding/shape (hit-area)
        super().__init__(path)
        col = QColor("#cc3b02") if critical else QColor("#6B7785")
        self.setPen(QPen(col, 1.3, Qt.SolidLine, Qt.FlatCap, Qt.MiterJoin))
        self.setBrush(QBrush(Qt.NoBrush))
        self.setZValue(2)
        # Datos para el menú contextual y drag
        self._tipo       = tipo
        self._lag        = lag
        self._pct        = pct
        self._tgt_pct    = tgt_pct
        self._source_pid = source_pid
        self._target_pid = target_pid
        self._on_context = on_context
        self._on_drag    = on_drag
        self._day_w      = day_w
        self._critical   = critical
        self._press_scene = None
        self._drag_active = False
        self._drag_label  = None
        self._related_highlight = False
        if on_context is not None and source_pid is not None and target_pid is not None:
            self.setAcceptHoverEvents(True)
            self.setCursor(Qt.PointingHandCursor)
            self.setFlag(QGraphicsItem.ItemIsSelectable, True)
            self.setFlag(QGraphicsItem.ItemIsFocusable, True)
            if pct:
                etiqueta = f"CC+{int(pct)}% (comienza cuando la pred. lleva {int(pct)}%)"
            elif tgt_pct:
                etiqueta = f"sucesor al {tgt_pct}% al fin de pred"
            else:
                from core.cronograma import _TIPO_ES
                etiqueta = _TIPO_ES.get(tipo, tipo) + (f"{lag:+d}" if lag else "")
            self.setToolTip(
                f"Dependencia: {etiqueta}\n"
                "Clic = seleccionar (Supr para eliminar) · arrastra para "
                "ajustar lag/% · clic derecho para más opciones."
            )

    def shape(self):
        """Hit-area más ancha que la línea para que sea fácil clicar."""
        stroker = QPainterPathStroker()
        stroker.setWidth(10)
        return stroker.createStroke(self.path())

    def boundingRect(self):
        # Margen extra para que las alas de la punta nunca se recorten.
        pad = getattr(self, '_head_pad', 6.0) + 2.0
        return super().boundingRect().adjusted(-pad, -pad, pad, pad)

    def paint(self, painter, option, widget=None):
        painter.setRenderHint(QPainter.Antialiasing, True)
        pen = self.pen()
        # Trazo del conector (sin relleno)
        painter.setPen(pen)
        painter.setBrush(Qt.NoBrush)
        poly = QPolygonF(self._elbow_pts)
        painter.drawPolyline(poly)
        # Punta de flecha rellena del mismo color que el trazo
        painter.setPen(QPen(pen.color(), 0.6))
        painter.setBrush(QBrush(pen.color()))
        painter.drawPolygon(self._head_poly)

    def _normal_pen(self):
        col = QColor("#cc3b02") if self._critical else QColor("#6B7785")
        return QPen(col, 1.3)

    def _hover_pen(self):
        col = QColor("#F37329") if not self._critical else QColor("#A10705")
        return QPen(col, 2.0)

    def _selected_pen(self):
        return QPen(QColor("#F37329"), 2.5, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin)

    def _related_pen(self):
        return QPen(QColor("#3689E6"), 2.0, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin)

    def set_related_highlight(self, on: bool):
        self._related_highlight = on
        if self.isSelected():
            return
        self.setPen(self._related_pen() if on else self._normal_pen())

    def hoverEnterEvent(self, event):
        if (self._on_context is not None and not self.isSelected()
                and not getattr(self, '_related_highlight', False)):
            self.setPen(self._hover_pen())
        super().hoverEnterEvent(event)

    def hoverLeaveEvent(self, event):
        if (self._on_context is not None and not self.isSelected()
                and not getattr(self, '_related_highlight', False)):
            self.setPen(self._normal_pen())
        super().hoverLeaveEvent(event)

    def itemChange(self, change, value):
        if change == QGraphicsItem.ItemSelectedHasChanged:
            if bool(value):
                self.setPen(self._selected_pen())
            elif getattr(self, '_related_highlight', False):
                self.setPen(self._related_pen())
            else:
                self.setPen(self._normal_pen())
        return super().itemChange(change, value)

    def contextMenuEvent(self, event):
        if self._on_context is not None and self._source_pid and self._target_pid:
            self._on_context(self._source_pid, self._target_pid,
                              self._tipo, self._lag, self._pct, self._tgt_pct,
                              event.screenPos())
            event.accept()
        else:
            super().contextMenuEvent(event)

    # ── Drag horizontal: ajusta lag (días) o pct (%) ─────────────────────
    # Si el mouse se mueve > 3px tras presionar, entra en modo drag;
    # si se suelta sin moverse, es un click → selecciona la flecha.
    def mousePressEvent(self, event):
        if (event.button() == Qt.LeftButton
                and self._source_pid and self._target_pid):
            self._press_scene = event.scenePos()
            self._drag_active = False
            self.setCursor(Qt.ClosedHandCursor)
            event.accept()
            return
        super().mousePressEvent(event)

    def _ensure_drag_label(self):
        if self._drag_label is None and self.scene() is not None:
            self._drag_label = QGraphicsTextItem()
            self._drag_label.setZValue(60)
            self.scene().addItem(self._drag_label)

    def _update_drag_label(self, scene_pos, dx_days):
        self._ensure_drag_label()
        if self._drag_label is None:
            return
        if self._pct:
            # Convertir dx_days a delta_% según duración de la pred
            txt = f"≈ {dx_days:+d}d"  # simple — el commit hace el cálculo real
        else:
            from core.cronograma import _TIPO_ES
            es = _TIPO_ES.get(self._tipo, self._tipo)   # FS→FC, SS→CC, …
            new_lag = self._lag + dx_days
            txt = f"{es}{new_lag:+d}d" if new_lag else es
        self._drag_label.setHtml(
            f"<div style='background:#F37329; color:white;"
            f" padding:1px 6px; border-radius:4px;"
            f" font-size:8pt; font-weight:700;'>{txt}</div>"
        )
        self._drag_label.setPos(scene_pos.x() + 10, scene_pos.y() - 22)

    def mouseMoveEvent(self, event):
        if self._press_scene is not None:
            dx = event.scenePos().x() - self._press_scene.x()
            # Activar drag solo si se movió suficiente — evita robar el click
            if not self._drag_active and abs(dx) > 3 and self._on_drag is not None:
                self._drag_active = True
            if self._drag_active:
                self.setTransform(QTransform().translate(dx, 0))
                dx_days = round(dx / max(1, self._day_w))
                self._update_drag_label(event.scenePos(), dx_days)
            event.accept()
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self._press_scene is None:
            return super().mouseReleaseEvent(event)
        was_drag = self._drag_active
        dx = event.scenePos().x() - self._press_scene.x()
        delta_days = round(dx / max(1, self._day_w))
        self._drag_active = False
        self._press_scene = None
        self.setTransform(QTransform())
        self.setCursor(Qt.PointingHandCursor)
        if self._drag_label is not None and self.scene() is not None:
            self.scene().removeItem(self._drag_label)
            self._drag_label = None
        if was_drag and delta_days != 0 and self._on_drag is not None:
            self._on_drag(self._source_pid, self._target_pid, delta_days,
                            self._tipo, self._lag, self._pct, self._tgt_pct)
        elif not was_drag:
            # Click puro → seleccionar (Supr la elimina). Ctrl+clic alterna
            # la selección para multi-eliminar varias dependencias a la vez.
            from PySide6.QtWidgets import QApplication as _QA
            mods = _QA.keyboardModifiers()
            sc = self.scene()
            if sc is not None and not (mods & Qt.ControlModifier):
                sc.clearSelection()
            if mods & Qt.ControlModifier:
                self.setSelected(not self.isSelected())
            else:
                self.setSelected(True)
            if sc is not None:
                for v in sc.views():
                    v.setFocus()
                    break
        event.accept()


class _GanttView(QGraphicsView):
    """QGraphicsView para el Gantt con soporte de Ctrl+wheel = zoom horizontal
    (acerca/aleja la escala de días sin afectar la altura de las barras)."""
    def __init__(self, scene, gantt_widget):
        super().__init__(scene)
        self._gantt = gantt_widget
        self.setFocusPolicy(Qt.StrongFocus)

    def wheelEvent(self, event):
        if event.modifiers() & Qt.ControlModifier:
            # Zoom horizontal: cambiar DAY_W del widget padre
            delta = event.angleDelta().y()
            if delta == 0:
                return
            step = 2 if abs(delta) < 240 else 4
            if delta > 0:
                self._gantt._zoom(+step)
            else:
                self._gantt._zoom(-step)
            event.accept()
            return
        super().wheelEvent(event)

    def keyPressEvent(self, event):
        """Supr / Backspace elimina la flecha (dependencia) seleccionada."""
        if event.key() in (Qt.Key_Delete, Qt.Key_Backspace):
            sc = self.scene()
            if sc is not None:
                for it in sc.selectedItems():
                    if isinstance(it, _GanttArrow) and it._source_pid and it._target_pid:
                        self._gantt._remove_dep(it._source_pid, it._target_pid)
                        event.accept()
                        return
        super().keyPressEvent(event)


class _GanttHeaderView(QGraphicsView):
    """Vista solo del header (días/meses). Fija verticalmente; sincroniza
    su scroll horizontal con la vista principal del Gantt para que las
    columnas siempre coincidan con las barras debajo."""
    def __init__(self, scene, gantt_widget):
        super().__init__(scene)
        self._gantt = gantt_widget
        self.setStyleSheet("background:white; border:none;")
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setRenderHint(QPainter.Antialiasing)
        self.setFrameShape(QGraphicsView.NoFrame)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

    def wheelEvent(self, event):
        # Replicar zoom y scroll vertical en la vista principal
        if event.modifiers() & Qt.ControlModifier:
            delta = event.angleDelta().y()
            if delta == 0:
                return
            step = 2 if abs(delta) < 240 else 4
            self._gantt._zoom(+step if delta > 0 else -step)
            event.accept()
            return
        # Forwardear el evento a la vista principal (para scroll vertical)
        self._gantt.view.wheelEvent(event)


class GanttWidget(QWidget):
    """Diagrama Gantt: tabla izquierda + barras a la derecha."""

    DAY_W = 2    # default Meses (2px/día) — Marco prefiere vista mensual al abrir
    ROW_H = 38   # ~2 líneas de texto a 11pt — descripciones largas se ven completas
    H_HDR_TOP = 18
    H_HDR_BOT = 16
    LEAD_DAYS = 1   # margen de días vacíos ANTES del inicio (respiro visual)

    def __init__(self, parent: CronogramaView):
        super().__init__(parent)
        self._cv = parent
        self._fit_pendiente = True   # auto-fit al primer render
        # Historial de dependencias (predecesoras) para deshacer/rehacer con
        # Ctrl+Z / Ctrl+Y. Cada entrada = {pid: predecesoras} de TODAS las
        # partidas antes de una mutación (_set_dep / _remove_dep).
        self._dep_undo: list[dict] = []
        self._dep_redo: list[dict] = []
        self._build_ui()
        from PySide6.QtGui import QShortcut, QKeySequence
        for seq, slot in ((QKeySequence.Undo, self._dep_deshacer),
                          (QKeySequence.Redo, self._dep_rehacer)):
            sc = QShortcut(seq, self)
            sc.setContext(Qt.WidgetWithChildrenShortcut)
            sc.activated.connect(slot)

    def _build_ui(self):
        vl = QVBoxLayout(self)
        vl.setContentsMargins(8, 8, 8, 8)
        vl.setSpacing(6)

        # Toolbar simple — escala/zoom
        tb = QFrame()
        tb.setStyleSheet("background:transparent; border:none;")
        tb_hl = QHBoxLayout(tb)
        tb_hl.setContentsMargins(0, 0, 0, 0)
        tb_hl.setSpacing(8)

        lbl = QLabel("Escala:")
        lbl.setStyleSheet(f"color:{SLATE_500}; font-size:11px;")
        tb_hl.addWidget(lbl)

        self._cmb_escala = QComboBox()
        # "Auto" = la unidad del encabezado se deriva del zoom (comportamiento
        # clásico). Las demás FIJAN la unidad del encabezado (días/semanas/meses)
        # independientemente del zoom — como la escala temporal de MS Project.
        self._cmb_escala.addItems(["Auto", "Días", "Semanas", "Meses"])
        from utils.tooltip import set_tooltip as _set_tt_esc
        _set_tt_esc(self._cmb_escala,
                    "Unidad del encabezado del Gantt.\n"
                    "Auto: se ajusta sola al zoom.\n"
                    "Días / Semanas / Meses: fija esa unidad (el zoom solo "
                    "cambia el tamaño, no la unidad).")
        # Restaurar la última unidad elegida por el usuario
        from PySide6.QtCore import QSettings as _QS_esc
        _saved_unit = str(_QS_esc("ingePresupuestos", "layout").value(
            "gantt_escala_unit", "auto"))
        _unit_idx = {"auto": 0, "dia": 1, "semana": 2, "mes": 3}.get(_saved_unit, 0)
        self._escala_unit = _saved_unit if _saved_unit in (
            "auto", "dia", "semana", "mes") else "auto"
        self._cmb_escala.setCurrentIndex(_unit_idx)
        self._cmb_escala.currentIndexChanged.connect(self._on_escala)
        self._cmb_escala.setStyleSheet(
            "QComboBox { min-height:0; padding:2px 8px; font-size:11px;"
            f" border:1px solid {SILVER_300}; border-radius:4px; }}"
        )
        tb_hl.addWidget(self._cmb_escala)

        tb_hl.addSpacing(20)
        lbl_z = QLabel("Zoom:")
        lbl_z.setStyleSheet(f"color:{SLATE_500}; font-size:11px;")
        tb_hl.addWidget(lbl_z)

        # − (zoom out)
        btn_zout = QPushButton("−")
        btn_zout.setFixedSize(26, 22)
        btn_zout.setCursor(Qt.PointingHandCursor)
        from utils.tooltip import set_tooltip as _set_tt
        _set_tt(btn_zout, "Alejar (Ctrl + rueda)")
        btn_zout.setStyleSheet(
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:14px; font-weight:700; min-height:0; padding:0; }}"
            f"QPushButton:hover {{ background:{SILVER_300}; }}"
        )
        btn_zout.clicked.connect(lambda: self._zoom(-2))
        tb_hl.addWidget(btn_zout)

        # Slider de zoom continuo (2..40 px/día)
        self._sl_zoom = QSlider(Qt.Horizontal)
        self._sl_zoom.setRange(2, 40)
        self._sl_zoom.setValue(int(self.DAY_W))
        self._sl_zoom.setFixedWidth(110)
        self._sl_zoom.setSingleStep(1)
        self._sl_zoom.setPageStep(4)
        self._sl_zoom.setTickInterval(4)
        _set_tt(self._sl_zoom, "Ancho de día (px) — arrastra para zoom continuo")
        self._sl_zoom.setStyleSheet(
            "QSlider::groove:horizontal {"
            f" background:{SILVER_300}; height:4px; border-radius:2px; }}"
            "QSlider::sub-page:horizontal {"
            " background:#F37329; height:4px; border-radius:2px; }"
            "QSlider::handle:horizontal {"
            " background:white; border:1.5px solid #F37329; width:12px;"
            " margin:-5px 0; border-radius:8px; }"
            "QSlider::handle:horizontal:hover { background:#FEF5EB; }"
        )
        self._sl_zoom.valueChanged.connect(self._on_zoom_slider)
        tb_hl.addWidget(self._sl_zoom)

        # + (zoom in)
        btn_zin = QPushButton("+")
        btn_zin.setFixedSize(26, 22)
        btn_zin.setCursor(Qt.PointingHandCursor)
        _set_tt(btn_zin, "Acercar (Ctrl + rueda)")
        btn_zin.setStyleSheet(
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:14px; font-weight:700; min-height:0; padding:0; }}"
            f"QPushButton:hover {{ background:{SILVER_300}; }}"
        )
        btn_zin.clicked.connect(lambda: self._zoom(+2))
        tb_hl.addWidget(btn_zin)

        self._lbl_zoom = QLabel(f"{self.DAY_W}px")
        self._lbl_zoom.setFixedWidth(40)
        self._lbl_zoom.setStyleSheet(f"color:{SLATE_500}; font-size:11px;")
        tb_hl.addWidget(self._lbl_zoom)

        # Botón "Ajustar" — calcula DAY_W para que el proyecto entero quepa
        btn_fit = QPushButton("⤢ Ajustar")
        btn_fit.setFixedHeight(22)
        btn_fit.setCursor(Qt.PointingHandCursor)
        _set_tt(btn_fit, "Calcular zoom para que el proyecto completo entre en pantalla")
        btn_fit.setStyleSheet(
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:2px 10px; min-height:0; }}"
            f"QPushButton:hover {{ background:{SILVER_300}; }}"
        )
        btn_fit.clicked.connect(self._zoom_fit)
        tb_hl.addWidget(btn_fit)

        # Botón "100%" — reset a DAY_W=16 (día completo grande)
        btn_reset = QPushButton("100%")
        btn_reset.setFixedHeight(22)
        btn_reset.setCursor(Qt.PointingHandCursor)
        _set_tt(btn_reset, "Restablecer zoom a tamaño normal (16 px/día)")
        btn_reset.setStyleSheet(
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:2px 8px; min-height:0; }}"
            f"QPushButton:hover {{ background:{SILVER_300}; }}"
        )
        btn_reset.clicked.connect(lambda: self._set_day_w(16))
        tb_hl.addWidget(btn_reset)

        # Toggle "Holgura" — barra gris detrás de cada tarea no crítica
        from PySide6.QtCore import QSettings as _QS
        s = _QS("ingePresupuestos", "layout")
        self._show_slack = bool(s.value("gantt_show_slack", False, type=bool))
        self._btn_slack = QPushButton("⏳ Holgura")
        self._btn_slack.setCheckable(True)
        self._btn_slack.setChecked(self._show_slack)
        self._btn_slack.setFixedHeight(22)
        self._btn_slack.setCursor(Qt.PointingHandCursor)
        _set_tt(self._btn_slack,
                  "Muestra/oculta la barra de holgura — cuánto puede atrasarse "
                  "cada tarea no crítica sin afectar el fin del proyecto.")
        self._btn_slack.setStyleSheet(
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:2px 10px; min-height:0; }}"
            f"QPushButton:hover {{ background:{SILVER_300}; }}"
            f"QPushButton:checked {{ background:#FEF5EB; border-color:#F37329;"
            f" color:#C0621A; }}"
        )
        self._btn_slack.toggled.connect(self._on_toggle_slack)
        tb_hl.addWidget(self._btn_slack)

        # ── Fecha de inicio + Feriados ──────────────────────────────────────
        tb_hl.addSpacing(20)
        from PySide6.QtWidgets import QDateEdit
        from PySide6.QtCore import QDate
        lbl_ini = QLabel("Inicio:")
        lbl_ini.setStyleSheet(f"color:{SLATE_500}; font-size:11px;")
        tb_hl.addWidget(lbl_ini)
        self._dp_inicio = QDateEdit()
        self._dp_inicio.setCalendarPopup(True)
        self._dp_inicio.setDisplayFormat("dd/MM/yyyy")
        # Cargar el valor actual del proyecto (o hoy)
        f_ini = self._project_start() or datetime.now()
        self._dp_inicio.setDate(QDate(f_ini.year, f_ini.month, f_ini.day))
        self._dp_inicio.setStyleSheet(
            "QDateEdit { min-height:0; padding:2px 4px; font-size:11px;"
            f" background:white; color:{SLATE_700};"
            f" border:1px solid {SILVER_300}; border-radius:4px; }}"
        )
        # Calendar popup styling — cubierto por `install_global_popup_styles(app)`
        self._dp_inicio.dateChanged.connect(self._on_fecha_inicio)
        tb_hl.addWidget(self._dp_inicio)

        btn_fer = QPushButton("Feriados…")
        btn_fer.setCursor(Qt.PointingHandCursor)
        from utils.tooltip import set_tooltip
        set_tooltip(btn_fer, "Días no laborables (se sombrean en el Gantt)")
        btn_fer.setStyleSheet(
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:3px 10px;"
            f" min-height:0; color:{SLATE_500}; }}"
            f"QPushButton:hover {{ background:{SILVER_100}; color:{SLATE_700}; }}"
        )
        btn_fer.clicked.connect(self._abrir_feriados)
        tb_hl.addWidget(btn_fer)

        tb_hl.addSpacing(20)
        btn_pdf = QPushButton("📄 PDF")
        btn_pdf.setCursor(Qt.PointingHandCursor)
        btn_pdf.setStyleSheet(
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:3px 10px;"
            f" min-height:0; color:{SLATE_500}; }}"
            f"QPushButton:hover {{ background:{SILVER_100}; color:{SLATE_700}; }}"
        )
        btn_pdf.clicked.connect(self._exportar_pdf)
        set_tooltip(btn_pdf, "Exportar el Gantt a PDF / imagen — con opciones "
                              "de papel, paginación y vista previa")
        tb_hl.addWidget(btn_pdf)

        btn_mpp = QPushButton("📊 MPP")
        btn_mpp.setCursor(Qt.PointingHandCursor)
        btn_mpp.setStyleSheet(
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:3px 10px;"
            f" min-height:0; color:{SLATE_500}; }}"
            f"QPushButton:hover {{ background:{SILVER_100}; color:{SLATE_700}; }}"
        )
        btn_mpp.clicked.connect(self._exportar_mpp)
        set_tooltip(btn_mpp, "Exportar a Microsoft Project (XML) — incluye "
                              "calendario, hitos, ruta crítica y notas con color")
        tb_hl.addWidget(btn_mpp)

        tb_hl.addStretch()

        lbl_leg = QLabel(
            "<span style='color:#3689E6'>▬</span> Tarea  "
            "<span style='color:#C6262E'>▬</span> Crítica  "
            "<span style='color:#7A36B1'>◆</span> Hito  "
            "<span style='color:#6B7785'>→</span> Dependencia  "
            "<span style='color:#C62828'>┊</span> Fin plazo"
        )
        lbl_leg.setStyleSheet(f"color:{SLATE_500}; font-size:10px;")
        tb_hl.addWidget(lbl_leg)

        vl.addWidget(tb)

        # Splitter: tabla izquierda | gantt derecha
        sp = QSplitter(Qt.Horizontal)
        # Handle ancho y visible — uno de 3px es casi imposible de agarrar
        # (a veces el clic caía dentro de la tabla o del gantt). Sin colapso.
        sp.setHandleWidth(8)
        sp.setChildrenCollapsible(False)
        sp.setOpaqueResize(True)
        sp.setStyleSheet(
            "QSplitter::handle { background:#D4D9E0; }"
            "QSplitter::handle:hover { background:#B7C0CC; }"
            "QSplitter::handle:pressed { background:#94A3B8; }"
        )

        # Tabla izquierda — 8 columnas (estilo MS Project: doble duración)
        self.tbl = QTableWidget(0, 8)
        self.tbl.setHorizontalHeaderLabels(
            ["id", "Ítem", "Descripción", "Días\ncal.", "Días\nlab.",
             "Inicio", "Fin", "Pred."]
        )
        # Tooltip explicativo de las dos columnas de duración.
        for _c, _tip in (
            (3, "Duración en días CALENDARIO (incluye domingos y feriados).\n"
                "Se calcula sola: Fin − Inicio + 1."),
            (4, "Duración en días LABORABLES (descuenta domingos/feriados).\n"
                "Es la que editas tú."),
        ):
            _h = self.tbl.horizontalHeaderItem(_c)
            if _h is not None:
                _h.setToolTip(_tip)
        # Tooltip del formato de predecesoras (notación MS Project en español)
        hdr_item = self.tbl.horizontalHeaderItem(7)
        if hdr_item is not None:
            hdr_item.setToolTip(
                "Predecesoras — notación MS Project (español):\n"
                "  5         → FC (default): inicia cuando 5 termina\n"
                "  5FC+3     → FC con 3 días de lag\n"
                "  5CC       → CC: inicia cuando 5 inicia\n"
                "  5CC+2     → CC con 2 días de lag\n"
                "  5FF       → FF: termina cuando 5 termina\n"
                "  5FF-1     → FF con 1 día de lead\n"
                "  5CF       → CF: termina cuando 5 inicia (raro)\n"
                "  5+50%     → arranca cuando 5 lleva 50% completado\n"
                "  3, 4CC    → varias separadas por coma\n"
                "  (también se aceptan FS/SS/SF en inglés)"
            )
        th = self.tbl.horizontalHeader()
        # Todas las columnas redimensionables a mano (como MS Project). La
        # Descripción arranca en modo Stretch para llenar el panel al abrir y
        # se «congela» a Interactive tras el primer showEvent (ver
        # _init_columns_resizable), de modo que también se pueda redimensionar.
        for c in (0, 1, 3, 4, 5, 6, 7):
            th.setSectionResizeMode(c, QHeaderView.Interactive)
        th.setSectionResizeMode(2, QHeaderView.Stretch)
        th.setStretchLastSection(False)
        th.setMinimumSectionSize(28)
        th.setCascadingSectionResizes(False)
        # Clic derecho sobre el encabezado → mostrar/ocultar columnas.
        th.setContextMenuPolicy(Qt.CustomContextMenu)
        th.customContextMenuRequested.connect(self._on_header_ctx_menu)
        self.tbl.setColumnWidth(0, 46)   # # — hasta 3-4 dígitos + padding
        self.tbl.setColumnWidth(1, 60)
        self.tbl.setColumnWidth(3, 44)   # Días calendario
        self.tbl.setColumnWidth(4, 44)   # Días laborables
        self.tbl.setColumnWidth(5, 80)   # Inicio: fecha dd/mm/yyyy
        self.tbl.setColumnWidth(6, 80)   # Fin:    fecha dd/mm/yyyy
        self.tbl.setColumnWidth(7, 80)
        self.tbl.verticalHeader().setVisible(False)
        self.tbl.verticalHeader().setDefaultSectionSize(self.ROW_H)
        # Wrap de texto en celdas — clave para que las descripciones largas
        # se vean en 2 líneas sin elidir
        self.tbl.setWordWrap(True)
        self.tbl.setTextElideMode(Qt.ElideRight)
        # Alinear el alto del header de la tabla con el header del Gantt
        # (H_HDR_TOP + H_HDR_BOT) para que la primera partida de la tabla
        # coincida con la primera fila del Gantt
        self.tbl.horizontalHeader().setFixedHeight(self.H_HDR_TOP + self.H_HDR_BOT)
        # Scroll en píxeles (no por fila) — necesario para sincronizar con el
        # QGraphicsView del Gantt que también es pixel-based
        self.tbl.setVerticalScrollMode(QAbstractItemView.ScrollPerPixel)
        self.tbl.setEditTriggers(
            QAbstractItemView.DoubleClicked | QAbstractItemView.SelectedClicked
        )
        self.tbl.setSelectionBehavior(QAbstractItemView.SelectRows)
        # Estilo unificado con Cronograma Valorizado y Adquisición de Insumos:
        # bordes silver, gridlines slate-200 sutiles, header slate-500 blanco,
        # padding 3×8. Consistencia visual entre las 3 vistas.
        self.tbl.setShowGrid(True)
        self.tbl.setStyleSheet(f"""
            QTableWidget {{ border:1px solid {SILVER_300}; font-size:11px;
                            gridline-color:#E0E5EC; background:white; }}
            QTableWidget::item {{ padding:3px 8px; }}
            QTableWidget::item:selected {{ background:#FEF0E0; color:#273445; }}
            QHeaderView::section {{
                background:{SLATE_500}; color:white; font-size:10px;
                font-weight:700; padding:4px 6px; border:none;
            }}
        """)
        # Bg delegate para que setBackground por item se pinte siempre
        # (la QSS global del proyecto a veces lo ignora).
        self.tbl.setItemDelegate(_BgFillDelegate(self.tbl))
        self.tbl.cellChanged.connect(self._on_cell_changed)
        self.tbl.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tbl.customContextMenuRequested.connect(self._on_table_ctx_menu)
        sp.addWidget(self.tbl)

        # Gantt derecha — header sticky arriba + barras abajo
        self.scene = QGraphicsScene()
        self.scene.selectionChanged.connect(self._on_selection_changed)
        self.view = _GanttView(self.scene, self)
        self.view.setRenderHint(QPainter.Antialiasing)
        # Anclar la escena ARRIBA-IZQUIERDA. Por defecto QGraphicsView CENTRA la
        # escena cuando es más baja que el viewport (proyecto corto + ventana
        # maximizada) → las barras se bajan media pantalla mientras la tabla
        # ancla sus filas arriba, y dejan de coincidir. Con AlignTop el Gantt
        # ancla igual que la tabla y las barras calzan a cualquier tamaño de
        # ventana. `[[project_gantt_align_top]]`
        self.view.setAlignment(Qt.AlignLeft | Qt.AlignTop)
        self.view.setStyleSheet(f"background:white; border:none;")
        self.view.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        # Una sola barra vertical: la de la tabla. La de la vista de gráficos se
        # oculta para que, al hacer zoom (que muestra/oculta el scroll HORIZONTAL
        # y cambia la altura de su viewport), su scroll vertical no se
        # desincronice de la tabla y las barras dejen de coincidir con sus filas.
        # La rueda sobre las barras sigue moviendo el scrollbar (oculto) y la
        # sincronización bidireccional lo replica en la tabla. Además, igualar la
        # vista al ancho del header (que ya no tiene barra vertical) mantiene las
        # columnas de días alineadas con las barras.
        self.view.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.view.setDragMode(QGraphicsView.NoDrag)

        self.hdr_scene = QGraphicsScene()
        self.hdr_view = _GanttHeaderView(self.hdr_scene, self)
        self.hdr_view.setFixedHeight(self.H_HDR_TOP + self.H_HDR_BOT + 1)
        # Mismo anclaje que la vista principal: si la escala de días es corta y
        # la ventana ancha, evita que el header centre y descuadre los días
        # respecto a las barras. `[[project_gantt_align_top]]`
        self.hdr_view.setAlignment(Qt.AlignLeft | Qt.AlignTop)

        # Sincronizar scroll horizontal del header con la vista principal
        self.view.horizontalScrollBar().valueChanged.connect(
            self.hdr_view.horizontalScrollBar().setValue
        )

        right_w = QWidget()
        right_vl = QVBoxLayout(right_w)
        right_vl.setContentsMargins(0, 0, 0, 0)
        right_vl.setSpacing(0)
        right_vl.addWidget(self.hdr_view)
        right_vl.addWidget(self.view, stretch=1)
        sp.addWidget(right_w)

        # Anchos mínimos para que el handle siempre tenga margen de movimiento
        # a ambos lados (evita que quede "trabado" en un extremo).
        self.tbl.setMinimumWidth(180)
        right_w.setMinimumWidth(160)
        sp.setStretchFactor(0, 40)
        sp.setStretchFactor(1, 60)
        sp.setSizes([400, 700])

        # Sincronizar scroll vertical
        self.tbl.verticalScrollBar().valueChanged.connect(
            self.view.verticalScrollBar().setValue
        )
        self.view.verticalScrollBar().valueChanged.connect(
            self.tbl.verticalScrollBar().setValue
        )

        vl.addWidget(sp, stretch=1)
        # Footer de hints a lo ancho (debajo de tabla + barras). Antes vivía solo
        # en la columna derecha y achicaba el viewport de las barras respecto al
        # de la tabla, lo que recortaba la última fila o dejaba scroll sobrante.
        vl.addWidget(self._build_footer_hints())

    def _on_escala(self, idx: int):
        # 0=Auto, 1=Días, 2=Semanas, 3=Meses
        self._escala_unit = ['auto', 'dia', 'semana', 'mes'][idx]
        from PySide6.QtCore import QSettings as _QS
        _QS("ingePresupuestos", "layout").setValue("gantt_escala_unit",
                                                      self._escala_unit)
        # Al fijar una unidad, llevar el zoom a un ancho cómodo para ella
        # (en "Auto" se respeta el zoom actual).
        defecto = {1: 16, 2: 6, 3: 2}.get(idx)
        if defecto is not None and defecto != self.DAY_W:
            self.DAY_W = defecto
            if hasattr(self, '_sl_zoom'):
                self._sl_zoom.blockSignals(True)
                self._sl_zoom.setValue(self.DAY_W)
                self._sl_zoom.blockSignals(False)
            if hasattr(self, '_lbl_zoom'):
                self._lbl_zoom.setText(f"{self.DAY_W}px")
        self._render_gantt()

    def _zoom(self, delta: int):
        self._set_day_w(self.DAY_W + delta)

    def _set_day_w(self, value: int):
        """Cambia DAY_W (clamp 2..40), sincroniza slider/label y re-renderiza."""
        new = max(2, min(40, int(value)))
        if new == self.DAY_W:
            return
        self.DAY_W = new
        if hasattr(self, '_lbl_zoom'):
            self._lbl_zoom.setText(f"{self.DAY_W}px")
        if hasattr(self, '_sl_zoom'):
            # Evitar feedback loop
            self._sl_zoom.blockSignals(True)
            self._sl_zoom.setValue(self.DAY_W)
            self._sl_zoom.blockSignals(False)
        self._render_gantt()

    def _on_zoom_slider(self, value: int):
        self._set_day_w(value)

    def _on_toggle_slack(self, on: bool):
        self._show_slack = bool(on)
        from PySide6.QtCore import QSettings as _QS
        _QS("ingePresupuestos", "layout").setValue("gantt_show_slack",
                                                      self._show_slack)
        self._render_gantt()

    # ── Columnas de la tabla: redimensionables a mano + persistentes ─────
    def showEvent(self, event):
        super().showEvent(event)
        if not getattr(self, '_cols_ready', False):
            self._cols_ready = True
            # Diferido para que el splitter ya haya dado ancho a la tabla y la
            # columna Descripción (Stretch) tenga su ancho real al congelarla.
            from PySide6.QtCore import QTimer as _QT
            _QT.singleShot(0, self._init_columns_resizable)

    def _init_columns_resizable(self):
        """Pasa Descripción de Stretch a Interactive conservando su ancho, y
        restaura los anchos que el usuario haya guardado antes."""
        th = self.tbl.horizontalHeader()
        w2 = self.tbl.columnWidth(2)
        th.setSectionResizeMode(2, QHeaderView.Interactive)
        self.tbl.setColumnWidth(2, w2 if w2 >= 80 else 200)
        self._restore_col_widths()
        self._restore_col_hidden()
        th.sectionResized.connect(self._on_col_resized)

    def _restore_col_widths(self):
        from PySide6.QtCore import QSettings as _QS
        raw = _QS("ingePresupuestos", "layout").value("gantt_tbl_col_widths", "")
        if not raw:
            return
        try:
            widths = [int(x) for x in str(raw).split(",")]
        except (ValueError, TypeError):
            return
        if len(widths) != self.tbl.columnCount():
            return
        self._restoring_cols = True
        for c, w in enumerate(widths):
            if w >= 24:
                self.tbl.setColumnWidth(c, w)
        self._restoring_cols = False

    def _on_col_resized(self, *_):
        if getattr(self, '_restoring_cols', False):
            return
        from PySide6.QtCore import QSettings as _QS
        widths = ",".join(str(self.tbl.columnWidth(c))
                            for c in range(self.tbl.columnCount()))
        _QS("ingePresupuestos", "layout").setValue("gantt_tbl_col_widths", widths)

    # ── Mostrar / ocultar columnas (clic derecho en el encabezado) ───────
    def _on_header_ctx_menu(self, pos):
        from PySide6.QtWidgets import QMenu
        menu = QMenu(self)
        hdr = menu.addAction("Mostrar / ocultar columnas")
        hdr.setEnabled(False)
        menu.addSeparator()
        for c in range(self.tbl.columnCount()):
            item = self.tbl.horizontalHeaderItem(c)
            label = (item.text() if item else f"Col {c}").replace("\n", " ").strip()
            act = menu.addAction(label)
            act.setCheckable(True)
            act.setChecked(not self.tbl.isColumnHidden(c))
            if c == 2:   # Descripción siempre visible (es la etiqueta principal)
                act.setEnabled(False)
            act.toggled.connect(
                lambda checked, col=c: self._set_col_visible(col, checked))
        menu.exec(self.tbl.horizontalHeader().mapToGlobal(pos))

    def _set_col_visible(self, col, visible):
        self.tbl.setColumnHidden(col, not visible)
        self._save_col_hidden()

    def _save_col_hidden(self):
        from PySide6.QtCore import QSettings as _QS
        hidden = ",".join(str(c) for c in range(self.tbl.columnCount())
                          if self.tbl.isColumnHidden(c))
        _QS("ingePresupuestos", "layout").setValue("gantt_tbl_hidden_cols", hidden)

    def _restore_col_hidden(self):
        from PySide6.QtCore import QSettings as _QS
        raw = _QS("ingePresupuestos", "layout").value("gantt_tbl_hidden_cols", "")
        if not raw:
            return
        try:
            hidden = {int(x) for x in str(raw).split(",") if x.strip() != ""}
        except (ValueError, TypeError):
            return
        for c in range(self.tbl.columnCount()):
            if c == 2:
                continue   # Descripción nunca se oculta
            self.tbl.setColumnHidden(c, c in hidden)

    def _zoom_fit(self):
        """Calcula DAY_W para que el proyecto completo entre en el viewport."""
        if not hasattr(self._cv, '_tasks'):
            return
        plazo = self._cv._proy.get('plazo') or 0
        tasks = self._cv._tasks
        max_ef = max((t['EF'] for t in tasks.values() if t['EF'] > 0),
                      default=plazo)
        n_dias = max(max_ef, plazo, 30)
        # Ancho útil del viewport (en px), restando barra de scroll y un margen
        vw = self.view.viewport().width() if hasattr(self, 'view') else 800
        vw = max(200, vw - 20)
        new_w = max(2, min(40, vw // max(1, n_dias)))
        self._set_day_w(int(new_w))

    def _build_footer_hints(self) -> QWidget:
        """Footer pequeño debajo del Gantt con un hint rotativo (predecesoras,
        atajos, IA, etc.) + botón pequeño "Ayúdame" que abre el chat de Tuxia.
        Reemplaza los popups intrusivos de Tuxia."""
        from PySide6.QtCore import QTimer
        bar = QFrame()
        bar.setFixedHeight(24)
        bar.setStyleSheet(
            f"QFrame {{ background:{SILVER_100};"
            f" border-top:1px solid {SILVER_300}; }}"
        )
        hl = QHBoxLayout(bar)
        hl.setContentsMargins(10, 0, 8, 0)
        hl.setSpacing(8)

        self._hints_idx = 0
        self._hints_lista = [
            "💡  <b>Arrastra una barra hacia otra</b> (hacia arriba/abajo) para "
            "crear una dependencia, como en MS Project. Sueltas al inicio = FC, "
            "al fin = FF, en el tercio central = CC+50% (el sucesor comienza "
            "cuando la pred. lleva el 50%)",
            "💡  <b>Arrastra</b> una flecha horizontalmente para ajustar el "
            "lag en días (o el % si la vinculaste a mitad de la pred.)",
            "💡  Haz <b>clic</b> sobre una flecha para seleccionarla y pulsa "
            "<b>Supr</b> para eliminar (Ctrl+clic agrega varias a la selección)",
            "💡  Haz <b>clic</b> sobre una <b>barra</b> y se resaltan en azul "
            "sus predecesoras y sucesoras",
            "💡  La <b style='color:#94A3B8'>barra gris</b> al final de una "
            "tarea es su <b>holgura</b> — días que puede atrasarse sin "
            "afectar el fin del proyecto",
            "💡  Clic derecho sobre una <b>flecha</b> para cambiar el tipo "
            "(FC/CC/FF/CF), editar lag/% o eliminarla",
            "💡  Predecesoras: <b>5</b> = FC · <b>5FC+3</b> = FC+lag · "
            "<b>5CC+2</b> = CC · <b>5FF-1</b> = FF · <b>5CF</b> = CF",
            "💡  <b>Ctrl + rueda</b> sobre el Gantt para hacer zoom horizontal "
            "en la línea de tiempo",
            "💡  Las barras <b style='color:#C6262E'>rojas</b> son la ruta "
            "crítica — su atraso retrasa el proyecto",
            "💡  Auto-programar ▾ → <b>Con IA</b> sugiere dependencias "
            "realistas según las descripciones",
            "💡  <b>F1</b> abre la lista completa de atajos de teclado",
            "💡  Clic derecho sobre una barra para dividirla, marcarla como "
            "hito o cambiarle el color",
        ]
        self.lbl_hint = QLabel(self._hints_lista[0])
        self.lbl_hint.setStyleSheet(
            f"color:{SLATE_500}; font-size:11px; background:transparent;"
        )
        self.lbl_hint.setTextFormat(Qt.RichText)
        hl.addWidget(self.lbl_hint, stretch=1)

        from utils.icons import icon as load_icon
        btn_help = QPushButton(" Ayúdame")
        btn_help.setIcon(load_icon("tuxia"))
        btn_help.setIconSize(QSize(16, 16))
        btn_help.setCursor(Qt.PointingHandCursor)
        btn_help.setFixedHeight(20)
        btn_help.setStyleSheet(
            "QPushButton { background:#A56DE2; color:white; border:none;"
            " border-radius:4px; padding:2px 10px; font-size:10px;"
            " font-weight:600; min-height:0; }"
            "QPushButton:hover { background:#7E3FCB; }"
        )
        btn_help.clicked.connect(self._abrir_chat_tuxia)
        hl.addWidget(btn_help)

        # Rotar el hint cada ~12 segundos
        self._hints_timer = QTimer(self)
        self._hints_timer.setInterval(12000)
        self._hints_timer.timeout.connect(self._rotar_hint)
        self._hints_timer.start()
        return bar

    def _rotar_hint(self):
        if not getattr(self, '_hints_lista', None):
            return
        self._hints_idx = (self._hints_idx + 1) % len(self._hints_lista)
        self.lbl_hint.setText(self._hints_lista[self._hints_idx])

    def _abrir_chat_tuxia(self):
        """Abre el chat de Tuxia (delegando al ProyectoView padre)."""
        # Subir hasta el ProyectoView para invocar su método
        w = self
        while w is not None:
            if hasattr(w, '_abrir_asistente_ia'):
                w._abrir_asistente_ia()
                return
            w = w.parent()

    def _on_fecha_inicio(self, qd):
        """Guarda fecha_inicio en proyecto y refresca tabla + Gantt."""
        nueva = qd.toString("yyyy-MM-dd")
        self._save_proy_field('fecha_inicio', nueva)
        self._cv._proy['fecha_inicio'] = nueva
        self._llenar_tabla()
        self._render_gantt()

    def _abrir_feriados(self):
        """Diálogo para editar feriados + toggle 'Saltar no laborables'."""
        from PySide6.QtWidgets import (QDialog, QVBoxLayout, QLabel, QPlainTextEdit,
                                          QDialogButtonBox, QCheckBox)
        from PySide6.QtGui import QPalette, QColor
        # No padre — para que el dialog no herede la paleta oscura del topbar
        dlg = QDialog()
        dlg.setWindowTitle("Feriados / días no laborables")
        dlg.resize(440, 380)
        dlg.setAttribute(Qt.WA_StyledBackground, True)
        dlg.setWindowModality(Qt.ApplicationModal)
        # Forzar paleta clara (sobrepone cualquier theme oscuro del sistema)
        pal = dlg.palette()
        pal.setColor(QPalette.Window, QColor("#FFFFFF"))
        pal.setColor(QPalette.WindowText, QColor("#1F2A38"))
        pal.setColor(QPalette.Base, QColor("#FFFFFF"))
        pal.setColor(QPalette.Text, QColor("#1F2A38"))
        pal.setColor(QPalette.Button, QColor("#F8F9FA"))
        pal.setColor(QPalette.ButtonText, QColor("#1F2A38"))
        pal.setColor(QPalette.Highlight, QColor("#F37329"))
        pal.setColor(QPalette.HighlightedText, QColor("#FFFFFF"))
        dlg.setPalette(pal)
        dlg.setStyleSheet(
            "QDialog { background:#FFFFFF; color:#1F2A38; }"
            "QLabel { color:#1F2A38; font-size:11px; background:transparent; }"
            "QPlainTextEdit { background:#FFFFFF; color:#1F2A38;"
            " border:1px solid #CBD5E1; border-radius:4px; padding:4px;"
            " font-size:11px; selection-background-color:#F37329;"
            " selection-color:#FFFFFF; }"
            "QCheckBox { color:#1F2A38; font-size:11px; padding:4px 0;"
            " background:transparent; }"
            "QCheckBox::indicator { width:14px; height:14px;"
            " border:1px solid #CBD5E1; background:#FFFFFF; border-radius:2px; }"
            "QCheckBox::indicator:checked { background:#F37329;"
            " border-color:#C0621A; image:none; }"
            "QDialogButtonBox QPushButton { min-width:88px; padding:5px 12px;"
            " font-size:11px; background:#F8F9FA; color:#1F2A38;"
            " border:1px solid #CBD5E1; border-radius:4px; }"
            "QDialogButtonBox QPushButton:hover { background:#FEF5EB;"
            " color:#C0621A; border-color:#F37329; }"
            "QDialogButtonBox QPushButton:default { background:#F37329;"
            " color:#FFFFFF; border-color:#C0621A; }"
        )
        vl = QVBoxLayout(dlg)

        # ── Toggle: saltar no laborables en cálculo del CPM ──
        chk_saltar = QCheckBox(
            "Saltar domingos y feriados\n"
            "Las tareas no inician ni terminan en estos días; la duración se "
            "cuenta solo en días laborables."
        )
        chk_saltar.setChecked(
            int(self._cv._proy.get('salta_no_laborables', 1) or 0) == 1
        )
        vl.addWidget(chk_saltar)

        vl.addWidget(QLabel(
            "\nFeriados — una fecha por línea o separadas por coma.\n"
            "Formato: AAAA-MM-DD. Ej: 2026-07-28, 2026-07-29 (Fiestas Patrias)"
        ))
        ed = QPlainTextEdit()
        actual = (self._cv._proy.get('feriados') or '').replace(',', '\n').strip()
        ed.setPlainText(actual)
        vl.addWidget(ed)

        # Botón: cargar feriados con IA (Perú, según años del proyecto)
        from PySide6.QtWidgets import QHBoxLayout, QPushButton, QApplication
        ia_row = QHBoxLayout()
        ia_row.addStretch()
        btn_ia = QPushButton("✨ Cargar feriados con IA")
        btn_ia.setCursor(Qt.PointingHandCursor)
        btn_ia.setStyleSheet(
            "QPushButton { background:#FEF5EB; color:#C0621A;"
            " border:1px solid #F37329; border-radius:4px;"
            " padding:5px 12px; font-size:11px; font-weight:600; }"
            "QPushButton:hover { background:#F37329; color:#FFFFFF; }"
        )

        def _cargar_ia():
            # Determinar años desde fecha_inicio + plazo, o usar año actual
            anios = set()
            f_ini = self._project_start()
            if f_ini:
                anios.add(f_ini.year)
                plazo = int(self._cv._proy.get('plazo') or 0)
                if plazo > 0:
                    f_fin = f_ini + timedelta(days=plazo)
                    anios.add(f_fin.year)
            if not anios:
                anios.add(datetime.now().year)
            QApplication.setOverrideCursor(Qt.WaitCursor)
            try:
                fechas, err = cargar_feriados_ia(sorted(anios))
            finally:
                QApplication.restoreOverrideCursor()
            if err:
                from PySide6.QtWidgets import QMessageBox
                QMessageBox.warning(dlg, "Cargar feriados con IA",
                    f"No se pudo obtener la lista:\n\n{err}")
                return
            # Combinar con las que ya están escritas (no perder ediciones)
            import re as _re
            existentes = set(_re.findall(r'\d{4}-\d{2}-\d{2}', ed.toPlainText()))
            existentes.update(fechas or [])
            ed.setPlainText('\n'.join(sorted(existentes)))

        btn_ia.clicked.connect(_cargar_ia)
        ia_row.addWidget(btn_ia)
        vl.addLayout(ia_row)
        bb = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)
        bb.button(QDialogButtonBox.Save).setText("Guardar")
        bb.button(QDialogButtonBox.Cancel).setText("Cancelar")
        bb.accepted.connect(dlg.accept)
        bb.rejected.connect(dlg.reject)
        vl.addWidget(bb)
        if dlg.exec() != QDialog.Accepted:
            return
        # Toggle saltar no laborables
        saltar = 1 if chk_saltar.isChecked() else 0
        self._save_proy_field('salta_no_laborables', saltar)
        self._cv._proy['salta_no_laborables'] = saltar
        # Normalizar: extraer todas las fechas YYYY-MM-DD válidas
        import re
        texto = ed.toPlainText()
        candidatos = re.findall(r'\d{4}-\d{2}-\d{2}', texto)
        validos = []
        for c in candidatos:
            try:
                datetime.strptime(c, '%Y-%m-%d')
                validos.append(c)
            except Exception:
                pass
        nuevo = ','.join(sorted(set(validos)))
        self._save_proy_field('feriados', nuevo)
        self._cv._proy['feriados'] = nuevo
        # Recalcular CPM y re-renderizar (afecta fechas si cambió el toggle)
        self._cv._calcular_cpm()
        self._llenar_tabla()
        self._render_gantt()

    def _save_proy_field(self, field: str, value):
        """Persiste un campo del proyecto a la BD."""
        from core.database import get_db
        conn = get_db()
        conn.execute(
            f"UPDATE proyectos SET {field}=? WHERE id=?",
            (value, self._cv.pid)
        )
        conn.commit()
        conn.close()

    # ──────────────────────────────────────────────────────────────────────
    # Carga + render
    # ──────────────────────────────────────────────────────────────────────

    def cargar(self):
        if not hasattr(self._cv, '_partidas'):
            return
        self._llenar_tabla()
        self._render_gantt()
        # Auto-fit la primera vez que se abre el Gantt: con la escala
        # default "Meses" (DAY_W=2) un proyecto largo igual se desborda,
        # y Marco no quiere tener que dar clic en "Ajustar" cada vez.
        # Diferimos 1 frame para que el viewport ya tenga ancho real.
        if self._fit_pendiente:
            self._fit_pendiente = False
            QTimer.singleShot(0, self._zoom_fit)

    def _llenar_tabla(self):
        from core.cronograma import contar_laborables, formatear_pred_es
        self.tbl.blockSignals(True)
        self.tbl.setRowCount(0)
        filas = self._cv.filas_con_hitos()   # proyecto + inicio + (sub+partidas)* + fin
        cmap = self._cv._cron_map
        tasks = self._cv._tasks
        plazo = self._cv._proy.get('plazo') or 0
        non_working = self._cv._non_working_set(plazo + 365)
        BLUE, GREEN = QColor("#3689E6"), QColor("#2EA043")
        # Estilo consistente con el árbol de Presupuesto:
        #  · Títulos / cabeceras (proyecto, subpresupuesto, título nivel 1)
        #    → banda más oscura.
        #  · Subtítulos (título nivel ≥2) y partidas → zebra; subtítulo en arándano.
        TITLE_BG = QColor("#E2E8F0")
        ZEBRA_BG = QColor("#F6F8FB")
        TITLE_FG = QColor("#B71C1C")   # rojo títulos (igual que Presupuesto)
        SUBT_FG  = QColor("#0D52BF")   # arándano (Blueberry 700)
        zebra_i = 0
        for p in filas:
            r = self.tbl.rowCount()
            self.tbl.insertRow(r)
            es_titulo = p['es_titulo']
            virtual = p.get('_virtual')
            cd = cmap.get(p['id'], {})
            t = tasks.get(p['id'], {})
            # "#" = posición 1-based en la lista (coincide con numerar_filas y
            # con las predecesoras).
            num_str = str(r + 1)

            if virtual:
                # Resumen del proyecto / subpresupuesto (span) o hitos (0 días).
                a = p.get('_ES') or 0
                b = p.get('_EF') or a
                if virtual in ('proyecto', 'subppto'):
                    cal_str = str(b - a + 1) if b >= a else ''
                    lab_str = str(contar_laborables(a, b, non_working))
                else:
                    cal_str = lab_str = "0"
                es_str = self._day_to_date_str(a)
                ef_str = self._day_to_date_str(b)
                pred_str = ""
            elif es_titulo:
                # Título: span real ef_max - es_min + 1 (considera paralelismos).
                es_min, ef_max = self._summary_for_title(p)
                if es_min and ef_max:
                    cal_str = str(ef_max - es_min + 1)
                    lab_str = str(contar_laborables(es_min, ef_max, non_working))
                    es_str = self._day_to_date_str(es_min)
                    ef_str = self._day_to_date_str(ef_max)
                else:
                    cal_str = lab_str = es_str = ef_str = ''
                pred_str = ''
            else:
                es_v, ef_v = t.get('ES'), t.get('EF')
                lab_str = str(cd.get('duracion', '') or '')
                cal_str = (str(ef_v - es_v + 1)
                           if (es_v and ef_v and ef_v >= es_v) else lab_str)
                es_str = self._day_to_date_str(es_v)
                ef_str = self._day_to_date_str(ef_v)
                pred_str = formatear_pred_es(cd.get('predecesoras', '') or '')

            editable = not (es_titulo or virtual)
            desc_txt = p['descripcion'] or ''
            if virtual in ('inicio', 'fin'):
                desc_txt = "◆ " + desc_txt
            cells = [
                (num_str,                Qt.AlignCenter,  False),
                (p.get('item', '') or '', Qt.AlignLeft,   False),
                (desc_txt,                Qt.AlignLeft,   False),
                (cal_str,                 Qt.AlignCenter,  False),     # Días cal: derivado
                (lab_str,                 Qt.AlignCenter,  editable),  # Días lab: editable
                (es_str,                  Qt.AlignCenter,  editable),
                (ef_str,                  Qt.AlignCenter,  editable),
                (pred_str,                Qt.AlignCenter,  editable),
            ]

            nivel_p   = p.get('nivel') or 1
            # Cabecera: proyecto/subpresupuesto (virtual) o título de nivel 1.
            es_header = (virtual in ('proyecto', 'subppto')) or \
                        (es_titulo and nivel_p <= 1)
            es_subt   = es_titulo and nivel_p >= 2 and not virtual
            es_hito_v = virtual in ('inicio', 'fin')
            # Zebra para subtítulos + partidas (no cabeceras, no hitos virtuales).
            row_bg = None
            if not es_header and not es_hito_v:
                zebra_i += 1
                if zebra_i % 2 == 0:
                    row_bg = ZEBRA_BG

            niv = nivel_p - 1
            for c, (txt, align, ed) in enumerate(cells):
                # Indentar descripción según nivel jerárquico
                if c == 2 and niv > 0 and not virtual:
                    txt = ('  ' * niv) + txt
                it = QTableWidgetItem(txt)
                it.setTextAlignment(align | Qt.AlignVCenter)
                if not ed:
                    it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                if es_header:
                    it.setBackground(QBrush(TITLE_BG))
                    f = QFont(); f.setBold(True)
                    # Subrayar solo títulos reales de nivel 1 (Ítem/Descripción).
                    if es_titulo and nivel_p == 1 and c in (1, 2):
                        f.setUnderline(True)
                    it.setFont(f)
                    it.setForeground(QBrush(TITLE_FG))
                elif es_hito_v:
                    f = QFont(); f.setBold(True); it.setFont(f)
                    it.setForeground(QBrush(BLUE if virtual == 'inicio' else GREEN))
                    if row_bg is not None:
                        it.setBackground(QBrush(row_bg))
                else:
                    # Subtítulo o partida — zebra; subtítulo en arándano + negrita.
                    if row_bg is not None:
                        it.setBackground(QBrush(row_bg))
                    if es_subt:
                        f = QFont(); f.setBold(True); it.setFont(f)
                        it.setForeground(QBrush(SUBT_FG))
                # Ruta crítica: solo resaltar la columna "Ítem" (c==1) en rojo
                # bold; el resto de la fila mantiene el color por defecto.
                if not es_titulo and not virtual and t.get('critical') and c == 1:
                    it.setForeground(QBrush(QColor(RED_500)))
                    f = QFont(); f.setBold(True); it.setFont(f)
                self.tbl.setItem(r, c, it)
            self.tbl.item(r, 0).setData(Qt.UserRole, p['id'])
            self.tbl.item(r, 0).setData(Qt.UserRole + 1, es_titulo)

        self.tbl.blockSignals(False)

    def _on_cell_changed(self, row, col):
        # Editables: 4=Días lab., 5=Inicio, 6=Fin, 7=Pred (col 3=Días cal. es derivada)
        if col not in (4, 5, 6, 7):
            return
        id_cell = self.tbl.item(row, 0)
        if not id_cell:
            return
        pid = id_cell.data(Qt.UserRole)
        es_t = id_cell.data(Qt.UserRole + 1)
        if not pid or pid < 0 or es_t:   # pid<0 → hito virtual Inicio/Fin
            return
        val = self.tbl.item(row, col).text().strip() if self.tbl.item(row, col) else ''
        cmap = self._cv._cron_map
        if pid not in cmap:
            cmap[pid] = {'partida_id': pid, 'duracion': 1, 'inicio_dia': 1,
                          'predecesoras': '', 'es_hito': 0, 'segmentos': ''}
        if col == 4:
            try: cmap[pid]['duracion'] = max(0, int(val))
            except ValueError: cmap[pid]['duracion'] = 1
        elif col == 5:
            # Inicio manual: acepta fecha (dd/mm/yyyy) o número de día corrido
            ini = self._date_str_to_day(val)
            if ini is not None:
                cmap[pid]['inicio_dia'] = ini
        elif col == 6:
            # Fin manual: convertir fecha a día corrido y derivar duración
            fin = self._date_str_to_day(val)
            if fin is not None:
                ini = int(cmap[pid].get('inicio_dia', 1) or 1)
                if fin >= ini:
                    cmap[pid]['duracion'] = fin - ini + 1
        elif col == 7:
            self._push_undo_dep()   # predecesoras editadas a mano → deshacible
            cmap[pid]['predecesoras'] = val
        # Recalcular CPM y refrescar Gantt + tabla
        self._cv._calcular_cpm()
        # Repintar barras y la fila editada para mostrar el nuevo ES/EF
        self._render_gantt()
        self._refrescar_filas_es_ef()

    def _refrescar_filas_es_ef(self):
        """Después de un cambio que afecta CPM, actualizar las celdas Días cal.,
        Inicio y Fin de todas las partidas con los valores recalculados (sin
        disparar otro cellChanged)."""
        self.tbl.blockSignals(True)
        try:
            tasks = self._cv._tasks
            for row in range(self.tbl.rowCount()):
                id_cell = self.tbl.item(row, 0)
                if not id_cell:
                    continue
                pid = id_cell.data(Qt.UserRole)
                es_t = id_cell.data(Qt.UserRole + 1)
                if not pid or pid < 0 or es_t:   # pid<0 → hito virtual
                    continue
                t = tasks.get(pid, {})
                es_v = t.get('ES') or 0
                ef_v = t.get('EF') or 0
                cal = (ef_v - es_v + 1) if (es_v and ef_v and ef_v >= es_v) else 0
                it_cal = self.tbl.item(row, 3)
                it_es = self.tbl.item(row, 5)
                it_ef = self.tbl.item(row, 6)
                if it_cal is not None and cal:
                    it_cal.setText(str(cal))
                if it_es is not None:
                    it_es.setText(self._day_to_date_str(es_v))
                if it_ef is not None:
                    it_ef.setText(self._day_to_date_str(ef_v))
        finally:
            self.tbl.blockSignals(False)

    def volcar_a_map(self):
        pass

    # ──────────────────────────────────────────────────────────────────────
    # Menú contextual de la tabla (dividir, hito)
    # ──────────────────────────────────────────────────────────────────────

    def _on_table_ctx_menu(self, pos):
        sel = self.tbl.currentRow()
        if sel < 0:
            return
        id_cell = self.tbl.item(sel, 0)
        if not id_cell:
            return
        pid = id_cell.data(Qt.UserRole)
        es_t = id_cell.data(Qt.UserRole + 1)
        if not pid or pid < 0 or es_t:   # pid<0 → hito virtual Inicio/Fin
            return
        global_pos = self.tbl.viewport().mapToGlobal(pos)
        self._show_partida_ctx_menu(pid, global_pos, parent=self.tbl)

    def _show_partida_ctx_menu(self, pid: int, global_pos, parent=None):
        """Muestra el menú contextual de una partida (Dividir / Unir / Hito).
        Reutilizable desde la tabla y desde una barra del Gantt."""
        cmap = self._cv._cron_map
        cd = cmap.get(pid, {})
        es_hito = int(cd.get('es_hito', 0) or 0)
        has_segs = bool(cd.get('segmentos', '') or '')

        menu = QMenu(parent or self)

        act_div = menu.addAction("✂ Dividir tarea")
        act_div.triggered.connect(lambda: self._dividir_tarea(pid))
        if has_segs:
            act_unir = menu.addAction("⊞ Unir segmentos")
            act_unir.triggered.connect(lambda: self._unir_segmentos(pid))
        menu.addSeparator()
        for label, val in [
            ("◯ Tarea normal", 0),
            ("◆ Marcar como hito (sin duración)", 1),
            ("▶ Marcar inicio de fase (mantiene duración)", 2),
            ("⏹ Marcar fin de fase (mantiene duración)", 3),
        ]:
            act = menu.addAction(label + (" ✓" if es_hito == val else ""))
            act.triggered.connect(lambda _=False, v=val: self._set_hito(pid, v))
        menu.addSeparator()
        # Color personalizado de la barra
        cur_color = (cd.get('color', '') or '').strip()
        marca = f" ({cur_color})" if cur_color else ""
        act_col = menu.addAction(f"🎨 Color de barra…{marca}")
        act_col.triggered.connect(lambda: self._elegir_color_barra(pid))
        if cur_color:
            act_clr = menu.addAction("✕ Quitar color personalizado")
            act_clr.triggered.connect(lambda: self._set_color_barra(pid, ''))
        menu.exec(global_pos)

    def _elegir_color_barra(self, pid: int):
        from PySide6.QtWidgets import QColorDialog
        cmap = self._cv._cron_map
        actual = (cmap.get(pid, {}).get('color', '') or '').strip()
        ini = QColor(actual) if actual else QColor("#F37329")
        col = QColorDialog.getColor(ini, self, "Color de la barra")
        if col.isValid():
            self._set_color_barra(pid, col.name())

    def _set_color_barra(self, pid: int, hex_color: str):
        cmap = self._cv._cron_map
        if pid not in cmap:
            cmap[pid] = {'partida_id': pid, 'duracion': 1, 'inicio_dia': 1,
                          'predecesoras': '', 'es_hito': 0, 'segmentos': '',
                          'color': ''}
        cmap[pid]['color'] = hex_color or ''
        self._cv._guardar_a_db()
        self._render_gantt()

    def _set_hito(self, pid: int, valor: int):
        cmap = self._cv._cron_map
        if pid not in cmap:
            cmap[pid] = {'partida_id': pid, 'duracion': 1, 'inicio_dia': 1,
                          'predecesoras': '', 'es_hito': 0, 'segmentos': ''}
        cmap[pid]['es_hito'] = valor
        # valor==1 → hito puro (sin duración)
        # valor==2 o 3 → marcador de inicio/fin de fase, mantiene la duración
        if valor == 1:
            cmap[pid]['duracion'] = 0
        elif (cmap[pid].get('duracion') or 0) == 0:
            cmap[pid]['duracion'] = 1
        self._cv._calcular_cpm()
        self._llenar_tabla()
        self._render_gantt()

    def _dividir_tarea(self, pid: int):
        cmap = self._cv._cron_map
        if pid not in cmap:
            return
        cd = cmap[pid]
        ini = int(cd.get('inicio_dia', 1) or 1)
        dur = int(cd.get('duracion', 0) or 0)
        if dur < 2:
            return
        segs_json = cd.get('segmentos', '') or ''
        try:
            segs = _json.loads(segs_json) if segs_json else []
        except Exception:
            segs = []
        if not segs:
            half = max(1, dur // 2)
            segs = [
                {'inicio_dia': ini, 'duracion': half},
                {'inicio_dia': ini + half + 2, 'duracion': dur - half},
            ]
        else:
            last = segs[-1]
            l_dur = last.get('duracion', 0) or 0
            if l_dur < 2:
                return
            half = max(1, l_dur // 2)
            l_ini = last.get('inicio_dia', 1) or 1
            last['duracion'] = half
            segs.append({'inicio_dia': l_ini + half + 2, 'duracion': l_dur - half})
        cd['segmentos'] = _json.dumps(segs)
        first = min(s.get('inicio_dia', 1) for s in segs)
        end = max(s.get('inicio_dia', 1) + s.get('duracion', 0) - 1 for s in segs)
        cd['inicio_dia'] = first
        cd['duracion'] = end - first + 1
        self._cv._calcular_cpm()
        self._render_gantt()

    def _unir_segmentos(self, pid: int):
        cmap = self._cv._cron_map
        if pid in cmap:
            cmap[pid]['segmentos'] = ''
        self._cv._calcular_cpm()
        self._render_gantt()

    # ──────────────────────────────────────────────────────────────────────
    # Exportar PDF
    # ──────────────────────────────────────────────────────────────────────

    # Mapa nombre de papel → QPageSize.PageSizeId
    _PAPEL_ENUM = {
        'A4': 'A4', 'A3': 'A3', 'A2': 'A2', 'A1': 'A1', 'A0': 'A0',
        'Letter': 'Letter', 'Tabloid': 'Tabloid',
    }

    def _aplicar_papel(self, printer, papel):
        """Configura el QPageSize del QPrinter según el label."""
        try:
            from PySide6.QtGui import QPageSize
            id_map = {
                'A4': QPageSize.A4, 'A3': QPageSize.A3, 'A2': QPageSize.A2,
                'A1': QPageSize.A1, 'A0': QPageSize.A0,
                'Letter': QPageSize.Letter, 'Tabloid': QPageSize.Tabloid,
            }
            ps = QPageSize(id_map.get(papel, QPageSize.A3))
            printer.setPageSize(ps)
        except Exception:
            pass

    def _dir_descargas_gantt(self) -> str:
        """Última carpeta usada para exportar el gantt (persistida en QSettings).
        Por defecto, la carpeta de Descargas del usuario."""
        import os
        from PySide6.QtCore import QSettings
        s = QSettings("ingePresupuestos", "exports")
        guardado = s.value("last_dir_gantt", "")
        if guardado and os.path.isdir(guardado):
            return guardado
        for c in (os.path.expanduser("~/Descargas"),
                   os.path.expanduser("~/Downloads"),
                   os.path.expanduser("~")):
            if os.path.isdir(c):
                return c
        return os.getcwd()

    def _save_dir_gantt(self, path: str):
        import os
        from PySide6.QtCore import QSettings
        if not path:
            return
        d = os.path.dirname(path)
        if d and os.path.isdir(d):
            QSettings("ingePresupuestos", "exports").setValue("last_dir_gantt", d)

    def _exportar_pdf(self):
        import os
        # 1. Diálogo previo: opciones
        opts = _DialogExportarGanttPdf.preguntar(
            self, on_preview=self._vista_previa_pdf,
            on_imagen=self._exportar_imagen,
        )
        if opts is None:
            return
        modo = opts['modo']
        orient = opts['orient']
        incluir_pred = opts['incluir_pred']
        escala = opts['escala']
        hojas_x = opts.get('hojas_x', 0)
        papel = opts.get('papel', 'A3')

        sugerido = os.path.join(self._dir_descargas_gantt(),
                                  f"gantt_{self._cv.pid}.pdf")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar Gantt a PDF", sugerido, "PDF (*.pdf)"
        )
        if not path:
            return
        try:
            self._render_pdf_completo(
                path, modo, orient, incluir_pred,
                escala, hojas_x, papel,
                incluir_header=opts.get('incluir_header', True),
                incluir_footer=opts.get('incluir_footer', True),
                incluir_legend=opts.get('incluir_legend', True),
                incluir_page=opts.get('incluir_page', True),
            )
            self._save_dir_gantt(path)
            QMessageBox.information(self, "Exportar PDF",
                                       f"Gantt exportado a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")

    def _vista_previa_pdf(self, opts):
        """Vista previa WYSIWYG: genera el PDF (mismo pipeline que el export) a
        un archivo temporal y lo muestra con QPdfView. Antes se usaba
        QPrintPreviewDialog, que dependía de la impresora física por defecto del
        sistema (su área imprimible / tamaños soportados), por lo que A3/A4 se
        veían mal aunque el PDF exportado fuera correcto."""
        import os, tempfile
        try:
            from PySide6.QtPdf import QPdfDocument
            from PySide6.QtPdfWidgets import QPdfView
        except ImportError:
            QMessageBox.warning(self, "Vista previa",
                                  "Módulo QtPdf/QtPdfWidgets no disponible.")
            return

        fd, tmp = tempfile.mkstemp(suffix='.pdf')
        os.close(fd)
        try:
            self._render_pdf_completo(
                tmp, opts['modo'], opts['orient'], opts['incluir_pred'],
                opts['escala'], opts.get('hojas_x', 0), opts.get('papel', 'A3'),
                incluir_header=opts.get('incluir_header', True),
                incluir_footer=opts.get('incluir_footer', True),
                incluir_legend=opts.get('incluir_legend', True),
                incluir_page=opts.get('incluir_page', True),
            )
        except Exception as e:
            import traceback; traceback.print_exc()
            try: os.unlink(tmp)
            except Exception: pass
            QMessageBox.warning(self, "Vista previa",
                                  f"No se pudo generar la vista previa: {e}")
            return

        dlg = QDialog(self)
        dlg.setWindowTitle("Vista previa — Gantt")
        dlg.setWindowModality(Qt.ApplicationModal)
        dlg.resize(1200, 820)
        vl = QVBoxLayout(dlg)
        vl.setContentsMargins(0, 0, 0, 0)
        vl.setSpacing(0)

        # Barra superior con acciones
        bar = QFrame()
        bar.setStyleSheet("background:#EEF1F5; border-bottom:1px solid #D4D4D4;")
        bhl = QHBoxLayout(bar)
        bhl.setContentsMargins(8, 4, 8, 4)
        lbl = QLabel("Vista previa del Gantt (igual al PDF que se exporta)")
        lbl.setStyleSheet("color:#485A6C; font-size:11px; border:none;")
        bhl.addWidget(lbl)
        bhl.addStretch()
        btn_img = QPushButton("🖼  Exportar imagen…")
        btn_pdf = QPushButton("Exportar PDF…")
        for b in (btn_img, btn_pdf):
            b.setCursor(Qt.PointingHandCursor)
            b.setStyleSheet(
                "QPushButton { background:white; border:1px solid #D4D4D4;"
                " border-radius:4px; padding:4px 10px; font-size:11px; }"
                "QPushButton:hover { border-color:#F37329; color:#C0621A; }")
        btn_img.clicked.connect(lambda: self._exportar_imagen(opts))
        bhl.addWidget(btn_img)
        bhl.addWidget(btn_pdf)
        vl.addWidget(bar)

        doc = QPdfDocument(dlg)
        doc.load(tmp)
        view = QPdfView(dlg)
        view.setDocument(doc)
        try:
            view.setPageMode(QPdfView.PageMode.MultiPage)
            view.setZoomMode(QPdfView.ZoomMode.FitToWidth)
        except Exception:
            pass
        vl.addWidget(view, stretch=1)

        def _export_pdf_desde_preview():
            sugerido = os.path.join(self._dir_descargas_gantt(),
                                      f"gantt_{self._cv.pid}.pdf")
            path, _ = QFileDialog.getSaveFileName(
                dlg, "Exportar Gantt a PDF", sugerido, "PDF (*.pdf)")
            if not path:
                return
            try:
                import shutil
                shutil.copyfile(tmp, path)
                self._save_dir_gantt(path)
                QMessageBox.information(dlg, "Exportar PDF",
                                          f"PDF exportado:\n{path}")
            except Exception as e:
                QMessageBox.warning(dlg, "Exportar PDF", f"No se pudo: {e}")
        btn_pdf.clicked.connect(_export_pdf_desde_preview)

        try:
            dlg.exec()
        finally:
            doc.close()
            try: os.unlink(tmp)
            except Exception: pass

    def _exportar_imagen(self, opts):
        """Exporta el Gantt como PNG o JPG. Genera primero un PDF temporal y
        lo rasteriza vía QPdfDocument. Multipágina ⇒ una imagen por página
        con sufijo _p1, _p2…"""
        import os, tempfile
        try:
            from PySide6.QtPdf import QPdfDocument
        except ImportError:
            QMessageBox.warning(self, "Exportar imagen",
                                  "Módulo QtPdf no disponible.")
            return

        sugerido = os.path.join(self._dir_descargas_gantt(),
                                  f"gantt_{self._cv.pid}.png")
        path, sel = QFileDialog.getSaveFileName(
            self, "Exportar Gantt como imagen", sugerido,
            "PNG (*.png);;JPEG (*.jpg *.jpeg)"
        )
        if not path:
            return
        # Asegurar extensión coherente con el filtro elegido
        ext = os.path.splitext(path)[1].lower()
        if 'jpg' in (sel or '').lower() or 'jpeg' in (sel or '').lower():
            if ext not in ('.jpg', '.jpeg'):
                path += '.jpg'
                ext = '.jpg'
            fmt = 'JPG'
        else:
            if ext != '.png':
                path += '.png'
                ext = '.png'
            fmt = 'PNG'

        # 1. Generar PDF temporal con los mismos parámetros
        tmp_fd, tmp_pdf = tempfile.mkstemp(suffix='.pdf', prefix='gantt_img_')
        os.close(tmp_fd)
        try:
            self._render_pdf_completo(
                tmp_pdf, opts['modo'], opts['orient'], opts['incluir_pred'],
                opts['escala'], opts.get('hojas_x', 0),
                opts.get('papel', 'A3'),
                incluir_header=opts.get('incluir_header', True),
                incluir_footer=opts.get('incluir_footer', True),
                incluir_legend=opts.get('incluir_legend', True),
                incluir_page=opts.get('incluir_page', True),
            )
            # 2. Rasterizar páginas a imagen
            doc = QPdfDocument()
            doc.load(tmp_pdf)
            n = doc.pageCount()
            if n <= 0:
                raise RuntimeError("PDF temporal vacío")

            from PySide6.QtGui import QImage
            dpi = 200  # buena calidad sin pesar demasiado
            base, real_ext = os.path.splitext(path)
            outs = []
            for i in range(n):
                page_pts = doc.pagePointSize(i)
                w_px = int(page_pts.width()  / 72.0 * dpi)
                h_px = int(page_pts.height() / 72.0 * dpi)
                rendered = doc.render(i, QSize(max(1, w_px), max(1, h_px)))
                # Componer sobre fondo blanco para evitar transparencias
                # (PNG queda con áreas transparentes y JPG no soporta alpha).
                canvas = QImage(rendered.size(), QImage.Format_RGB32)
                canvas.fill(QColor("white"))
                pp = QPainter(canvas)
                try:
                    pp.drawImage(0, 0, rendered)
                finally:
                    pp.end()
                out = path if n == 1 else f"{base}_p{i+1}{real_ext}"
                ok = canvas.save(out, fmt, 92 if fmt == 'JPG' else -1)
                if not ok:
                    raise RuntimeError(f"No se pudo guardar {out}")
                outs.append(out)
            self._save_dir_gantt(path)
            msg = "Imagen exportada:\n" + outs[0] if n == 1 \
                else f"{n} páginas exportadas:\n{outs[0]}\n…\n{outs[-1]}"
            QMessageBox.information(self, "Exportar imagen", msg)
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Exportar imagen",
                                  f"No se pudo exportar: {e}")
        finally:
            try: os.unlink(tmp_pdf)
            except Exception: pass

    # ──────────────────────────────────────────────────────────────────────
    # Pipeline PDF: tabla de partidas + gantt (1 hoja o varias)
    # ──────────────────────────────────────────────────────────────────────

    def _render_pdf_completo(self, path, modo, orient, incluir_pred, escala='auto',
                                hojas_x=0, papel='A3',
                                incluir_header=True, incluir_footer=True,
                                incluir_legend=True, incluir_page=True,
                                pie_offset: int = 0,
                                pie_total: int | None = None):
        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(path)
        self._aplicar_papel(printer, papel)
        printer.setPageOrientation(
            QPageLayout.Landscape if orient == 'landscape' else QPageLayout.Portrait
        )
        printer.setPageMargins(QMarginsF(10, 10, 10, 10), QPageLayout.Millimeter)

        painter = QPainter(printer)
        painter.setFont(QFont('Inter', 10))   # font base — Inter consistente
        try:
            # Pasamos pie_offset/pie_total como attrs temporales — leídos por
            # `_render_gantt_a_painter` para el footer continuo del Reporte
            # Completo.
            self._pie_offset_render = int(pie_offset or 0)
            self._pie_total_render  = int(pie_total) if pie_total else None
            self._render_gantt_a_painter(painter, printer, modo, incluir_pred,
                                            escala, hojas_x,
                                            incluir_header=incluir_header,
                                            incluir_footer=incluir_footer,
                                            incluir_legend=incluir_legend,
                                            incluir_page=incluir_page)
        finally:
            painter.end()
            self._pie_offset_render = 0
            self._pie_total_render  = None

    def _render_gantt_a_painter(self, painter, printer, modo, incluir_pred, escala='auto',
                                  hojas_x=0, incluir_header=True,
                                  incluir_footer=True, incluir_legend=True,
                                  incluir_page=True):
        """Pipeline puro de dibujo (sin crear/cerrar painter): se reusa entre
        export a archivo y QPrintPreviewDialog."""
        partidas = self._cv.filas_con_hitos()   # [Inicio] + partidas + [Fin]
        tasks    = self._cv._tasks
        cmap     = self._cv._cron_map
        proy     = self._cv._proy
        plazo    = proy.get('plazo') or 0
        max_ef   = max((t['EF'] for t in tasks.values() if t['EF'] > 0),
                        default=plazo)
        n_dias   = max(max_ef, plazo, 30)
        f_ini    = self._project_start()
        lead     = self.LEAD_DAYS   # margen vacío antes del día 1 (1ª hoja)

        if True:
            dpi = printer.resolution()
            mm = lambda v: v * dpi / 25.4
            # viewport() devuelve QRect en device coords (entero); usar QRectF
            vp = painter.viewport()
            PG_W = vp.width()
            PG_H = vp.height()

            # ── Layout vertical ──────────────────────────────────────────
            HEADER_H   = mm(18) if incluir_header else mm(0)
            # El footer puede contener: textos del formato (row 1) y/o
            # leyenda (row 2). Si solo uno de los dos está activo, lo
            # comprimimos a mm(5.5); si están ambos, mm(10).
            _row1_on = incluir_footer or incluir_page
            _row2_on = incluir_legend
            if _row1_on and _row2_on:
                FOOTER_H = mm(10)
            elif _row1_on or _row2_on:
                FOOTER_H = mm(5.5)
            else:
                FOOTER_H = mm(0)
            # Cabecera tabla y franja de tiempo del gantt deben coincidir en
            # altura para que las filas queden alineadas.
            HDR_TIME_H = mm(9)
            COL_HDR_H  = HDR_TIME_H
            body_top   = HEADER_H
            body_bottom = PG_H - FOOTER_H
            body_h     = body_bottom - body_top

            # ── Columnas de la tabla ─────────────────────────────────────
            #   #, Ítem, Descripción, Cal, Lab, Inicio, Fin  (+ Pred opcional)
            col_defs = [
                ('id',           mm(9),  Qt.AlignCenter),
                ('Ítem',         mm(14), Qt.AlignLeft | Qt.AlignVCenter),
                ('Descripción',  mm(50), Qt.AlignLeft | Qt.AlignVCenter),
                ('Días\ncal.',   mm(10), Qt.AlignCenter),
                ('Días\nlab.',   mm(10), Qt.AlignCenter),
                ('Inicio',       mm(17), Qt.AlignCenter),
                ('Fin',          mm(17), Qt.AlignCenter),
            ]
            if incluir_pred:
                col_defs.append(('Pred.', mm(18), Qt.AlignCenter))
            TABLE_W = sum(w for _, w, _ in col_defs)

            gantt_x = TABLE_W
            gantt_w = PG_W - TABLE_W
            if gantt_w < mm(40):
                # tabla demasiado ancha; achicar Descripción
                extra = mm(40) - gantt_w
                d_idx = 2
                _h, _w, _a = col_defs[d_idx]
                col_defs[d_idx] = (_h, max(mm(20), _w - extra), _a)
                TABLE_W = sum(w for _, w, _ in col_defs)
                gantt_x = TABLE_W
                gantt_w = PG_W - TABLE_W

            # ── Planificación de páginas ─────────────────────────────────
            # En modo 'fit' el alto de la franja de tiempo y la cabecera de la
            # tabla se calculan proporcionalmente al alto de fila — así todo
            # queda escalado coherentemente cuando hay muchas partidas.
            if modo == 'fit':
                rows_total = max(1, len(partidas))
                # iterar: HDR_TIME_H depende de row_h_pdf que depende de HDR_TIME_H
                for _ in range(4):
                    rows_area_h = body_h - HDR_TIME_H
                    rh = rows_area_h / rows_total
                    rh = max(mm(1.5), min(mm(7.0), rh))
                    ideal_hdr = max(mm(4.0), min(mm(11.0), rh * 2.0))
                    if abs(ideal_hdr - HDR_TIME_H) < mm(0.3):
                        HDR_TIME_H = ideal_hdr
                        break
                    HDR_TIME_H = ideal_hdr
                COL_HDR_H = HDR_TIME_H
                rows_area_h = body_h - HDR_TIME_H
                row_h_pdf = max(mm(1.5), min(mm(7.0), rows_area_h / rows_total))
                rows_chunks = [(0, len(partidas))]
                day_w_pdf  = gantt_w / (n_dias + lead)
                days_per_page = [n_dias]
            else:
                # Múltiples hojas: escala "real" y se pagina
                rows_area_h = body_h - HDR_TIME_H
                row_h_pdf  = mm(5.0)
                rows_per_page = max(1, int(rows_area_h // row_h_pdf))
                rows_chunks = []
                idx = 0
                while idx < len(partidas):
                    rows_chunks.append((idx, min(idx + rows_per_page, len(partidas))))
                    idx += rows_per_page

                if hojas_x and hojas_x > 0:
                    # Reparto controlado por el usuario: dividir n_dias en exactamente
                    # `hojas_x` hojas, lo más parejas posibles.
                    base = n_dias // hojas_x
                    rem  = n_dias - base * hojas_x
                    days_per_page = []
                    for i in range(hojas_x):
                        days_per_page.append(base + (1 if i < rem else 0))
                    days_per_page = [d for d in days_per_page if d > 0]
                    if not days_per_page:
                        days_per_page = [n_dias]
                    day_w_pdf = gantt_w / (max(days_per_page) + lead)
                else:
                    # Escala legible + reparto que LLENA el ancho en cada hoja.
                    # 1) día_w base legible (clamp); 2) días que caben por hoja;
                    # 3) nº MÍNIMO de hojas; 4) recalcular día_w para que las
                    # columnas (incluido el margen `lead` de la 1ª hoja) llenen
                    # el ancho del gantt en TODAS las hojas. Antes, con proyectos
                    # que entraban en ~1 hoja, quedaba un hueco a la derecha
                    # (notorio en A3) y la última hoja dispersa.
                    target_days = max(20, min(70, n_dias))
                    day_w0 = max(mm(1.2), min(mm(6.0), gantt_w / target_days))
                    days_per_full = max(1, int(gantt_w // day_w0))
                    cols_total = n_dias + lead          # columnas totales (con margen)
                    n_x = max(1, -(-cols_total // days_per_full))   # ceil → mín. hojas
                    cols_per_page = -(-cols_total // n_x)           # ceil → reparto parejo
                    day_w_pdf = gantt_w / cols_per_page             # llena el ancho
                    days_per_page = []
                    rem = n_dias
                    _first = True
                    while rem > 0:
                        # En la 1ª hoja reservamos `lead` días para el margen.
                        cap = max(1, cols_per_page - (lead if _first else 0))
                        take = min(cap, rem)
                        days_per_page.append(take)
                        rem -= take
                        _first = False

            x_pages = len(days_per_page)
            y_pages = len(rows_chunks)
            total_pages = x_pages * y_pages

            # ── Renderizar páginas ───────────────────────────────────────
            page_num = 0
            day_cursor = 1
            for xi, days_in_page in enumerate(days_per_page):
                # Margen inicial (lead) solo en la 1ª hoja: deja aire a la
                # izquierda para que se vean las líneas de predecesoras que
                # salen por ese lado (igual que en pantalla).
                pg_lead = lead if xi == 0 else 0
                d_ini = day_cursor - pg_lead
                d_fin = day_cursor + days_in_page - 1
                day_cursor += days_in_page
                page_gantt_w = (d_fin - d_ini + 1) * day_w_pdf
                for yi, (r0, r1) in enumerate(rows_chunks):
                    if page_num > 0:
                        printer.newPage()
                    page_num += 1

                    if incluir_header:
                        self._pdf_paint_header(
                            painter, 0, 0, PG_W, HEADER_H,
                            proy, page_num, total_pages, xi + 1, x_pages,
                            yi + 1, y_pages, d_ini, d_fin, f_ini,
                        )

                    # Cabeceras de columnas
                    self._pdf_paint_tabla_header(
                        painter, 0, body_top, COL_HDR_H, col_defs,
                    )
                    # Tabla de partidas (filas r0..r1)
                    self._pdf_paint_tabla_filas(
                        painter, 0, body_top + COL_HDR_H,
                        TABLE_W, rows_area_h - COL_HDR_H,
                        col_defs, partidas, tasks, cmap, r0, r1, row_h_pdf,
                    )

                    # Cabecera de tiempo del gantt
                    self._pdf_paint_gantt_header(
                        painter, gantt_x, body_top, page_gantt_w, HDR_TIME_H,
                        d_ini, d_fin, f_ini, day_w_pdf, escala,
                    )
                    # Cuerpo del gantt (slice)
                    self._pdf_paint_gantt_body(
                        painter, gantt_x, body_top + HDR_TIME_H,
                        page_gantt_w, rows_area_h - 0,
                        partidas, tasks, cmap, r0, r1, d_ini, d_fin,
                        row_h_pdf, day_w_pdf, f_ini,
                    )

                    # Marco general de la zona de cuerpo
                    painter.save()
                    painter.setPen(QPen(QColor("#94A3B8"), max(1, mm(0.15))))
                    painter.setBrush(Qt.NoBrush)
                    painter.drawRect(QRectF(0, body_top, TABLE_W + page_gantt_w,
                                              rows_area_h))
                    painter.restore()

                    if _row1_on or _row2_on:
                        self._pdf_paint_footer(
                            painter, 0, PG_H - FOOTER_H, PG_W, FOOTER_H,
                            page_num, total_pages,
                            incluir_footer=incluir_footer,
                            incluir_legend=incluir_legend,
                            incluir_page=incluir_page,
                        )

    # — Helpers de pintura —————————————————————————————————————

    def _pdf_paint_header(self, p, x, y, w, h, proy, pg, tot,
                            xpg, xtot, ypg, ytot, d_ini, d_fin, f_ini):
        """Header estilo Centro de Reportes: línea naranja + 3 columnas
        (empresa/logo · título + nombre proyecto · costo al + modalidad)."""
        p.save()
        dpi = p.device().logicalDpiX()
        mm = lambda v: v * dpi / 25.4

        # Paleta unificada con pdf_reports
        ORANGE     = "#F37329"
        ORANGE_DK  = "#C0621A"
        SLATE_900  = "#1F2A38"
        SLATE_700  = "#2E3C52"
        SLATE_500  = "#485A6C"
        SLATE_300  = "#94A3B8"
        SLATE_100  = "#E2E8F0"

        # Fallback de marca respeta el toggle 'Reportes sobrios' — si el
        # usuario configuró un color custom (rep_color_marca) éste gana.
        try:
            from utils.theme import accent_reportes
            _o_def, _od_def, _ = accent_reportes()
        except Exception:
            _o_def, _od_def = ORANGE, ORANGE_DK

        fmt_ = {}
        try:
            fmt_ = _get_formato_reportes() or {}
        except Exception:
            pass
        color_marca    = QColor(fmt_.get('rep_color_marca') or _o_def)
        color_marca_dk = QColor(fmt_.get('rep_color_marca_dk') or _od_def)
        empresa        = fmt_.get('rep_empresa_nombre') or 'ingePresupuestos'
        sub_empresa    = fmt_.get('rep_empresa_subtitulo') or 'Presupuestos de Obra Pública'

        # Sin línea de marca arriba (eliminada por consistencia con
        # pdf_reports y para ahorrar tinta — el usuario reportó que la
        # línea de color se ve "sucia" en entregas formales).
        line_h = 0
        # Geometría: 3 columnas (izq, centro, der)
        margin = mm(2)
        col_top  = y + mm(2)
        col_h    = h - mm(2.5)
        left_w   = mm(50)
        right_w  = mm(55)
        center_x = x + margin + left_w + mm(3)
        center_w = w - 2 * margin - left_w - right_w - mm(6)

        # ── Columna izquierda: logo (si hay) o empresa + subtitulo ──────
        logo_b64 = (fmt_.get('rep_logo_b64') or '').strip()
        drew_logo = False
        if logo_b64:
            try:
                from PySide6.QtCore import QByteArray
                from PySide6.QtGui import QImage
                ba = QByteArray.fromBase64(logo_b64.encode('ascii'))
                img = QImage()
                if img.loadFromData(ba):
                    img_h = col_h
                    img_w = img.width() * img_h / max(1, img.height())
                    img_w = min(img_w, left_w)
                    p.drawImage(QRectF(x + margin, col_top, img_w, img_h), img)
                    drew_logo = True
            except Exception:
                drew_logo = False

        if not drew_logo:
            # Empresa
            f1 = p.font(); f1.setPointSizeF(11.5); f1.setBold(True); p.setFont(f1)
            p.setPen(color_marca_dk)
            p.drawText(QRectF(x + margin, col_top, left_w, col_h * 0.42),
                        Qt.AlignLeft | Qt.AlignVCenter, empresa)
            # Sub-empresa
            f2 = p.font(); f2.setPointSizeF(7.0); f2.setBold(False); p.setFont(f2)
            p.setPen(QColor(SLATE_300))
            p.drawText(QRectF(x + margin, col_top + col_h * 0.42, left_w, col_h * 0.30),
                        Qt.AlignLeft | Qt.AlignVCenter, sub_empresa)
            # Línea cliente (si existe)
            cliente = (proy.get('cliente') or '').strip()
            if cliente:
                f2b = p.font(); f2b.setPointSizeF(6.8); f2b.setItalic(True); p.setFont(f2b)
                p.setPen(QColor(SLATE_500))
                p.drawText(QRectF(x + margin, col_top + col_h * 0.70, left_w, col_h * 0.28),
                            Qt.AlignLeft | Qt.AlignVCenter, cliente)
                f2b.setItalic(False); p.setFont(f2b)

        # ── Columna central: "Diagrama de Gantt" + nombre proyecto (3 líneas) ──
        # Título del reporte
        f3 = p.font(); f3.setPointSizeF(11.0); f3.setBold(True); f3.setItalic(False); p.setFont(f3)
        p.setPen(QColor(SLATE_900))
        rect_titulo = QRectF(center_x, col_top, center_w, col_h * 0.34)
        p.drawText(rect_titulo, Qt.AlignCenter | Qt.AlignVCenter, "Diagrama de Gantt")

        # Nombre proyecto — wrap greedy hasta 3 líneas
        from PySide6.QtGui import QFontMetrics
        nom = (proy.get('nombre') or 'Proyecto sin nombre').strip()
        f4 = p.font(); f4.setPointSizeF(7.8); f4.setItalic(True); f4.setBold(False); p.setFont(f4)
        p.setPen(QColor(SLATE_500))
        fm = QFontMetrics(f4)
        MAX_LINES = 3
        line_sp = fm.lineSpacing()
        words = nom.split()
        lines: list[list[str]] = [[]]
        for word in words:
            test = (' '.join(lines[-1] + [word])).strip()
            if fm.horizontalAdvance(test) <= int(center_w):
                lines[-1].append(word)
            else:
                if len(lines) < MAX_LINES:
                    lines.append([word])
                else:
                    last = ' '.join(lines[-1] + [word])
                    lines[-1] = [fm.elidedText(last, Qt.ElideRight, int(center_w))]
                    break
        texto = '\n'.join(' '.join(ws) for ws in lines if ws)
        rect_proy = QRectF(center_x, col_top + col_h * 0.34,
                            center_w, col_h * 0.62)
        p.drawText(rect_proy, Qt.AlignHCenter | Qt.AlignTop, texto)
        f4.setItalic(False); p.setFont(f4)

        # ── Columna derecha: costo al, plazo, tramo, modalidad, hojas ──
        from datetime import timedelta
        right_x = x + w - margin - right_w
        f5 = p.font(); f5.setPointSizeF(8.0); f5.setBold(True); p.setFont(f5)
        p.setPen(QColor(SLATE_700))
        plazo = proy.get('plazo') or 0
        from core.pdf_reports import _clean_costo_al
        costo_al = _clean_costo_al(proy.get('costo_al'))
        p.drawText(QRectF(right_x, col_top, right_w, col_h * 0.22),
                    Qt.AlignRight | Qt.AlignVCenter,
                    f"Costo al: {costo_al}")

        # Detalles secundarios
        f6 = p.font(); f6.setPointSizeF(7.0); f6.setBold(False); f6.setItalic(False); p.setFont(f6)
        p.setPen(QColor(SLATE_500))
        det = [f"Plazo: {plazo} días"]
        if f_ini:
            fi = (f_ini + timedelta(days=d_ini - 1)).strftime('%d/%m/%Y')
            ff = (f_ini + timedelta(days=d_fin - 1)).strftime('%d/%m/%Y')
            det.append(f"Tramo: {fi} → {ff}")
        if xtot > 1 or ytot > 1:
            det.append(f"Hoja {pg}/{tot} (col {xpg}/{xtot} · fila {ypg}/{ytot})")
        modalidad = (proy.get('modalidad') or '').strip()
        if modalidad:
            det.append(modalidad)
        det_h_total = col_h * 0.74
        det_line_h = det_h_total / max(1, len(det))
        for i, txt in enumerate(det):
            yy = col_top + col_h * 0.22 + i * det_line_h
            p.drawText(QRectF(right_x, yy, right_w, det_line_h),
                        Qt.AlignRight | Qt.AlignVCenter, txt)

        # Línea separadora inferior del header
        sep_y = y + h - mm(1.5)
        p.setPen(QPen(QColor(SLATE_100), max(1, mm(0.18))))
        p.drawLine(QPointF(x + margin, sep_y), QPointF(x + w - margin, sep_y))
        p.restore()

    def _pdf_paint_tabla_header(self, p, x, y, h, col_defs):
        """Header de la tabla estilo Centro de Reportes: fondo sutil silver,
        texto slate-900 (casi negro), bordes slate arriba/abajo, separadores
        suaves."""
        p.save()
        dpi = p.device().logicalDpiX()
        mm = lambda v: v * dpi / 25.4

        SLATE_900 = "#1F2A38"
        SLATE_700 = "#2E3C52"
        SLATE_300 = "#94A3B8"
        SILVER_BG = "#F1F5F9"  # fondo sutil (slate-100)

        total_w = sum(w for _, w, _ in col_defs)
        # Fondo sutil
        p.setPen(Qt.NoPen)
        p.setBrush(QBrush(QColor(SILVER_BG)))
        p.drawRect(QRectF(x, y, total_w, h))

        # Textos — slate-900 bold (casi negro)
        h_pt = h * 72.0 / dpi
        hdr_pt = max(4.5, min(10.0, h_pt * 0.45))
        p.setPen(QColor(SLATE_900))
        f = p.font(); f.setPointSizeF(hdr_pt); f.setBold(True); p.setFont(f)
        cx = x
        for label, cw, _align in col_defs:
            # AlignCenter + WordWrap: respeta los '\n' de encabezados de 2
            # líneas (p.ej. "Días\ncal.") y evita que se recorten.
            p.drawText(QRectF(cx + 2, y, cw - 4, h),
                        Qt.AlignCenter | Qt.TextWordWrap, label)
            cx += cw

        # Borde superior + inferior (1.5pt arriba / 0.8pt abajo, estilo reportes)
        top_w = max(1, mm(0.40))
        bot_w = max(1, mm(0.22))
        p.setPen(QPen(QColor(SLATE_700), top_w))
        p.drawLine(QPointF(x, y), QPointF(x + total_w, y))
        p.setPen(QPen(QColor(SLATE_700), bot_w))
        p.drawLine(QPointF(x, y + h), QPointF(x + total_w, y + h))

        # Separadores verticales suaves
        p.setPen(QPen(QColor(SLATE_300), max(1, mm(0.10))))
        cx = x
        for _, cw, _ in col_defs[:-1]:
            cx += cw
            p.drawLine(QPointF(cx, y + h * 0.18), QPointF(cx, y + h * 0.82))
        p.restore()

    def _pdf_paint_tabla_filas(self, p, x, y, w, max_h,
                                  col_defs, partidas, tasks, cmap,
                                  r0, r1, row_h):
        from core.cronograma import contar_laborables, formatear_pred_es
        p.save()
        dpi = p.device().logicalDpiX()
        mm = lambda v: v * dpi / 25.4

        # "#" = posición 1-based en la lista (coincide con numerar_filas).
        _plazo = self._cv._proy.get('plazo') or 0
        non_working = self._cv._non_working_set(_plazo + 365)

        cur_y = y
        # Escalar font al alto de fila (1pt = dpi/72 dots; altura visual ≈ 1.25*pt)
        row_h_pt = row_h * 72.0 / dpi
        body_pt  = max(4.5, min(8.5, row_h_pt * 0.55))
        title_pt = max(5.0, min(9.0, row_h_pt * 0.60))
        f_body = p.font();  f_body.setPointSizeF(body_pt);   f_body.setBold(False)
        f_title = p.font(); f_title.setPointSizeF(title_pt); f_title.setBold(True)

        for r in range(r0, r1):
            pt = partidas[r]
            es_titulo = bool(pt['es_titulo'])
            cd = cmap.get(pt['id'], {})
            t = tasks.get(pt['id'], {})

            # Datos por columna
            virtual = pt.get('_virtual')
            num_str = str(r + 1)
            if virtual:
                a = pt.get('_ES') or 0
                b = pt.get('_EF') or a
                if virtual in ('proyecto', 'subppto'):
                    cal_str = str(b - a + 1) if b >= a else ''
                    lab_str = str(contar_laborables(a, b, non_working))
                else:
                    cal_str = lab_str = '0'
                es_str = self._day_to_date_str(a)
                ef_str = self._day_to_date_str(b)
                pred_str = ''
            elif es_titulo:
                es_min, ef_max = self._summary_for_title(pt)
                if es_min and ef_max:
                    cal_str = str(ef_max - es_min + 1)
                    lab_str = str(contar_laborables(es_min, ef_max, non_working))
                    es_str = self._day_to_date_str(es_min)
                    ef_str = self._day_to_date_str(ef_max)
                else:
                    cal_str = lab_str = es_str = ef_str = ''
                pred_str = ''
            else:
                es_v, ef_v = t.get('ES'), t.get('EF')
                lab_str = str(cd.get('duracion', '') or '')
                cal_str = (str(ef_v - es_v + 1)
                           if (es_v and ef_v and ef_v >= es_v) else lab_str)
                es_str  = self._day_to_date_str(es_v)
                ef_str  = self._day_to_date_str(ef_v)
                pred_str = formatear_pred_es((cd.get('predecesoras', '') or '').strip())

            niv = (pt.get('nivel') or 1) - 1
            desc = pt['descripcion'] or ''
            if virtual in ('inicio', 'fin'):
                desc = "◆ " + desc
            elif niv > 0 and not virtual:
                desc = ('  ' * niv) + desc

            row_vals = [num_str, pt['item'] or '', desc, cal_str, lab_str,
                        es_str, ef_str]
            if len(col_defs) >= 8:
                row_vals.append(pred_str)

            # Fondo
            p.setPen(Qt.NoPen)
            if es_titulo:
                p.setBrush(QBrush(QColor("#F1F5F9")))
            elif (r - r0) % 2 == 1:
                p.setBrush(QBrush(QColor("#FAFBFC")))
            else:
                p.setBrush(QBrush(QColor("white")))
            p.drawRect(QRectF(x, cur_y, w, row_h))

            # Texto por celda
            p.setFont(f_title if es_titulo else f_body)
            critical = bool(t.get('critical'))
            # Color por defecto: títulos casi negro, cuerpo gris-oscuro.
            # Para críticas NO pintamos toda la fila en rojo: solo la
            # columna "Ítem" (índice 1) se resalta en rojo + bold.
            if es_titulo:
                default_pen = QColor("#1E2635")
            else:
                default_pen = QColor("#273445")
            p.setPen(default_pen)

            # Subrayar SOLO los títulos de nivel 1 (no subtítulos ni partidas).
            is_main_title = es_titulo and (pt.get('nivel') or 0) == 1
            cx = x
            for idx, (val, (label, cw, align)) in enumerate(zip(row_vals, col_defs)):
                rect_cell = QRectF(cx + 2, cur_y, cw - 4, row_h)
                # Fuente fresca por celda (evita arrastrar estado): subrayado en
                # el código/nombre de los títulos de nivel 1; bold en críticas.
                cell_critical = critical and not es_titulo and idx == 1
                cf = QFont(f_title if es_titulo else f_body)
                if is_main_title and idx in (1, 2):
                    cf.setUnderline(True)
                if cell_critical:
                    cf.setBold(True)
                p.setFont(cf)
                p.setPen(QColor("#C6262E") if cell_critical else default_pen)
                # Elide manual si el texto excede
                fm = p.fontMetrics()
                txt = val
                if fm.horizontalAdvance(txt) > cw - 4:
                    txt = fm.elidedText(txt, Qt.ElideRight, int(cw - 4))
                # Clip al rect de celda para evitar cualquier desbordamiento
                p.save()
                p.setClipRect(rect_cell)
                p.drawText(rect_cell, int(align), txt)
                p.restore()
                cx += cw

            # Borde inferior de fila
            p.setPen(QPen(QColor("#E8EAED"), max(1, mm(0.1))))
            p.drawLine(QPointF(x, cur_y + row_h), QPointF(x + w, cur_y + row_h))

            cur_y += row_h
            if cur_y > y + max_h:
                break

        # Separadores verticales de columnas
        p.setPen(QPen(QColor("#D4D4D4"), max(1, mm(0.12))))
        cx = x
        for _, cw, _ in col_defs[:-1]:
            cx += cw
            p.drawLine(QPointF(cx, y), QPointF(cx, min(cur_y, y + max_h)))
        # Borde derecho de la tabla
        p.drawLine(QPointF(x + w, y), QPointF(x + w, min(cur_y, y + max_h)))
        p.restore()

    def _pdf_paint_gantt_header(self, p, x, y, w, h, d_ini, d_fin, f_ini,
                                 day_w_pdf, escala='auto'):
        """Franja de escala temporal. `escala`: 'auto'|'dias'|'semanas'|'meses'."""
        p.save()
        dpi = p.device().logicalDpiX()
        mm = lambda v: v * dpi / 25.4

        h_top = h * 0.55
        h_bot = h - h_top

        SLATE_900 = "#1F2A38"
        SLATE_700 = "#2E3C52"
        SLATE_500 = "#485A6C"
        SLATE_300 = "#94A3B8"
        SILVER_BG = "#F1F5F9"  # fondo sutil (idéntico al col header)
        SILVER_50 = "#FAFBFC"

        # Fondo sutil (combina con el header de la tabla)
        p.setPen(Qt.NoPen)
        p.setBrush(QBrush(QColor(SILVER_BG)))
        p.drawRect(QRectF(x, y, w, h_top))
        p.setBrush(QBrush(QColor(SILVER_50)))
        p.drawRect(QRectF(x, y + h_top, w, h_bot))

        from datetime import timedelta
        MESES = ["Ene", "Feb", "Mar", "Abr", "May", "Jun",
                  "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]

        # Tamaño tipográfico escalado al alto de la franja
        h_top_pt = h_top * 72.0 / dpi
        h_bot_pt = h_bot * 72.0 / dpi
        # Meses en el tope pueden ser grandes; las fechas en la franja inferior
        # se aprietan al mínimo para que entren muchas etiquetas.
        top_pt = max(4.5, min(11.0, h_top_pt * 0.70))
        bot_pt = max(3.0, min(5.5, h_bot_pt * 0.42))

        day_w_mm = day_w_pdf / mm(1.0)

        # Resolver escala 'auto' a una concreta para la fila inferior.
        # Preferimos 'dias' (con paso adaptativo) hasta que las etiquetas
        # serían demasiado densas; entonces caemos a semanas y, en último
        # extremo, a meses.
        if escala == 'auto':
            if day_w_mm >= 0.45:
                detail = 'dias'
            elif day_w_mm * 7 >= 5:
                detail = 'semanas'
            else:
                detail = 'meses'
        else:
            detail = escala

        # ── Fila superior: meses (siempre) ────────────────────────────────
        f1 = p.font(); f1.setPointSizeF(top_pt); f1.setBold(True); p.setFont(f1)
        p.setPen(QColor(SLATE_900))

        if f_ini is not None:
            cur_mk = None
            band_start_day = d_ini
            for d in range(d_ini, d_fin + 2):
                if d > d_fin or (cur_mk is not None and
                                  (f_ini + timedelta(days=d - 1)).month != cur_mk[1]):
                    # flush banda anterior
                    if cur_mk is not None:
                        x_s = x + (band_start_day - d_ini) * day_w_pdf
                        x_e = x + (d - d_ini) * day_w_pdf
                        ww = x_e - x_s
                        year, month = cur_mk
                        label_full  = f"{MESES[month - 1]} {year}"
                        label_short = MESES[month - 1]
                        label_micro = label_short[0]
                        p.setPen(QColor(SLATE_900))
                        fm = p.fontMetrics()
                        if fm.horizontalAdvance(label_full) <= ww - 2:
                            txt = label_full
                        elif fm.horizontalAdvance(label_short) <= ww - 2:
                            txt = label_short
                        elif ww >= 4:
                            txt = label_micro
                        else:
                            txt = ''
                        if txt:
                            p.save()
                            p.setClipRect(QRectF(x_s, y, ww, h_top))
                            p.drawText(QRectF(x_s, y, ww, h_top),
                                          Qt.AlignCenter, txt)
                            p.restore()
                        # Línea divisoria entre meses
                        p.setPen(QPen(QColor(SLATE_300), max(1, mm(0.12))))
                        p.drawLine(QPointF(x_s, y), QPointF(x_s, y + h))
                    if d > d_fin:
                        break
                    cur_mk = ((f_ini + timedelta(days=d - 1)).year,
                              (f_ini + timedelta(days=d - 1)).month)
                    band_start_day = d
                    continue
                date = f_ini + timedelta(days=d - 1)
                if cur_mk is None:
                    cur_mk = (date.year, date.month)
                    band_start_day = d

        # ── Fila inferior según `detail` ──────────────────────────────────
        p.setPen(QColor("#485A6C"))
        f2 = p.font(); f2.setPointSizeF(bot_pt); f2.setBold(False); p.setFont(f2)
        fm2 = p.fontMetrics()

        if detail == 'dias':
            # Mostrar día (1, 5, 10, ...) — paso adaptativo para que no se sobrepongan
            sample = fm2.horizontalAdvance("00")
            min_gap = sample + max(1, mm(0.3))
            step = max(1, int((min_gap + day_w_pdf - 1) // max(1.0, day_w_pdf)))
            # redondear a divisores agradables
            for cand in (1, 2, 3, 5, 7, 10, 14, 15, 20, 30, 45, 60, 90, 120, 180, 365):
                if cand >= step:
                    step = cand
                    break
            for d in range(d_ini, d_fin + 1):
                rel = d - d_ini
                if rel % step != 0 and d != d_ini:
                    continue
                xx = x + rel * day_w_pdf
                if f_ini is not None:
                    fd = f_ini + timedelta(days=d - 1)
                    # Si el step es <=7 mostramos solo día; si es mayor, dd/mm
                    if step <= 7:
                        lbl = str(fd.day)
                    else:
                        lbl = fd.strftime("%d/%m")
                else:
                    lbl = str(d)
                tw = max(day_w_pdf * step, fm2.horizontalAdvance(lbl) + mm(0.5))
                p.save()
                p.setClipRect(QRectF(xx - tw / 2, y + h_top, tw, h_bot))
                p.drawText(QRectF(xx - tw / 2, y + h_top, tw, h_bot),
                              Qt.AlignCenter, lbl)
                p.restore()
                # Tick
                p.setPen(QPen(QColor("#C5CFDB"), max(1, mm(0.1))))
                p.drawLine(QPointF(xx, y + h - mm(0.7)), QPointF(xx, y + h))
                p.setPen(QColor("#485A6C"))

        elif detail == 'semanas':
            # cada N semanas según ancho
            week_w = day_w_pdf * 7
            sample = fm2.horizontalAdvance("S99")
            step_w = max(1, int((sample + mm(1.0)) // max(1.0, week_w)))
            d = d_ini
            wnum_base = (d_ini - 1) // 7 + 1
            i = 0
            while d <= d_fin:
                if i % step_w == 0:
                    xx = x + (d - d_ini) * day_w_pdf
                    w_no = (d - 1) // 7 + 1
                    lbl = f"S{w_no}"
                    tw = week_w * step_w
                    p.save()
                    p.setClipRect(QRectF(xx, y + h_top, tw, h_bot))
                    p.drawText(QRectF(xx, y + h_top, tw, h_bot),
                                  Qt.AlignCenter, lbl)
                    p.restore()
                i += 1
                d += 7

        elif detail == 'meses':
            # Resaltar mes/año en la franja inferior (texto más grande + fecha 1)
            f3 = p.font(); f3.setPointSizeF(max(bot_pt, top_pt * 0.85))
            f3.setBold(False); p.setFont(f3)
            fm3 = p.fontMetrics()
            if f_ini is not None:
                cur_mk = None
                band_start_day = d_ini
                for d in range(d_ini, d_fin + 2):
                    date_d = (f_ini + timedelta(days=d - 1)) if d <= d_fin else None
                    cur = (date_d.year, date_d.month) if date_d else None
                    if cur_mk is None and cur is not None:
                        cur_mk = cur
                        band_start_day = d
                        continue
                    if d > d_fin or cur != cur_mk:
                        # flush
                        x_s = x + (band_start_day - d_ini) * day_w_pdf
                        x_e = x + (d - d_ini) * day_w_pdf
                        ww = x_e - x_s
                        year, month = cur_mk
                        # Primer día del mes en formato dd/mm/yy si entra
                        first = f_ini + timedelta(days=band_start_day - 1)
                        label = first.strftime("%d/%m/%y")
                        if fm3.horizontalAdvance(label) > ww - 2:
                            label = f"{MESES[month - 1]} {str(year)[-2:]}"
                        if fm3.horizontalAdvance(label) > ww - 2:
                            label = MESES[month - 1]
                        if ww > 4 and fm3.horizontalAdvance(label) <= ww - 1:
                            p.save()
                            p.setClipRect(QRectF(x_s, y + h_top, ww, h_bot))
                            p.drawText(QRectF(x_s, y + h_top, ww, h_bot),
                                          Qt.AlignCenter, label)
                            p.restore()
                        if d > d_fin:
                            break
                        cur_mk = cur
                        band_start_day = d

        # Separadores: arriba/abajo gruesos (slate-700) — coinciden con la
        # tabla de partidas. División mes/día interna más suave.
        top_w = max(1, mm(0.40))
        bot_w = max(1, mm(0.22))
        p.setPen(QPen(QColor(SLATE_700), top_w))
        p.drawLine(QPointF(x, y), QPointF(x + w, y))
        p.setPen(QPen(QColor(SLATE_300), max(1, mm(0.10))))
        p.drawLine(QPointF(x, y + h_top), QPointF(x + w, y + h_top))
        p.setPen(QPen(QColor(SLATE_700), bot_w))
        p.drawLine(QPointF(x, y + h), QPointF(x + w, y + h))
        p.restore()

    def _pdf_paint_gantt_body(self, p, x, y, w, max_h,
                                 partidas, tasks, cmap, r0, r1,
                                 d_ini, d_fin, row_h, day_w_pdf, f_ini):
        """Pinta el slice del cuerpo del gantt para filas [r0,r1) y
        días [d_ini,d_fin]."""
        p.save()
        dpi = p.device().logicalDpiX()
        mm = lambda v: v * dpi / 25.4

        days_total = d_fin - d_ini + 1
        h_total = (r1 - r0) * row_h
        # Clip al área del gantt
        p.setClipRect(QRectF(x, y, w, h_total))

        from datetime import timedelta

        # Sombreado domingos / feriados (capa por encima)
        feriados = self._feriados_set()
        for d in range(d_ini, d_fin + 1):
            fecha_dia = (f_ini + timedelta(days=d - 1)) if f_ini else None
            es_domingo = fecha_dia.weekday() == 6 if fecha_dia else False
            es_feriado = (fecha_dia.strftime('%Y-%m-%d') in feriados) if fecha_dia else False
            if not (es_domingo or es_feriado):
                continue
            xx = x + (d - d_ini) * day_w_pdf
            p.setPen(Qt.NoPen)
            if es_feriado:
                p.setBrush(QBrush(QColor(243, 115, 41, 38)))
            else:
                p.setBrush(QBrush(QColor(148, 163, 184, 30)))
            p.drawRect(QRectF(xx, y, day_w_pdf, h_total))

        # Grilla vertical — líneas guía PUNTEADAS (entrecortadas) estilo MS
        # Project: una por día cuando el día es ancho, o por semana si es muy
        # angosto. Las de inicio de semana, un poco más marcadas.
        step_g = 1 if day_w_pdf >= mm(2.4) else 7
        pen_day  = QPen(QColor("#CDD3DB"), max(1, mm(0.08)), Qt.DotLine)
        pen_week = QPen(QColor("#AEB6C0"), max(1, mm(0.09)), Qt.DotLine)
        for d in range(d_ini, d_fin + 1):
            if d == d_ini:
                continue
            is_week = (d % 7 == 1)
            if step_g == 1:
                p.setPen(pen_week if is_week else pen_day)
            elif is_week:
                p.setPen(pen_week)
            else:
                continue
            xx = x + (d - d_ini) * day_w_pdf
            p.drawLine(QPointF(xx, y), QPointF(xx, y + h_total))

        # Grilla horizontal — líneas guía PUNTEADAS por fila (en lugar de la
        # zebra), para seguir cada barra hasta su partida. Mismo estilo que las
        # verticales. Van bajo las barras.
        p.setPen(QPen(QColor("#CDD3DB"), max(1, mm(0.08)), Qt.DotLine))
        for r in range(r0, r1):
            ly = y + (r - r0 + 1) * row_h
            p.drawLine(QPointF(x, ly), QPointF(x + w, ly))

        # ── Cuerpo ────────────────────────────────────────────────────────
        row_info = {}
        for r in range(r0, r1):
            pt = partidas[r]
            es_titulo = bool(pt['es_titulo'])
            yy = y + (r - r0) * row_h

            virtual = pt.get('_virtual')
            # Resumen del proyecto / subpresupuesto: barra resumen del rango.
            if virtual in ('proyecto', 'subppto'):
                a0 = max(pt.get('_ES') or 1, d_ini)
                b0 = min(pt.get('_EF') or 1, d_fin)
                if b0 >= a0:
                    bx = x + (a0 - d_ini) * day_w_pdf
                    bw = (b0 - a0 + 1) * day_w_pdf
                    by = yy + row_h / 2 - max(1, mm(0.5))
                    bh = max(1.8, mm(1.1))
                    p.setPen(Qt.NoPen)
                    p.setBrush(QBrush(QColor("#0F1419")))
                    p.drawRect(QRectF(bx, by, bw, bh))
                continue

            # Hitos virtuales (Inicio/Fin): diamante azul/verde.
            if virtual:
                dia = pt.get('_ES') or 1
                cy = yy + row_h / 2
                # Registrar el hito en row_info ANTES del chequeo de slice — así
                # las flechas de predecesoras que salen del hito de Inicio (o
                # entran al de Fin) se dibujan (espejo de la UI en pantalla).
                row_info[pt['id']] = {'y_center': cy, 'ES': dia, 'EF': dia,
                                        'critical': False}
                if dia < d_ini or dia > d_fin:
                    continue
                s = max(2.2, min(row_h * 0.38, mm(2.2)))
                if virtual == 'inicio':
                    cx = x + (dia - d_ini) * day_w_pdf + s          # borde de arranque
                elif virtual == 'fin':
                    cx = x + (dia - d_ini + 1) * day_w_pdf - s      # borde de cierre
                else:
                    cx = x + (dia - d_ini) * day_w_pdf + day_w_pdf / 2
                p.setRenderHint(QPainter.Antialiasing, True)
                poly = QPolygonF([
                    QPointF(cx, cy - s), QPointF(cx + s, cy),
                    QPointF(cx, cy + s), QPointF(cx - s, cy),
                ])
                col = "#3689E6" if virtual == 'inicio' else "#2EA043"
                p.setPen(Qt.NoPen)
                p.setBrush(QBrush(QColor(col)))
                p.drawPolygon(poly)
                continue

            if es_titulo:
                es_min, ef_max = self._summary_for_title(pt)
                if es_min and ef_max:
                    # recortar al slice visible
                    a = max(es_min, d_ini)
                    b = min(ef_max, d_fin)
                    if b >= a:
                        bx = x + (a - d_ini) * day_w_pdf
                        bw = (b - a + 1) * day_w_pdf
                        by = yy + row_h / 2 - max(1, mm(0.5))
                        bh = max(1.8, mm(1.1))
                        p.setRenderHint(QPainter.Antialiasing, True)
                        # Gradiente sutil slate-900 → slate-700 (look MS Project)
                        grad_t = QLinearGradient(bx, by, bx, by + bh)
                        grad_t.setColorAt(0.0, QColor("#2E3C52"))
                        grad_t.setColorAt(1.0, QColor("#0F1419"))
                        p.setPen(Qt.NoPen)
                        p.setBrush(QBrush(grad_t))
                        p.drawRect(QRectF(bx, by, bw, bh))
                        # triángulos solo si los extremos reales caen en el slice
                        tri_size = max(2.5, mm(1.1))
                        p.setBrush(QBrush(QColor("#0F1419")))
                        if es_min >= d_ini and es_min <= d_fin:
                            poly = QPolygonF([
                                QPointF(bx, by),
                                QPointF(bx + tri_size, by),
                                QPointF(bx, by + bh + tri_size),
                            ])
                            p.drawPolygon(poly)
                        if ef_max >= d_ini and ef_max <= d_fin:
                            ex = x + (ef_max - d_ini + 1) * day_w_pdf
                            poly = QPolygonF([
                                QPointF(ex, by),
                                QPointF(ex - tri_size, by),
                                QPointF(ex, by + bh + tri_size),
                            ])
                            p.drawPolygon(poly)
                continue

            t = tasks.get(pt['id'], {})
            es = t.get('ES', 0)
            ef = t.get('EF', 0)
            dur = t.get('dur', 0)
            cd = cmap.get(pt['id'], {})
            es_hito = int(cd.get('es_hito', 0) or 0)
            critical = bool(t.get('critical'))
            if es <= 0:
                continue
            row_info[pt['id']] = {'y_center': yy + row_h / 2,
                                    'ES': es, 'EF': ef, 'critical': critical}

            # Hito puro (sin duración): solo diamante, sin barra
            if es_hito == 1 or dur == 0:
                if es < d_ini or es > d_fin:
                    continue
                cx = x + (es - d_ini) * day_w_pdf + day_w_pdf / 2
                cy = yy + row_h / 2
                s = max(2.2, min(row_h * 0.38, mm(2.2)))
                p.setRenderHint(QPainter.Antialiasing, True)
                poly = QPolygonF([
                    QPointF(cx, cy - s),
                    QPointF(cx + s, cy),
                    QPointF(cx, cy + s),
                    QPointF(cx - s, cy),
                ])
                # Sombra debajo
                shadow_poly = QPolygonF([
                    QPointF(cx, cy - s + mm(0.2)),
                    QPointF(cx + s, cy + mm(0.2)),
                    QPointF(cx, cy + s + mm(0.2)),
                    QPointF(cx - s, cy + mm(0.2)),
                ])
                p.setPen(Qt.NoPen)
                p.setBrush(QBrush(QColor(0, 0, 0, 30)))
                p.drawPolygon(shadow_poly)
                # Hito con borde — morado puro
                p.setBrush(QBrush(QColor("#7A36B1")))
                p.setPen(QPen(QColor("#4F1F76"), max(0.6, mm(0.2))))
                p.drawPolygon(poly)
                continue

            # Segmentos
            segs = []
            segs_json = cd.get('segmentos', '') or ''
            if segs_json:
                try:
                    segs = _json.loads(segs_json)
                except Exception:
                    segs = []

            # Tipografía escalada al alto de fila (para etiquetas dentro de barras)
            row_h_pt = row_h * 72.0 / dpi
            bar_lbl_pt = max(4.5, min(7.5, row_h_pt * 0.45))

            # Color de barra: prioridad custom > crítica > default — mirror de
            # _GanttBar._colors() en la UI.
            color_custom = (cd.get('color') or '').strip()
            if color_custom:
                _c = QColor(color_custom)
                if not _c.isValid():
                    _c = QColor("#F37329")
                bar_fill   = _c
                bar_top    = _c.lighter(130)
                bar_border = _c.darker(135)
            elif critical:
                bar_fill   = QColor("#C6262E")
                bar_top    = QColor("#ED5353")
                bar_border = QColor("#8B161B")
            else:
                bar_fill   = QColor(GANTT_BAR)
                bar_top    = QColor(GANTT_BAR_TOP)
                bar_border = QColor(GANTT_BAR_BORDER)

            # Radio de redondeo proporcional al alto de fila (~2.5pt en A4)
            radius = max(1.2, min(3.5, mm(0.8)))

            def _draw_bar(d_s, d_e, seg_idx=0, label=''):
                if d_e < d_ini or d_s > d_fin:
                    return
                a = max(d_s, d_ini)
                b = min(d_e, d_fin)
                bx = x + (a - d_ini) * day_w_pdf
                bw = (b - a + 1) * day_w_pdf
                by = yy + row_h * 0.18
                bh = row_h * 0.64
                p.setRenderHint(QPainter.Antialiasing, True)
                rect_bar = QRectF(bx, by, bw, bh)
                # Sombra muy sutil debajo (acentúa la barra contra fondo blanco)
                sh = max(0.5, mm(0.25))
                p.setPen(Qt.NoPen)
                p.setBrush(QBrush(QColor(0, 0, 0, 25)))
                p.drawRoundedRect(QRectF(bx, by + sh, bw, bh), radius, radius)
                # Gradiente vertical (más claro arriba, base abajo)
                grad = QLinearGradient(bx, by, bx, by + bh)
                grad.setColorAt(0.0, bar_top)
                grad.setColorAt(0.55, bar_fill)
                grad.setColorAt(1.0, bar_fill.darker(112))
                p.setBrush(QBrush(grad))
                # Borde 1px del color del fill oscurecido
                p.setPen(QPen(bar_border, max(0.6, mm(0.18))))
                p.drawRoundedRect(rect_bar, radius, radius)
                # Highlight superior (1px más claro, da efecto glossy)
                highlight = bar_top.lighter(120)
                hh = max(0.6, bh * 0.22)
                p.setPen(Qt.NoPen)
                p.setBrush(QBrush(QColor(255, 255, 255, 50)))
                hi_rect = QRectF(bx + 0.5, by + 0.5, bw - 1, hh)
                p.save()
                clip = QPainterPath()
                clip.addRoundedRect(rect_bar, radius, radius)
                p.setClipPath(clip)
                p.drawRect(hi_rect)
                p.restore()
                # Etiqueta dentro si cabe (texto blanco bold)
                if label and bw >= mm(8) and seg_idx == 0:
                    p.save()
                    p.setClipRect(rect_bar)
                    f = p.font(); f.setPointSizeF(bar_lbl_pt); f.setBold(True); p.setFont(f)
                    # Sombra de texto suave para legibilidad sobre cualquier color
                    p.setPen(QColor(0, 0, 0, 90))
                    p.drawText(QRectF(bx + 3, by + 0.5, bw - 6, bh),
                                  Qt.AlignLeft | Qt.AlignVCenter, label)
                    p.setPen(QColor("white"))
                    p.drawText(QRectF(bx + 3, by, bw - 6, bh),
                                  Qt.AlignLeft | Qt.AlignVCenter, label)
                    p.restore()

            if segs:
                offset = es - (segs[0].get('inicio_dia', es) if segs else es)
                first = True
                seg_spans_pdf = []
                for seg_idx, s in enumerate(segs):
                    s_ini = (s.get('inicio_dia', 1) or 1) + offset
                    s_dur = s.get('duracion', 0) or 0
                    if s_dur <= 0:
                        continue
                    _draw_bar(s_ini, s_ini + s_dur - 1, seg_idx,
                                label=pt['item'] or '' if first else '')
                    first = False
                    a = max(s_ini, d_ini)
                    b = min(s_ini + s_dur - 1, d_fin)
                    if a <= b:
                        seg_spans_pdf.append((
                            x + (a - d_ini) * day_w_pdf,
                            x + (b - d_ini + 1) * day_w_pdf,
                        ))
                # Línea entrecortada uniendo segmentos consecutivos
                seg_spans_pdf.sort()
                if len(seg_spans_pdf) >= 2:
                    cy = yy + row_h / 2
                    link_col = QColor("#A10705") if critical else QColor("#94A3B8")
                    pen_link = QPen(link_col, max(0.5, mm(0.15)), Qt.DashLine)
                    pen_link.setDashPattern([2.5, 2.5])
                    p.setPen(pen_link)
                    p.setBrush(Qt.NoBrush)
                    for (_, x_end), (x_start, _) in zip(seg_spans_pdf, seg_spans_pdf[1:]):
                        if x_start > x_end:
                            p.drawLine(QPointF(x_end, cy), QPointF(x_start, cy))
            else:
                _draw_bar(es, ef, 0, label=pt['item'] or '')

            # ── Barra de holgura (float) — espejo de la UI ────────────────
            if (getattr(self, '_show_slack', False) and not critical
                    and (segs == [])):
                slack_d = int(t.get('float', 0) or 0)
                if slack_d > 0:
                    fa = max(ef + 1, d_ini)
                    fb = min(ef + slack_d, d_fin)
                    if fa <= fb:
                        fx = x + (fa - d_ini) * day_w_pdf
                        fw = (fb - fa + 1) * day_w_pdf
                        fy = yy + row_h / 2 - max(0.6, mm(0.25))
                        p.setPen(Qt.NoPen)
                        p.setBrush(QBrush(QColor(148, 163, 184, 140)))
                        p.drawRect(QRectF(fx, fy, fw, max(1.2, mm(0.6))))
                        if ef + slack_d <= d_fin:
                            tip = QPolygonF([
                                QPointF(fx + fw, fy - mm(0.3)),
                                QPointF(fx + fw + mm(0.8), fy + max(0.6, mm(0.3))),
                                QPointF(fx + fw, fy + max(1.2, mm(0.6)) + mm(0.3)),
                            ])
                            p.setBrush(QBrush(QColor(148, 163, 184, 200)))
                            p.drawPolygon(tip)

            # Marcador de inicio/fin de fase (es_hito 2 o 3) — espejo del UI
            if es_hito in (2, 3):
                m_day = es if es_hito == 2 else ef + 1
                if d_ini <= m_day <= d_fin + 1:
                    m_x = x + (m_day - d_ini) * day_w_pdf
                    m_y = yy + row_h / 2
                    sm = max(2.2, min(row_h * 0.36, mm(2.0)))
                    fill_m = QColor("#3689E6") if es_hito == 2 else QColor("#3A9104")
                    border_m = QColor("#1E5DA8") if es_hito == 2 else QColor("#206700")
                    poly_m = QPolygonF([
                        QPointF(m_x, m_y - sm),
                        QPointF(m_x + sm, m_y),
                        QPointF(m_x, m_y + sm),
                        QPointF(m_x - sm, m_y),
                    ])
                    # Sombra
                    shadow = QPolygonF([
                        QPointF(m_x, m_y - sm + mm(0.2)),
                        QPointF(m_x + sm, m_y + mm(0.2)),
                        QPointF(m_x, m_y + sm + mm(0.2)),
                        QPointF(m_x - sm, m_y + mm(0.2)),
                    ])
                    p.setPen(Qt.NoPen)
                    p.setBrush(QBrush(QColor(0, 0, 0, 30)))
                    p.drawPolygon(shadow)
                    p.setBrush(QBrush(fill_m))
                    p.setPen(QPen(border_m, max(0.6, mm(0.18))))
                    p.drawPolygon(poly_m)

        # ── Flechas de dependencia ────────────────────────────────────────
        # Se pintan solo si ambos extremos caen en filas visibles del slice.
        # Pre-pass: agrupa flechas con mismo predecesor (≥2 sucesores) Z-
        # elegibles desde el stub default → comparten exactamente el mismo
        # stub_x para dibujar un tronco visual único (espejo del bundling
        # en _GanttArrow). Las flechas que no son Z-elegibles quedan fuera
        # del bundle y conservan ruteo individual.
        from collections import defaultdict as _dd_pdf
        from core.cronograma import INICIO_PID as _INI_pdf
        _HL_pdf = max(mm(0.5), min(mm(1.6), day_w_pdf * 0.45))
        _STUB_pdf = max(mm(0.6), min(mm(1.8), day_w_pdf * 0.45))
        _outgoing_pdf = _dd_pdf(list)
        for _r in range(r0, r1):
            _ptp = partidas[_r]
            if _ptp['es_titulo']: continue
            _ttp = tasks.get(_ptp['id'], {})
            for _prp in _ttp.get('preds', []):
                _pidp = _prp['pid']
                if _pidp not in row_info or _ptp['id'] not in row_info:
                    continue
                _pip = row_info[_pidp]
                _cip = row_info[_ptp['id']]
                _ti = _prp.get('tipo', 'FS') or 'FS'
                _pp = int(round(_prp.get('pct', 0) or 0))
                _tgp = int(round(_prp.get('tgt_pct', 0) or 0))
                if _tgp > 0:
                    _sf = True; _tf = False
                    _x1df = _pip['EF']
                    _cdurp = _cip['EF'] - _cip['ES'] + 1
                    _x2df = (_cip['ES'] - 1) + _cdurp * _tgp / 100.0
                elif _pp > 0:
                    _sf = False; _tf = False
                    _pdurp = _pip['EF'] - _pip['ES'] + 1
                    _x1df = (_pip['ES'] - 1) + _pdurp * _pp / 100.0
                    _x2df = _cip['ES'] - 1
                else:
                    _sf = _ti in ('FS', 'FF')
                    _tf = _ti in ('FF', 'SF')
                    _x1df = _pip['EF'] if _sf else _pip['ES'] - 1
                    _x2df = _cip['EF'] if _tf else _cip['ES'] - 1
                if _pidp == _INI_pdf:
                    _x1df = _pip['ES'] - 1
                if max(_x1df, _x2df) < d_ini - 1 or min(_x1df, _x2df) > d_fin:
                    continue
                _x1s = x + (_x1df - d_ini + 1) * day_w_pdf
                _x2s = x + (_x2df - d_ini + 1) * day_w_pdf
                _td_p = 1 if not _tf else -1
                _bx_p = _x2s - _td_p * _HL_pdf
                _y1pdf = _pip['y_center']
                _y2pdf = _cip['y_center']
                _crit_pdf = bool(_pip.get('critical')) and bool(_cip.get('critical'))
                _outgoing_pdf[(_pidp, _sf)].append(
                    (_bx_p, _td_p, _x1s, _ptp['id'], _y1pdf, _y2pdf, _crit_pdf)
                )
        # Lookup por flecha individual: (source_pid, target_pid) → stub_x
        # _bundle_trunks_pdf acumula tronco compartido por bundle. Mismo
        # criterio que en pantalla — usamos x2 (la punta) y no base_x, así
        # FS+0 adyacentes tambien bundlean con stub=x1 (vertical pura).
        _shared_stub_pdf = {}
        _bundle_trunks_pdf = []
        for _kk, _ee in _outgoing_pdf.items():
            if len(_ee) < 2: continue
            _sat = _kk[1]
            _ed = 1 if _sat else -1
            _tdirs_pdf = {e[1] for e in _ee}
            if len(_tdirs_pdf) > 1: continue
            _td_pdf = _tdirs_pdf.pop()
            _x1ref = _ee[0][2]
            _def_s = _x1ref + _ed * _STUB_pdf
            _x2s_pdf = [e[0] + e[1] * _HL_pdf for e in _ee]
            if _ed == _td_pdf:
                if _td_pdf == 1:
                    if min(_x2s_pdf) < _x1ref: continue
                    _sx_pdf = max(_x1ref, min(_def_s, min(_x2s_pdf)))
                else:
                    if max(_x2s_pdf) > _x1ref: continue
                    _sx_pdf = min(_x1ref, max(_def_s, max(_x2s_pdf)))
            else:
                continue
            _y1ref = _ee[0][4]
            _y2_list = [e[5] for e in _ee]
            _y_lo_pdf = min(_y1ref, min(_y2_list))
            _y_hi_pdf = max(_y1ref, max(_y2_list))
            _has_crit_pdf = any(e[6] for e in _ee)
            _bundle_trunks_pdf.append(
                (_x1ref, _y1ref, _sx_pdf, _y_lo_pdf, _y_hi_pdf, _has_crit_pdf)
            )
            for _bx, _tdir, _x1e, _tgt_pid, _y1e, _y2e, _crit_e in _ee:
                _shared_stub_pdf[(_kk[0], _tgt_pid)] = _sx_pdf
        for r in range(r0, r1):
            pt = partidas[r]
            if pt['es_titulo']:
                continue
            t = tasks.get(pt['id'], {})
            for pr in t.get('preds', []):
                pid_pred = pr['pid']
                if pid_pred not in row_info or pt['id'] not in row_info:
                    continue
                pred_info = row_info[pid_pred]
                cur_info  = row_info[pt['id']]
                tipo  = pr.get('tipo', 'FS') or 'FS'
                pct   = int(round(pr.get('pct', 0) or 0))
                tgt_p = int(round(pr.get('tgt_pct', 0) or 0))
                # Coordenadas espejo del UI: tgt_pct y pct cambian los puntos.
                if tgt_p > 0:
                    x1_day_f = pred_info['EF']
                    cur_dur = cur_info['EF'] - cur_info['ES'] + 1
                    x2_day_f = (cur_info['ES'] - 1) + cur_dur * tgt_p / 100.0
                elif pct > 0:
                    pred_dur = pred_info['EF'] - pred_info['ES'] + 1
                    x1_day_f = (pred_info['ES'] - 1) + pred_dur * pct / 100.0
                    x2_day_f = cur_info['ES'] - 1
                else:
                    # FS/SS/FF/SF: lado source/target según tipo
                    src_finish = tipo in ('FS', 'FF')
                    tgt_finish = tipo in ('FF', 'SF')
                    x1_day_f = pred_info['EF'] if src_finish else pred_info['ES'] - 1
                    x2_day_f = cur_info['EF'] if tgt_finish else cur_info['ES'] - 1
                # El hito «Inicio de Obra» marca el ARRANQUE del día 1: su línea
                # sale por la izquierda (borde de arranque), como MS Project.
                from core.cronograma import INICIO_PID as _INI
                if pid_pred == _INI:
                    x1_day_f = pred_info['ES'] - 1
                # Si la flecha cae por completo fuera del slice, omitir
                if max(x1_day_f, x2_day_f) < d_ini - 1 or min(x1_day_f, x2_day_f) > d_fin:
                    continue
                x1 = x + (x1_day_f - d_ini + 1) * day_w_pdf
                y1 = pred_info['y_center']
                x2 = x + (x2_day_f - d_ini + 1) * day_w_pdf
                y2 = cur_info['y_center']
                col = (QColor("#C6262E") if (pred_info['critical'] and cur_info['critical'])
                       else QColor("#6B7785"))
                # Ruteo ortogonal Z/C + punta rellena — espejo de _GanttArrow
                # en pantalla, para que la impresión se vea igual de prolija.
                if tgt_p > 0:
                    src_finish, tgt_finish = True, False
                elif pct > 0:
                    src_finish, tgt_finish = False, False
                else:
                    src_finish = tipo in ('FS', 'FF')
                    tgt_finish = tipo in ('FF', 'SF')
                exit_dir = 1 if src_finish else -1
                tip_dir  = 1 if not tgt_finish else -1
                # Codo/punta escalados al ancho de día: así, aun pegados al
                # borde izquierdo (hito de Inicio), el trazo no se sale del
                # margen ni se recorta.
                HEAD_LEN = max(mm(0.5), min(mm(1.6), day_w_pdf * 0.45))
                HEAD_W   = max(mm(0.5), min(mm(1.4), day_w_pdf * 0.70))
                STUB     = max(mm(0.6), min(mm(1.8), day_w_pdf * 0.45))
                base_x  = x2 - tip_dir * HEAD_LEN
                _stub_ovr = _shared_stub_pdf.get((pid_pred, pt['id']))
                if _stub_ovr is not None:
                    sx_stub = _stub_ovr
                    _force_z_pdf = True
                    _skip_trunk_pdf = True
                else:
                    sx_stub = x1 + exit_dir * STUB
                    _force_z_pdf = False
                    _skip_trunk_pdf = False
                if _skip_trunk_pdf:
                    # Solo lateral; el tronco lo pinta el item compartido
                    # despues del loop.
                    pts = [QPointF(sx_stub, y2), QPointF(base_x, y2)]
                else:
                    pts = [QPointF(x1, y1), QPointF(sx_stub, y1)]
                    if _force_z_pdf or (base_x - sx_stub) * tip_dir >= 0:
                        pts.append(QPointF(sx_stub, y2))
                        pts.append(QPointF(base_x, y2))
                    else:
                        mid_y  = (y1 + y2) / 2.0
                        turn_x = base_x - tip_dir * STUB
                        pts.append(QPointF(sx_stub, mid_y))
                        pts.append(QPointF(turn_x, mid_y))
                        pts.append(QPointF(turn_x, y2))
                        pts.append(QPointF(base_x, y2))
                p.setRenderHint(QPainter.Antialiasing, True)
                p.setPen(QPen(col, max(0.8, mm(0.18)), Qt.SolidLine, Qt.FlatCap, Qt.MiterJoin))
                p.setBrush(Qt.NoBrush)
                p.drawPolyline(QPolygonF(pts))
                head = QPolygonF([
                    QPointF(x2, y2),
                    QPointF(base_x, y2 - HEAD_W / 2.0),
                    QPointF(base_x, y2 + HEAD_W / 2.0),
                ])
                p.setPen(QPen(col, max(0.4, mm(0.06))))
                p.setBrush(QBrush(col))
                p.drawPolygon(head)
                mid_x = (x1 + x2) / 2.0   # ancla para la etiqueta del codo
                # Etiqueta del tipo cerca del codo (notación MS Project español)
                if pct > 0:
                    etiqueta = f"CC+{int(pct)}%"
                elif tgt_p > 0:
                    etiqueta = f"→{tgt_p}%"
                else:
                    lag_v = int(round(pr.get('lag', 0) or 0))
                    if tipo != 'FS' or lag_v != 0:
                        from core.cronograma import _TIPO_ES
                        etiqueta = _TIPO_ES.get(tipo, tipo) + (f"{lag_v:+d}" if lag_v else "")
                    else:
                        etiqueta = ''
                if etiqueta:
                    lbl_pt = max(4.0, min(6.5, row_h * 72.0 / dpi * 0.35))
                    f_lbl = p.font(); f_lbl.setPointSizeF(lbl_pt); f_lbl.setBold(True)
                    p.setFont(f_lbl)
                    fm_lbl = p.fontMetrics()
                    tw = fm_lbl.horizontalAdvance(etiqueta) + mm(1.0)
                    th = fm_lbl.height()
                    lx = mid_x - tw / 2
                    ly = (y1 + y2) / 2 - th / 2 - mm(0.3)
                    p.setPen(Qt.NoPen)
                    p.setBrush(QBrush(QColor(255, 255, 255, 220)))
                    p.drawRoundedRect(QRectF(lx, ly, tw, th), 1.5, 1.5)
                    p.setPen(QPen(QColor("#D4D4D4"), max(0.5, mm(0.05))))
                    p.drawRoundedRect(QRectF(lx, ly, tw, th), 1.5, 1.5)
                    p.setPen(QColor("#485A6C"))
                    p.drawText(QRectF(lx, ly, tw, th), Qt.AlignCenter, etiqueta)

        # ── Troncos compartidos (bundling MS Project) ─────────────────────
        for _x1t, _y1t, _sxt, _y_lo_t, _y_hi_t, _crit_t in _bundle_trunks_pdf:
            _tcol_pdf = QColor("#C6262E") if _crit_t else QColor("#6B7785")
            p.setPen(QPen(_tcol_pdf, max(0.8, mm(0.18)),
                          Qt.SolidLine, Qt.FlatCap, Qt.MiterJoin))
            p.setBrush(Qt.NoBrush)
            p.drawLine(QPointF(_x1t, _y1t), QPointF(_sxt, _y1t))
            p.drawLine(QPointF(_sxt, _y_lo_t), QPointF(_sxt, _y_hi_t))

        # ── Línea "hoy" ───────────────────────────────────────────────────
        if f_ini:
            from datetime import datetime as _dt
            try:
                hoy = _dt.now()
                day_n = (hoy - f_ini).days + 1
                if d_ini <= day_n <= d_fin:
                    xx = x + (day_n - d_ini) * day_w_pdf
                    p.setPen(QPen(QColor("#C6262E"), max(1, mm(0.25)), Qt.DashLine))
                    p.drawLine(QPointF(xx, y), QPointF(xx, y + h_total))
            except Exception:
                pass

        # ── Fin del plazo programado (línea roja entrecortada) ──────────────
        # Igual que en pantalla: marca el último día del plazo; si alguna barra
        # la cruza, la obra se pasa del tiempo previsto.
        _plazo = self._cv._proy.get('plazo') or 0
        if _plazo and (d_ini - 1) <= _plazo <= d_fin:
            xf = x + (_plazo - d_ini + 1) * day_w_pdf
            p.setPen(QPen(QColor("#C62828"), max(1, mm(0.28)), Qt.DashLine))
            p.drawLine(QPointF(xf, y), QPointF(xf, y + h_total))

        p.restore()

    def _pdf_paint_footer(self, p, x, y, w, h, page_num, total,
                              incluir_footer=True, incluir_legend=True,
                              incluir_page=True):
        """Footer estilo Centro de Reportes: separador + 3 columnas (cliente
        / fecha / página). Bajo eso una franja con la leyenda coloreada.
        Cada sección se puede excluir vía los flags `incluir_*`."""
        p.save()
        dpi = p.device().logicalDpiX()
        mm = lambda v: v * dpi / 25.4

        SLATE_500 = "#485A6C"
        SLATE_300 = "#94A3B8"
        SLATE_100 = "#E2E8F0"

        row1_on = incluir_footer or incluir_page
        row2_on = incluir_legend

        # Separador superior (solo si hay algo que pintar)
        p.setPen(QPen(QColor(SLATE_100), max(1, mm(0.18))))
        p.drawLine(QPointF(x + mm(4), y), QPointF(x + w - mm(4), y))

        # Cargar formato (pie izq/cen/der custom)
        fmt_ = {}
        try:
            fmt_ = _get_formato_reportes() or {}
        except Exception:
            pass

        if row1_on and row2_on:
            row1_h = h * 0.50
            row2_h = h - row1_h - mm(1.5)
            row1_y = y + mm(1.5)
            row2_y = y + row1_h + mm(0.5)
        elif row1_on:
            row1_h = h - mm(1.5)
            row1_y = y + mm(1.5)
            row2_h = 0
            row2_y = 0
        else:  # solo leyenda
            row2_h = h - mm(1.5)
            row2_y = y + mm(1.5)
            row1_h = 0
            row1_y = 0

        # ── Fila 1: pie izq | pie central | página X de Y ─────────────
        if row1_on:
            from datetime import datetime as _dt
            f1 = p.font(); f1.setPointSizeF(7.0); f1.setBold(False); f1.setItalic(False); p.setFont(f1)
            p.setPen(QColor(SLATE_300))

            proy = self._cv._proy

            # Solo pintar pie izq/cen si incluir_footer está activo
            if incluir_footer:
                pie_izq = (fmt_.get('rep_pie_izquierdo') or '').strip()
                if not pie_izq:
                    cliente = (proy.get('cliente') or '').strip()
                    pie_izq = f"Cliente: {cliente}" if cliente else ''
                pie_cen = (fmt_.get('rep_pie_central') or '').strip()
                if not pie_cen:
                    pie_cen = _dt.now().strftime('%d de %B de %Y, %H:%M')
                if pie_izq:
                    p.drawText(QRectF(x + mm(4), row1_y, w * 0.40, row1_h),
                                Qt.AlignLeft | Qt.AlignVCenter, pie_izq)
                if pie_cen:
                    p.drawText(QRectF(x + w * 0.30, row1_y, w * 0.40, row1_h),
                                Qt.AlignCenter | Qt.AlignVCenter, pie_cen)

            # Lado derecho: pie der personalizado o "Página X de Y" según flag.
            # Si hay pie_total_render (Reporte Completo), usar numeración global.
            if incluir_footer:
                pie_der_custom = (fmt_.get('rep_pie_derecho') or '').strip()
            else:
                pie_der_custom = ''
            global_total = getattr(self, '_pie_total_render', None)
            global_off   = getattr(self, '_pie_offset_render', 0) or 0
            if pie_der_custom:
                pie_der = pie_der_custom
            elif global_total and incluir_page:
                pie_der = f"Página {page_num + global_off} de {global_total}"
            elif incluir_page:
                pie_der = f"Página {page_num} de {total}"
            else:
                pie_der = ''
            if pie_der:
                p.drawText(QRectF(x + w - mm(4) - w * 0.30, row1_y, w * 0.30, row1_h),
                            Qt.AlignRight | Qt.AlignVCenter, pie_der)

        # ── Fila 2: leyenda con mini-iconos reales (centrada) ─────────
        if row2_on:
            leg_pt = max(4.5, min(8.0, row2_h * 72.0 / dpi * 0.55))
            f2 = p.font(); f2.setPointSizeF(leg_pt); f2.setBold(False); p.setFont(f2)
            fm = p.fontMetrics()
            cy = row2_y + row2_h / 2.0
            sw_h  = max(mm(1.1), row2_h * 0.42)   # alto del swatch
            bar_w = sw_h * 2.0
            dia_r = sw_h * 0.55
            arr_w = sw_h * 2.1
            dash_w = mm(1.4)

            items = [
                ('bar',     "#3689E6", "Tarea"),
                ('bar',     "#C6262E", "Crítica"),
                ('diamond', "#7A36B1", "Hito"),
                ('arrow',   "#6B7785", "Dependencia"),
                ('dash',    "#C6262E", "Hoy"),
            ]
            if (self._cv._proy.get('plazo') or 0):
                items.append(('dash', "#C62828", "Fin plazo"))

            def _sw_w(kind):
                return {'bar': bar_w, 'arrow': arr_w,
                        'diamond': dia_r * 2}.get(kind, dash_w)

            def _draw_sw(kind, col, sx):
                qc = QColor(col)
                p.setRenderHint(QPainter.Antialiasing, True)
                if kind == 'bar':
                    r = QRectF(sx, cy - sw_h / 2, bar_w, sw_h)
                    grad = QLinearGradient(sx, cy - sw_h / 2, sx, cy + sw_h / 2)
                    grad.setColorAt(0.0, qc.lighter(135))
                    grad.setColorAt(0.6, qc)
                    grad.setColorAt(1.0, qc.darker(112))
                    p.setPen(QPen(qc.darker(135), max(0.4, mm(0.12))))
                    p.setBrush(QBrush(grad))
                    p.drawRoundedRect(r, sw_h * 0.35, sw_h * 0.35)
                elif kind == 'diamond':
                    poly = QPolygonF([
                        QPointF(sx + dia_r, cy - dia_r),
                        QPointF(sx + dia_r * 2, cy),
                        QPointF(sx + dia_r, cy + dia_r),
                        QPointF(sx, cy),
                    ])
                    p.setPen(QPen(qc.darker(140), max(0.4, mm(0.1))))
                    p.setBrush(QBrush(qc))
                    p.drawPolygon(poly)
                elif kind == 'arrow':
                    hl = sw_h * 0.85
                    p.setPen(QPen(qc, max(0.6, mm(0.14)), Qt.SolidLine,
                                   Qt.FlatCap, Qt.MiterJoin))
                    p.setBrush(Qt.NoBrush)
                    p.drawLine(QPointF(sx, cy), QPointF(sx + arr_w - hl, cy))
                    head = QPolygonF([
                        QPointF(sx + arr_w, cy),
                        QPointF(sx + arr_w - hl, cy - hl * 0.5),
                        QPointF(sx + arr_w - hl, cy + hl * 0.5),
                    ])
                    p.setPen(Qt.NoPen); p.setBrush(QBrush(qc))
                    p.drawPolygon(head)
                else:  # dash (línea vertical entrecortada)
                    p.setPen(QPen(qc, max(0.6, mm(0.18)), Qt.DashLine))
                    mx = sx + dash_w / 2
                    p.drawLine(QPointF(mx, cy - sw_h * 0.75),
                                QPointF(mx, cy + sw_h * 0.75))

            gap = mm(3.0)
            sym_gap = mm(1.0)
            total_w = -gap
            for kind, _c, lbl in items:
                total_w += _sw_w(kind) + sym_gap + fm.horizontalAdvance(lbl) + gap
            cur_x = x + (w - total_w) / 2
            for kind, col, lbl in items:
                _draw_sw(kind, col, cur_x)
                cur_x += _sw_w(kind) + sym_gap
                lw = fm.horizontalAdvance(lbl)
                p.setPen(QColor(SLATE_500))
                p.setBrush(Qt.NoBrush)
                p.drawText(QRectF(cur_x, row2_y, lw + mm(0.5), row2_h),
                            Qt.AlignLeft | Qt.AlignVCenter, lbl)
                cur_x += lw + gap
        p.restore()

    def _exportar_mpp(self):
        """Exporta el cronograma a XML compatible con Microsoft Project (MSPDI).
        Incluye: metadata del proyecto, calendario con domingos no laborables
        y feriados como excepciones, milestones, ruta crítica, dependencias
        con lag, y notas con color personalizado de barra + spec resumida."""
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os
        sugerido = os.path.join(self._dir_descargas_gantt(),
                                  f"cronograma_{self._cv.pid}.xml")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar a Microsoft Project XML",
            sugerido, "MS Project XML (*.xml)"
        )
        if not path:
            return
        try:
            import xml.etree.ElementTree as ET
            import re as _re
            partidas = self._cv._partidas
            tasks = self._cv._tasks
            cmap = self._cv._cron_map
            proy = self._cv._proy

            # Hitos virtuales de proyecto (Inicio/Fin) derivados del CPM, para
            # cerrar la red al abrir en MS Project (sin tareas sueltas).
            _filas = self._cv.filas_con_hitos()
            hito_ini = next((f for f in _filas if f.get('_virtual') == 'inicio'), None)
            hito_fin = next((f for f in _filas if f.get('_virtual') == 'fin'), None)
            sin_pred = set(hito_ini['_succ_ids']) if hito_ini else set()
            terminales = hito_fin['_pred_ids'] if hito_fin else []

            # Cargar empresa desde el formato compartido de reportes
            fmt_ = {}
            try:
                fmt_ = _get_formato_reportes() or {}
            except Exception:
                pass
            empresa = (fmt_.get('rep_empresa_nombre') or '').strip()

            # Fecha de inicio del proyecto
            f_ini = self._project_start() or datetime.now()
            f_ini_str = f_ini.strftime("%Y-%m-%dT08:00:00")
            plazo = proy.get('plazo') or 0
            max_ef = max((t['EF'] for t in tasks.values() if t['EF'] > 0),
                          default=plazo)
            f_fin = f_ini + timedelta(days=max(max_ef, plazo))
            f_fin_str = f_fin.strftime("%Y-%m-%dT17:00:00")

            ns = "http://schemas.microsoft.com/project"
            ET.register_namespace('', ns)
            root = ET.Element(f"{{{ns}}}Project")

            # ── Metadata del proyecto (header) ────────────────────────────
            ET.SubElement(root, f"{{{ns}}}SaveVersion").text = "14"
            ET.SubElement(root, f"{{{ns}}}Name").text = proy.get('nombre', 'Proyecto')
            ET.SubElement(root, f"{{{ns}}}Title").text = proy.get('nombre', 'Proyecto')
            ET.SubElement(root, f"{{{ns}}}Subject").text = (proy.get('ubicacion') or '').strip()
            ET.SubElement(root, f"{{{ns}}}Category").text = (proy.get('modalidad') or '').strip()
            ET.SubElement(root, f"{{{ns}}}Company").text = empresa
            ET.SubElement(root, f"{{{ns}}}Manager").text = (proy.get('cliente') or '').strip()
            ET.SubElement(root, f"{{{ns}}}Author").text = empresa or 'ingePresupuestos'
            ET.SubElement(root, f"{{{ns}}}CreationDate").text = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
            ET.SubElement(root, f"{{{ns}}}StartDate").text = f_ini_str
            ET.SubElement(root, f"{{{ns}}}FinishDate").text = f_fin_str
            ET.SubElement(root, f"{{{ns}}}MinutesPerDay").text = "480"
            ET.SubElement(root, f"{{{ns}}}MinutesPerWeek").text = "2400"
            ET.SubElement(root, f"{{{ns}}}DaysPerMonth").text = "30"
            ET.SubElement(root, f"{{{ns}}}DefaultStartTime").text = "08:00:00"
            ET.SubElement(root, f"{{{ns}}}DefaultFinishTime").text = "17:00:00"
            ET.SubElement(root, f"{{{ns}}}DefaultFixedCostAccrual").text = "3"
            ET.SubElement(root, f"{{{ns}}}CalendarUID").text = "1"
            # Programación automática a nivel proyecto (espejo de captura/mpp.xml):
            # sin esto Project puede importar las tareas como «manuales» y
            # mostrar duración 0.
            ET.SubElement(root, f"{{{ns}}}ScheduleFromStart").text = "1"
            ET.SubElement(root, f"{{{ns}}}NewTasksAreManual").text = "0"
            ET.SubElement(root, f"{{{ns}}}DefaultTaskType").text = "0"

            # ── Calendario "Standard" con domingos no-lab + feriados ──────
            cals = ET.SubElement(root, f"{{{ns}}}Calendars")
            cal = ET.SubElement(cals, f"{{{ns}}}Calendar")
            ET.SubElement(cal, f"{{{ns}}}UID").text = "1"
            ET.SubElement(cal, f"{{{ns}}}Name").text = "Standard"
            ET.SubElement(cal, f"{{{ns}}}IsBaseCalendar").text = "1"
            ET.SubElement(cal, f"{{{ns}}}BaseCalendarUID").text = "-1"

            # WeekDays: MSPDI DayType 1=Domingo … 7=Sábado
            weeks = ET.SubElement(cal, f"{{{ns}}}WeekDays")
            for dn in range(1, 8):
                wd = ET.SubElement(weeks, f"{{{ns}}}WeekDay")
                ET.SubElement(wd, f"{{{ns}}}DayType").text = str(dn)
                # Domingo (1) no laborable; Lun-Sáb laborable
                working = "0" if dn == 1 else "1"
                ET.SubElement(wd, f"{{{ns}}}DayWorking").text = working
                if working == "1":
                    wts = ET.SubElement(wd, f"{{{ns}}}WorkingTimes")
                    wt = ET.SubElement(wts, f"{{{ns}}}WorkingTime")
                    ET.SubElement(wt, f"{{{ns}}}FromTime").text = "08:00:00"
                    ET.SubElement(wt, f"{{{ns}}}ToTime").text = "12:00:00"
                    wt2 = ET.SubElement(wts, f"{{{ns}}}WorkingTime")
                    ET.SubElement(wt2, f"{{{ns}}}FromTime").text = "13:00:00"
                    ET.SubElement(wt2, f"{{{ns}}}ToTime").text = "17:00:00"

            # Excepciones (feriados)
            feriados = sorted(self._feriados_set())
            if feriados:
                exc = ET.SubElement(cal, f"{{{ns}}}Exceptions")
                for fer in feriados:
                    e = ET.SubElement(exc, f"{{{ns}}}Exception")
                    ET.SubElement(e, f"{{{ns}}}EnteredByOccurrences").text = "0"
                    rng = ET.SubElement(e, f"{{{ns}}}TimePeriod")
                    ET.SubElement(rng, f"{{{ns}}}FromDate").text = f"{fer}T00:00:00"
                    ET.SubElement(rng, f"{{{ns}}}ToDate").text = f"{fer}T23:59:00"
                    ET.SubElement(e, f"{{{ns}}}Occurrences").text = "1"
                    ET.SubElement(e, f"{{{ns}}}Name").text = f"Feriado {fer}"
                    ET.SubElement(e, f"{{{ns}}}Type").text = "7"
                    ET.SubElement(e, f"{{{ns}}}DayWorking").text = "0"

            # ── Extended Attributes (campos personalizados) ───────────────
            ext_attrs = ET.SubElement(root, f"{{{ns}}}ExtendedAttributes")
            # Text30 = color personalizado de barra.
            ea = ET.SubElement(ext_attrs, f"{{{ns}}}ExtendedAttribute")
            ET.SubElement(ea, f"{{{ns}}}FieldID").text = "188744029"
            ET.SubElement(ea, f"{{{ns}}}FieldName").text = "Text30"
            ET.SubElement(ea, f"{{{ns}}}Alias").text = "ColorBarra"
            # Text29 = ID interno de la partida en ingePresupuestos. Permite
            # reimportar el XML editado en Project mapeando cada tarea a su
            # partida sin ambigüedad (el importador futuro lee este campo).
            ea_id = ET.SubElement(ext_attrs, f"{{{ns}}}ExtendedAttribute")
            ET.SubElement(ea_id, f"{{{ns}}}FieldID").text = "188744028"
            ET.SubElement(ea_id, f"{{{ns}}}FieldName").text = "Text29"
            ET.SubElement(ea_id, f"{{{ns}}}Alias").text = "IngeID"

            # ── Tasks ─────────────────────────────────────────────────────
            tasks_el = ET.SubElement(root, f"{{{ns}}}Tasks")
            uid_map = {}
            uid = 1
            row_n = 0

            def _spec_a_texto(html_or_text: str, max_len=400) -> str:
                """Pasa especificaciones HTML a texto plano truncado."""
                if not html_or_text:
                    return ''
                txt = str(html_or_text)
                # Quitar bloques <style>/<script> COMPLETOS (etiqueta + contenido);
                # si solo se quitaran las etiquetas, las reglas CSS de QTextDocument
                # se colarían como texto en la nota del Gantt MSPDI.
                txt = _re.sub(r'(?is)<(style|script)\b[^>]*>.*?</\1>', ' ', txt)
                txt = _re.sub(r'<[^>]+>', ' ', txt)
                txt = _re.sub(r'\s+', ' ', txt).strip()
                if len(txt) > max_len:
                    txt = txt[:max_len].rstrip() + '…'
                return txt

            def _emit_hito(nombre, dia, pred_uids):
                """Emite un Task hito (Milestone, 0 días) y devuelve su UID."""
                nonlocal uid, row_n
                row_n += 1
                el = ET.SubElement(tasks_el, f"{{{ns}}}Task")
                ET.SubElement(el, f"{{{ns}}}UID").text = str(uid)
                ET.SubElement(el, f"{{{ns}}}ID").text = str(row_n)
                ET.SubElement(el, f"{{{ns}}}Name").text = nombre
                ET.SubElement(el, f"{{{ns}}}OutlineLevel").text = "2"  # bajo el resumen
                f_h = f_ini + timedelta(days=max(0, (dia or 1) - 1))
                fh_str = f_h.strftime("%Y-%m-%dT08:00:00")
                ET.SubElement(el, f"{{{ns}}}Type").text = "1"
                ET.SubElement(el, f"{{{ns}}}Manual").text = "0"   # auto-programada
                ET.SubElement(el, f"{{{ns}}}Duration").text = "PT0H0M0S"
                ET.SubElement(el, f"{{{ns}}}DurationFormat").text = "7"
                ET.SubElement(el, f"{{{ns}}}Start").text = fh_str
                ET.SubElement(el, f"{{{ns}}}Finish").text = fh_str
                ET.SubElement(el, f"{{{ns}}}ConstraintType").text = "0"
                ET.SubElement(el, f"{{{ns}}}Summary").text = "0"
                ET.SubElement(el, f"{{{ns}}}Milestone").text = "1"
                for puid in pred_uids:
                    link = ET.SubElement(el, f"{{{ns}}}PredecessorLink")
                    ET.SubElement(link, f"{{{ns}}}PredecessorUID").text = str(puid)
                    ET.SubElement(link, f"{{{ns}}}Type").text = "1"  # FS
                cur = uid
                uid += 1
                return cur

            # Resumen del proyecto (#1) — contiene toda la obra (OutlineLevel 1).
            row_n += 1
            ps = ET.SubElement(tasks_el, f"{{{ns}}}Task")
            ET.SubElement(ps, f"{{{ns}}}UID").text = str(uid)
            ET.SubElement(ps, f"{{{ns}}}ID").text = str(row_n)
            ET.SubElement(ps, f"{{{ns}}}Name").text = proy.get('nombre', 'Proyecto')
            ET.SubElement(ps, f"{{{ns}}}OutlineLevel").text = "1"
            ET.SubElement(ps, f"{{{ns}}}Summary").text = "1"
            ET.SubElement(ps, f"{{{ns}}}Type").text = "1"
            ET.SubElement(ps, f"{{{ns}}}Manual").text = "0"   # auto-programada
            ET.SubElement(ps, f"{{{ns}}}Start").text = f_ini_str
            ET.SubElement(ps, f"{{{ns}}}Finish").text = f_fin_str
            uid += 1

            # Hito «Inicio de Obra» (#2) — sin predecesoras.
            inicio_uid = (_emit_hito(hito_ini['descripcion'], 1, [])
                          if hito_ini else None)
            # Mapear el id centinela del hito para que las predecesoras
            # explícitas "1" (que parse_predecesoras resuelve a INICIO_PID)
            # generen su PredecessorLink al hito de inicio.
            if inicio_uid is not None:
                from core.cronograma import INICIO_PID
                uid_map[INICIO_PID] = inicio_uid

            # Cabeceras de subpresupuesto (Summary, OutlineLevel 2) emitidas al
            # cambiar de grupo; las partidas cuelgan de ellas (OutlineLevel +2).
            sub_nombres = getattr(self._cv, '_sub_nombres', {None: 'Principal'})
            _prev_sub = object()

            for p in partidas:
                g = p.get('sub_presupuesto_id')
                if g != _prev_sub:
                    _prev_sub = g
                    row_n += 1
                    se = ET.SubElement(tasks_el, f"{{{ns}}}Task")
                    ET.SubElement(se, f"{{{ns}}}UID").text = str(uid)
                    ET.SubElement(se, f"{{{ns}}}ID").text = str(row_n)
                    ET.SubElement(se, f"{{{ns}}}Name").text = (sub_nombres.get(g) or 'Subpresupuesto')
                    ET.SubElement(se, f"{{{ns}}}OutlineLevel").text = "2"
                    ET.SubElement(se, f"{{{ns}}}Summary").text = "1"
                    ET.SubElement(se, f"{{{ns}}}Type").text = "1"
                    ET.SubElement(se, f"{{{ns}}}Manual").text = "0"   # auto-programada
                    ET.SubElement(se, f"{{{ns}}}Start").text = f_ini_str
                    ET.SubElement(se, f"{{{ns}}}Finish").text = f_fin_str
                    uid += 1

                row_n += 1
                t_el = ET.SubElement(tasks_el, f"{{{ns}}}Task")
                ET.SubElement(t_el, f"{{{ns}}}UID").text = str(uid)
                ET.SubElement(t_el, f"{{{ns}}}ID").text = str(row_n)
                ET.SubElement(t_el, f"{{{ns}}}Name").text = (p.get('item', '') + ' ' +
                                                              (p.get('descripcion') or '')).strip()
                # +2: cuelga del resumen del proyecto (1) y del subpresupuesto (2).
                ET.SubElement(t_el, f"{{{ns}}}OutlineLevel").text = str((p.get('nivel') or 1) + 2)
                ET.SubElement(t_el, f"{{{ns}}}WBS").text = p.get('item', '')
                # ID interno de la partida (Text29 / «IngeID») para reimportar.
                ea_pid = ET.SubElement(t_el, f"{{{ns}}}ExtendedAttribute")
                ET.SubElement(ea_pid, f"{{{ns}}}FieldID").text = "188744028"
                ET.SubElement(ea_pid, f"{{{ns}}}Value").text = str(p['id'])

                cd_full = cmap.get(p['id'], {})
                color_custom = (cd_full.get('color') or '').strip()
                spec_txt = _spec_a_texto(p.get('especificaciones') or '')

                if p['es_titulo']:
                    ET.SubElement(t_el, f"{{{ns}}}Summary").text = "1"
                    ET.SubElement(t_el, f"{{{ns}}}Type").text = "1"
                    ET.SubElement(t_el, f"{{{ns}}}Manual").text = "0"   # auto-programada
                    ET.SubElement(t_el, f"{{{ns}}}Duration").text = "PT0H0M0S"
                    ET.SubElement(t_el, f"{{{ns}}}Start").text = f_ini_str
                    ET.SubElement(t_el, f"{{{ns}}}Finish").text = f_ini_str
                else:
                    tinfo = tasks.get(p['id'], {})
                    es = tinfo.get('ES', 1)
                    ef = tinfo.get('EF', 1)
                    dur = tinfo.get('dur', 1) or 1
                    es_hito = int(cd_full.get('es_hito', 0) or 0)
                    # Solo es_hito==1 cuenta como milestone en MS Project;
                    # los valores 2/3 son marcadores de inicio/fin de fase que
                    # mantienen la duración como tarea normal.
                    is_hito = es_hito == 1 or dur == 0
                    f_es = f_ini + timedelta(days=es - 1)
                    f_ef = f_ini + timedelta(days=ef - 1) if ef > 0 else f_es
                    dur_str    = f"PT{dur*8}H0M0S"
                    start_str  = f_es.strftime("%Y-%m-%dT08:00:00")
                    finish_str = f_ef.strftime("%Y-%m-%dT17:00:00")
                    # Espejo de la estructura de tarea hoja que emite el propio
                    # MS Project (ver captura/mpp.xml): Manual=0 (auto), Type=0
                    # (Fixed Units, default de Project), y el triplete Manual*
                    # + Work + RemainingDuration. Sin estos campos, Project
                    # importa la tarea como manual sin duración → 0 días.
                    # IMPORTANTE: NO emitimos Finish/ManualFinish. Con tarea
                    # auto-programada (Manual=0, ASAP), si damos Start+Finish+
                    # Duration los tres, Project mueve el Start según los enlaces
                    # pero conserva nuestro Finish → duración = Finish-Start_nuevo
                    # = fraccional (0.75 d, etc.). Dando solo Start + Duration,
                    # Project deriva Finish = Start + Duration → duración entera.
                    ET.SubElement(t_el, f"{{{ns}}}Active").text = "1"
                    ET.SubElement(t_el, f"{{{ns}}}Manual").text = "0"
                    ET.SubElement(t_el, f"{{{ns}}}Type").text = "0"
                    ET.SubElement(t_el, f"{{{ns}}}EffortDriven").text = "0"
                    ET.SubElement(t_el, f"{{{ns}}}Start").text = start_str
                    ET.SubElement(t_el, f"{{{ns}}}Duration").text = dur_str
                    ET.SubElement(t_el, f"{{{ns}}}ManualStart").text = start_str
                    ET.SubElement(t_el, f"{{{ns}}}ManualDuration").text = dur_str
                    ET.SubElement(t_el, f"{{{ns}}}DurationFormat").text = "7"  # días
                    ET.SubElement(t_el, f"{{{ns}}}Work").text = "PT0H0M0S"
                    ET.SubElement(t_el, f"{{{ns}}}Estimated").text = "0"
                    ET.SubElement(t_el, f"{{{ns}}}RemainingDuration").text = dur_str
                    # Restricción de programación, igual que MS Project:
                    #  - CON predecesora → ASAP (0): la tarea fluye por enlaces.
                    #  - SIN predecesora → «Comenzar no antes del» (SNET=4) con
                    #    la fecha de inicio que el usuario fijó en ingePresupuestos,
                    #    SIN inventar predecesora. Así Project respeta la fecha
                    #    exacta (antes la mandaba al día 1 / la colgaba del hito).
                    if tinfo.get('preds'):
                        ET.SubElement(t_el, f"{{{ns}}}ConstraintType").text = "0"
                    else:
                        ET.SubElement(t_el, f"{{{ns}}}ConstraintType").text = "4"
                        ET.SubElement(t_el, f"{{{ns}}}ConstraintDate").text = start_str
                    ET.SubElement(t_el, f"{{{ns}}}Summary").text = "0"
                    if is_hito:
                        ET.SubElement(t_el, f"{{{ns}}}Milestone").text = "1"
                    if tinfo.get('critical'):
                        ET.SubElement(t_el, f"{{{ns}}}Critical").text = "1"
                    # Predecesoras — soporta FS/SS/FF/SF + lag, y aproxima
                    # pct/tgt_pct como FS con lag negativo (MS Project no
                    # tiene esos modos nativos).
                    # Códigos MPP Type: 0=FF, 1=FS, 2=SF, 3=SS
                    _TIPO_MPP = {'FF': '0', 'FS': '1', 'SF': '2', 'SS': '3'}
                    import math as _math
                    for pr in tinfo.get('preds', []):
                        ppid = pr.get('pid')
                        if ppid not in uid_map:
                            continue
                        link = ET.SubElement(t_el, f"{{{ns}}}PredecessorLink")
                        ET.SubElement(link, f"{{{ns}}}PredecessorUID").text = str(uid_map[ppid])
                        tipo_pr = (pr.get('tipo', 'FS') or 'FS').upper()
                        pct_pr = float(pr.get('pct', 0) or 0)
                        tgt_pct_pr = float(pr.get('tgt_pct', 0) or 0)
                        lag_days = int(pr.get('lag', 0) or 0)
                        if tgt_pct_pr > 0:
                            # B llega al X% cuando A termina → equivale a FS
                            # con lag negativo de dur_B * X / 100.
                            t_b = tasks.get(p['id'], {})
                            dur_b = max(1, (t_b.get('EF', 0) - t_b.get('ES', 0) + 1))
                            mpp_type = '1'  # FS
                            lag_days = -int(_math.ceil(dur_b * tgt_pct_pr / 100))
                        elif pct_pr > 0:
                            # B arranca cuando A lleva X% → equivale a SS con
                            # lag positivo de dur_A * X / 100.
                            t_a = tasks.get(ppid, {})
                            dur_a = max(1, (t_a.get('EF', 0) - t_a.get('ES', 0) + 1))
                            mpp_type = '3'  # SS
                            lag_days = int(_math.ceil(dur_a * pct_pr / 100))
                        else:
                            mpp_type = _TIPO_MPP.get(tipo_pr, '1')
                        ET.SubElement(link, f"{{{ns}}}Type").text = mpp_type
                        if lag_days:
                            # LinkLag en décimas de minuto; un día = 480*10 = 4800
                            ET.SubElement(link, f"{{{ns}}}LinkLag").text = str(lag_days * 4800)
                            ET.SubElement(link, f"{{{ns}}}LagFormat").text = "7"  # días

                    # (Ya NO colgamos del hito «Inicio de Obra» las tareas sin
                    # predecesora: ahora llevan restricción SNET con su fecha
                    # real, sin predecesora inventada — ver bloque ConstraintType.)

                # Notas (color personalizado + spec)
                note_lines = []
                if color_custom:
                    note_lines.append(f"Color de barra: {color_custom}")
                if spec_txt:
                    note_lines.append("Especificación: " + spec_txt)
                if note_lines:
                    ET.SubElement(t_el, f"{{{ns}}}Notes").text = "\n".join(note_lines)

                # Extended attribute: ColorBarra (Text30)
                if color_custom:
                    ev = ET.SubElement(t_el, f"{{{ns}}}ExtendedAttribute")
                    ET.SubElement(ev, f"{{{ns}}}FieldID").text = "188744029"
                    ET.SubElement(ev, f"{{{ns}}}Value").text = color_custom

                uid_map[p['id']] = uid
                uid += 1

            # Hito «Termino de Obra» — cuelga SOLO de la(s) tarea(s) que
            # terminan más tarde (mayor EF), no de todas las terminales, para
            # un Gantt más limpio. Si hay empate en el EF máximo, se enlazan
            # todas las que comparten ese fin.
            if hito_fin:
                cand = [pid for pid in terminales if pid in uid_map]
                if cand:
                    _ef = lambda pid: (tasks.get(pid, {}) or {}).get('EF', 0) or 0
                    max_ef = max(_ef(pid) for pid in cand)
                    cand = [pid for pid in cand if _ef(pid) == max_ef]
                term_uids = [uid_map[pid] for pid in cand]
                _emit_hito(hito_fin['descripcion'], self._cv._proj_end(), term_uids)

            # Indentar y guardar
            try:
                ET.indent(root)
            except Exception:
                pass
            tree = ET.ElementTree(root)
            tree.write(path, encoding='utf-8', xml_declaration=True)
            self._save_dir_gantt(path)

            msg = (
                f"Cronograma exportado a:\n{path}\n\n"
                "Incluye:\n"
                "  • Metadata: empresa, cliente, ubicación, modalidad\n"
                "  • Calendario: domingos no laborables\n"
                f"  • {len(feriados)} feriados como excepciones del calendario\n"
                "  • Hitos, ruta crítica y dependencias con lag\n"
                "  • Colores personalizados → notas + columna ColorBarra (Text30)\n\n"
                "Formato estándar abierto (XML de MS Project / MSPDI):\n"
                "ábrelo en MS Project o, SIN necesidad de licencia, en\n"
                "ProjectLibre (gratis — Windows, Mac y Linux). También\n"
                "compatible con GanttProject."
            )
            QMessageBox.information(self, "Exportar a MS Project", msg)
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")

    # ──────────────────────────────────────────────────────────────────────
    # Render del Gantt (QGraphicsScene)
    # ──────────────────────────────────────────────────────────────────────

    def _on_bar_change(self, partida_id, new_inicio, new_dur, mode, seg_idx):
        """Callback cuando el usuario suelta una barra después de drag/resize."""
        cmap = self._cv._cron_map
        if partida_id not in cmap:
            cmap[partida_id] = {'partida_id': partida_id, 'duracion': 1,
                                  'inicio_dia': 1, 'predecesoras': '',
                                  'es_hito': 0, 'segmentos': ''}
        cd = cmap[partida_id]
        # Si hay segmentos, modificar el segmento correspondiente
        segs_json = cd.get('segmentos', '') or ''
        if segs_json and mode != 'move-milestone':
            try:
                segs = _json.loads(segs_json)
                if 0 <= seg_idx < len(segs):
                    if mode == 'move':
                        segs[seg_idx]['inicio_dia'] = new_inicio
                    elif mode == 'resize-l':
                        delta = new_inicio - segs[seg_idx].get('inicio_dia', 1)
                        segs[seg_idx]['inicio_dia'] = new_inicio
                        segs[seg_idx]['duracion'] = max(1, segs[seg_idx].get('duracion', 1) - delta)
                    elif mode == 'resize-r':
                        segs[seg_idx]['duracion'] = new_dur
                    cd['segmentos'] = _json.dumps(segs)
                    # Recomputar inicio_dia y duracion como envoltura
                    if segs:
                        first = min(s.get('inicio_dia', 1) for s in segs)
                        last = max(s.get('inicio_dia', 1) + s.get('duracion', 0) - 1
                                    for s in segs)
                        cd['inicio_dia'] = first
                        cd['duracion'] = last - first + 1
            except Exception:
                pass
        else:
            if mode in ('move', 'resize-l'):
                cd['inicio_dia'] = new_inicio
            if mode in ('resize-l', 'resize-r'):
                cd['duracion'] = new_dur
        # Recalcular CPM y re-renderizar
        self._cv._calcular_cpm()
        self._llenar_tabla()
        self._render_gantt()

    # ──────────────────────────────────────────────────────────────────────
    # Dependencias gráficas (drag desde handles + menú sobre flechas)
    # ──────────────────────────────────────────────────────────────────────

    def _rownum_for_pid(self, pid: int):
        """# de fila (estilo MS Project Id: cuenta TODAS las filas) de una partida."""
        from core.cronograma import numerar_filas
        return numerar_filas(self._cv._partidas).get(pid)

    def _pid_for_rownum(self, n: int):
        from core.cronograma import numerar_filas
        for pid, rn in numerar_filas(self._cv._partidas).items():
            if rn == n:
                return pid
        return None

    def _bar_at_scene(self, scene_pos):
        """Encuentra _GanttBar bajo el puntero (ignora hitos y flechas)."""
        for it in self.scene.items(scene_pos):
            if isinstance(it, _GanttBar):
                return it
        return None

    def _dep_drag_start(self, source_pid, source_side, scene_pos):
        self._dep_drag = {'src_pid': source_pid, 'src_side': source_side,
                          'src_pos': QPointF(scene_pos)}
        # Línea fantasma
        pen = QPen(QColor("#3689E6"), 1.5, Qt.DashLine)
        line = QGraphicsLineItem(scene_pos.x(), scene_pos.y(),
                                  scene_pos.x(), scene_pos.y())
        line.setPen(pen)
        line.setZValue(50)
        self.scene.addItem(line)
        self._dep_ghost = line

    def _dep_drag_update(self, scene_pos):
        if not getattr(self, '_dep_drag', None) or not getattr(self, '_dep_ghost', None):
            return
        src = self._dep_drag['src_pos']
        self._dep_ghost.setLine(src.x(), src.y(), scene_pos.x(), scene_pos.y())

    def _dep_drag_cancel(self):
        """Quita la línea fantasma sin crear dependencia (al alternar entre
        enlazar y mover dentro del mismo arrastre)."""
        if getattr(self, '_dep_ghost', None):
            self.scene.removeItem(self._dep_ghost)
            self._dep_ghost = None
        self._dep_drag = None

    def _dep_drag_end(self, scene_pos):
        drag = getattr(self, '_dep_drag', None)
        if not drag:
            return
        # Limpiar ghost
        if getattr(self, '_dep_ghost', None):
            self.scene.removeItem(self._dep_ghost)
            self._dep_ghost = None
        self._dep_drag = None
        # Resolver target (debe ser una partida real, no una fila virtual)
        tgt_bar = self._bar_at_scene(scene_pos)
        if (tgt_bar is None or tgt_bar._pid == drag['src_pid']
                or tgt_bar._pid < 0):
            return
        # Posición relativa del drop dentro de la barra destino (0..1)
        r = tgt_bar.rect()
        rel = (scene_pos.x() - r.x()) / max(1, r.width())
        rel = max(0.0, min(1.0, rel))
        # Soltar en el tercio central = vínculo de porcentaje estilo MS Project:
        # el sucesor comienza cuando la pred. lleva el 50% (CC+50%).
        if 0.33 <= rel <= 0.66:
            self._set_dep(drag['src_pid'], tgt_bar._pid,
                            tipo='SS', lag=0, pct=50, tgt_pct=0)
            return
        # Lado del target según mitad de la barra donde se soltó
        on_finish = rel > 0.5
        target_side = 'finish' if on_finish else 'start'
        # Mapear lados a tipo MS Project
        tipo_map = {
            ('finish', 'start'):  'FS',
            ('finish', 'finish'): 'FF',
            ('start',  'start'):  'SS',
            ('start',  'finish'): 'SF',
        }
        tipo = tipo_map.get((drag['src_side'], target_side), 'FS')
        self._set_dep(drag['src_pid'], tgt_bar._pid, tipo=tipo, lag=0)

    def _parse_preds(self, preds_str: str):
        """Parsea 'predecesoras' en lista
        [(base_str, tipo, lag, pct, tgt_pct, raw_token)].
        Conserva referencias por item (no-numéricas) tal cual."""
        import re as _re
        from core.cronograma import _PAT_PRED, _TIPO_NORM
        out = []
        if not preds_str:
            return out
        pat = _PAT_PRED   # acepta notación española (FC/CC/CF) e inglesa
        for tok in _re.split(r'[,;]+', preds_str):
            tok = tok.strip()
            if not tok:
                continue
            m = pat.match(tok)
            if not m:
                out.append((tok, 'FS', 0, 0, 0, tok))
                continue
            base = (m.group(1) or '').strip()
            tipo = _TIPO_NORM.get((m.group(2) or 'FS').upper(), 'FS')
            sig  = m.group(3)
            num  = m.group(4)
            pctm = m.group(5)
            num_tgt = m.group(6)
            lag = 0
            pct = 0
            tgt_pct = 0
            if num:
                v = float(num)
                if sig == '-':
                    v = -v
                if pctm:
                    pct = int(round(v))
                else:
                    lag = int(round(v))
            elif num_tgt:
                tgt_pct = int(round(float(num_tgt)))
            out.append((base, tipo, lag, pct, tgt_pct, tok))
        return out

    def _build_pred_token(self, source_pid, tipo='FS', lag=0, pct=0, tgt_pct=0):
        """Construye el token de predecesora.
        Precedencia: tgt_pct > pct > tipo+lag."""
        rn = self._rownum_for_pid(source_pid)
        if rn is None:
            return None
        token = str(rn)
        if tgt_pct and tgt_pct > 0:
            # Formato '{rn}T{tgt_pct}%' — extensión propia (sin equivalente MSP),
            # conservada solo para leer datos antiguos.
            token += f"T{int(tgt_pct)}%"
            return token
        if pct and pct > 0:
            # MS Project: 'comenzar cuando la pred. lleva pct%' = Comienzo-a-
            # Comienzo con lag en porcentaje → '{rn}CC+{pct}%'.
            token += f"CC+{int(pct)}%"
            return token
        if tipo != 'FS':
            from core.cronograma import _TIPO_ES
            token += _TIPO_ES.get(tipo, tipo)   # notación española (CC/FF/CF)
        if lag != 0:
            token += ('+' if lag >= 0 else '-') + str(abs(int(lag)))
        return token

    def _replace_pred(self, target_pid, source_pid, new_token):
        """Reemplaza o agrega el token correspondiente al source en target.
        Si new_token es None, elimina la entrada (sirve para eliminar dep)."""
        cmap = self._cv._cron_map
        if target_pid not in cmap:
            cmap[target_pid] = {'partida_id': target_pid, 'duracion': 1,
                                  'inicio_dia': 1, 'predecesoras': '',
                                  'es_hito': 0, 'segmentos': '', 'color': ''}
        actual = (cmap[target_pid].get('predecesoras', '') or '').strip()
        items = self._parse_preds(actual)
        rn_src = self._rownum_for_pid(source_pid)
        nuevos = []
        reemplazado = False
        for base, _t, _l, _p, _tp, raw in items:
            # Misma fuente: comparar como número
            try:
                if rn_src is not None and int(base) == rn_src:
                    if new_token is not None and not reemplazado:
                        nuevos.append(new_token)
                        reemplazado = True
                    continue
            except (ValueError, TypeError):
                pass
            nuevos.append(raw)
        if new_token is not None and not reemplazado:
            nuevos.append(new_token)
        cmap[target_pid]['predecesoras'] = ', '.join(nuevos)

    def _evita_ciclo(self, source_pid, target_pid) -> bool:
        """¿Crear dep source→target generaría un ciclo? (target ya alcanza a source)."""
        if source_pid == target_pid:
            return True
        cmap = self._cv._cron_map
        # BFS desde target siguiendo sus predecesoras hacia atrás — si
        # alcanzamos source, existe ciclo invertido (target ya depende de algo
        # que depende de source si añadimos esta arista). Simplificación:
        # comprobar si source depende de target directa o indirectamente.
        visit = set()
        stack = [source_pid]
        while stack:
            cur = stack.pop()
            if cur in visit:
                continue
            visit.add(cur)
            preds = (cmap.get(cur, {}).get('predecesoras', '') or '')
            for base, _t, _l, _p, _tp, _raw in self._parse_preds(preds):
                try:
                    pid_pred = self._pid_for_rownum(int(base))
                except (ValueError, TypeError):
                    pid_pred = None
                if pid_pred == target_pid:
                    return True
                if pid_pred is not None and pid_pred not in visit:
                    stack.append(pid_pred)
        return False

    # ── Deshacer / rehacer dependencias (Ctrl+Z / Ctrl+Y) ──────────────────
    def _snapshot_preds(self) -> dict:
        """Copia {pid: predecesoras} de todas las partidas del cron_map."""
        return {pid: (d.get('predecesoras', '') or '')
                for pid, d in self._cv._cron_map.items()}

    def _push_undo_dep(self):
        """Guarda el estado actual de predecesoras antes de una mutación.
        Una acción nueva invalida la pila de rehacer."""
        self._dep_undo.append(self._snapshot_preds())
        if len(self._dep_undo) > 100:
            self._dep_undo.pop(0)
        self._dep_redo.clear()

    def _restaurar_preds(self, snap: dict):
        """Aplica un snapshot: cada partida del cron_map toma su predecesoras
        del snapshot (las creadas después del snapshot quedan sin pred.)."""
        for pid, d in self._cv._cron_map.items():
            d['predecesoras'] = snap.get(pid, '')
        self._cv._calcular_cpm()
        self._cv._guardar_a_db()
        self._llenar_tabla()
        self._render_gantt()

    def _dep_deshacer(self):
        if not self._dep_undo:
            return
        self._dep_redo.append(self._snapshot_preds())
        self._restaurar_preds(self._dep_undo.pop())

    def _dep_rehacer(self):
        if not self._dep_redo:
            return
        self._dep_undo.append(self._snapshot_preds())
        self._restaurar_preds(self._dep_redo.pop())

    def _set_dep(self, source_pid, target_pid, tipo='FS', lag=0, pct=0,
                  tgt_pct=0):
        if source_pid == target_pid:
            return
        if self._evita_ciclo(source_pid, target_pid):
            QMessageBox.warning(self, "Dependencia circular",
                                  "Esa dependencia crearía un ciclo entre partidas. "
                                  "Revisa la cadena de predecesoras.")
            return
        token = self._build_pred_token(source_pid, tipo, lag, pct, tgt_pct)
        if token is None:
            return
        self._push_undo_dep()
        self._replace_pred(target_pid, source_pid, token)
        self._cv._calcular_cpm()
        self._cv._guardar_a_db()
        self._llenar_tabla()
        self._render_gantt()

    def _remove_dep(self, source_pid, target_pid):
        self._push_undo_dep()
        self._replace_pred(target_pid, source_pid, None)
        self._cv._calcular_cpm()
        self._cv._guardar_a_db()
        self._llenar_tabla()
        self._render_gantt()

    def _arrow_ctx_menu(self, source_pid, target_pid, tipo_actual, lag_actual,
                          pct_actual, tgt_pct_actual, global_pos):
        menu = QMenu(self)
        # Título informativo
        rn_src = self._rownum_for_pid(source_pid)
        rn_tgt = self._rownum_for_pid(target_pid)
        if pct_actual:
            etiqueta = f"CC+{int(pct_actual)}%"
        elif tgt_pct_actual:
            etiqueta = f"sucesor al {tgt_pct_actual}% cuando pred. termina"
        else:
            from core.cronograma import _TIPO_ES
            etiqueta = _TIPO_ES.get(tipo_actual, tipo_actual) + (f"{lag_actual:+d}" if lag_actual else "")
        hdr = menu.addAction(f"#{rn_src} → #{rn_tgt}    {etiqueta}")
        hdr.setEnabled(False)
        menu.addSeparator()
        modo_especial = bool(pct_actual or tgt_pct_actual)
        # Cambiar tipo (no aplica si está en modo % o tgt_pct)
        for t_lbl, t_val, t_desc in [
            ('FC — Fin a Comienzo (default)', 'FS',
              'El sucesor inicia cuando la pred. termina'),
            ('CC — Comienzo a Comienzo',      'SS',
              'Ambas inician a la vez'),
            ('FF — Fin a Fin',                'FF',
              'Ambas terminan a la vez'),
            ('CF — Comienzo a Fin (raro)',    'SF',
              'El sucesor termina cuando la pred. inicia'),
        ]:
            marca = ' ✓' if (t_val == tipo_actual and not modo_especial) else ''
            act = menu.addAction(t_lbl + marca)
            act.setToolTip(t_desc)
            act.triggered.connect(
                lambda _=False, v=t_val: self._set_dep(
                    source_pid, target_pid, tipo=v,
                    lag=lag_actual if not modo_especial else 0,
                    pct=0, tgt_pct=0)
            )
        menu.addSeparator()
        # Editar lag (días)
        lbl_lag = (f"⏱ Editar lag (días)…   actual: {lag_actual:+d}"
                    if lag_actual and not modo_especial else "⏱ Editar lag (días)…")
        act_lag = menu.addAction(lbl_lag)
        act_lag.triggered.connect(
            lambda: self._edit_dep_lag(source_pid, target_pid, tipo_actual,
                                         lag_actual)
        )
        # Lag en porcentaje, estilo MS Project: CC+X% (B comienza cuando A
        # lleva X% de avance).
        lbl_pct = (f"◐ Lag en % — CC+X%…   actual: CC+{int(pct_actual)}%"
                    if pct_actual else "◐ Lag en % — CC+X%…")
        act_pct = menu.addAction(lbl_pct)
        act_pct.setToolTip("Comenzar B cuando A lleva X% de avance — equivale a "
                            "CC+X% en MS Project (p.ej. 50% del encofrado para "
                            "empezar a vaciar concreto).")
        act_pct.triggered.connect(
            lambda: self._edit_dep_pct(source_pid, target_pid, pct_actual)
        )
        # Extensión propia (sin equivalente MSP): solo se ofrece para editar/
        # quitar relaciones antiguas que ya la usen.
        if tgt_pct_actual:
            lbl_tgt = f"◑ Llegar al X% del sucesor (al fin de pred)…   actual: {tgt_pct_actual}%"
            act_tgt = menu.addAction(lbl_tgt)
            act_tgt.setToolTip("Relación heredada sin equivalente en MS Project. "
                                "Para cuando A termine, B ya debe estar al X%.")
            act_tgt.triggered.connect(
                lambda: self._edit_dep_tgt_pct(source_pid, target_pid, tgt_pct_actual)
            )
        if modo_especial:
            act_clr = menu.addAction("⊘ Quitar vínculo especial (volver a FS)")
            act_clr.triggered.connect(
                lambda: self._set_dep(source_pid, target_pid, tipo='FS',
                                        lag=0, pct=0, tgt_pct=0)
            )
        menu.addSeparator()
        act_del = menu.addAction("✕ Eliminar dependencia")
        act_del.triggered.connect(lambda: self._remove_dep(source_pid, target_pid))
        menu.exec(global_pos)

    def _edit_dep_lag(self, source_pid, target_pid, tipo, lag_actual):
        val, ok = QInputDialog.getInt(
            self, "Lag de la dependencia",
            "Días de adelanto (negativo) o retraso (positivo):",
            int(lag_actual), -365, 365, 1
        )
        if not ok:
            return
        self._set_dep(source_pid, target_pid, tipo=tipo, lag=int(val), pct=0)

    def _edit_dep_pct(self, source_pid, target_pid, pct_actual):
        val, ok = QInputDialog.getInt(
            self, "Lag en % (CC+X%)",
            "Porcentaje de avance de la pred. para iniciar este sucesor\n"
            "— equivale a CC+X% en MS Project (ej. 50 = a la mitad):",
            int(pct_actual) if pct_actual else 50, 1, 100, 5
        )
        if not ok:
            return
        # CC+X% = Comienzo-a-Comienzo con lag en porcentaje.
        self._set_dep(source_pid, target_pid, tipo='SS', lag=0, pct=int(val),
                        tgt_pct=0)

    def _edit_dep_tgt_pct(self, source_pid, target_pid, tgt_pct_actual):
        val, ok = QInputDialog.getInt(
            self, "Llegar al % del sucesor",
            "% que debe haber avanzado el SUCESOR cuando la pred. termine\n"
            "(ej. 50 = sucesor a la mitad, 75 = a 3/4 cuando A termine):",
            int(tgt_pct_actual) if tgt_pct_actual else 50, 1, 100, 5
        )
        if not ok:
            return
        self._set_dep(source_pid, target_pid, tipo='FS', lag=0, pct=0,
                        tgt_pct=int(val))

    def _delete_selected_arrow_if_any(self) -> bool:
        """Si hay una o más flechas seleccionadas en la escena del Gantt,
        las elimina (en batch) y devuelve True. Disparado desde el QShortcut
        Supr de ProyectoView."""
        sc = getattr(self, 'scene', None)
        if sc is None:
            return False
        pairs = [(it._source_pid, it._target_pid) for it in sc.selectedItems()
                  if isinstance(it, _GanttArrow)
                  and it._source_pid and it._target_pid]
        if not pairs:
            return False
        self._push_undo_dep()   # un solo paso de deshacer para todo el batch
        for source_pid, target_pid in pairs:
            self._replace_pred(target_pid, source_pid, None)
        self._cv._calcular_cpm()
        self._cv._guardar_a_db()
        self._llenar_tabla()
        self._render_gantt()
        return True

    # ── Resaltado de relaciones (pred/sucesor) ──────────────────────────
    def _related_pids(self, pid: int):
        """Devuelve (preds, sucs) — sets de pids relacionados con `pid`."""
        cmap = self._cv._cron_map
        preds = set()
        for base, _t, _l, _p, _tp, _raw in self._parse_preds(
                cmap.get(pid, {}).get('predecesoras', '') or ''):
            try:
                pp = self._pid_for_rownum(int(base))
                if pp:
                    preds.add(pp)
            except (ValueError, TypeError):
                pass
        sucs = set()
        for p in self._cv._partidas:
            if p['es_titulo']:
                continue
            for base, _t, _l, _p2, _tp, _raw in self._parse_preds(
                    cmap.get(p['id'], {}).get('predecesoras', '') or ''):
                try:
                    if self._pid_for_rownum(int(base)) == pid:
                        sucs.add(p['id'])
                        break
                except (ValueError, TypeError):
                    pass
        return preds, sucs

    def _on_selection_changed(self):
        """Resalta en azul las barras pred/sucesoras (y las flechas entre
        ellas) cuando hay un ítem seleccionado en la escena."""
        sc = self.scene
        related_bars   = set()
        selected_pids  = set()
        selected_arrows = []
        for it in sc.selectedItems():
            if isinstance(it, _GanttArrow):
                selected_arrows.append(it)
                if it._source_pid:
                    related_bars.add(it._source_pid)
                if it._target_pid:
                    related_bars.add(it._target_pid)
            elif isinstance(it, _GanttBar):
                selected_pids.add(it._pid)
                preds, sucs = self._related_pids(it._pid)
                related_bars |= preds | sucs
        # Bars no seleccionadas pero relacionadas
        for it in sc.items():
            if isinstance(it, _GanttBar):
                it.set_highlighted(
                    it._pid in related_bars and not it.isSelected()
                )
            elif isinstance(it, _GanttArrow):
                if it.isSelected():
                    continue
                rel = False
                if selected_pids and (it._source_pid in selected_pids
                                       or it._target_pid in selected_pids):
                    rel = True
                it.set_related_highlight(rel)

    def _on_arrow_drag(self, source_pid, target_pid, delta_days,
                          tipo, lag_actual, pct_actual, tgt_pct_actual=0):
        """Ajusta la dependencia al arrastrar la flecha horizontalmente.
        - modo tgt_pct: mueve el % del sucesor proporcional a los días.
        - modo pct:     mueve el % de la pred proporcional a los días.
        - modo normal:  suma/resta lag (días)."""
        tasks = self._cv._tasks
        if tgt_pct_actual and tgt_pct_actual > 0:
            t = tasks.get(target_pid, {})
            tgt_dur = max(1, (t.get('EF', 0) or 0) - (t.get('ES', 0) or 0) + 1)
            # Si A termina más tarde, B avanza más → tgt_pct sube
            delta_pct = int(round(delta_days / tgt_dur * 100))
            if delta_pct == 0:
                delta_pct = 1 if delta_days > 0 else -1
            new_pct = max(1, min(100, int(tgt_pct_actual) + delta_pct))
            if new_pct == tgt_pct_actual:
                return
            self._set_dep(source_pid, target_pid, tipo='FS', lag=0, pct=0,
                            tgt_pct=new_pct)
        elif pct_actual and pct_actual > 0:
            t = tasks.get(source_pid, {})
            pred_dur = max(1, (t.get('EF', 0) or 0) - (t.get('ES', 0) or 0) + 1)
            delta_pct = int(round(delta_days / pred_dur * 100))
            if delta_pct == 0:
                delta_pct = 1 if delta_days > 0 else -1
            new_pct = max(1, min(100, int(pct_actual) + delta_pct))
            if new_pct == pct_actual:
                return
            self._set_dep(source_pid, target_pid, tipo='FS', lag=0, pct=new_pct,
                            tgt_pct=0)
        else:
            new_lag = int(lag_actual) + int(delta_days)
            self._set_dep(source_pid, target_pid, tipo=tipo,
                            lag=new_lag, pct=0, tgt_pct=0)

    # ──────────────────────────────────────────────────────────────────────
    # Helpers de fechas y header
    # ──────────────────────────────────────────────────────────────────────

    def _project_start(self):
        """Devuelve datetime para el día 1 del proyecto. Usa el campo
        `fecha_inicio` si está configurado; cae a `costo_al` (legacy) y
        finalmente a la fecha actual."""
        for key in ('fecha_inicio', 'costo_al'):
            fi = (self._cv._proy.get(key) or '').strip()
            if fi:
                try:
                    return datetime.strptime(fi, '%Y-%m-%d')
                except Exception:
                    pass
        return datetime.now()

    def _day_to_date_str(self, d) -> str:
        """Convierte un número de día corrido (1=día 1 del proyecto) a una
        fecha 'dd/mm/yyyy'. Si no hay fecha de inicio o d es vacío, devuelve
        la representación del día como string."""
        if not d:
            return ''
        try:
            d = int(d)
        except (TypeError, ValueError):
            return str(d) if d else ''
        f_ini = self._project_start()
        if f_ini is None:
            return str(d)
        return (f_ini + timedelta(days=d - 1)).strftime('%d/%m/%Y')

    def _date_str_to_day(self, s) -> int | None:
        """Convierte 'dd/mm/yyyy' (o variantes) a número de día corrido del
        proyecto. Devuelve None si no se puede parsear."""
        if not s:
            return None
        s = str(s).strip()
        if not s:
            return None
        # Si es solo número, ya es día corrido
        if s.isdigit():
            return max(1, int(s))
        f_ini = self._project_start()
        if f_ini is None:
            return None
        for fmt in ('%d/%m/%Y', '%d-%m-%Y', '%Y-%m-%d', '%d/%m/%y', '%d-%m-%y'):
            try:
                fecha = datetime.strptime(s, fmt)
                delta = (fecha - f_ini).days + 1
                return max(1, delta)
            except ValueError:
                continue
        return None

    def _feriados_set(self) -> set:
        """Conjunto de fechas (YYYY-MM-DD) marcadas como feriados."""
        raw = (self._cv._proy.get('feriados') or '').strip()
        if not raw:
            return set()
        out = set()
        for tok in raw.replace(';', ',').split(','):
            tok = tok.strip()
            if tok:
                # Validar formato y normalizar
                try:
                    d = datetime.strptime(tok, '%Y-%m-%d')
                    out.add(d.strftime('%Y-%m-%d'))
                except Exception:
                    pass
        return out

    def _draw_header(self, n_dias, W, f_ini, target_scene=None, lead=0):
        """Dibuja el header de 2 filas: período arriba, día/semana abajo.
        Si `target_scene` se pasa, dibuja ahí; si no, usa self.scene.
        `lead` = px de margen vacío a la izquierda (antes del día 1)."""
        sc = target_scene if target_scene is not None else self.scene
        H_TOP = self.H_HDR_TOP
        H_BOT = self.H_HDR_BOT
        H_HDR = H_TOP + H_BOT

        # Fondos (se extienden al margen izquierdo para que sea continuo)
        rect_t = QGraphicsRectItem(-lead, 0, W + lead, H_TOP)
        rect_t.setBrush(QBrush(QColor("#E3E9F1")))
        rect_t.setPen(QPen(Qt.NoPen))
        sc.addItem(rect_t)
        rect_b = QGraphicsRectItem(-lead, H_TOP, W + lead, H_BOT)
        rect_b.setBrush(QBrush(QColor("#F2F5F9")))
        rect_b.setPen(QPen(Qt.NoPen))
        sc.addItem(rect_b)

        # ── Unidad del encabezado: explícita (días/semanas/meses) o derivada
        #    del zoom cuando está en "auto". ──
        unit = getattr(self, '_escala_unit', 'auto')
        if unit == 'auto':
            bottom = ('dia' if self.DAY_W >= 8
                      else 'semana' if self.DAY_W >= 4 else 'mes')
        else:
            bottom = unit
        top_is_months = (bottom != 'mes')

        # ── Top: meses (cuando abajo van días/semanas) o años (cuando meses) ──
        if top_is_months:
            self._draw_months_band(0, H_TOP, n_dias, f_ini, target_scene=sc)
        else:
            self._draw_years_band(0, H_TOP, n_dias, f_ini, target_scene=sc)

        # ── Bottom: días/semanas/meses ──
        if bottom == 'dia':
            # Días — etiquetar con el día-del-mes real (no el día corrido)
            # para que cada mes arranque su numeración en 1. El paso evita que
            # las etiquetas se encimen cuando el día es angosto.
            step = max(1, int(round(22 / max(1, self.DAY_W))))
            for d in range(1, n_dias + 1):
                x = (d - 1) * self.DAY_W
                if d % step == 0 or d == 1:
                    if f_ini is not None:
                        fecha_dia = f_ini + timedelta(days=d - 1)
                        label = str(fecha_dia.day)
                    else:
                        label = str(d)
                    txt = QGraphicsTextItem(label)
                    f = QFont(); f.setPointSize(7); txt.setFont(f)
                    txt.setDefaultTextColor(QColor(SLATE_500))
                    txt.setPos(x + 2, H_TOP + 2)
                    sc.addItem(txt)
                # Tick
                tick_y = H_TOP + H_BOT - 3
                line = QGraphicsLineItem(x, tick_y, x, H_TOP + H_BOT)
                line.setPen(QPen(QColor(SLATE_300), 0.5))
                sc.addItem(line)
        elif bottom == 'semana':
            # Semanas — paso para no encimar etiquetas con semanas angostas.
            wk_w = 7 * self.DAY_W
            wk_step = max(1, int(round(26 / max(1, wk_w))))
            for w in range((n_dias + 6) // 7):
                x = w * 7 * self.DAY_W
                if w % wk_step == 0:
                    txt = QGraphicsTextItem(f"S{w+1}")
                    f = QFont(); f.setPointSize(7); txt.setFont(f)
                    txt.setDefaultTextColor(QColor(SLATE_500))
                    txt.setPos(x + 2, H_TOP + 2)
                    sc.addItem(txt)
                line = QGraphicsLineItem(x, H_TOP, x, H_TOP + H_BOT)
                line.setPen(QPen(QColor(SLATE_300), 0.5))
                sc.addItem(line)
        else:
            # Meses (escala mensual)
            self._draw_months_band(H_TOP, H_BOT, n_dias, f_ini, short=True,
                                    target_scene=sc)

        # Línea separadora entre top y bottom
        sep1 = QGraphicsLineItem(-lead, H_TOP, W, H_TOP)
        sep1.setPen(QPen(QColor("#D0D8E4"), 0.5))
        sc.addItem(sep1)
        # Línea separadora bajo el header
        sep2 = QGraphicsLineItem(-lead, H_HDR, W, H_HDR)
        sep2.setPen(QPen(QColor(SILVER_300), 1))
        sc.addItem(sep2)

        return H_HDR

    def _draw_months_band(self, y, h, n_dias, f_ini, short=False, target_scene=None):
        """Dibuja bandas mensuales en el rango [y, y+h]."""
        sc = target_scene if target_scene is not None else self.scene
        MESES = ["Ene", "Feb", "Mar", "Abr", "May", "Jun",
                 "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
        if not f_ini:
            for i in range((n_dias + 29) // 30):
                x = i * 30 * self.DAY_W
                w = min(30 * self.DAY_W, (n_dias - i * 30) * self.DAY_W)
                line = QGraphicsLineItem(x, y, x, y + h)
                line.setPen(QPen(QColor("#D0D8E4"), 0.5))
                sc.addItem(line)
                txt = QGraphicsTextItem(f"Mes {i+1}")
                f = QFont(); f.setPointSize(7 if short else 8); f.setBold(not short)
                txt.setFont(f)
                txt.setDefaultTextColor(QColor(SLATE_700))
                txt.document().setDocumentMargin(0)
                th = txt.boundingRect().height()
                tw = txt.boundingRect().width()
                txt.setPos(x + max(2, (w - tw) / 2), y + max(0.0, (h - th) / 2.0))
                sc.addItem(txt)
            return

        cur_month = (f_ini.year, f_ini.month)
        band_start = 1
        for d in range(1, n_dias + 2):
            if d > n_dias:
                self._flush_month_band(y, h, band_start, n_dias, cur_month, short, sc)
                break
            date = f_ini + timedelta(days=d - 1)
            mk = (date.year, date.month)
            if mk != cur_month:
                self._flush_month_band(y, h, band_start, d - 1, cur_month, short, sc)
                cur_month = mk
                band_start = d

    def _flush_month_band(self, y, h, d_start, d_end, month_key, short, sc=None):
        if sc is None:
            sc = self.scene
        MESES = ["Ene", "Feb", "Mar", "Abr", "May", "Jun",
                 "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
        x_start = (d_start - 1) * self.DAY_W
        x_end = d_end * self.DAY_W
        w = x_end - x_start
        line = QGraphicsLineItem(x_start, y, x_start, y + h)
        line.setPen(QPen(QColor("#D0D8E4"), 0.5))
        sc.addItem(line)
        year, month = month_key
        if short:
            label = MESES[month - 1]
        else:
            label = f"{MESES[month - 1]} {year}"
        txt = QGraphicsTextItem(label)
        f = QFont(); f.setPointSize(7 if short else 8); f.setBold(not short)
        txt.setFont(f)
        txt.setDefaultTextColor(QColor(SLATE_700))
        # Sin margen interno del documento + centrado vertical: evita que la
        # banda inferior de meses (solo ~16px) corte los nombres por abajo.
        txt.document().setDocumentMargin(0)
        th = txt.boundingRect().height()
        ty = y + max(0.0, (h - th) / 2.0)
        tw = txt.boundingRect().width()
        if w >= tw + 4:
            txt.setPos(x_start + (w - tw) / 2, ty)
        elif w >= 20:
            short_lbl = MESES[month - 1] if not short else label[:3]
            txt.setPlainText(short_lbl)
            tw2 = txt.boundingRect().width()
            txt.setPos(x_start + max(1, (w - tw2) / 2), ty)
        else:
            return
        sc.addItem(txt)

    def _draw_years_band(self, y, h, n_dias, f_ini, target_scene=None):
        sc = target_scene if target_scene is not None else self.scene
        if not f_ini:
            txt = QGraphicsTextItem("Días 1 → " + str(n_dias))
            f = QFont(); f.setPointSize(9); f.setBold(True); txt.setFont(f)
            txt.setDefaultTextColor(QColor(SLATE_700))
            txt.setPos(8, y + 2)
            sc.addItem(txt)
            return
        cur_year = f_ini.year
        band_start = 1
        for d in range(1, n_dias + 2):
            if d > n_dias:
                self._flush_year_band(y, h, band_start, n_dias, cur_year, sc)
                break
            date = f_ini + timedelta(days=d - 1)
            if date.year != cur_year:
                self._flush_year_band(y, h, band_start, d - 1, cur_year, sc)
                cur_year = date.year
                band_start = d

    def _flush_year_band(self, y, h, d_start, d_end, year, sc=None):
        if sc is None:
            sc = self.scene
        x_start = (d_start - 1) * self.DAY_W
        x_end = d_end * self.DAY_W
        w = x_end - x_start
        line = QGraphicsLineItem(x_start, y, x_start, y + h)
        line.setPen(QPen(QColor("#D0D8E4"), 0.5))
        sc.addItem(line)
        if w < 30: return
        txt = QGraphicsTextItem(str(year))
        f = QFont(); f.setPointSize(9); f.setBold(True); txt.setFont(f)
        txt.setDefaultTextColor(QColor(SLATE_700))
        tw = txt.boundingRect().width()
        txt.setPos(x_start + (w - tw) / 2, y + 2)
        sc.addItem(txt)

    def _summary_for_title(self, title_p):
        """Calcula (es_min, ef_max) para los hijos de un título.

        Recibe el dict de partida del título (no un índice), para ser inmune al
        offset que introducen las filas virtuales de hitos en el render."""
        partidas = self._cv._partidas
        tasks = self._cv._tasks
        title_item = title_p['item'] or ''
        title_nivel = title_p['nivel'] or 1
        if not title_item:
            return None, None
        try:
            title_idx = next(i for i, p in enumerate(partidas)
                             if p['id'] == title_p['id'])
        except StopIteration:
            return None, None
        es_l, ef_l = [], []
        for j in range(title_idx + 1, len(partidas)):
            p = partidas[j]
            niv = p['nivel'] or 1
            it = p['item'] or ''
            # Detener si llegamos a otro título de igual o menor nivel
            if p['es_titulo'] and niv <= title_nivel:
                break
            if p['es_titulo']:
                continue
            # Solo descendientes
            if not (it == title_item or it.startswith(title_item + '.')):
                continue
            t = tasks.get(p['id'], {})
            es = t.get('ES', 0); ef = t.get('EF', 0)
            if es > 0 and ef > 0:
                es_l.append(es); ef_l.append(ef)
        if not es_l:
            return None, None
        return min(es_l), max(ef_l)

    def _render_gantt(self):
        self.scene.clear()
        self.hdr_scene.clear()
        if not hasattr(self._cv, '_partidas'):
            return
        partidas = self._cv.filas_con_hitos()   # [Inicio] + partidas + [Fin]
        tasks    = self._cv._tasks
        cmap     = self._cv._cron_map
        plazo    = self._cv._proy.get('plazo') or 0
        max_ef   = max((t['EF'] for t in tasks.values() if t['EF'] > 0),
                        default=plazo)
        n_dias   = max(max_ef, plazo, 30)

        f_ini = self._project_start()
        W = n_dias * self.DAY_W
        # Margen visible ANTES del día 1 (respiro a la izquierda, estilo MS
        # Project). La escena arranca en x negativo; las barras/flechas siguen
        # en x≥0 sin cambios — solo se extienden fondos y filas hacia la izq.
        lead = self.LEAD_DAYS * self.DAY_W
        # El header va a su propia escena (sticky); en el contenido principal
        # las barras arrancan en y=0 (sin offset)
        self._draw_header(n_dias, W, f_ini, target_scene=self.hdr_scene, lead=lead)
        H_HDR = 0
        H_total = len(partidas) * self.ROW_H
        self.hdr_scene.setSceneRect(-lead, 0, W + lead, self.H_HDR_TOP + self.H_HDR_BOT)
        self.scene.setSceneRect(-lead, 0, W + lead, H_total)

        # ── Sombreado de domingos y feriados ──────────────────────────────
        # z=-2 para quedar SOBRE las filas alt (z=-3) pero debajo de líneas/barras
        feriados = self._feriados_set()
        if f_ini is not None:
            from datetime import timedelta
            for d in range(1, n_dias + 1):
                fecha_dia = f_ini + timedelta(days=d - 1)
                es_domingo = fecha_dia.weekday() == 6  # Mon=0 ... Sun=6
                es_feriado = fecha_dia.strftime('%Y-%m-%d') in feriados
                if not (es_domingo or es_feriado):
                    continue
                x = (d - 1) * self.DAY_W
                rect = QGraphicsRectItem(x, H_HDR, self.DAY_W, H_total - H_HDR)
                if es_feriado:
                    rect.setBrush(QBrush(QColor(243, 115, 41, 38)))   # naranja claro
                else:
                    rect.setBrush(QBrush(QColor(148, 163, 184, 30)))  # slate-300 claro
                rect.setPen(QPen(Qt.NoPen))
                rect.setZValue(-2)
                self.scene.addItem(rect)

        # ── Grilla punteada (entrecortada), estilo MS Project ──────────────
        # Vertical: una por día si el día es ancho, o por semana si es angosto;
        # inicio de semana un poco más marcado.
        _grid_pen = QPen(QColor("#CDD3DB"), 0.7); _grid_pen.setStyle(Qt.DotLine)
        _week_pen = QPen(QColor("#AEB6C0"), 0.8); _week_pen.setStyle(Qt.DotLine)
        step_v = 1 if self.DAY_W >= 6 else 7
        for d in range(2, n_dias + 1):
            is_week = (d % 7 == 1)
            if step_v == 1:
                pen = _week_pen if is_week else _grid_pen
            elif is_week:
                pen = _week_pen
            else:
                continue
            x = (d - 1) * self.DAY_W
            line = QGraphicsLineItem(x, H_HDR, x, H_total)
            line.setPen(pen)
            line.setZValue(-1)
            self.scene.addItem(line)
        # Horizontal: una línea punteada al pie de cada fila (reemplaza la zebra).
        for r in range(len(partidas)):
            ly = H_HDR + (r + 1) * self.ROW_H
            line = QGraphicsLineItem(-lead, ly, W, ly)
            line.setPen(_grid_pen)
            line.setZValue(-1)
            self.scene.addItem(line)

        # ── Filas y barras ─────────────────────────────────────────────────
        row_info = {}

        for r, p in enumerate(partidas):
            y = H_HDR + r * self.ROW_H
            virtual = p.get('_virtual')
            # Resumen del proyecto / subpresupuesto: barra resumen del rango.
            if virtual in ('proyecto', 'subppto'):
                a = p.get('_ES') or 1
                b = p.get('_EF') or a
                x = (a - 1) * self.DAY_W
                w = (b - a + 1) * self.DAY_W
                by = y + self.ROW_H / 2 - 2
                bh = 5
                bar = QGraphicsRectItem(x, by, w, bh)
                bar.setBrush(QBrush(QColor("#1E2635")))
                bar.setPen(QPen(Qt.NoPen))
                bar.setZValue(1)
                self.scene.addItem(bar)
                for poly in (
                    QPolygonF([QPointF(x, by), QPointF(x + 4, by),
                               QPointF(x, by + bh + 3)]),
                    QPolygonF([QPointF(x + w, by), QPointF(x + w - 4, by),
                               QPointF(x + w, by + bh + 3)]),
                ):
                    tri = QGraphicsPolygonItem(poly)
                    tri.setBrush(QBrush(QColor("#1E2635")))
                    tri.setPen(QPen(Qt.NoPen))
                    tri.setZValue(1)
                    self.scene.addItem(tri)
                continue

            # Hitos virtuales (Inicio/Fin): diamante azul/verde.
            if virtual:
                dia = p.get('_ES') or 1
                s2 = _GanttMilestone.SIZE / 2
                if virtual == 'inicio':
                    # Borde de arranque del día (no centrado): inicio de obra.
                    cx = (dia - 1) * self.DAY_W + s2
                elif virtual == 'fin':
                    # Borde de cierre del último día.
                    cx = dia * self.DAY_W - s2
                else:
                    cx = (dia - 1) * self.DAY_W + self.DAY_W / 2
                cy = y + self.ROW_H / 2
                ms = _GanttMilestone(cx, cy, self.DAY_W, p['id'],
                                      lambda *a: None, p.get('es_hito', 2))
                ms.setToolTip(f"{p['descripcion']} · día {dia}")
                self.scene.addItem(ms)
                row_info[p['id']] = {'y_center': cy, 'ES': dia, 'EF': dia,
                                      'critical': False}
                continue

            if p['es_titulo']:
                # Barra de resumen MS Project (negra delgada con extremos triangulares)
                es_min, ef_max = self._summary_for_title(p)
                if es_min and ef_max:
                    x = (es_min - 1) * self.DAY_W
                    w = (ef_max - es_min + 1) * self.DAY_W
                    by = y + self.ROW_H / 2 - 2
                    bh = 5
                    bar = QGraphicsRectItem(x, by, w, bh)
                    bar.setBrush(QBrush(QColor("#1E2635")))
                    bar.setPen(QPen(Qt.NoPen))
                    bar.setZValue(1)
                    self.scene.addItem(bar)
                    # Triángulos en los extremos
                    poly_l = QPolygonF([
                        QPointF(x, by),
                        QPointF(x + 4, by),
                        QPointF(x, by + bh + 3),
                    ])
                    poly_r = QPolygonF([
                        QPointF(x + w, by),
                        QPointF(x + w - 4, by),
                        QPointF(x + w, by + bh + 3),
                    ])
                    for poly in (poly_l, poly_r):
                        tri = QGraphicsPolygonItem(poly)
                        tri.setBrush(QBrush(QColor("#1E2635")))
                        tri.setPen(QPen(Qt.NoPen))
                        tri.setZValue(1)
                        self.scene.addItem(tri)
                continue

            t = tasks.get(p['id'], {})
            es = t.get('ES', 0)
            ef = t.get('EF', 0)
            dur = t.get('dur', 0)
            cd = cmap.get(p['id'], {})
            es_hito = int(cd.get('es_hito', 0) or 0)
            critical = t.get('critical', False)

            if es <= 0:
                continue

            row_info[p['id']] = {'y_center': y + self.ROW_H / 2,
                                   'ES': es, 'EF': ef,
                                   'critical': critical}

            # Hito puro (sin duración): solo diamante, sin barra
            if es_hito == 1 or dur == 0:
                cx = (es - 1) * self.DAY_W + self.DAY_W / 2
                cy = y + self.ROW_H / 2
                ms = _GanttMilestone(cx, cy, self.DAY_W, p['id'],
                                      self._on_bar_change, max(es_hito, 1),
                                      on_context=self._show_partida_ctx_menu)
                ms.setToolTip(f"{p['item']} {p['descripcion']}\nHito · día {es}")
                self.scene.addItem(ms)
                continue

            # Segmentos
            segs_json = cd.get('segmentos', '') or ''
            segs = []
            if segs_json:
                try:
                    segs = _json.loads(segs_json)
                except Exception:
                    segs = []

            if segs:
                offset = es - (segs[0].get('inicio_dia', es) if segs else es)
                seg_spans = []   # [(x_inicio, x_fin)] para conectar con línea
                for seg_idx, s in enumerate(segs):
                    s_ini = (s.get('inicio_dia', 1) or 1) + offset
                    s_dur = s.get('duracion', 0) or 0
                    if s_dur <= 0: continue
                    x = (s_ini - 1) * self.DAY_W
                    w = s_dur * self.DAY_W
                    seg_spans.append((x, x + w))
                    bar = _GanttBar(x, y + 5, w, self.ROW_H - 10,
                                      self.DAY_W, p['id'],
                                      self._on_bar_change, critical, seg_idx,
                                      on_context=self._show_partida_ctx_menu,
                                      color=cd.get('color', ''),
                                      gantt_widget=self)
                    bar.setToolTip(
                        f"{p['item']} {p['descripcion']}\n"
                        f"Segmento {seg_idx+1}: día {s_ini} → {s_ini + s_dur - 1}"
                        + ("\n⚠ Ruta crítica" if critical else "")
                    )
                    self.scene.addItem(bar)
                    if w >= 30 and seg_idx == 0:
                        lbl = QGraphicsTextItem(p['item'] or '')
                        f = QFont(); f.setPointSize(7); f.setBold(True); lbl.setFont(f)
                        lbl.setDefaultTextColor(QColor("white"))
                        lbl.setPos(x + 4, y + 5)
                        self.scene.addItem(lbl)
                # Línea entrecortada que conecta segmentos consecutivos
                # (estilo MS Project: indica que pertenecen a la misma tarea)
                seg_spans.sort()
                if len(seg_spans) >= 2:
                    cy = y + self.ROW_H / 2
                    link_color = QColor("#A10705") if critical else QColor("#94A3B8")
                    for (_, x_end), (x_start, _) in zip(seg_spans, seg_spans[1:]):
                        if x_start <= x_end:
                            continue
                        ln = QGraphicsLineItem(x_end, cy, x_start, cy)
                        pen = QPen(link_color, 1.0, Qt.DashLine)
                        pen.setDashPattern([3, 3])
                        ln.setPen(pen)
                        ln.setZValue(0)
                        self.scene.addItem(ln)
            else:
                x = (es - 1) * self.DAY_W
                # Ancho = span calendario real (incluye días no laborables que
                # el CPM saltó). Así la barra termina exactamente donde arranca
                # la flecha de dependencia (en el borde derecho del día EF).
                w = max(1, ef - es + 1) * self.DAY_W
                bar = _GanttBar(x, y + 5, w, self.ROW_H - 10,
                                  self.DAY_W, p['id'],
                                  self._on_bar_change, critical, 0,
                                  on_context=self._show_partida_ctx_menu,
                                  color=cd.get('color', ''),
                                  gantt_widget=self)
                slack = int(t.get('float', 0) or 0)
                tip_extra = ("\n⚠ Ruta crítica" if critical
                              else (f"\nHolgura: {slack}d" if slack > 0 else ""))
                bar.setToolTip(
                    f"{p['item']} {p['descripcion']}\n"
                    f"Inicio: día {es}  ·  Fin: día {ef}  ·  Dur: {dur}d"
                    + tip_extra
                )
                self.scene.addItem(bar)
                # ── Barra de holgura (float) — solo no críticas y si está activado ──
                if not critical and slack > 0 and getattr(self, '_show_slack', False):
                    fx = ef * self.DAY_W
                    fw = slack * self.DAY_W
                    fy = y + self.ROW_H / 2 - 1
                    slack_bar = QGraphicsRectItem(fx, fy, fw, 3)
                    slack_bar.setBrush(QBrush(QColor(148, 163, 184, 130)))
                    slack_bar.setPen(QPen(Qt.NoPen))
                    slack_bar.setZValue(0)
                    slack_bar.setToolTip(
                        f"Holgura libre: {slack}d — la tarea puede atrasarse "
                        "ese tiempo sin afectar el fin del proyecto."
                    )
                    self.scene.addItem(slack_bar)
                    # Triángulo de cierre en la punta derecha (marca el LF)
                    tip = QPolygonF([
                        QPointF(fx + fw, fy - 2),
                        QPointF(fx + fw + 4, fy + 1.5),
                        QPointF(fx + fw, fy + 5),
                    ])
                    tri = QGraphicsPolygonItem(tip)
                    tri.setBrush(QBrush(QColor(148, 163, 184, 180)))
                    tri.setPen(QPen(Qt.NoPen))
                    tri.setZValue(0)
                    self.scene.addItem(tri)
                if w >= 30:
                    lbl = QGraphicsTextItem(p['item'] or '')
                    f = QFont(); f.setPointSize(7); f.setBold(True); lbl.setFont(f)
                    lbl.setDefaultTextColor(QColor("white"))
                    lbl.setPos(x + 4, y + 5)
                    self.scene.addItem(lbl)

            # Marcador de inicio/fin de fase (es_hito 2 o 3) — diamante extra
            # sobre la barra, sin alterar la duración.
            if es_hito in (2, 3):
                m_x = (es - 1) * self.DAY_W if es_hito == 2 else ef * self.DAY_W
                m_y = y + self.ROW_H / 2
                m_color = QColor("#3689E6") if es_hito == 2 else QColor("#3A9104")
                m_border = QColor("#1E5DA8") if es_hito == 2 else QColor("#206700")
                s = 6
                poly = QPolygonF([
                    QPointF(m_x, m_y - s),
                    QPointF(m_x + s, m_y),
                    QPointF(m_x, m_y + s),
                    QPointF(m_x - s, m_y),
                ])
                marker = QGraphicsPolygonItem(poly)
                marker.setBrush(QBrush(m_color))
                marker.setPen(QPen(m_border, 1.5))
                marker.setZValue(4)
                marker.setToolTip(
                    f"{p['item']} {p['descripcion']}\n"
                    + ("Marcador de inicio de fase" if es_hito == 2
                        else "Marcador de fin de fase")
                )
                self.scene.addItem(marker)

        # ── Flechas de dependencia ────────────────────────────────────────
        # Pre-pass: cuando un mismo predecesor tiene ≥2 sucesores que pueden
        # Z-rutear desde el stub default (9px), marcamos esas flechas para
        # compartir exactamente el mismo stub_x → un solo tronco visual
        # estilo MS Project. Las flechas cuyo target queda DEMASIADO CERCA
        # del stub default (necesitarían ruta C) quedan fuera del bundle y
        # conservan su ruteo individual — así evitamos encoger el stub para
        # forzar Z y arruinar el resto del grupo.
        from collections import defaultdict as _dd
        from core.cronograma import INICIO_PID as _INI_pre
        _HEAD_LEN_GANTT = 7.0
        _STUB_GANTT = 9.0
        _outgoing = _dd(list)  # (src_pid, src_finish) -> [(base_x, tip_dir, x1, target_pid)]
        for _pp in partidas:
            if _pp['es_titulo']: continue
            _tt = tasks.get(_pp['id'], {})
            for _pr in _tt.get('preds', []):
                _pid_pred = _pr['pid']
                if _pid_pred not in row_info or _pp['id'] not in row_info:
                    continue
                _pre_info = row_info[_pid_pred]
                _cur_info_pre = row_info[_pp['id']]
                _tipo = _pr.get('tipo', 'FS') or 'FS'
                _pct  = int(round(_pr.get('pct', 0) or 0))
                _tgt_pct = int(round(_pr.get('tgt_pct', 0) or 0))
                if _tgt_pct > 0:
                    _src_f = True; _tgt_f = False
                    _x1p = _pre_info['EF'] * self.DAY_W
                    _cdur = _cur_info_pre['EF'] - _cur_info_pre['ES'] + 1
                    _x2p = (_cur_info_pre['ES'] - 1 + _cdur * _tgt_pct / 100.0) * self.DAY_W
                elif _pct > 0:
                    _src_f = False; _tgt_f = False
                    _pdur = _pre_info['EF'] - _pre_info['ES'] + 1
                    _x1p = (_pre_info['ES'] - 1 + _pdur * _pct / 100.0) * self.DAY_W
                    _x2p = (_cur_info_pre['ES'] - 1) * self.DAY_W
                else:
                    _src_f = _tipo in ('FS', 'FF')
                    _tgt_f = _tipo in ('FF', 'SF')
                    _x1p = (_pre_info['EF'] * self.DAY_W) if _src_f \
                           else ((_pre_info['ES'] - 1) * self.DAY_W)
                    _x2p = (_cur_info_pre['EF'] * self.DAY_W) if _tgt_f \
                           else ((_cur_info_pre['ES'] - 1) * self.DAY_W)
                if _pid_pred == _INI_pre:
                    _x1p = (_pre_info['ES'] - 1) * self.DAY_W
                _td = 1 if not _tgt_f else -1
                _bx = _x2p - _td * _HEAD_LEN_GANTT
                _y1g = _pre_info['y_center']
                _y2g = _cur_info_pre['y_center']
                _crit_g = bool(_cur_info_pre.get('critical')) and bool(_pre_info.get('critical'))
                _outgoing[(_pid_pred, _src_f)].append(
                    (_bx, _td, _x1p, _pp['id'], _y1g, _y2g, _crit_g)
                )
        # _shared_stub mapea por flecha individual: (source_pid, target_pid) → stub_x
        # Calculamos el stub considerando el x2 real (la punta) y no el base_x,
        # para que casos FS+0 adyacentes (x2 = x1) tambien bundleen — el codo
        # se colapsa a una vertical pura en ese caso, igual que MS Project.
        # _bundle_trunks acumula el dibujo del tronco compartido por bundle.
        _shared_stub = {}
        _bundle_trunks = []
        for _key, _entries in _outgoing.items():
            if len(_entries) < 2: continue
            _sat = _key[1]
            _ed = 1 if _sat else -1
            _tdirs = {e[1] for e in _entries}
            if len(_tdirs) > 1: continue  # mix de direcciones — no bundleamos
            _td = _tdirs.pop()
            _x1r = _entries[0][2]
            _def_stub = _x1r + _ed * _STUB_GANTT
            # x2 = base_x + tip_dir * HEAD_LEN
            _x2s = [e[0] + e[1] * _HEAD_LEN_GANTT for e in _entries]
            # Caso comun: salida y llegada por el MISMO lado (FS, SS).
            if _ed == _td:
                if _td == 1:
                    # No bundleamos si algun sucesor empieza ANTES del pred (lead)
                    if min(_x2s) < _x1r: continue
                    _sx = max(_x1r, min(_def_stub, min(_x2s)))
                else:
                    if max(_x2s) > _x1r: continue
                    _sx = min(_x1r, max(_def_stub, max(_x2s)))
            else:
                # FF/SF: salida y llegada por lados opuestos — topologia mas
                # rara, mantenemos ruteo individual por ahora.
                continue
            _y1r = _entries[0][4]
            _y2s = [e[5] for e in _entries]
            _y_lo = min(_y1r, min(_y2s))
            _y_hi = max(_y1r, max(_y2s))
            _has_crit = any(e[6] for e in _entries)
            _bundle_trunks.append((_x1r, _y1r, _sx, _y_lo, _y_hi, _has_crit))
            for _bx, _td_e, _x1_e, _tgt_pid, _y1e, _y2e, _crit_e in _entries:
                _shared_stub[(_key[0], _tgt_pid)] = _sx
        for p in partidas:
            if p['es_titulo']: continue
            t = tasks.get(p['id'], {})
            for pr in t.get('preds', []):
                pid_pred = pr['pid']
                if pid_pred not in row_info or p['id'] not in row_info:
                    continue
                pred_info = row_info[pid_pred]
                cur_info  = row_info[p['id']]
                tipo = pr.get('tipo', 'FS') or 'FS'
                lag  = int(round(pr.get('lag', 0) or 0))
                pct  = int(round(pr.get('pct', 0) or 0))
                tgt_pct = int(round(pr.get('tgt_pct', 0) or 0))
                if tgt_pct > 0:
                    # Modo tgt_pct: la flecha sale del Fin de la pred y entra
                    # al punto tgt_pct% dentro de la barra del sucesor.
                    x1 = pred_info['EF'] * self.DAY_W
                    cur_es = cur_info['ES']
                    cur_ef = cur_info['EF']
                    cur_dur = cur_ef - cur_es + 1
                    x2 = (cur_es - 1 + cur_dur * tgt_pct / 100.0) * self.DAY_W
                    source_at_finish = True
                    target_at_finish = False
                elif pct > 0:
                    # Modo %: la flecha sale del punto pct% dentro de la barra
                    # de la pred y entra al inicio del sucesor (que es donde
                    # el motor CPM calcula su ES).
                    pred_es = pred_info['ES']
                    pred_ef = pred_info['EF']
                    pred_dur = pred_ef - pred_es + 1
                    x1 = (pred_es - 1 + pred_dur * pct / 100.0) * self.DAY_W
                    x2 = (cur_info['ES'] - 1) * self.DAY_W
                    source_at_finish = False
                    target_at_finish = False
                else:
                    # Lado de la pred (source): FS/FF parten del Fin; SS/SF del Inicio
                    source_at_finish = tipo in ('FS', 'FF')
                    # Lado del sucesor (target): FS/SS llegan al Inicio; FF/SF al Fin
                    target_at_finish = tipo in ('FF', 'SF')
                    if source_at_finish:
                        x1 = pred_info['EF'] * self.DAY_W
                    else:
                        x1 = (pred_info['ES'] - 1) * self.DAY_W
                    if target_at_finish:
                        x2 = cur_info['EF'] * self.DAY_W
                    else:
                        x2 = (cur_info['ES'] - 1) * self.DAY_W
                # El hito «Inicio de Obra» marca el ARRANQUE del día 1 (borde
                # izquierdo): su línea de dependencia debe salir por la izquierda,
                # como en MS Project (no por el borde derecho del día).
                from core.cronograma import INICIO_PID as _INI
                if pid_pred == _INI:
                    x1 = (pred_info['ES'] - 1) * self.DAY_W
                y1 = pred_info['y_center']
                y2 = cur_info['y_center']
                arrow = _GanttArrow(
                    x1, y1, x2, y2,
                    critical=cur_info['critical'] and pred_info['critical'],
                    tipo=tipo, lag=lag, pct=pct, tgt_pct=tgt_pct,
                    source_pid=pid_pred, target_pid=p['id'],
                    on_context=self._arrow_ctx_menu,
                    on_drag=self._on_arrow_drag,
                    day_w=self.DAY_W,
                    source_at_finish=source_at_finish,
                    target_at_finish=target_at_finish,
                    stub_x_override=_shared_stub.get((pid_pred, p['id'])),
                )
                self.scene.addItem(arrow)
                # Etiqueta cerca del codo (notación MS Project en español)
                if pct > 0:
                    etiqueta = f"CC+{int(pct)}%"
                elif tgt_pct > 0:
                    etiqueta = f"→{tgt_pct}%"   # legacy (sin equivalente MSP)
                elif tipo != 'FS' or lag != 0:
                    from core.cronograma import _TIPO_ES
                    etiqueta = _TIPO_ES.get(tipo, tipo) + (f"{lag:+d}" if lag else "")
                else:
                    etiqueta = ''
                if etiqueta:
                    lbl = QGraphicsTextItem(etiqueta)
                    f = QFont(); f.setPointSize(7); f.setBold(True); lbl.setFont(f)
                    lbl.setDefaultTextColor(QColor("#485A6C"))
                    # Pequeño fondo blanco semi-opaco para que se lea sobre la
                    # zebra (lo logramos vía HTML inline)
                    lbl.setHtml(
                        f"<div style='background:rgba(255,255,255,0.85);"
                        f" padding:0 2px; border:1px solid #D4D4D4;"
                        f" border-radius:2px; color:#485A6C;"
                        f" font-size:7pt; font-weight:700;'>{etiqueta}</div>"
                    )
                    mid_x = (x1 + x2) / 2
                    mid_y = (y1 + y2) / 2 - 8
                    lbl.setPos(mid_x - 14, mid_y)
                    lbl.setZValue(3)
                    self.scene.addItem(lbl)

        # ── Troncos compartidos (bundling MS Project) ─────────────────────
        # Un solo QGraphicsPathItem por grupo: horizontal desde el predecesor
        # hasta sx_stub, más vertical cubriendo todo el rango de filas
        # destino. Color crítico si alguna flecha del grupo lo es. Se dibuja
        # bajo las flechas individuales para que el tronco no oculte sus
        # ramas laterales ni interfiera con su hit-area.
        for _x1t, _y1t, _sxt, _y_lo_t, _y_hi_t, _crit_t in _bundle_trunks:
            _tp = QPainterPath()
            _tp.moveTo(QPointF(_x1t, _y1t))
            _tp.lineTo(QPointF(_sxt, _y1t))
            _tp.moveTo(QPointF(_sxt, _y_lo_t))
            _tp.lineTo(QPointF(_sxt, _y_hi_t))
            _tit = QGraphicsPathItem(_tp)
            _tcol = QColor("#cc3b02") if _crit_t else QColor("#6B7785")
            _tit.setPen(QPen(_tcol, 1.3, Qt.SolidLine, Qt.FlatCap, Qt.MiterJoin))
            _tit.setBrush(Qt.NoBrush)
            _tit.setZValue(1.9)
            self.scene.addItem(_tit)

        # ── Línea "hoy" ────────────────────────────────────────────────────
        if f_ini:
            try:
                hoy = datetime.now()
                day_n = (hoy - f_ini).days + 1
                if 1 <= day_n <= n_dias:
                    x = (day_n - 1) * self.DAY_W
                    line = QGraphicsLineItem(x, H_HDR, x, H_total)
                    line.setPen(QPen(QColor("#C6262E"), 1.5, Qt.DashLine))
                    line.setZValue(3)
                    self.scene.addItem(line)
                    lbl = QGraphicsTextItem("HOY")
                    f = QFont(); f.setPointSize(7); f.setBold(True); lbl.setFont(f)
                    lbl.setDefaultTextColor(QColor("#C6262E"))
                    lbl.setPos(x + 2, H_HDR - 12)
                    self.scene.addItem(lbl)
            except Exception:
                pass

        # ── Fin del plazo programado (línea roja entrecortada) ──────────────
        # Marca el último día del plazo previsto: si alguna barra la cruza, la
        # obra se está pasando del tiempo programado.
        if plazo and plazo > 0:
            xf = plazo * self.DAY_W
            line = QGraphicsLineItem(xf, H_HDR, xf, H_total)
            line.setPen(QPen(QColor("#C62828"), 1.6, Qt.DashLine))
            line.setZValue(4)
            self.scene.addItem(line)
            lbl = QGraphicsTextItem()
            lbl.setHtml(
                "<div style='background:rgba(255,255,255,0.9); color:#C62828;"
                " padding:0 3px; border:1px solid #E0A0A0; border-radius:2px;"
                " font-size:7pt; font-weight:700;'>Fin plazo</div>"
            )
            lbl.setPos(xf + 3, 2)
            lbl.setZValue(4)
            self.scene.addItem(lbl)

        self.scene.setSceneRect(-lead, 0, W + lead, H_total)


# ════════════════════════════════════════════════════════════════════════════
# 2. CRONOGRAMA VALORIZADO
# ════════════════════════════════════════════════════════════════════════════


class _DialogExportarReportePdf(QDialog):
    """Diálogo genérico para exportar reportes a PDF con papel/orientación
    configurables + vista previa + exportar como imagen. Comparte estilo
    con `_DialogExportarGanttPdf` pero más simple (no aplica paginación
    fit/multi ni escala temporal — esos son del Gantt)."""

    preview_solicitado = Signal(dict)
    imagen_solicitada  = Signal(dict)

    PAPEL_OPCIONES = [
        ("A4 (210 × 297 mm)", "A4"),
        ("A3 (297 × 420 mm)", "A3"),
        ("A2 (420 × 594 mm)", "A2"),
        ("A1 (594 × 841 mm)", "A1"),
        ("A0 (841 × 1189 mm)", "A0"),
        ("Carta / Letter (216 × 279 mm)", "Letter"),
        ("Tabloide / Ledger (279 × 432 mm)", "Tabloid"),
    ]

    def __init__(self, parent=None, titulo: str = "Exportar reporte"):
        super().__init__(parent)
        self.setWindowTitle(titulo)
        self.setMinimumWidth(520)
        self.setMinimumHeight(420)
        self.resize(540, 460)
        self.setModal(True)

        from PySide6.QtGui import QPalette
        pal = self.palette()
        pal.setColor(QPalette.Window, QColor("white"))
        pal.setColor(QPalette.Base, QColor("white"))
        pal.setColor(QPalette.WindowText, QColor("#273445"))
        pal.setColor(QPalette.Text, QColor("#273445"))
        self.setPalette(pal)
        self.setAutoFillBackground(True)
        self.setStyleSheet(
            "QDialog { background:white; }"
            "QLabel { color:#273445; font-size:12px; background:transparent; }"
            "QRadioButton { color:#273445; font-size:12px; padding:4px 0;"
            "  background:transparent; spacing:8px; }"
            "QRadioButton::indicator { width:16px; height:16px;"
            "  background:transparent; border:none; }"
            "QRadioButton::indicator:unchecked {"
            "  image: url(resources/icons/radio_orange_off.svg); }"
            "QRadioButton::indicator:checked {"
            "  image: url(resources/icons/radio_orange_on.svg); }"
            "QCheckBox { color:#273445; font-size:12px; padding:2px 0;"
            "  background:transparent; spacing:8px; }"
            "QCheckBox::indicator { width:16px; height:16px;"
            "  background:transparent; border:none; }"
            "QCheckBox::indicator:unchecked {"
            "  image: url(resources/icons/check_orange_off.svg); }"
            "QCheckBox::indicator:checked {"
            "  image: url(resources/icons/check_orange_on.svg); }"
        )

        vl = QVBoxLayout(self)
        vl.setContentsMargins(18, 16, 18, 14)
        vl.setSpacing(10)

        ttl = QLabel("Opciones de exportación")
        f = ttl.font(); f.setPointSize(13); f.setBold(True); ttl.setFont(f)
        ttl.setStyleSheet("color:#1E2635;")
        vl.addWidget(ttl)

        # Tamaño de papel
        hh_pap = QHBoxLayout()
        hh_pap.addWidget(QLabel("Tamaño de papel:"))
        self.cmb_papel = QComboBox()
        for label, _ in self.PAPEL_OPCIONES:
            self.cmb_papel.addItem(label)
        self.cmb_papel.setCurrentIndex(0)  # A4
        self.cmb_papel.setStyleSheet(
            "QComboBox { padding:3px 6px; border:1px solid #D4D4D4;"
            " border-radius:4px; background:white; min-width:200px; }"
        )
        hh_pap.addWidget(self.cmb_papel, 1)
        vl.addLayout(hh_pap)

        # Orientación
        vl.addWidget(QLabel("Orientación:"))
        self.rb_port = QRadioButton("Vertical (recomendado)")
        self.rb_land = QRadioButton("Horizontal")
        self.rb_port.setChecked(True)
        grp_or = QButtonGroup(self)
        grp_or.addButton(self.rb_port)
        grp_or.addButton(self.rb_land)
        hh = QHBoxLayout()
        hh.addWidget(self.rb_port)
        hh.addWidget(self.rb_land)
        hh.addStretch()
        vl.addLayout(hh)

        # Cuántos períodos (semanas/meses) entran por hoja
        from PySide6.QtWidgets import QSpinBox
        vl.addWidget(QLabel("Períodos por hoja:"))
        self.rb_pag_una = QRadioButton(
            "Todos los períodos en una hoja (comprime las columnas)")
        self.rb_pag_auto = QRadioButton(
            "Automático (la app decide cuántos caben según el papel)")
        self.rb_pag_auto.setChecked(True)
        self.rb_pag_man = QRadioButton("Manual:")
        grp_pag = QButtonGroup(self)
        for rb in (self.rb_pag_una, self.rb_pag_auto, self.rb_pag_man):
            grp_pag.addButton(rb)
        vl.addWidget(self.rb_pag_una)
        vl.addWidget(self.rb_pag_auto)

        hh_pag = QHBoxLayout()
        hh_pag.setContentsMargins(0, 0, 0, 0)
        hh_pag.addWidget(self.rb_pag_man)
        self.sp_paginas = QSpinBox()
        self.sp_paginas.setRange(2, 52)
        self.sp_paginas.setValue(12)
        self.sp_paginas.setStyleSheet(
            "QSpinBox { padding:3px 6px; border:1px solid #D4D4D4;"
            " border-radius:4px; background:white; min-width:60px; }"
        )
        hh_pag.addWidget(self.sp_paginas)
        hh_pag.addWidget(QLabel("períodos por hoja"))
        hh_pag.addStretch()
        vl.addLayout(hh_pag)

        # Habilitar spinbox solo cuando "Manual" esté seleccionado
        def _on_pag_mode():
            self.sp_paginas.setEnabled(self.rb_pag_man.isChecked())
        for rb in (self.rb_pag_una, self.rb_pag_auto, self.rb_pag_man):
            rb.toggled.connect(_on_pag_mode)
        _on_pag_mode()

        # Checkbox: mostrar fechas en columnas (igual que en UI)
        from PySide6.QtWidgets import QCheckBox as _QChk
        self.chk_fechas = _QChk("Mostrar fechas de inicio/fin en cada columna")
        self.chk_fechas.setChecked(True)
        vl.addWidget(self.chk_fechas)

        hint = QLabel(
            "Logo, nombre de empresa, colores y textos del encabezado/pie "
            "se toman del formato compartido con el Centro de Reportes."
        )
        hint.setStyleSheet("color:#667885; font-size:10px;")
        hint.setWordWrap(True)
        vl.addWidget(hint)

        vl.addStretch()

        # Botones
        bar = QHBoxLayout()
        bar.setSpacing(8)
        btn_cancel = QPushButton("Cancelar")
        btn_cancel.setMinimumWidth(96); btn_cancel.setMinimumHeight(32)
        btn_cancel.setStyleSheet(
            "QPushButton { background:white; color:#485A6C; border:1px solid #D4D4D4;"
            " border-radius:4px; padding:6px 18px; }"
            "QPushButton:hover { background:#F1F5F9; }"
        )
        btn_cancel.clicked.connect(self.reject)

        btn_img = QPushButton("🖼  Imagen…")
        btn_img.setMinimumWidth(120); btn_img.setMinimumHeight(32)
        btn_img.setStyleSheet(
            "QPushButton { background:white; color:#273445; border:1px solid #D4D4D4;"
            " border-radius:4px; padding:6px 14px; font-weight:600; }"
            "QPushButton:hover { background:#FEF5EB; border-color:#F37329; color:#C0621A; }"
        )
        btn_img.clicked.connect(self._emitir_imagen)

        btn_prev = QPushButton("👁  Vista previa")
        btn_prev.setMinimumWidth(130); btn_prev.setMinimumHeight(32)
        btn_prev.setStyleSheet(
            "QPushButton { background:white; color:#273445; border:1px solid #D4D4D4;"
            " border-radius:4px; padding:6px 14px; font-weight:600; }"
            "QPushButton:hover { background:#FEF5EB; border-color:#F37329; color:#C0621A; }"
        )
        btn_prev.clicked.connect(self._emitir_preview)

        btn_ok = QPushButton("Exportar PDF…")
        btn_ok.setMinimumWidth(130); btn_ok.setMinimumHeight(32)
        btn_ok.setDefault(True)
        btn_ok.setStyleSheet(
            "QPushButton { background:#F37329; color:white; border:none;"
            " border-radius:4px; padding:6px 18px; font-weight:600; }"
            "QPushButton:hover { background:#C0621A; }"
        )
        btn_ok.clicked.connect(self.accept)

        bar.addWidget(btn_cancel)
        bar.addStretch()
        bar.addWidget(btn_img)
        bar.addWidget(btn_prev)
        bar.addWidget(btn_ok)
        vl.addLayout(bar)

    def _opts(self) -> dict:
        idx = max(0, self.cmb_papel.currentIndex())
        # `periodos_por_pagina`:
        #   -1 → todo en una sola hoja (compresión elástica)
        #    0 → auto (la app calcula según papel)
        #   N>0 → N períodos por hoja
        if self.rb_pag_una.isChecked():
            periodos = -1
        elif self.rb_pag_man.isChecked():
            periodos = int(self.sp_paginas.value())
        else:
            periodos = 0
        return {
            'papel':              self.PAPEL_OPCIONES[idx][1],
            'orient':             'landscape' if self.rb_land.isChecked() else 'portrait',
            'with_cover':         False,
            'periodos_por_pagina': periodos,
            'show_fechas':        self.chk_fechas.isChecked(),
        }

    def _emitir_preview(self):
        self.preview_solicitado.emit(self._opts())

    def _emitir_imagen(self):
        self.imagen_solicitada.emit(self._opts())

    @classmethod
    def preguntar(cls, parent=None, titulo="Exportar reporte",
                    on_preview=None, on_imagen=None):
        dlg = cls(parent, titulo=titulo)
        if on_preview is not None:
            dlg.preview_solicitado.connect(on_preview)
        if on_imagen is not None:
            dlg.imagen_solicitada.connect(on_imagen)
        if dlg.exec() != QDialog.Accepted:
            return None
        return dlg._opts()


class _BgFillDelegate(QStyledItemDelegate):
    """Delegate que pinta el background del item leyendo Qt.BackgroundRole,
    aún en estilos QSS que normalmente lo ignorarían. Necesario porque la
    QSS global del proyecto (`QTableWidget::item:selected/:hover`) hace que
    Qt entre en modo QSS-rendering y descarte el BackgroundRole por
    defecto."""

    def paint(self, painter, option, index):
        # Si el item está seleccionado dejamos el comportamiento default
        # (la regla `::item:selected` de QSS pinta el resaltado).
        if option.state & QStyle.State_Selected:
            return super().paint(painter, option, index)
        bg = index.data(Qt.BackgroundRole)
        if bg is not None:
            br = bg if isinstance(bg, QBrush) else QBrush(bg)
            painter.fillRect(option.rect, br)
            # Anular la backgroundBrush del style option para que el super
            # no vuelva a pintar encima.
            opt = QStyleOptionViewItem(option)
            self.initStyleOption(opt, index)
            opt.backgroundBrush = QBrush(Qt.transparent)
            return super().paint(painter, opt, index)
        return super().paint(painter, option, index)


class ValorizadoWidget(QWidget):
    """Tabla con distribución valorizada por semanas/meses con totales,
    % por período, acumulados, %-distribución por partida y resaltado del
    período actual."""

    def __init__(self, parent: CronogramaView):
        super().__init__(parent)
        self._cv = parent
        self._period_days = 30  # default Mensual (Marco prefiere mes al abrir)
        self._build_ui()

    def _build_ui(self):
        vl = QVBoxLayout(self)
        vl.setContentsMargins(8, 8, 8, 8)
        vl.setSpacing(6)

        # (KPI strip removido por pedido de Marco — no se veía consistente.)

        # ── Toolbar ────────────────────────────────────────────────────────
        tb = QFrame()
        tb.setStyleSheet("background:transparent; border:none;")
        tb_hl = QHBoxLayout(tb)
        tb_hl.setContentsMargins(0, 0, 0, 0)
        tb_hl.setSpacing(8)

        lbl = QLabel("Período:")
        lbl.setStyleSheet(f"color:{SLATE_500}; font-size:11px;")
        tb_hl.addWidget(lbl)

        cmb = QComboBox()
        cmb.addItems(["Semanal", "Mensual"])
        cmb.setCurrentIndex(1)  # default Mensual (espejo de _period_days=30)
        cmb.currentIndexChanged.connect(self._on_periodo)
        cmb.setStyleSheet(
            "QComboBox { min-height:0; padding:2px 8px; font-size:11px;"
            f" border:1px solid {SILVER_300}; border-radius:4px; }}"
        )
        tb_hl.addWidget(cmb)

        # Toggle: mostrar fechas de inicio/fin de cada período en el header
        from PySide6.QtWidgets import QCheckBox
        self._show_fechas = True
        chk_fechas = QCheckBox("Mostrar fechas")
        chk_fechas.setChecked(self._show_fechas)
        chk_fechas.setStyleSheet(
            f"QCheckBox {{ color:{SLATE_500}; font-size:11px; padding:0 8px; }}"
        )
        chk_fechas.toggled.connect(self._on_toggle_fechas)
        tb_hl.addWidget(chk_fechas)

        tb_hl.addStretch()

        # Botones de exportación — mismo estilo que en GanttWidget
        from utils.tooltip import set_tooltip
        btn_style = (
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:3px 10px;"
            f" min-height:0; color:{SLATE_500}; }}"
            f"QPushButton:hover {{ background:{SILVER_100}; color:{SLATE_700}; }}"
        )
        btn_pdf = QPushButton("📄 PDF")
        btn_pdf.setCursor(Qt.PointingHandCursor)
        btn_pdf.setStyleSheet(btn_style)
        btn_pdf.clicked.connect(self._exportar_pdf)
        set_tooltip(btn_pdf,
                     "Exportar el cronograma valorizado a PDF / imagen "
                     "(con vista previa y selección de papel)")
        tb_hl.addWidget(btn_pdf)

        btn_xls = QPushButton("📊 Excel")
        btn_xls.setCursor(Qt.PointingHandCursor)
        btn_xls.setStyleSheet(btn_style)
        btn_xls.clicked.connect(self._exportar_excel)
        set_tooltip(btn_xls, "Exportar a hoja de cálculo (.xlsx)")
        tb_hl.addWidget(btn_xls)

        btn_ods = QPushButton("📑 ODS")
        btn_ods.setCursor(Qt.PointingHandCursor)
        btn_ods.setStyleSheet(btn_style)
        btn_ods.clicked.connect(self._exportar_ods)
        set_tooltip(btn_ods,
                     "Exportar a OpenDocument Spreadsheet (.ods) — "
                     "requiere LibreOffice instalado")
        tb_hl.addWidget(btn_ods)

        vl.addWidget(tb)

        # ── Splitter horizontal con dos paneles ───────────────────────────
        # Panel izquierdo: cols frozen (Ítem · Descripción · Parcial · % Total)
        # Panel derecho: períodos + Total (scroll horizontal independiente)
        self.split = QSplitter(Qt.Horizontal)
        self.split.setHandleWidth(6)
        self.split.setStyleSheet(
            "QSplitter::handle { background:#CBD5E1; }"
            "QSplitter::handle:horizontal { width:6px; margin:0; }"
            "QSplitter::handle:hover { background:#F37329; }"
        )
        self.split.setChildrenCollapsible(False)

        def _make_panel(is_left):
            """Crea (frame, tabla_principal, tabla_footer) — patrón espejo
            con scrolls sincronizables verticalmente y header consistente."""
            fr = QFrame()
            v = QVBoxLayout(fr)
            v.setContentsMargins(0, 0, 0, 0)
            v.setSpacing(0)
            tbl = QTableWidget(0, 0)
            tbl.verticalHeader().setVisible(False)
            tbl.setEditTriggers(QAbstractItemView.NoEditTriggers)
            tbl.setSelectionBehavior(QAbstractItemView.SelectRows)
            # Mostramos gridlines suaves para que las filas se distingan
            # incluso sin zebra striping muy contrastado.
            tbl.setShowGrid(True)
            tbl.setAlternatingRowColors(False)
            tbl.setStyleSheet(f"""
                QTableWidget {{ border:1px solid {SILVER_300};
                                border-bottom:none; font-size:11px;
                                gridline-color:#E0E5EC; background:white; }}
                QTableWidget::item {{ padding:3px 8px; }}
                QHeaderView::section {{
                    background:{SLATE_500}; color:white; font-size:10px;
                    font-weight:700; padding:4px 6px; border:none;
                }}
            """)
            # En el panel izquierdo escondemos la scrollbar horizontal — su
            # ancho lo controla el splitter; igualmente la vertical va al
            # panel derecho para evitar duplicarla.
            if is_left:
                tbl.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                tbl.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            v.addWidget(tbl, stretch=1)

            ftr = QTableWidget(0, 0)
            ftr.verticalHeader().setVisible(False)
            ftr.horizontalHeader().setVisible(False)
            ftr.setEditTriggers(QAbstractItemView.NoEditTriggers)
            ftr.setSelectionMode(QAbstractItemView.NoSelection)
            ftr.setFocusPolicy(Qt.NoFocus)
            ftr.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            ftr.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            # Cuadrícula sutil en el footer también — ayuda a leer columna a
            # columna en TOTAL/% PERÍODO/ACUMULADO/% ACUMULADO.
            ftr.setShowGrid(True)
            ftr.setStyleSheet(f"""
                QTableWidget {{ border:1px solid {SILVER_300};
                                border-top:2px solid {SLATE_500};
                                font-size:11px; background:white;
                                gridline-color:#C5D1E0; }}
                QTableWidget::item {{ padding:4px 8px; }}
            """)
            v.addWidget(ftr)
            return fr, tbl, ftr

        left_fr, self.tbl_l, self.tbl_lf = _make_panel(is_left=True)
        right_fr, self.tbl_r, self.tbl_rf = _make_panel(is_left=False)

        # Cabecera de 2 niveles: período (Mes/Sem + fecha) abarcando sus 2
        # sub-columnas (Metrado | Valorización) — estilo Delphin Express.
        self._period_header = _PeriodGroupHeader(('Metrado', 'Valoriz.'))
        self.tbl_r.setHorizontalHeader(self._period_header)

        # Aplicar delegate de bg (cruza el bypass de QSS para zebra y footers)
        self._bg_delegate = _BgFillDelegate(self)
        for tbl_ in (self.tbl_l, self.tbl_r, self.tbl_lf, self.tbl_rf):
            tbl_.setItemDelegate(self._bg_delegate)

        # Padding mínimo (sin exceso) — el ancho dinámico de columna ya da
        # respiración visual cuando hay pocos períodos. Con muchas semanas,
        # cada columna se queda al MIN y el padding extra cortaría texto.
        right_pad_qss = (
            "QTableWidget::item { padding:3px 8px; }"
        )
        self.tbl_r.setStyleSheet(self.tbl_r.styleSheet() + right_pad_qss)
        self.tbl_rf.setStyleSheet(self.tbl_rf.styleSheet() + right_pad_qss)

        # Click en el header de una columna del panel derecho → selecciona
        # toda la columna (en main + footer) para inspeccionar un período.
        self.tbl_r.horizontalHeader().setSectionsClickable(True)
        self.tbl_r.horizontalHeader().sectionClicked.connect(
            self._on_period_header_clicked)
        # Al clickar cualquier celda (no-header) se debe deselectar la
        # columna del footer, que mi selectionModel().select dejaba pegada
        # incluso al cambiar de fila.
        self.tbl_r.cellClicked.connect(self._on_cell_clicked_clear_col)
        self.tbl_l.cellClicked.connect(self._on_cell_clicked_clear_col)
        # ESC deselecciona — instalamos eventFilter a nivel de QApplication
        # para capturar la tecla con seguridad. El QShortcut y eventFilter
        # locales no eran suficientes (QAbstractItemView y/o QSplitter
        # consumían el evento antes).
        from PySide6.QtWidgets import QApplication
        QApplication.instance().installEventFilter(self)
        self._col_selected = None

        # Alias por compatibilidad (en caso de referencias externas)
        self.tbl = self.tbl_r
        self.tbl_footer = self.tbl_rf

        self.split.addWidget(left_fr)
        self.split.addWidget(right_fr)
        self.split.setStretchFactor(0, 0)
        self.split.setStretchFactor(1, 1)
        # Anchura inicial del panel izquierdo: lo justo para los 4 cols
        # Ítem(60) + Descripción(280) + Und(50) + Cantidad(80) + Precio(95) + %Total(65)
        self.split.setSizes([60 + 280 + 50 + 80 + 95 + 65, 800])
        vl.addWidget(self.split, stretch=1)

        # ── Sincronizar scrolls ───────────────────────────────────────────
        # Vertical: panel izquierdo ↔ panel derecho (selección y scroll
        # del usuario en cualquiera de los dos se replica en el otro).
        self.tbl_r.verticalScrollBar().valueChanged.connect(
            self.tbl_l.verticalScrollBar().setValue)
        self.tbl_l.verticalScrollBar().valueChanged.connect(
            self.tbl_r.verticalScrollBar().setValue)
        # Horizontal: tabla derecha ↔ su footer
        self.tbl_r.horizontalScrollBar().valueChanged.connect(
            self.tbl_rf.horizontalScrollBar().setValue)
        # Selección de fila sincronizada
        self.tbl_l.itemSelectionChanged.connect(
            lambda: self._sync_selection(self.tbl_l, self.tbl_r))
        self.tbl_r.itemSelectionChanged.connect(
            lambda: self._sync_selection(self.tbl_r, self.tbl_l))

    def _sync_selection(self, src, dst):
        if dst.signalsBlocked():
            return
        dst.blockSignals(True)
        try:
            rows = {ix.row() for ix in src.selectedIndexes()}
            dst.clearSelection()
            for r in rows:
                if 0 <= r < dst.rowCount():
                    dst.selectRow(r)
        finally:
            dst.blockSignals(False)

    def _on_periodo(self, idx: int):
        self._period_days = 7 if idx == 0 else 30
        self.cargar()

    def _on_toggle_fechas(self, checked: bool):
        self._show_fechas = bool(checked)
        self.cargar()

    def _on_period_header_clicked(self, col: int):
        """Selecciona toda la columna en el panel derecho cuando el usuario
        hace clic en el header de una semana/mes. Sincroniza también con el
        footer para que las filas TOTAL/%/ACUMULADO/% se resalten."""
        if col < 0 or col >= self.tbl_r.columnCount():
            return
        from PySide6.QtCore import QItemSelection, QItemSelectionModel
        # Guardar scroll vertical para restaurarlo: la selección de un rango
        # grande hace que Qt auto-scrolle al último item (bottom_right).
        v_r = self.tbl_r.verticalScrollBar().value()
        v_l = self.tbl_l.verticalScrollBar().value()

        def _sel_col(tbl, c):
            n_rows = tbl.rowCount()
            if n_rows <= 0:
                return
            mdl = tbl.model()
            tl = mdl.index(0, c)
            br = mdl.index(n_rows - 1, c)
            tbl.selectionModel().select(
                QItemSelection(tl, br),
                QItemSelectionModel.ClearAndSelect,
            )
        _sel_col(self.tbl_r, col)
        _sel_col(self.tbl_rf, col)
        self._col_selected = col
        # Restaurar scroll en ambos paneles (sincronizados)
        self.tbl_r.verticalScrollBar().setValue(v_r)
        self.tbl_l.verticalScrollBar().setValue(v_l)

    def _clear_column_selection(self):
        """Limpia la selección que dejó la pulsación de un header de columna."""
        self.tbl_l.clearSelection()
        self.tbl_r.clearSelection()
        self.tbl_rf.clearSelection()
        self._col_selected = None

    def _on_cell_clicked_clear_col(self, *_args):
        """Al hacer clic en una celda cualquiera (no en el header), limpia
        la columna previamente seleccionada (footer queda colgado si no)."""
        if getattr(self, '_col_selected', None) is None:
            return
        self.tbl_rf.clearSelection()
        self._col_selected = None

    def eventFilter(self, obj, event):
        from PySide6.QtCore import QEvent
        from PySide6.QtWidgets import QApplication
        # ESC en cualquier sub-widget de Valorizado → deseleccionar todo
        if event.type() == QEvent.KeyPress and event.key() == Qt.Key_Escape:
            fw = QApplication.focusWidget()
            # Solo actuar si el foco está en este widget o uno de sus hijos
            if fw is not None and (fw is self or self.isAncestorOf(fw)):
                self._clear_column_selection()
                return True
        return super().eventFilter(obj, event)

    def _apply_period_widths(self):
        """Ajusta el ancho de las columnas de período según el viewport:
        - Si caben holgadamente → cada una toma viewport_w / n (sin scroll).
        - Si no caben → cada una toma el MIN (80px) y aparece scroll."""
        n = getattr(self, '_last_n_subcols', 0)
        total_w = getattr(self, '_last_total_w', 110)
        if n <= 0:
            return
        viewport_w = max(200, self.tbl_r.viewport().width())
        avail = viewport_w - total_w - 4  # -4 para borde
        # Mínimo por sub-columna: 78px por defecto, pero ampliado dinámicamente
        # si las filas TOTAL/ACUMULADO traen montos grandes (calculado en
        # `cargar()` con QFontMetrics) para que no se recorten.
        MIN_PER = getattr(self, '_min_per_w', 78)   # cada sub-columna (Met./Valor.)
        per_w = max(MIN_PER, avail // n)
        for c in range(n):
            self.tbl_r.setColumnWidth(c, per_w)
            self.tbl_rf.setColumnWidth(c, per_w)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        # Reajustar anchos de períodos al redimensionar la ventana / splitter
        if hasattr(self, '_last_n_periods'):
            QTimer.singleShot(0, self._apply_period_widths)

    def _project_start(self):
        """Devuelve datetime para el día 1 del proyecto (usa fecha_inicio,
        cae a costo_al, finalmente a None)."""
        for key in ('fecha_inicio', 'costo_al'):
            fi = (self._cv._proy.get(key) or '').strip()
            if fi:
                try:
                    return datetime.strptime(fi, '%Y-%m-%d')
                except Exception:
                    pass
        return None

    def _periodo_label(self, unidad: str, i: int) -> str:
        """Construye el label de un período. Si _show_fechas está activo y
        hay fecha de inicio configurada, añade el rango en una 2ª línea
        (mismo mes: '01-07 May'; cruza mes: '30 Abr - 06 May'; cruza año:
        '30 Dic 25 - 05 Ene 26')."""
        base = f"{unidad} {i + 1}"
        if not self._show_fechas:
            return base
        f_ini = self._project_start()
        if f_ini is None:
            return base
        MESES = ["Ene", "Feb", "Mar", "Abr", "May", "Jun",
                  "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
        start_day = i * self._period_days + 1
        end_day   = (i + 1) * self._period_days
        d_ini = f_ini + timedelta(days=start_day - 1)
        d_fin = f_ini + timedelta(days=end_day - 1)
        if d_ini.month == d_fin.month and d_ini.year == d_fin.year:
            rng = f"{d_ini.day:02d}–{d_fin.day:02d} {MESES[d_ini.month - 1]}"
        elif d_ini.year == d_fin.year:
            rng = (f"{d_ini.day:02d} {MESES[d_ini.month - 1]} – "
                    f"{d_fin.day:02d} {MESES[d_fin.month - 1]}")
        else:
            rng = (f"{d_ini.day:02d} {MESES[d_ini.month - 1]} {str(d_ini.year)[-2:]} – "
                    f"{d_fin.day:02d} {MESES[d_fin.month - 1]} {str(d_fin.year)[-2:]}")
        return f"{base}\n{rng}"

    def _periodo_actual_idx(self, n_periods: int) -> int:
        """Devuelve el índice (0-based) del período en curso usando hoy vs
        f_ini del proyecto. -1 si está fuera de rango."""
        try:
            f_ini = self._cv.gantt._project_start() if hasattr(self._cv, 'gantt') \
                    else None
        except Exception:
            f_ini = None
        # Fallback: buscar GanttWidget vía atributo
        if f_ini is None:
            for attr in ('_gantt_w',):
                w = getattr(self._cv, attr, None)
                if w is not None and hasattr(w, '_project_start'):
                    f_ini = w._project_start()
                    break
        if f_ini is None:
            return -1
        try:
            delta_dias = (datetime.now() - f_ini).days + 1
        except Exception:
            return -1
        if delta_dias < 1:
            return -1
        idx = (delta_dias - 1) // self._period_days
        return idx if 0 <= idx < n_periods else -1

    def cargar(self):
        if not hasattr(self._cv, '_partidas'):
            return
        partidas = self._cv._partidas
        tasks = self._cv._tasks
        cmap = self._cv._cron_map
        plazo = self._cv._proy.get('plazo') or 0
        max_ef = max((t['EF'] for t in tasks.values() if t['EF'] > 0),
                      default=plazo)
        n_dias = max(max_ef, plazo, self._period_days)
        n_periods = (n_dias + self._period_days - 1) // self._period_days

        unidad = "Sem" if self._period_days == 7 else "Mes"
        moneda = self._cv._moneda
        per_actual = self._periodo_actual_idx(n_periods)

        # ── Estructura de columnas ──────────────────────────────────────
        # Panel izquierdo (frozen): Ítem · Descripción · Und · Cantidad · Precio · % Total
        # Panel derecho (scrollable): [períodos con fechas opcionales] · Total
        left_headers = ["Ítem", "Descripción", "Und", "Cantidad", "Precio", "Total"]
        # Texto de header del MODELO (lo lee el Excel); en pantalla el header
        # personalizado de 2 niveles dibuja el período combinado.
        period_labels = [self._periodo_label(unidad, i) for i in range(n_periods)]
        right_headers = []
        for plab in period_labels:
            right_headers.append(f"{plab}\nMetrado")
            right_headers.append(f"{plab}\nValoriz.")
        right_headers.append("Total")
        N_RIGHT = len(right_headers)
        COL_TOTAL_R = N_RIGHT - 1  # índice de "Total" en panel derecho
        n_sub = 2 * n_periods      # sub-columnas de período

        # Header de 2 niveles: período (1-2 líneas con fecha) + sub-columna.
        hdr_h = 50 if self._show_fechas else 34
        self.tbl_r.horizontalHeader().setFixedHeight(hdr_h)
        self.tbl_l.horizontalHeader().setFixedHeight(hdr_h)
        if hasattr(self, '_period_header'):
            self._period_header.set_periodos(period_labels, hdr_h)
        # Word-wrap habilitado en los headers del panel derecho
        self.tbl_r.horizontalHeader().setDefaultAlignment(Qt.AlignCenter)
        self.tbl_l.horizontalHeader().setDefaultAlignment(Qt.AlignCenter)

        for tbl_ in (self.tbl_l, self.tbl_r, self.tbl_lf, self.tbl_rf):
            tbl_.setUpdatesEnabled(False)
            tbl_.setRowCount(0)
            tbl_.clearSpans()
        self.tbl_l.setColumnCount(len(left_headers))
        self.tbl_l.setHorizontalHeaderLabels(left_headers)
        self.tbl_r.setColumnCount(N_RIGHT)
        self.tbl_r.setHorizontalHeaderLabels(right_headers)
        self.tbl_lf.setColumnCount(len(left_headers))
        self.tbl_rf.setColumnCount(N_RIGHT)

        # Anchos
        th_l = self.tbl_l.horizontalHeader()
        th_r = self.tbl_r.horizontalHeader()
        th_lf = self.tbl_lf.horizontalHeader()
        th_rf = self.tbl_rf.horizontalHeader()
        th_lf.setVisible(False); th_rf.setVisible(False)

        # Izquierdo — Ítem · Descripción · Und · Cantidad · Precio · Precio Total
        for c, w in enumerate([60, None, 50, 80, 95, 100]):
            if w is None:
                th_l.setSectionResizeMode(c, QHeaderView.Stretch)
                th_lf.setSectionResizeMode(c, QHeaderView.Stretch)
            else:
                th_l.setSectionResizeMode(c, QHeaderView.Interactive)
                self.tbl_l.setColumnWidth(c, w)
                th_lf.setSectionResizeMode(c, QHeaderView.Interactive)
                self.tbl_lf.setColumnWidth(c, w)
        # Cuando el usuario redimensiona en uno, replicar en el otro
        th_l.sectionResized.connect(
            lambda i, _o, n: self.tbl_lf.setColumnWidth(i, n))

        # Derecho — Períodos en modo Interactive con ancho calculado
        # dinámicamente: si caben todos en el viewport, se distribuyen
        # uniformemente; si no, cada uno toma el MIN y aparece scroll
        # horizontal. Total con ancho fijo 110px.
        TOTAL_W = 110
        for c in range(N_RIGHT - 1):
            th_r.setSectionResizeMode(c, QHeaderView.Interactive)
            th_rf.setSectionResizeMode(c, QHeaderView.Interactive)
        th_r.setSectionResizeMode(COL_TOTAL_R, QHeaderView.Fixed)
        self.tbl_r.setColumnWidth(COL_TOTAL_R, TOTAL_W)
        th_rf.setSectionResizeMode(COL_TOTAL_R, QHeaderView.Fixed)
        self.tbl_rf.setColumnWidth(COL_TOTAL_R, TOTAL_W)
        th_r.setStretchLastSection(False)
        th_rf.setStretchLastSection(False)
        # Mínimo por sección — si N períodos × este min > viewport, scroll.
        # 100px asegura que "S/ 1,234.56" (típico) entre sin elide.
        th_r.setMinimumSectionSize(78)
        th_rf.setMinimumSectionSize(78)
        self._last_n_periods = n_periods
        self._last_n_subcols = n_sub
        self._last_total_w = TOTAL_W
        # Mirror al footer cuando el usuario redimensiona manualmente
        try:
            th_r.sectionResized.disconnect()
        except Exception:
            pass
        th_r.sectionResized.connect(
            lambda i, _o, n: self.tbl_rf.setColumnWidth(i, n))
        # Aplicar anchos dinámicos (deferred — viewport puede no estar listo)
        QTimer.singleShot(0, self._apply_period_widths)

        totales_per = [0.0] * n_periods
        total_general = 0.0

        # Estilo consistente con el árbol de Presupuesto / tabla del Gantt:
        #  · Títulos (nivel 1) → banda más oscura + texto rojo.
        #  · Subtítulos (nivel ≥2) → zebra + texto arándano.
        #  · Partidas → zebra. Zebra sutil, solo partidas+subtítulos.
        ZEBRA_BG = QColor("#F6F8FB")
        TITLE_BG = QColor("#E2E8F0")
        TITLE_FG = QColor("#B71C1C")   # rojo títulos (igual que Presupuesto)
        SUBT_FG  = QColor("#0D52BF")   # arándano (Blueberry 700)
        zebra_i  = 0
        # Word-wrap en descripción para no perder texto largo
        self.tbl_l.setWordWrap(True)

        for p in partidas:
            r = self.tbl_l.rowCount()
            # Insertar fila en AMBAS tablas para mantener alineación
            self.tbl_l.insertRow(r)
            self.tbl_r.insertRow(r)
            es_titulo = p['es_titulo']
            nivel_p   = p.get('nivel') or 1
            es_header = es_titulo and nivel_p <= 1     # título de capítulo
            es_subt   = es_titulo and nivel_p >= 2     # subtítulo
            # Zebra para subtítulos + partidas (las cabeceras llevan su banda).
            row_bg = None
            if not es_header:
                zebra_i += 1
                if zebra_i % 2 == 0:
                    row_bg = ZEBRA_BG

            if es_titulo:
                # Cabecera (nivel 1): banda oscura + rojo. Subtítulo: zebra +
                # arándano. Solo se ve Ítem+Descripción en el panel izquierdo.
                bg_t = TITLE_BG if es_header else row_bg
                fg_t = TITLE_FG if es_header else SUBT_FG
                for c in range(len(left_headers)):
                    txt = p['item'] if c == 0 else (p['descripcion'] if c == 1 else '')
                    it = QTableWidgetItem(txt)
                    if bg_t is not None:
                        it.setBackground(QBrush(bg_t))
                    f = QFont(); f.setBold(True)
                    if es_header and nivel_p == 1 and c in (0, 1):
                        f.setUnderline(True)
                    it.setFont(f)
                    it.setForeground(QBrush(fg_t))
                    it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                    self.tbl_l.setItem(r, c, it)
                # Mismas filas del derecho con el mismo fondo
                for c in range(N_RIGHT):
                    it = QTableWidgetItem('')
                    if bg_t is not None:
                        it.setBackground(QBrush(bg_t))
                    it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                    self.tbl_r.setItem(r, c, it)
                continue

            cd = cmap.get(p['id'], {})
            ini = cd.get('inicio_dia', 1) or 1
            dur = cd.get('duracion', 0) or 0
            t = tasks.get(p['id'], {})
            ini = t.get('ES', ini)

            parcial = parcial_wysiwyg(p['metrado'], p['precio_unitario'])
            segs_ = cd.get('segmentos', '') or ''
            weekly = distribuir_periodos(
                segs_, ini, dur, parcial, n_periods, self._period_days
            )
            metrado_tot = p['metrado'] or 0
            weekly_met = (distribuir_periodos(
                segs_, ini, dur, metrado_tot, n_periods, self._period_days)
                if metrado_tot else [0.0] * n_periods)

            # Panel izquierdo: Ítem · Descripción · Und · Cantidad · Precio · Precio Total
            metrado_v = p['metrado'] or 0
            pu_v       = p['precio_unitario'] or 0
            und_v      = p.get('unidad') or p.get('und') or ''
            # Cantidad (= metrado): número sin símbolo de moneda, con los
            # decimales de metrado configurados y coma como separador de miles
            _dm = get_decimales_metrado()
            cant_txt = (f"{metrado_v:,.{_dm}f}") if metrado_v else ''
            cells_left = [
                (p['item'] or '',           Qt.AlignLeft  | Qt.AlignVCenter),
                (p['descripcion'] or '',     Qt.AlignLeft  | Qt.AlignVCenter),
                (str(und_v),                 Qt.AlignCenter),
                (cant_txt,                   Qt.AlignRight | Qt.AlignVCenter),
                (fmt(pu_v, moneda) if pu_v else '',
                                              Qt.AlignRight | Qt.AlignVCenter),
                (fmt(parcial, moneda) if parcial else '',
                                              Qt.AlignRight | Qt.AlignVCenter),
            ]
            for c, (val, align) in enumerate(cells_left):
                it = QTableWidgetItem(val)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                it.setTextAlignment(align)
                # Tooltip con texto completo en descripción (por si el texto
                # excede el ancho de columna y queda elided).
                if c == 1 and (p['descripcion'] or ''):
                    it.setToolTip(p['descripcion'])
                if row_bg is not None:
                    it.setBackground(QBrush(row_bg))
                self.tbl_l.setItem(r, c, it)

            # Panel derecho: por período → 2 sub-columnas (Metrado | Valoriz.) + Total
            row_total = 0.0
            for i in range(n_periods):
                mval = weekly_met[i] if i < len(weekly_met) else 0.0
                vval = weekly[i] if i < len(weekly) else 0.0
                # Metrado (sub-col 2i)
                itm = QTableWidgetItem(f"{mval:,.2f}" if mval > 0 else '')
                itm.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                itm.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                if mval > 0:
                    itm.setForeground(QBrush(QColor(SLATE_500)))
                if row_bg is not None:
                    itm.setBackground(QBrush(row_bg))
                self.tbl_r.setItem(r, 2 * i, itm)
                # Valorización (sub-col 2i+1)
                itv = QTableWidgetItem(fmt(vval, moneda) if vval > 0 else '')
                itv.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                itv.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                if vval > 0:
                    itv.setForeground(QBrush(QColor("#1F2A38")))
                if row_bg is not None:
                    itv.setBackground(QBrush(row_bg))
                self.tbl_r.setItem(r, 2 * i + 1, itv)
                row_total += vval
                totales_per[i] += vval

            it = QTableWidgetItem(fmt(row_total, moneda))
            it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
            it.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            f = QFont(); f.setBold(True); it.setFont(f)
            if row_bg is not None:
                it.setBackground(QBrush(row_bg))
            self.tbl_r.setItem(r, COL_TOTAL_R, it)
            total_general += row_total

        # ── Filas de resumen en el FOOTER (sticky bottom) ─────────────────
        BG_MONTO = "#E2E8F0"   # slate-100
        BG_PCT   = "#F1F5F9"   # silver más claro
        FG_TEXT  = "#1F2A38"   # slate-900
        FG_LBL   = "#273445"   # slate-700

        def _add_summary_row(label, vals_per, val_total, *, bg, pct=False):
            bg_brush = QBrush(QColor(bg))
            fg_brush = QBrush(QColor(FG_TEXT))
            lbl_brush = QBrush(QColor(FG_LBL))

            # Inserta fila en AMBOS footers
            r = self.tbl_lf.rowCount()
            self.tbl_lf.insertRow(r); self.tbl_rf.insertRow(r)
            self.tbl_lf.setRowHeight(r, 28); self.tbl_rf.setRowHeight(r, 28)

            # Footer izquierdo: label que abarca las 4 columnas
            it_lbl = QTableWidgetItem(label)
            it_lbl.setBackground(bg_brush)
            it_lbl.setForeground(lbl_brush)
            f = QFont(); f.setBold(True); f.setPointSize(10)
            it_lbl.setFont(f)
            it_lbl.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            it_lbl.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
            self.tbl_lf.setItem(r, 0, it_lbl)
            for c in range(1, len(left_headers)):
                it = QTableWidgetItem('')
                it.setBackground(bg_brush)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_lf.setItem(r, c, it)
            self.tbl_lf.setSpan(r, 0, 1, len(left_headers))

            # Footer derecho: por período el valor va en la sub-columna de
            # Valorización (2i+1); la de Metrado (2i) queda vacía (sumar
            # metrados de distintas unidades no tiene sentido).
            for i, v in enumerate(vals_per):
                # Sub-col Metrado vacía (solo fondo)
                itm = QTableWidgetItem('')
                itm.setBackground(bg_brush)
                itm.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_rf.setItem(r, 2 * i, itm)
                # Sub-col Valorización
                if pct:
                    txt = f"{v:.1f}%" if v > 0 else ''
                else:
                    txt = fmt(v, moneda) if v > 0 else ''
                it = QTableWidgetItem(txt)
                it.setBackground(bg_brush)
                it.setForeground(fg_brush)
                it.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                f = QFont(); f.setBold(True); it.setFont(f)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_rf.setItem(r, 2 * i + 1, it)

            if val_total is not None:
                txt = f"{val_total:.1f}%" if pct else fmt(val_total, moneda)
                it = QTableWidgetItem(txt)
                it.setBackground(bg_brush)
                it.setForeground(fg_brush)
                it.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                f = QFont(); f.setBold(True); it.setFont(f)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_rf.setItem(r, COL_TOTAL_R, it)

        _add_summary_row("TOTAL", totales_per, total_general, bg=BG_MONTO)
        pct_per = [(v / total_general * 100.0 if total_general else 0.0)
                    for v in totales_per]
        _add_summary_row("% PERÍODO", pct_per, 100.0 if total_general else 0.0,
                          bg=BG_PCT, pct=True)
        acumulado = []
        s = 0.0
        for v in totales_per:
            s += v
            acumulado.append(s)
        _add_summary_row("ACUMULADO", acumulado, total_general, bg=BG_MONTO)
        pct_acum = [(v / total_general * 100.0 if total_general else 0.0)
                     for v in acumulado]
        _add_summary_row("% ACUMULADO", pct_acum,
                          100.0 if total_general else 0.0,
                          bg=BG_PCT, pct=True)

        # Anchos dinámicos para que TOTAL/ACUMULADO (montos grandes) no se
        # corten: medir el valor más ancho que aparecerá en una sub-columna de
        # Valorización (= acumulado final = total del proyecto) y en la columna
        # Total, y ampliar el mínimo de columna en consecuencia (con tope para
        # no forzar scroll horizontal excesivo).
        from PySide6.QtGui import QFontMetrics
        _fb = QFont(); _fb.setBold(True)
        _fm = QFontMetrics(_fb)
        _widest = fmt(total_general, moneda) if total_general else ''
        _need = (_fm.horizontalAdvance(_widest) + 22) if _widest else 78
        self._min_per_w  = max(78, min(_need, 170))
        self._last_total_w = max(110, min(_need + 6, 185))
        self.tbl_r.setColumnWidth(COL_TOTAL_R, self._last_total_w)
        self.tbl_rf.setColumnWidth(COL_TOTAL_R, self._last_total_w)

        for tbl_ in (self.tbl_l, self.tbl_r, self.tbl_lf, self.tbl_rf):
            tbl_.setUpdatesEnabled(True)
        QTimer.singleShot(0, self._apply_period_widths)
        # Auto-fit alturas del panel izquierdo (descripción puede wrap a 2-3
        # líneas) y replicar al panel derecho para mantener alineación de filas.
        self.tbl_l.resizeRowsToContents()
        for r in range(self.tbl_l.rowCount()):
            self.tbl_r.setRowHeight(r, self.tbl_l.rowHeight(r))
        # Footers con altura fija al contenido (4 filas)
        h_total = sum(self.tbl_lf.rowHeight(i) for i in range(self.tbl_lf.rowCount())) + 4
        self.tbl_lf.setFixedHeight(h_total)
        self.tbl_rf.setFixedHeight(h_total)

    # ──────────────────────────────────────────────────────────────────────
    # Exportación
    # ──────────────────────────────────────────────────────────────────────

    def _exportar_pdf(self):
        """Abre un diálogo de opciones (papel, orientación, cover, vista
        previa, imagen) y exporta a PDF usando `core.pdf_reports`."""
        opts = _DialogExportarReportePdf.preguntar(
            self, "Exportar cronograma valorizado",
            on_preview=self._vista_previa_pdf,
            on_imagen=self._exportar_imagen,
        )
        if opts is None:
            return
        self._render_pdf_a_archivo_con_dialog(opts)

    def _render_pdf_a_archivo_con_dialog(self, opts):
        import os
        try:
            from core.pdf_reports import generar_pdf_archivo
        except Exception as e:
            QMessageBox.warning(self, "Exportar PDF",
                                  f"No se pudo cargar el generador de PDF:\n{e}")
            return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"valorizado_{self._cv.pid}.pdf")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar cronograma valorizado a PDF",
            sugerido, "PDF (*.pdf)"
        )
        if not path:
            return
        try:
            generar_pdf_archivo(
                'cronograma_valorizado', self._cv.pid, path,
                with_cover=opts.get('with_cover', True),
                paper=opts.get('papel', 'A4'),
                orient=opts.get('orient', 'portrait'),
            )
            self._save_dir(path)
            QMessageBox.information(self, "Exportar PDF",
                                       f"Cronograma valorizado exportado a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")

    def _vista_previa_pdf(self, opts):
        """Genera un PDF temporal y lo muestra en un diálogo con QPdfView
        (visor nativo de PDF) — más confiable que el QPrintPreviewDialog
        cuando se trata de mostrar el contenido directamente."""
        import os, tempfile
        try:
            from core.pdf_reports import generar_pdf_archivo
            from PySide6.QtPdf import QPdfDocument
            from PySide6.QtPdfWidgets import QPdfView
        except Exception as e:
            QMessageBox.warning(self, "Vista previa",
                                  f"No se pudo abrir vista previa:\n{e}")
            return
        tmp_fd, tmp_pdf = tempfile.mkstemp(suffix='.pdf', prefix='valz_prev_')
        os.close(tmp_fd)
        try:
            # Asegurar que el reporte respete escala (semana/mes) y páginas
            # horizontales elegidas en el diálogo
            from core import pdf_reports as _pr
            _pr._BUILD_HTML_PAPER['escala'] = (
                'semana' if self._period_days == 7 else 'mes')
            _pr._BUILD_HTML_PAPER['periodos_por_pagina'] = int(
                opts.get('periodos_por_pagina', 0) or 0)
            _pr._BUILD_HTML_PAPER['show_fechas'] = bool(opts.get('show_fechas', False))
            generar_pdf_archivo(
                'cronograma_valorizado', self._cv.pid, tmp_pdf,
                with_cover=opts.get('with_cover', False),
                paper=opts.get('papel', 'A4'),
                orient=opts.get('orient', 'portrait'),
            )
            if not os.path.exists(tmp_pdf) or os.path.getsize(tmp_pdf) < 500:
                QMessageBox.warning(self, "Vista previa", "El PDF generado está vacío.")
                return

            dlg = QDialog(self)
            dlg.setWindowTitle("Vista previa — Cronograma Valorizado")
            dlg.setWindowModality(Qt.ApplicationModal)
            dlg.resize(1100, 800)

            v = QVBoxLayout(dlg)
            v.setContentsMargins(8, 8, 8, 8)
            v.setSpacing(6)

            # Toolbar superior
            tb = QHBoxLayout()
            btn_img = QPushButton("🖼  Exportar imagen…")
            btn_img.setStyleSheet(
                f"QPushButton {{ background:white; color:{SLATE_700};"
                f" border:1px solid {SILVER_300}; border-radius:4px;"
                f" padding:6px 14px; }}"
                "QPushButton:hover { background:#FEF5EB; border-color:#F37329; color:#C0621A; }"
            )
            btn_img.clicked.connect(lambda: self._exportar_imagen(opts))
            btn_pdf = QPushButton("📄  Guardar PDF…")
            btn_pdf.setStyleSheet(btn_img.styleSheet())
            btn_pdf.clicked.connect(lambda: self._render_pdf_a_archivo_con_dialog(opts))
            btn_close = QPushButton("Cerrar")
            btn_close.setStyleSheet(
                f"QPushButton {{ background:{SLATE_500}; color:white; border:none;"
                f" border-radius:4px; padding:6px 18px; }}"
                f"QPushButton:hover {{ background:{SLATE_700}; }}"
            )
            btn_close.clicked.connect(dlg.accept)

            tb.addWidget(btn_img)
            tb.addWidget(btn_pdf)
            tb.addStretch()
            tb.addWidget(btn_close)
            v.addLayout(tb)

            # Visor de PDF nativo
            doc = QPdfDocument(dlg)
            doc.load(tmp_pdf)
            view = QPdfView(dlg)
            view.setDocument(doc)
            view.setPageMode(QPdfView.PageMode.MultiPage)
            view.setZoomMode(QPdfView.ZoomMode.FitToWidth)
            v.addWidget(view, stretch=1)

            dlg.exec()
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Vista previa",
                                  f"No se pudo abrir vista previa: {e}")
        finally:
            try: os.unlink(tmp_pdf)
            except Exception: pass

    def _exportar_imagen(self, opts):
        """Genera un PDF temporal y lo rasteriza a PNG/JPG (una imagen por
        página si hay varias). Compuesto sobre fondo blanco."""
        import os, tempfile
        try:
            from core.pdf_reports import generar_pdf_archivo
            from PySide6.QtPdf import QPdfDocument
            from PySide6.QtGui import QImage
        except Exception as e:
            QMessageBox.warning(self, "Exportar imagen",
                                  f"Módulos necesarios no disponibles:\n{e}")
            return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"valorizado_{self._cv.pid}.png")
        path, sel = QFileDialog.getSaveFileName(
            self, "Exportar como imagen", sugerido,
            "PNG (*.png);;JPEG (*.jpg *.jpeg)"
        )
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        if 'jpg' in (sel or '').lower() or 'jpeg' in (sel or '').lower():
            if ext not in ('.jpg', '.jpeg'):
                path += '.jpg'
            fmt = 'JPG'
        else:
            if ext != '.png':
                path += '.png'
            fmt = 'PNG'

        tmp_fd, tmp_pdf = tempfile.mkstemp(suffix='.pdf', prefix='valz_img_')
        os.close(tmp_fd)
        try:
            generar_pdf_archivo(
                'cronograma_valorizado', self._cv.pid, tmp_pdf,
                with_cover=opts.get('with_cover', True),
                paper=opts.get('papel', 'A4'),
                orient=opts.get('orient', 'portrait'),
            )
            doc = QPdfDocument()
            doc.load(tmp_pdf)
            n = doc.pageCount()
            if n <= 0:
                raise RuntimeError("PDF vacío")
            dpi = 200
            base, real_ext = os.path.splitext(path)
            outs = []
            for i in range(n):
                page_pts = doc.pagePointSize(i)
                w_px = int(page_pts.width()  / 72.0 * dpi)
                h_px = int(page_pts.height() / 72.0 * dpi)
                rendered = doc.render(i, QSize(max(1, w_px), max(1, h_px)))
                canvas = QImage(rendered.size(), QImage.Format_RGB32)
                canvas.fill(QColor("white"))
                pp = QPainter(canvas)
                try:
                    pp.drawImage(0, 0, rendered)
                finally:
                    pp.end()
                out = path if n == 1 else f"{base}_p{i+1}{real_ext}"
                if not canvas.save(out, fmt, 92 if fmt == 'JPG' else -1):
                    raise RuntimeError(f"No se pudo guardar {out}")
                outs.append(out)
            self._save_dir(path)
            msg = ("Imagen exportada:\n" + outs[0]) if n == 1 \
                else f"{n} páginas exportadas:\n{outs[0]}\n…\n{outs[-1]}"
            QMessageBox.information(self, "Exportar imagen", msg)
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Exportar imagen", f"No se pudo exportar: {e}")
        finally:
            try: os.unlink(tmp_pdf)
            except Exception: pass

    def _exportar_excel(self):
        """Exporta el cronograma valorizado a Excel (.xlsx). Mismo layout
        que el reporte PDF: mismas columnas, fondos sutiles para títulos y
        las 4 filas de resumen al pie, bordes finos en todas las celdas,
        freeze panes después de % Total."""
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os
        sugerido = os.path.join(self._dir_descargas(),
                                  f"valorizado_{self._cv.pid}.xlsx")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar cronograma valorizado a Excel",
            sugerido, "Excel (*.xlsx)"
        )
        if not path:
            return
        try:
            self._build_xlsx_valorizado(path)
            self._save_dir(path)
            QMessageBox.information(self, "Exportar Excel",
                                       f"Cronograma valorizado exportado a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")

    def _exportar_ods(self):
        """Exporta a OpenDocument Spreadsheet (.ods). Genera primero .xlsx
        y lo convierte vía `libreoffice --headless --convert-to ods`."""
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os, subprocess, tempfile
        from core.soffice import find_soffice, mensaje_instalacion
        soffice = find_soffice()
        if not soffice:
            QMessageBox.warning(self, "Exportar ODS", mensaje_instalacion())
            return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"valorizado_{self._cv.pid}.ods")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar cronograma valorizado a ODS",
            sugerido, "OpenDocument Spreadsheet (*.ods)"
        )
        if not path:
            return
        tmp_fd, tmp_xlsx = tempfile.mkstemp(suffix='.xlsx', prefix='valz_ods_')
        os.close(tmp_fd)
        try:
            self._build_xlsx_valorizado(tmp_xlsx)
            out_dir = os.path.dirname(os.path.abspath(path))
            subprocess.run(
                [soffice, '--headless', '--convert-to', 'ods',
                 '--outdir', out_dir, tmp_xlsx],
                check=True, capture_output=True, timeout=60,
            )
            generated = os.path.join(out_dir, os.path.splitext(os.path.basename(tmp_xlsx))[0] + '.ods')
            if os.path.exists(generated) and generated != path:
                os.replace(generated, path)
            if not os.path.exists(path):
                raise RuntimeError("LibreOffice no generó el archivo .ods")
            self._save_dir(path)
            QMessageBox.information(self, "Exportar ODS",
                                       f"Cronograma valorizado exportado a:\n{path}")
        except subprocess.CalledProcessError as e:
            QMessageBox.warning(self, "Exportar ODS",
                f"LibreOffice falló al convertir:\n{e.stderr.decode()[:300]}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")
        finally:
            try: os.unlink(tmp_xlsx)
            except Exception: pass

    def _build_xlsx_valorizado(self, path: str):
        """Construye el .xlsx — espejo VISUAL del PDF de Cronograma Valorizado.

        Sigue el patrón "PDF visible, no PDF CSS" del proyecto:
        - Header tripartito vía `_xlsx_header_pdf_style` (empresa | título |
          costo) + nombre proyecto full-width + separador slate-300.
        - h2 izquierdo "Cronograma Valorizado — por {label}" 13pt color
          `accent_reportes()[1]` (slate-900 en sobrio, naranja-700 en marca).
        - Sin gridlines de Excel (`showGridLines=False`); las separaciones
          visibles son bordes silver-100 0.6pt (igual que el PDF, que usa
          `border:0.6pt solid #CBD5E1` inline en cada `<td>`).
        - Header de columnas bg `#F1F5F9` + border-bottom 1.5pt accent `o`
          (espejo de `border-bottom:1.5pt solid {o}` del PDF).
        - Títulos N1: blanco + uppercase + border-top/bottom slate-700.
          Títulos N2+: color accent `od` + border-top/bottom 0.6pt accent `o`.
        - Filas resumen: TOTAL/ACUMULADO bg `#F1F5F9`, %PERÍODO/%ACUMULADO
          bg `#FBFCFD`. La fila TOTAL lleva `border-top` 1.5pt slate-700.

        Lee de los QTableWidget en pantalla (poblados por `cargar()`) — eso
        respeta período semanal/mensual y toggle "Mostrar fechas".
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from core.exporter import _xlsx_header_pdf_style
        from utils.theme import accent_reportes

        wb = Workbook()
        ws = wb.active
        ws.title = "Valorizado"

        proy = self._cv._proy

        # ── Paleta — espejo del PDF (`_html_cronograma_valorizado`) ──────
        o_hex, od_hex, _ = accent_reportes()
        ACC_O    = o_hex.lstrip('#').upper()        # accent primario
        ACC_OD   = od_hex.lstrip('#').upper()       # accent dark (h2)
        GRID     = 'CBD5E1'                          # silver-100 (PDF `GRID`)
        HDR_BG   = 'F1F5F9'                          # header bg
        BG_MONTO = 'F1F5F9'                          # TOTAL/ACUMULADO
        BG_PCT   = 'FBFCFD'                          # %PERÍODO/%ACUMULADO
        ALT_BG   = 'FBFCFD'                          # zebra muy sutil
        TXT_DK   = '1F2A38'                          # SLATE_900
        TXT_700  = '273445'                          # SLATE_700
        TXT_500  = '485A6C'                          # SLATE_500

        side_slate7  = Side(style='medium', color=TXT_700)     # ~1.5pt slate-700

        # Sin bordes en headers/datos/títulos — diferenciación por peso, color
        # y bg sutil (espejo del Presupuesto). Solo persiste el border-top
        # slate-700 de la fila TOTAL para separar el pie del cuerpo.
        border_grid    = Border()
        border_head    = Border()
        border_title1  = Border()
        border_title2  = Border()
        border_total   = Border(top=side_slate7)

        align_left   = Alignment(horizontal='left',   vertical='center',
                                  wrap_text=True)
        align_right  = Alignment(horizontal='right',  vertical='center')
        align_center = Alignment(horizontal='center', vertical='center',
                                  wrap_text=True)

        # SPLIT de Descripción en B+C — patrón espejo de Presupuesto/ACU/Insumos.
        # En cada fila de datos B+C se mergean → user ve 1 sola col Descripción.
        n_left_logical = self.tbl_l.columnCount()   # 6 cols de tbl_l
        n_left  = n_left_logical + 1                 # 7 cols físicas tras split
        n_right = self.tbl_r.columnCount()
        n_cols  = n_left + n_right

        # Helper: logical col idx (0-based, tbl_l) → physical Excel col (1-based)
        def _phys(c_logical: int) -> int:
            if c_logical <= 0:
                return 1                # Ítem
            if c_logical == 1:
                return 2                # Descripción (merged con col 3)
            return c_logical + 2        # resto shifted por el split

        def _merge_desc(row_idx: int):
            ws.merge_cells(start_row=row_idx, start_column=2,
                              end_row=row_idx, end_column=3)

        # ── Anchos de columna ANTES del header (necesario para wrap calc) ─
        ws.column_dimensions['A'].width = 11   # Ítem
        ws.column_dimensions['B'].width = 18   # Descripción parte 1
        ws.column_dimensions['C'].width = 24   # Descripción parte 2
        ws.column_dimensions['D'].width = 7    # Und
        ws.column_dimensions['E'].width = 13   # Cantidad
        ws.column_dimensions['F'].width = 14   # Precio
        ws.column_dimensions['G'].width = 14   # Precio Total
        for c in range(n_left + 1, n_cols + 1):
            ws.column_dimensions[get_column_letter(c)].width = 14

        # Quitar gridlines nativas — los bordes visibles los pintamos nosotros.
        ws.sheet_view.showGridLines = False
        ws.print_options.gridLines = False
        ws.print_options.gridLinesSet = True

        # ── Fila 1-3: header tripartito espejo del PDF ────────────────────
        # cols_partition=(2, 7, n_cols-1) → 3-tuple con hueco intencional:
        # Left = A-B (Ítem+Desc1), Center = C-G (Desc2+Und+Cantidad+Precio+
        # %Total) para el título y nombre del proyecto, Right = last 2 cols
        # (costo + modalidad). Cols intermedias (períodos) quedan sin
        # contenido en el header → solo los datos.
        if n_cols >= 9:
            row = _xlsx_header_pdf_style(
                ws, proy, 'Cronograma Valorizado', n_cols,
                cols_partition=(2, 7, n_cols - 1),
            )
        else:
            # Proyecto muy chico — fallback al cálculo automático.
            row = _xlsx_header_pdf_style(
                ws, proy, 'Cronograma Valorizado', n_cols,
            )

        # ── h2 izquierdo + subtítulo ─────────────────────────────────────
        label = 'Semana' if self._period_days == 7 else 'Mes'
        label_plural = 'semanas' if self._period_days == 7 else 'meses'
        plazo = int(proy.get('plazo') or 0)
        n_periods = max(0, n_right - 1)  # n_right incluye col "Total"

        ws.merge_cells(start_row=row, start_column=1,
                          end_row=row, end_column=n_cols)
        c_h2 = ws.cell(row=row, column=1, value=f"Cronograma Valorizado — por {label}")
        c_h2.font      = Font(name='Inter', bold=True, size=13, color=ACC_OD)
        c_h2.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[row].height = 20
        row += 1

        ws.merge_cells(start_row=row, start_column=1,
                          end_row=row, end_column=n_cols)
        c_sub = ws.cell(row=row, column=1,
                          value=f"Plazo: {plazo} días  ·  "
                                f"Períodos: {n_periods} {label_plural}")
        c_sub.font      = Font(name='Inter', size=9, color=TXT_500)
        c_sub.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[row].height = 14
        row += 1

        # Espacio antes de la tabla
        ws.row_dimensions[row].height = 6
        row += 1

        # ── Header de la tabla — 2 niveles (período sobre sus 2 sub-cols) ──
        head_fill = PatternFill('solid', fgColor=HDR_BG)
        head_font = Font(name='Inter', bold=True, size=10, color=TXT_DK)
        n_periods_h = (n_right - 1) // 2
        hdr1 = row
        hdr2 = row + 1
        # Fill del bloque de header en ambas filas
        for col_ in range(1, n_cols + 1):
            ws.cell(hdr1, col_).fill = head_fill
            ws.cell(hdr2, col_).fill = head_fill
        # Columnas izquierdas: rótulo merged verticalmente (hdr1:hdr2)
        for c in range(n_left_logical):
            pc = _phys(c)
            txt = (self.tbl_l.horizontalHeaderItem(c).text() or '').replace('\n', ' ')
            if c == 5:   # «% Total» → «Total» (precio total = pu × metrado)
                txt = 'Total'
            cell = ws.cell(hdr1, pc, txt)
            cell.font = head_font; cell.alignment = align_center
            if c == 1:   # Descripción ocupa B+C, ambas filas
                ws.merge_cells(start_row=hdr1, start_column=2, end_row=hdr2, end_column=3)
            else:
                ws.merge_cells(start_row=hdr1, start_column=pc, end_row=hdr2, end_column=pc)
        # Períodos: nombre (Mes/Sem + fecha) abarcando las 2 sub-cols en hdr1,
        # sub-etiquetas (Metrado | Valoriz.) en hdr2.
        for i in range(n_periods_h):
            mc = n_left + 1 + 2 * i
            vc = mc + 1
            raw_m = (self.tbl_r.horizontalHeaderItem(2 * i).text() or '')
            lines = raw_m.split('\n')
            plabel = '\n'.join(lines[:-1]) if len(lines) > 1 else raw_m
            sub_m  = lines[-1] if lines else 'Metrado'
            sub_v  = (self.tbl_r.horizontalHeaderItem(2 * i + 1).text() or 'Valoriz.').split('\n')[-1]
            ws.merge_cells(start_row=hdr1, start_column=mc, end_row=hdr1, end_column=vc)
            cp = ws.cell(hdr1, mc, plabel); cp.font = head_font; cp.alignment = align_center
            cm = ws.cell(hdr2, mc, sub_m);  cm.font = head_font; cm.alignment = align_center
            cvv= ws.cell(hdr2, vc, sub_v);  cvv.font = head_font; cvv.alignment = align_center
        # Total: merged vertical
        tc = n_cols
        ws.merge_cells(start_row=hdr1, start_column=tc, end_row=hdr2, end_column=tc)
        ct = ws.cell(hdr1, tc, (self.tbl_r.horizontalHeaderItem(n_right - 1).text() or 'Total'))
        ct.font = head_font; ct.alignment = align_center
        head_row = hdr1
        ws.row_dimensions[hdr1].height = 26 if self._show_fechas else 16
        ws.row_dimensions[hdr2].height = 14
        row += 2

        # ── Filas de datos ───────────────────────────────────────────────
        def _txt(item):
            return (item.text() if item is not None else '') or ''

        def _num(s):
            """Convierte string a float (quita símbolos de moneda)."""
            if not s:
                return None
            t = (str(s).replace('S/', '').replace('US$', '').replace('€', '')
                       .strip().replace(',', '').rstrip('%').strip())
            try:
                return float(t)
            except ValueError:
                return s

        # Acceso a las partidas para detectar títulos por índice (las filas
        # del QTableWidget mantienen el mismo orden que `_cv._partidas`).
        partidas = getattr(self._cv, '_partidas', None) or []
        # Colores de título por nivel = espejo del PDF (_NIVEL_COL) y del
        # programa: N1 rojo, N2 arándano, N3 morado, N4 rosa, default marrón.
        NIVEL_COL_X = {1: 'B71C1C', 2: '0D52BF', 3: '6A1B9A', 4: 'AD1457'}
        # Profundidad para la sangría (tabs) de la columna Descripción =
        # item.count('.') - min_dots (igual que el PDF / Presupuesto).
        _dots = [(p.get('item') or '').count('.') for p in partidas if p.get('item')]
        _min_dots = min(_dots) if _dots else 0
        def _depth_of(p):
            return max(0, (p.get('item') or '').count('.') - _min_dots)
        def _desc_align(depth):
            return (Alignment(horizontal='left', vertical='center',
                               wrap_text=True, indent=depth) if depth else align_left)

        def _alto_desc(texto: str, depth: int, bold: bool = False) -> float:
            """Alto de fila (pts) para que la Descripción mergeada B+C con
            wrap_text NO se corte. Excel no auto-ajusta el alto en celdas
            MERGEADAS con wrap, así que lo estimamos: ancho útil ≈ B+C menos
            la sangría, y contamos cuántas líneas ocupa el texto.

            `bold=True` (títulos/subtítulos) usa un ancho menor: el texto en
            negrita es más ancho → caben menos caracteres por línea."""
            ancho = (38 if bold else 42) - depth * 3   # B(18)+C(24) útiles − sangría
            if ancho < 12:
                ancho = 12
            lineas = 0
            for seg in (texto or '').split('\n'):
                lineas += max(1, -(-len(seg) // ancho))   # ceil div
            return max(15.0, lineas * 15.0)
        cell_font_data  = Font(name='Inter', size=9, color=TXT_DK)
        cell_font_alt   = Font(name='Inter', size=9, color=TXT_DK)
        cell_font_total = Font(name='Inter', size=9, color=TXT_DK, bold=True)

        # Contador propio de filas-partida para zebra striping CONTINUO:
        # los títulos no se cuentan, así que al volver a una partida después
        # de un título la alternancia sigue el patrón visual esperado.
        partida_idx = 0
        for r in range(self.tbl_l.rowCount()):
            es_titulo = (r < len(partidas) and bool(partidas[r].get('es_titulo')))
            if es_titulo:
                zebra = False
            else:
                zebra = (partida_idx % 2 == 1)
                partida_idx += 1

            if es_titulo:
                p = partidas[r]
                niv = int(p.get('nivel') or 1)
                # Color por nivel (espejo del PDF). N1: rojo + uppercase +
                # subrayado; N2+: color de nivel, sin subrayado.
                if niv <= 1:
                    t_color  = NIVEL_COL_X[1]
                    t_size   = 10
                    t_bdr    = border_title1
                    text_tx  = lambda s: (s or '').upper()
                else:
                    t_color  = NIVEL_COL_X.get(niv, '92400E')
                    t_size   = 10
                    t_bdr    = border_title2
                    text_tx  = lambda s: (s or '')
                t_font = Font(name='Inter', bold=True, size=t_size, color=t_color)
                # Subrayar solo títulos de nivel 1 (Ítem/Descripción).
                t_font_u = (Font(name='Inter', bold=True, size=t_size, color=t_color,
                                  underline='single') if niv <= 1 else t_font)
                # Col 1 = Ítem; col 2 = Descripción (merged con col 3); resto
                # vacío; col n_cols = parcial del título.
                cell = ws.cell(row=row, column=1, value=text_tx(p.get('item') or ''))
                cell.font = t_font_u; cell.alignment = align_left; cell.border = t_bdr
                _desc_txt = text_tx(p.get('descripcion') or '')
                cell = ws.cell(row=row, column=2, value=_desc_txt)
                cell.font = t_font_u; cell.alignment = _desc_align(_depth_of(p))
                cell.border = t_bdr
                ws.row_dimensions[row].height = _alto_desc(_desc_txt, _depth_of(p), bold=True)
                # Cols 3..n_cols vacías (col 3 ya parte del merge B+C)
                for c in range(3, n_cols + 1):
                    cell = ws.cell(row=row, column=c, value='')
                    cell.font = t_font; cell.alignment = align_right; cell.border = t_bdr
                _merge_desc(row)
                # En títulos/subtítulos la columna Total va vacía (solo las
                # partidas llevan total) — espejo del PDF. La col n_cols ya
                # quedó vacía en el loop de arriba.
                row += 1
                continue

            # Fila de partida ordinaria
            zebra_fill = (PatternFill('solid', fgColor=ALT_BG) if zebra else None)
            # PRE-aplicar fill+border en TODAS las cols del left panel
            # (incluye col 3 = Desc2, MergedCell post-merge)
            for col_ in range(1, n_left + 1):
                cell = ws.cell(row=row, column=col_)
                cell.font = cell_font_data
                cell.border = border_grid
                if zebra_fill:
                    cell.fill = zebra_fill
            for c in range(n_left_logical):
                val = _txt(self.tbl_l.item(r, c))
                cell = ws.cell(row=row, column=_phys(c))
                if c in (3, 4):                   # Cantidad / Precio
                    n = _num(val)
                    cell.value = n if isinstance(n, (int, float)) else val
                    if isinstance(n, (int, float)):
                        cell.number_format = '#,##0.00'
                    cell.alignment = align_right
                elif c == 5:                      # Precio Total = pu × metrado
                    p_row = partidas[r] if r < len(partidas) else {}
                    precio_total = parcial_wysiwyg(p_row.get('metrado') or 0,
                                                    p_row.get('precio_unitario') or 0)
                    cell.value = precio_total if precio_total else ''
                    if precio_total:
                        cell.number_format = '#,##0.00'
                    cell.alignment = align_right
                elif c == 2:                      # Und
                    cell.value = val
                    cell.alignment = align_center
                elif c == 1:                      # Descripción — con sangría
                    cell.value = val
                    p_row = partidas[r] if r < len(partidas) else {}
                    cell.alignment = _desc_align(_depth_of(p_row))
                else:
                    cell.value = val
                    cell.alignment = align_left
            _merge_desc(row)

            for c in range(n_right):
                val  = _txt(self.tbl_r.item(r, c))
                n    = _num(val)
                cell = ws.cell(row=row, column=n_left + c + 1)
                cell.value = n if isinstance(n, (int, float)) else val
                if isinstance(n, (int, float)):
                    cell.number_format = '#,##0.00'
                cell.alignment = align_right
                cell.font   = cell_font_total if c == n_right - 1 else cell_font_data
                cell.border = border_grid
                if zebra_fill:
                    cell.fill = zebra_fill
            # Alto de fila para que la descripción larga (B+C mergeada con
            # wrap) no se corte — Excel no lo auto-ajusta en celdas mergeadas.
            _p_row = partidas[r] if r < len(partidas) else {}
            ws.row_dimensions[row].height = _alto_desc(
                _txt(self.tbl_l.item(r, 1)), _depth_of(_p_row))
            row += 1

        # ── Filas de resumen (footer) ─────────────────────────────────────
        ftr_font = Font(name='Inter', bold=True, size=10, color=TXT_DK)
        lbl_font = Font(name='Inter', bold=True, size=10, color=TXT_700)
        for r in range(self.tbl_lf.rowCount()):
            is_pct  = (r in (1, 3))                  # %PERÍODO, %ACUMULADO
            is_total = (r == 0)                       # TOTAL — top bord. slate-700
            bg      = BG_PCT if is_pct else BG_MONTO
            fill    = PatternFill('solid', fgColor=bg)
            row_bdr = border_total if is_total else border_grid

            # Etiqueta merged sobre las n_left cols (espejo del PDF colspan=6)
            lbl = (_txt(self.tbl_lf.item(r, 0)) or '').strip()
            if not lbl:
                # fallback por si la celda está en col 1 vacía
                for c in range(n_left):
                    t = _txt(self.tbl_lf.item(r, c))
                    if t.strip():
                        lbl = t.strip()
                        break
            # PRE-estilar todas las cols del label ANTES del merge_cells —
            # tocar .fill/.border en un MergedCell ya creado lanza error en
            # openpyxl ≥ 3.1. El orden importa.
            for c in range(n_left):
                cell = ws.cell(row=row, column=c + 1)
                if c == 0:
                    cell.value = lbl
                cell.font      = lbl_font
                cell.fill      = fill
                cell.border    = row_bdr
                cell.alignment = align_right
            ws.merge_cells(start_row=row, start_column=1,
                              end_row=row, end_column=n_left)

            for c in range(n_right):
                val  = _txt(self.tbl_rf.item(r, c))
                n    = _num(val)
                cell = ws.cell(row=row, column=n_left + c + 1)
                if isinstance(n, (int, float)):
                    cell.value = n
                    cell.number_format = '0.0"%"' if '%' in val else '#,##0.00'
                else:
                    cell.value = val
                cell.font      = ftr_font
                cell.fill      = fill
                cell.border    = row_bdr
                cell.alignment = align_right
            ws.row_dimensions[row].height = 18
            row += 1

        # ── Líneas divisorias (estilo Delphin / hoja de Marco) ────────────
        # Verticales: medias en borde izq. de la tabla, tras el bloque
        # izquierdo (% Total) y en el borde derecho; finas entre cada mes.
        # NO se dibujan horizontales entre partidas — la separación es la
        # zebra. El header lleva recuadro medio (top en hdr1, bottom en hdr2).
        last_row = row - 1
        MED_CLR  = '94A3B8'      # slate-400 — divisoria fuerte
        THIN_CLR = 'D5DCE5'      # slate-200 — divisoria sutil entre meses

        def _side(cell, edge, style, color):
            b = cell.border
            kw = {s: getattr(b, s) for s in ('top', 'bottom', 'left', 'right')}
            kw[edge] = Side(style=style, color=color)
            cell.border = Border(**kw)

        med_cols = {1: 'left', n_left: 'right', n_cols: 'right'}   # físicas
        # límites finos: borde derecho de la valoriz. de cada mes (= sep. mes)
        thin_right = {n_left + 2 * (m + 1) for m in range(n_periods_h)}
        thin_right.discard(n_cols - 1)   # ese límite ya queda junto a Total
        for r in range(hdr1, last_row + 1):
            for col_, edge in med_cols.items():
                _side(ws.cell(r, col_), edge, 'medium', MED_CLR)
            for col_ in thin_right:
                _side(ws.cell(r, col_), 'right', 'thin', THIN_CLR)
        # Recuadro del header (top en hdr1, bottom en hdr2) y cierre inferior
        for col_ in range(1, n_cols + 1):
            _side(ws.cell(hdr1, col_), 'top', 'medium', MED_CLR)
            _side(ws.cell(hdr2, col_), 'bottom', 'medium', MED_CLR)
            _side(ws.cell(last_row, col_), 'bottom', 'medium', MED_CLR)

        # Freeze panes después de % Total (col F) y bajo el header de 2 filas
        # Coordenada como STRING — pasar Cell falla si la celda destino acaba
        # como MergedCell por colisión con merges previos (openpyxl 3.x).
        ws.freeze_panes = f'{get_column_letter(n_left + 1)}{head_row + 2}'

        # Repetir filas del encabezado (tripartito + h2 + subtítulo + header
        # de columnas de 2 filas) en cada página impresa. LibreOffice mapea
        # esto a `<table:table-header-rows>` al convertir a ODS.
        ws.print_title_rows = f'1:{head_row + 1}'
        # Repetir las cols frozen (Ítem · Descripción · Und · Cantidad ·
        # Precio · % Total) en cada página horizontal — sin esto la 2ª
        # página de períodos sale sin contexto de partida.
        ws.print_title_cols = f'A:{get_column_letter(n_left)}'

        # Orientación apaisada para imprimir; ajustar al ancho
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.print_options.horizontalCentered = True
        # Márgenes compactos espejo de `core/exporter.py::_setup_impresion`.
        from openpyxl.worksheet.page import PageMargins as _PM
        ws.page_margins = _PM(left=0.4, right=0.4, top=0.35, bottom=0.5,
                                 header=0.15, footer=0.2)

        # Pie tripartito espejo del PDF: Cliente | Fecha | Página X de N.
        # `&K{HEX}&{SIZE}` setea color+tamaño; `&P`/`&N` = página actual / total.
        from datetime import datetime as _dt_p
        cliente = (proy.get('cliente') or '').strip()
        fecha   = _dt_p.now().strftime('%d/%m/%Y')
        # OJO: el código de tamaño (&7) debe ir ANTES del color (&K). Si va
        # último y el texto empieza con dígito (la fecha), Excel lee `&7` +
        # los dígitos del texto como un solo tamaño (p.ej. &701 → 701pt). Con
        # `&K` al final lee exactamente 6 hex y el dígito siguiente ya es texto.
        prefix  = '&7&K485A6C'  # 7pt + slate-500
        ws.oddFooter.left.text   = f"{prefix}Cliente: {cliente}" if cliente else f"{prefix} "
        ws.oddFooter.center.text = f"{prefix}{fecha}"
        ws.oddFooter.right.text  = f"{prefix}Página &P de &N"

        wb.save(path)

    def _dir_descargas(self) -> str:
        import os
        from PySide6.QtCore import QSettings
        s = QSettings("ingePresupuestos", "exports")
        guardado = s.value("last_dir_valorizado", "")
        if guardado and os.path.isdir(guardado):
            return guardado
        for c in (os.path.expanduser("~/Descargas"),
                   os.path.expanduser("~/Downloads"),
                   os.path.expanduser("~")):
            if os.path.isdir(c):
                return c
        return os.getcwd()

    def _save_dir(self, path: str):
        import os
        from PySide6.QtCore import QSettings
        if not path:
            return
        d = os.path.dirname(path)
        if d and os.path.isdir(d):
            QSettings("ingePresupuestos", "exports").setValue("last_dir_valorizado", d)


# ════════════════════════════════════════════════════════════════════════════
# 3. CURVA S
# ════════════════════════════════════════════════════════════════════════════

class CurvaSWidget(QWidget):
    """Curva S — avance acumulado por semana/mes con tabla + gráfico
    sincronizados. Mismo patrón visual que Valorizado/Insumos:
    KPI strip + toolbar + splitter (tabla zebra | chart) + footer + exportes.
    """

    def __init__(self, parent: CronogramaView):
        super().__init__(parent)
        self._cv = parent
        self._period_days = 30  # default Mensual (Marco prefiere mes al abrir)
        self._show_fechas = True
        self._suavizar = True   # default: curva suave (Bézier)
        self._show_barras = True  # default: muestra barras de avance
        self._last_data = None  # cached para exportes
        self._build_ui()

    # ── UI ───────────────────────────────────────────────────────────────
    def _build_ui(self):
        vl = QVBoxLayout(self)
        vl.setContentsMargins(8, 8, 8, 8)
        vl.setSpacing(6)

        # (KPI strip removido por pedido de Marco — no se veía consistente.)

        # ── Toolbar ──────────────────────────────────────────────────────
        tb = QFrame()
        tb.setStyleSheet("background:transparent; border:none;")
        tb_hl = QHBoxLayout(tb)
        tb_hl.setContentsMargins(0, 0, 0, 0); tb_hl.setSpacing(8)
        lbl = QLabel("Período:")
        lbl.setStyleSheet(f"color:{SLATE_500}; font-size:11px;")
        tb_hl.addWidget(lbl)
        cmb = QComboBox()
        cmb.addItems(["Semanal", "Mensual"])
        cmb.setCurrentIndex(1)  # default Mensual (espejo de _period_days=30)
        cmb.currentIndexChanged.connect(self._on_periodo)
        cmb.setStyleSheet(
            "QComboBox { min-height:0; padding:2px 8px; font-size:11px;"
            f" border:1px solid {SILVER_300}; border-radius:4px; }}"
        )
        tb_hl.addWidget(cmb)

        from PySide6.QtWidgets import QCheckBox
        chk_fechas = QCheckBox("Mostrar fechas")
        chk_fechas.setChecked(self._show_fechas)
        chk_fechas.setStyleSheet(
            f"QCheckBox {{ color:{SLATE_500}; font-size:11px; padding:0 8px; }}"
        )
        chk_fechas.toggled.connect(self._on_toggle_fechas)
        tb_hl.addWidget(chk_fechas)

        chk_smooth = QCheckBox("Suavizar curva")
        chk_smooth.setChecked(self._suavizar)
        chk_smooth.setStyleSheet(
            f"QCheckBox {{ color:{SLATE_500}; font-size:11px; padding:0 8px; }}"
        )
        chk_smooth.toggled.connect(self._on_toggle_suavizar)
        from utils.tooltip import set_tooltip as _set_tt_sm
        _set_tt_sm(chk_smooth,
                     "Dibuja la curva con un suavizado Bézier en vez de "
                     "segmentos rectos entre puntos.")
        tb_hl.addWidget(chk_smooth)

        chk_bars = QCheckBox("Mostrar barras de avance")
        chk_bars.setChecked(self._show_barras)
        chk_bars.setStyleSheet(
            f"QCheckBox {{ color:{SLATE_500}; font-size:11px; padding:0 8px; }}"
        )
        chk_bars.toggled.connect(self._on_toggle_barras)
        _set_tt_sm(chk_bars,
                     "Muestra/oculta las barras naranjas que indican el % de "
                     "avance de cada período (la curva azul queda igual).")
        tb_hl.addWidget(chk_bars)

        tb_hl.addStretch()

        from utils.tooltip import set_tooltip
        btn_style = (
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:3px 10px;"
            f" min-height:0; color:{SLATE_500}; }}"
            f"QPushButton:hover {{ background:{SILVER_100}; color:{SLATE_700}; }}"
        )
        btn_pdf = QPushButton("📄 PDF")
        btn_pdf.setCursor(Qt.PointingHandCursor)
        btn_pdf.setStyleSheet(btn_style)
        btn_pdf.clicked.connect(self._exportar_pdf)
        set_tooltip(btn_pdf,
                      "Exportar la Curva S a PDF (con vista previa, "
                      "tamaño de papel y orientación configurables)")
        tb_hl.addWidget(btn_pdf)

        btn_word = QPushButton("📝 Word")
        btn_word.setCursor(Qt.PointingHandCursor)
        btn_word.setStyleSheet(btn_style)
        btn_word.clicked.connect(self._exportar_word)
        set_tooltip(btn_word,
                      "Exportar reporte editable .docx con gráfico embebido "
                      "+ tabla resumen + notas")
        tb_hl.addWidget(btn_word)

        btn_odt = QPushButton("📰 ODT")
        btn_odt.setCursor(Qt.PointingHandCursor)
        btn_odt.setStyleSheet(btn_style)
        btn_odt.clicked.connect(self._exportar_odt)
        set_tooltip(btn_odt,
                      "Exportar reporte editable .odt (LibreOffice / "
                      "OpenOffice). Requiere LibreOffice instalado.")
        tb_hl.addWidget(btn_odt)

        vl.addWidget(tb)

        # ── Splitter vertical: gráfico (arriba) | tabla (abajo) ──────────
        # La curva S es ante todo un gráfico; va arriba a todo el ancho y la
        # tabla de respaldo, compacta, debajo (antes lado a lado dejaba mucho
        # espacio en blanco y el gráfico se veía perdido).
        self.split = QSplitter(Qt.Vertical)
        self.split.setHandleWidth(6)
        self.split.setStyleSheet(
            "QSplitter::handle { background:#CBD5E1; }"
            "QSplitter::handle:vertical { height:6px; margin:0; }"
            "QSplitter::handle:hover { background:#F37329; }"
        )
        self.split.setChildrenCollapsible(False)

        # ── Panel izquierdo: tabla + footer sticky ──
        left_fr = QFrame()
        lv = QVBoxLayout(left_fr)
        lv.setContentsMargins(0, 0, 0, 0); lv.setSpacing(0)
        self.tbl = QTableWidget(0, 0)
        self.tbl.verticalHeader().setVisible(False)
        self.tbl.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.tbl.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tbl.setShowGrid(True)
        self.tbl.setAlternatingRowColors(False)
        self.tbl.setStyleSheet(f"""
            QTableWidget {{ border:1px solid {SILVER_300};
                            border-bottom:none; font-size:11px;
                            gridline-color:#E0E5EC; background:white; }}
            QTableWidget::item {{ padding:3px 8px; }}
            QHeaderView::section {{
                background:{SLATE_500}; color:white; font-size:10px;
                font-weight:700; padding:4px 6px; border:none;
            }}
        """)
        lv.addWidget(self.tbl, stretch=1)

        # Footer sticky con TOTAL · 100%
        self.tbl_ftr = QTableWidget(0, 0)
        self.tbl_ftr.verticalHeader().setVisible(False)
        self.tbl_ftr.horizontalHeader().setVisible(False)
        self.tbl_ftr.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.tbl_ftr.setSelectionMode(QAbstractItemView.NoSelection)
        self.tbl_ftr.setFocusPolicy(Qt.NoFocus)
        self.tbl_ftr.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.tbl_ftr.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.tbl_ftr.setShowGrid(True)
        self.tbl_ftr.setStyleSheet(f"""
            QTableWidget {{ border:1px solid {SILVER_300};
                            border-top:2px solid {SLATE_500};
                            font-size:11px; background:white;
                            gridline-color:#C5D1E0; }}
            QTableWidget::item {{ padding:4px 8px; }}
        """)
        lv.addWidget(self.tbl_ftr)

        # Delegate de bg para zebra robusta
        self._bg_delegate = _BgFillDelegate(self)
        for tbl_ in (self.tbl, self.tbl_ftr):
            tbl_.setItemDelegate(self._bg_delegate)

        # ── Panel derecho: gráfico de la curva ──
        right_fr = QFrame()
        rv = QVBoxLayout(right_fr)
        rv.setContentsMargins(0, 0, 0, 0); rv.setSpacing(0)
        self.scene = QGraphicsScene()
        self.view = QGraphicsView(self.scene)
        self.view.setRenderHint(QPainter.Antialiasing)
        self.view.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.view.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.view.setStyleSheet(f"background:white; border:1px solid {SILVER_300}; border-radius:6px;")
        rv.addWidget(self.view, stretch=1)

        # Gráfico ARRIBA (grande), tabla ABAJO (compacta).
        self.split.addWidget(right_fr)
        self.split.addWidget(left_fr)
        self.split.setStretchFactor(0, 7)
        self.split.setStretchFactor(1, 3)
        self.split.setSizes([460, 220])
        vl.addWidget(self.split, stretch=1)

    # ── Acciones de toolbar ─────────────────────────────────────────────
    def _on_periodo(self, idx: int):
        self._period_days = 7 if idx == 0 else 30
        self.cargar()

    def _on_toggle_fechas(self, on: bool):
        self._show_fechas = bool(on)
        self.cargar()

    def _on_toggle_suavizar(self, on: bool):
        self._suavizar = bool(on)
        self._dibujar_chart()

    def _on_toggle_barras(self, on: bool):
        self._show_barras = bool(on)
        self._dibujar_chart()

    # ── Helpers ─────────────────────────────────────────────────────────
    def _project_start(self):
        for key in ('fecha_inicio', 'costo_al'):
            fi = (self._cv._proy.get(key) or '').strip()
            if fi:
                try:
                    return datetime.strptime(fi, '%Y-%m-%d')
                except Exception:
                    pass
        return datetime.now()

    def _periodo_label(self, unidad: str, i: int) -> str:
        """Etiqueta como 'Sem 1\n01/01–07/01' si _show_fechas y hay fecha."""
        base = f"{unidad} {i + 1}"
        if not self._show_fechas:
            return base
        f_ini = self._project_start()
        if f_ini is None:
            return base
        ini = f_ini + timedelta(days=i * self._period_days)
        fin = ini + timedelta(days=self._period_days - 1)
        return f"{base}\n{ini.strftime('%d/%m')}–{fin.strftime('%d/%m')}"

    def _moneda_simbolo(self):
        try:
            from core.config import MONEDAS
            return MONEDAS.get(self._cv._moneda, {}).get('simbolo', 'S/')
        except Exception:
            return 'S/'

    # ── Cálculo y render ────────────────────────────────────────────────
    def cargar(self):
        if not hasattr(self._cv, '_partidas'):
            return
        partidas = self._cv._partidas
        tasks    = self._cv._tasks
        cmap     = self._cv._cron_map
        plazo    = self._cv._proy.get('plazo') or 0
        max_ef   = max((t['EF'] for t in tasks.values() if t['EF'] > 0),
                        default=plazo)
        n_dias   = max(max_ef, plazo, self._period_days)
        n_periods = max(1, (n_dias + self._period_days - 1) // self._period_days)
        unidad   = "Sem" if self._period_days == 7 else "Mes"
        sym      = self._moneda_simbolo()
        dec      = get_decimales_ppto()

        # Distribución por período
        total_per = [0.0] * n_periods
        for p in partidas:
            if p['es_titulo']:
                continue
            cd = cmap.get(p['id'], {})
            ini = cd.get('inicio_dia', 1) or 1
            dur = cd.get('duracion', 0) or 0
            t = tasks.get(p['id'], {})
            ini = t.get('ES', ini)
            parcial = parcial_wysiwyg(p['metrado'], p['precio_unitario'])
            weekly = distribuir_periodos(
                cd.get('segmentos', '') or '',
                ini, dur, parcial, n_periods, self._period_days
            )
            for i, v in enumerate(weekly):
                total_per[i] += v

        total_general = sum(total_per)
        acum = [0.0] * n_periods
        run = 0.0
        for i in range(n_periods):
            run += total_per[i]
            acum[i] = run
        per_pct  = [(v / total_general * 100 if total_general > 0 else 0) for v in total_per]
        acum_pct = [(v / total_general * 100 if total_general > 0 else 0) for v in acum]
        pico_idx = (max(range(n_periods), key=lambda i: total_per[i])
                     if total_general > 0 else 0)

        self._last_data = {
            'unidad': unidad, 'n_periods': n_periods, 'total_per': total_per,
            'acum': acum, 'per_pct': per_pct, 'acum_pct': acum_pct,
            'total_general': total_general, 'pico_idx': pico_idx,
        }

        # ── Llenar tabla ──
        self.tbl.blockSignals(True)
        self.tbl.clear()
        self.tbl.setColumnCount(6)
        self.tbl.setHorizontalHeaderLabels(
            ["#", "Período", f"Avance ({sym})", "% Período",
             f"Acumulado ({sym})", "% Acumulado"]
        )
        # Anchos cómodos
        self.tbl.setColumnWidth(0, 36)
        self.tbl.setColumnWidth(1, 130)
        self.tbl.setColumnWidth(2, 110)
        self.tbl.setColumnWidth(3, 80)
        self.tbl.setColumnWidth(4, 130)
        self.tbl.setColumnWidth(5, 88)
        self.tbl.horizontalHeader().setStretchLastSection(False)
        self.tbl.setRowCount(n_periods)

        zebra_bg = QColor("#F8FAFC")
        pico_bg  = QColor("#FEF5EB")
        for i in range(n_periods):
            row_color = pico_bg if (i == pico_idx and total_general > 0) else (
                zebra_bg if i % 2 == 1 else None)
            vals = [
                (str(i + 1), Qt.AlignCenter),
                (self._periodo_label(unidad, i).replace("\n", " · "), Qt.AlignLeft),
                (fmt(total_per[i], self._cv._moneda, dec), Qt.AlignRight),
                (f"{per_pct[i]:.2f}%", Qt.AlignRight),
                (fmt(acum[i], self._cv._moneda, dec), Qt.AlignRight),
                (f"{acum_pct[i]:.2f}%", Qt.AlignRight),
            ]
            for c, (txt, al) in enumerate(vals):
                it = QTableWidgetItem(txt)
                it.setTextAlignment(int(al) | Qt.AlignVCenter)
                if row_color is not None:
                    it.setBackground(QBrush(row_color))
                if i == pico_idx and total_general > 0 and c in (2, 3):
                    f = QFont(); f.setBold(True); it.setFont(f)
                    it.setForeground(QBrush(QColor("#C0621A")))
                self.tbl.setItem(i, c, it)
        self.tbl.verticalHeader().setDefaultSectionSize(24)
        self.tbl.blockSignals(False)

        # ── Footer sticky con TOTAL ──
        self.tbl_ftr.setColumnCount(6)
        for c, w in enumerate([36, 130, 110, 80, 130, 88]):
            self.tbl_ftr.setColumnWidth(c, w)
        self.tbl_ftr.setRowCount(1)
        ftr_vals = [
            "", "TOTAL", fmt(total_general, self._cv._moneda, dec),
            "100.00%" if total_general > 0 else "—",
            fmt(total_general, self._cv._moneda, dec),
            "100.00%" if total_general > 0 else "—",
        ]
        for c, txt in enumerate(ftr_vals):
            it = QTableWidgetItem(txt)
            al = Qt.AlignRight if c in (2, 3, 4, 5) else (
                Qt.AlignCenter if c == 0 else Qt.AlignLeft)
            it.setTextAlignment(int(al) | Qt.AlignVCenter)
            it.setBackground(QBrush(QColor("#E2E8F0")))
            f = QFont(); f.setBold(True); it.setFont(f)
            it.setForeground(QBrush(QColor(SLATE_700)))
            self.tbl_ftr.setItem(0, c, it)
        self.tbl_ftr.setFixedHeight(self.tbl_ftr.verticalHeader().defaultSectionSize() + 2)

        # Sincronizar scroll horizontal de tabla y footer
        self.tbl.horizontalScrollBar().valueChanged.connect(
            self.tbl_ftr.horizontalScrollBar().setValue)

        # ── Dibujar gráfico ──
        self._dibujar_chart()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        # Redibujar la curva para que llene el ancho/alto disponible.
        if getattr(self, '_last_data', None):
            QTimer.singleShot(0, self._dibujar_chart)

    def showEvent(self, event):
        super().showEvent(event)
        if getattr(self, '_last_data', None):
            QTimer.singleShot(0, self._dibujar_chart)

    def _dibujar_chart(self):
        """Pinta la curva S en self.scene usando _last_data."""
        self.scene.clear()
        data = self._last_data or {}
        n_periods = data.get('n_periods', 0)
        acum_pct = data.get('acum_pct', [])
        per_pct  = data.get('per_pct', [])
        unidad   = data.get('unidad', 'Sem')
        total_general = data.get('total_general', 0)

        if total_general <= 0 or n_periods == 0:
            txt = QGraphicsTextItem(
                "Sin datos. Configura duraciones y precios para ver la curva."
            )
            f = QFont(); f.setPointSize(11); txt.setFont(f)
            txt.setDefaultTextColor(QColor(SLATE_300))
            txt.setPos(40, 40)
            self.scene.addItem(txt)
            self.scene.setSceneRect(0, 0, 600, 100)
            return

        # Curva acumulada arranca en 0 (antes del período 1)
        pts_acum = [0.0] + list(acum_pct)
        # El gráfico llena el viewport (se redibuja en resizeEvent). Fallback a
        # tamaño fijo si la vista aún no tiene dimensiones (1ª construcción).
        vp = self.view.viewport().size()
        W = max(620, vp.width() - 4) if vp.width() > 60 else max(620, n_periods * 60 + 100)
        H = max(300, vp.height() - 4) if vp.height() > 60 else 380
        margin_l, margin_r, margin_t, margin_b = 64, 30, 36, 64
        plot_w = W - margin_l - margin_r
        plot_h = H - margin_t - margin_b

        # Líneas horizontales 0/25/50/75/100
        for y_pct in (0, 25, 50, 75, 100):
            y = margin_t + plot_h * (1 - y_pct / 100)
            line = QGraphicsLineItem(margin_l, y, margin_l + plot_w, y)
            line.setPen(QPen(QColor("#E8EAED"), 0.5, Qt.DashLine))
            self.scene.addItem(line)
            txt = QGraphicsTextItem(f"{y_pct}%")
            f = QFont(); f.setPointSize(8); txt.setFont(f)
            txt.setDefaultTextColor(QColor(SLATE_500))
            txt.setPos(margin_l - 32, y - 8)
            self.scene.addItem(txt)

        # Barras de avance por período (naranja translúcido) — opcional
        if getattr(self, '_show_barras', True):
            bar_w = plot_w / max(n_periods, 1) * 0.7
            for i in range(n_periods):
                x_c = margin_l + plot_w * (i + 0.5) / n_periods
                h = plot_h * (per_pct[i] / 100)
                rect = QGraphicsRectItem(x_c - bar_w / 2,
                                           margin_t + plot_h - h,
                                           bar_w, h)
                rect.setBrush(QBrush(QColor(243, 115, 41, 70)))
                rect.setPen(QPen(QColor("#F37329"), 0.4))
                rect.setZValue(0)
                self.scene.addItem(rect)

        # Eje X — ticks y labels
        for i in range(n_periods + 1):
            x = margin_l + plot_w * i / n_periods
            line = QGraphicsLineItem(x, margin_t + plot_h,
                                       x, margin_t + plot_h + 4)
            line.setPen(QPen(QColor(SLATE_300), 1))
            self.scene.addItem(line)
            step = max(1, n_periods // 10)
            if i % step == 0 or i == n_periods:
                lbl = f"{unidad} {i}" if i > 0 else "0"
                txt = QGraphicsTextItem(lbl)
                f = QFont(); f.setPointSize(8); txt.setFont(f)
                txt.setDefaultTextColor(QColor(SLATE_500))
                tw = txt.boundingRect().width()
                txt.setPos(x - tw / 2, margin_t + plot_h + 8)
                self.scene.addItem(txt)

        # Curva acumulada (azul)
        path_pts = []
        for i, pct in enumerate(pts_acum):
            x = margin_l + plot_w * i / n_periods
            y = margin_t + plot_h * (1 - pct / 100)
            path_pts.append((x, y))

        path = QPainterPath()
        if path_pts:
            path.moveTo(*path_pts[0])
        if getattr(self, '_suavizar', True) and len(path_pts) >= 2:
            # Curva suave Catmull-Rom → Bézier cúbico. Para cada segmento
            # P1→P2 calculamos control points usando P0 y P3 (puntos vecinos)
            # con tensión 0.5; los extremos repiten el primer/último punto.
            ext = [path_pts[0]] + path_pts + [path_pts[-1]]
            for i in range(1, len(ext) - 2):
                p0x, p0y = ext[i - 1]
                p1x, p1y = ext[i]
                p2x, p2y = ext[i + 1]
                p3x, p3y = ext[i + 2]
                c1x = p1x + (p2x - p0x) / 6.0
                c1y = p1y + (p2y - p0y) / 6.0
                c2x = p2x - (p3x - p1x) / 6.0
                c2y = p2y - (p3y - p1y) / 6.0
                # Para evitar overshoot (la curva acumulada es monótona
                # creciente — y nunca debe bajar visualmente), clamp Y
                # de control points al rango entre P1.y y P2.y.
                ymin, ymax = min(p1y, p2y), max(p1y, p2y)
                c1y = max(ymin, min(ymax, c1y))
                c2y = max(ymin, min(ymax, c2y))
                path.cubicTo(c1x, c1y, c2x, c2y, p2x, p2y)
        else:
            for (x, y) in path_pts[1:]:
                path.lineTo(x, y)

        path_item = QGraphicsPathItem(path)
        path_item.setPen(QPen(QColor("#3689E6"), 2.5))
        path_item.setBrush(QBrush(Qt.NoBrush))
        path_item.setZValue(2)
        self.scene.addItem(path_item)

        # Dots en los puntos reales (siempre, suavice o no)
        for (x, y) in path_pts:
            dot = QGraphicsRectItem(x - 3, y - 3, 6, 6)
            dot.setBrush(QBrush(QColor("#3689E6")))
            dot.setPen(QPen(QColor("white"), 1))
            dot.setZValue(3)
            self.scene.addItem(dot)

        # Etiquetas % en cada punto (cada N o si es 100)
        step_lbl = max(1, n_periods // 8)
        for i, pct in enumerate(pts_acum):
            if i == 0:
                continue
            if i % step_lbl != 0 and i != n_periods:
                continue
            x = margin_l + plot_w * i / n_periods
            y = margin_t + plot_h * (1 - pct / 100)
            txt = QGraphicsTextItem(f"{pct:.0f}%")
            f = QFont(); f.setPointSize(7); f.setBold(True); txt.setFont(f)
            txt.setDefaultTextColor(QColor("#1E5DA8"))
            txt.setPos(x - 10, y - 18)
            txt.setZValue(4)
            self.scene.addItem(txt)

        # Ejes
        line_y = QGraphicsLineItem(margin_l, margin_t,
                                     margin_l, margin_t + plot_h)
        line_y.setPen(QPen(QColor(SLATE_500), 1.5))
        self.scene.addItem(line_y)
        line_x = QGraphicsLineItem(margin_l, margin_t + plot_h,
                                     margin_l + plot_w, margin_t + plot_h)
        line_x.setPen(QPen(QColor(SLATE_500), 1.5))
        self.scene.addItem(line_x)

        # Título
        ttl = QGraphicsTextItem("CURVA S — Avance financiero acumulado")
        f = QFont(); f.setPointSize(11); f.setBold(True); ttl.setFont(f)
        ttl.setDefaultTextColor(QColor(SLATE_700))
        ttl.setPos(margin_l, 6)
        self.scene.addItem(ttl)

        # Leyenda — solo entradas relevantes
        leg_y = H - 22
        f = QFont(); f.setPointSize(8)
        cur_lx = margin_l
        if getattr(self, '_show_barras', True):
            sw = QGraphicsRectItem(cur_lx, leg_y, 10, 10)
            sw.setBrush(QBrush(QColor(243, 115, 41, 130)))
            sw.setPen(QPen(QColor("#F37329"), 0.6))
            self.scene.addItem(sw)
            t1 = QGraphicsTextItem("Avance del período (%)")
            t1.setFont(f); t1.setDefaultTextColor(QColor(SLATE_500))
            t1.setPos(cur_lx + 14, leg_y - 4)
            self.scene.addItem(t1)
            cur_lx += 170
        ln = QGraphicsLineItem(cur_lx, leg_y + 5, cur_lx + 18, leg_y + 5)
        ln.setPen(QPen(QColor("#3689E6"), 2.5))
        self.scene.addItem(ln)
        t2 = QGraphicsTextItem("Acumulado (%)")
        t2.setFont(f); t2.setDefaultTextColor(QColor(SLATE_500))
        t2.setPos(cur_lx + 26, leg_y - 4)
        self.scene.addItem(t2)

        self.scene.setSceneRect(0, 0, W, H)

    # ── Exportes ────────────────────────────────────────────────────────
    def _dir_descargas(self) -> str:
        import os
        from PySide6.QtCore import QSettings
        s = QSettings("ingePresupuestos", "exports")
        guardado = s.value("last_dir_curvas", "")
        if guardado and os.path.isdir(guardado):
            return guardado
        for c in (os.path.expanduser("~/Descargas"),
                   os.path.expanduser("~/Downloads"),
                   os.path.expanduser("~")):
            if os.path.isdir(c):
                return c
        return os.getcwd()

    def _save_dir(self, path: str):
        import os
        from PySide6.QtCore import QSettings
        if not path:
            return
        d = os.path.dirname(path)
        if d and os.path.isdir(d):
            QSettings("ingePresupuestos", "exports").setValue("last_dir_curvas", d)

    def _exportar_pdf(self):
        """Abre un diálogo con tamaño de papel, orientación, vista previa y
        exporta el reporte de Curva S a PDF."""
        if not self._last_data or (self._last_data.get('total_general', 0) <= 0):
            QMessageBox.information(self, "Exportar PDF",
                "No hay datos para exportar (configura duraciones y precios).")
            return
        opts = self._dialog_pdf_opciones()
        if opts is None:
            return
        if opts.get('_action') == 'preview':
            return  # ya se mostró
        self._exportar_pdf_a_archivo(opts)

    def _exportar_pdf_a_archivo(self, opts: dict):
        import os
        sugerido = os.path.join(self._dir_descargas(),
                                  f"curva_s_{self._cv.pid}.pdf")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar Curva S a PDF", sugerido, "PDF (*.pdf)"
        )
        if not path:
            return
        try:
            self._render_pdf(path, paper=opts.get('papel', 'A4'),
                              orient=opts.get('orient', 'portrait'))
            self._save_dir(path)
            QMessageBox.information(self, "Exportar PDF",
                                       f"Curva S exportada a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")

    def _dialog_pdf_opciones(self):
        """Diálogo simple: papel + orientación + vista previa. Devuelve
        un dict con `papel`/`orient` o None si se canceló."""
        from PySide6.QtGui import QPalette as _QPal
        dlg = QDialog(self)
        dlg.setWindowTitle("Exportar Curva S a PDF")
        dlg.setModal(True)
        dlg.resize(440, 280)
        pal = dlg.palette()
        pal.setColor(_QPal.Window, QColor("white"))
        pal.setColor(_QPal.WindowText, QColor("#273445"))
        dlg.setPalette(pal); dlg.setAutoFillBackground(True)
        dlg.setStyleSheet(
            "QDialog { background:white; }"
            "QLabel { color:#273445; font-size:12px; background:transparent; }"
            "QComboBox, QRadioButton { font-size:12px; color:#273445; }"
            "QComboBox { padding:3px 6px; border:1px solid #D4D4D4;"
            " border-radius:4px; background:white; min-width:200px; }"
        )
        v = QVBoxLayout(dlg)
        v.setContentsMargins(18, 16, 18, 14); v.setSpacing(10)
        ttl = QLabel("Opciones de exportación")
        f = ttl.font(); f.setPointSize(13); f.setBold(True); ttl.setFont(f)
        v.addWidget(ttl)

        # Papel
        h1 = QHBoxLayout()
        h1.addWidget(QLabel("Tamaño de papel:"))
        cmb = QComboBox()
        PAPEL_OPS = [
            ("A4 (210 × 297 mm)", "A4"),
            ("A3 (297 × 420 mm)", "A3"),
            ("A2 (420 × 594 mm)", "A2"),
            ("A1 (594 × 841 mm)", "A1"),
            ("A0 (841 × 1189 mm)", "A0"),
            ("Carta (216 × 279 mm)", "Letter"),
            ("Tabloide (279 × 432 mm)", "Tabloid"),
        ]
        for lbl, _ in PAPEL_OPS:
            cmb.addItem(lbl)
        cmb.setCurrentIndex(0)
        h1.addWidget(cmb, 1)
        v.addLayout(h1)

        # Orientación
        v.addWidget(QLabel("Orientación:"))
        rb_p = QRadioButton("Vertical (Portrait) — recomendado")
        rb_l = QRadioButton("Horizontal (Landscape)")
        rb_p.setChecked(True)
        grp = QButtonGroup(dlg); grp.addButton(rb_p); grp.addButton(rb_l)
        h2 = QHBoxLayout(); h2.addWidget(rb_p); h2.addWidget(rb_l); h2.addStretch()
        v.addLayout(h2)
        v.addStretch()

        # Botones
        bar = QHBoxLayout(); bar.setSpacing(8)
        btn_cancel = QPushButton("Cancelar")
        btn_cancel.setMinimumWidth(96); btn_cancel.setMinimumHeight(32)
        btn_cancel.setStyleSheet(
            "QPushButton { background:white; color:#485A6C;"
            " border:1px solid #D4D4D4; border-radius:4px; padding:6px 18px; }"
            "QPushButton:hover { background:#F1F5F9; }"
        )
        btn_cancel.clicked.connect(dlg.reject)
        btn_prev = QPushButton("👁  Vista previa")
        btn_prev.setMinimumHeight(32)
        btn_prev.setStyleSheet(
            "QPushButton { background:white; color:#273445;"
            " border:1px solid #D4D4D4; border-radius:4px;"
            " padding:6px 14px; font-weight:600; }"
            "QPushButton:hover { background:#FEF5EB;"
            " border-color:#F37329; color:#C0621A; }"
        )
        btn_ok = QPushButton("Exportar PDF…")
        btn_ok.setDefault(True); btn_ok.setMinimumHeight(32)
        btn_ok.setStyleSheet(
            "QPushButton { background:#F37329; color:white; border:none;"
            " border-radius:4px; padding:6px 18px; font-weight:600; }"
            "QPushButton:hover { background:#C0621A; }"
        )
        btn_ok.clicked.connect(dlg.accept)

        def _opts():
            return {
                'papel':  PAPEL_OPS[cmb.currentIndex()][1],
                'orient': 'landscape' if rb_l.isChecked() else 'portrait',
            }
        btn_prev.clicked.connect(lambda: self._vista_previa_pdf(_opts()))
        bar.addWidget(btn_cancel); bar.addStretch()
        bar.addWidget(btn_prev); bar.addWidget(btn_ok)
        v.addLayout(bar)

        if dlg.exec() != QDialog.Accepted:
            return None
        return _opts()

    def _vista_previa_pdf(self, opts: dict):
        """Genera un PDF temporal y lo muestra en un diálogo con QPdfView."""
        import os, tempfile
        try:
            from PySide6.QtPdf import QPdfDocument
            from PySide6.QtPdfWidgets import QPdfView
        except Exception as e:
            QMessageBox.warning(self, "Vista previa",
                                  f"No se pudo cargar el visor de PDF:\n{e}")
            return
        tmp_fd, tmp_pdf = tempfile.mkstemp(suffix='.pdf', prefix='curvaS_prev_')
        os.close(tmp_fd)
        try:
            self._render_pdf(tmp_pdf, paper=opts.get('papel', 'A4'),
                              orient=opts.get('orient', 'portrait'))
            if not os.path.exists(tmp_pdf) or os.path.getsize(tmp_pdf) < 500:
                QMessageBox.warning(self, "Vista previa",
                                      "El PDF generado está vacío.")
                return
            dlg = QDialog(self)
            dlg.setWindowTitle("Vista previa — Curva S")
            dlg.setWindowModality(Qt.ApplicationModal)
            dlg.resize(1100, 800)
            vv = QVBoxLayout(dlg)
            vv.setContentsMargins(8, 8, 8, 8); vv.setSpacing(6)
            tb = QHBoxLayout()
            btn_save = QPushButton("📄  Guardar PDF…")
            btn_save.setStyleSheet(
                f"QPushButton {{ background:white; color:{SLATE_700};"
                f" border:1px solid {SILVER_300}; border-radius:4px;"
                f" padding:6px 14px; }}"
                "QPushButton:hover { background:#FEF5EB;"
                " border-color:#F37329; color:#C0621A; }"
            )
            btn_save.clicked.connect(lambda: (
                self._exportar_pdf_a_archivo(opts), dlg.accept()
            ))
            btn_close = QPushButton("Cerrar")
            btn_close.setStyleSheet(btn_save.styleSheet())
            btn_close.clicked.connect(dlg.accept)
            tb.addStretch(); tb.addWidget(btn_save); tb.addWidget(btn_close)
            vv.addLayout(tb)
            doc = QPdfDocument(dlg)
            doc.load(tmp_pdf)
            pdfv = QPdfView(dlg)
            pdfv.setDocument(doc)
            try:
                pdfv.setPageMode(QPdfView.PageMode.MultiPage)
                pdfv.setZoomMode(QPdfView.ZoomMode.FitToWidth)
            except Exception:
                pass
            vv.addWidget(pdfv, stretch=1)
            dlg.exec()
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Vista previa",
                                  f"No se pudo generar la vista previa:\n{e}")
        finally:
            try: os.unlink(tmp_pdf)
            except Exception: pass

    def _render_pdf(self, path: str, paper: str = 'A4',
                      orient: str = 'portrait',
                      pie_offset: int = 0, pie_total: int | None = None):
        """Genera el PDF de Curva S con el papel/orientación dados.
        Header naranja + KPI cards + gráfico + tabla resumen + footer."""
        from PySide6.QtGui import QPdfWriter, QPageSize
        from PySide6.QtCore import QMarginsF as _QM
        from datetime import datetime as _dt
        # Mapa de papel → QPageSize.PageSizeId
        PAPER_MAP = {
            'A0':      QPageSize.A0,     'A1': QPageSize.A1,
            'A2':      QPageSize.A2,     'A3': QPageSize.A3,
            'A4':      QPageSize.A4,
            'Letter':  QPageSize.Letter, 'Tabloid': QPageSize.Tabloid,
        }
        writer = QPdfWriter(path)
        writer.setResolution(150)
        writer.setPageSize(QPageSize(PAPER_MAP.get(paper, QPageSize.A4)))
        layout = QPageLayout(writer.pageLayout())
        layout.setOrientation(
            QPageLayout.Landscape if (orient or '').lower() == 'landscape'
            else QPageLayout.Portrait
        )
        layout.setMargins(_QM(14, 14, 14, 14))
        writer.setPageLayout(layout)
        painter = QPainter(writer)
        painter.setRenderHint(QPainter.Antialiasing, True)
        # Inter como font base del PDF — consistente cross-platform.
        # La fuente se registra en main.py via QFontDatabase al arranque.
        painter.setFont(QFont('Inter', 10))
        try:
            dpi = writer.logicalDpiX()
            mm = lambda v: v * dpi / 25.4
            page_w = writer.pageLayout().paintRect(QPageLayout.Inch).width() * dpi
            page_h = writer.pageLayout().paintRect(QPageLayout.Inch).height() * dpi

            proy = self._cv._proy
            data = self._last_data or {}
            sym = self._moneda_simbolo()
            dec = get_decimales_ppto()
            unidad = data.get('unidad', 'Sem')

            # ── Header tripartite estilo Resumen Ejecutivo ────────────────
            # 3 zonas: IZQ empresa + subtítulo · CENTRO título + nombre del
            # proyecto wrapped · DER "Costo al:" + modalidad. Línea fina
            # slate-100 como separador. Espejo de `_draw_header` del PDF
            # principal.
            try:
                from utils.theme import accent_reportes
                _o_r, _od_r, _ = accent_reportes()
            except Exception:
                _o_r, _od_r = "#F37329", "#C0621A"
            try:
                from core.pdf_reports import get_formato, _clean_costo_al
                _fmt = get_formato()
            except Exception:
                _fmt = {}; _clean_costo_al = lambda v: str(v or '—')
            empresa  = _fmt.get('rep_empresa_nombre')    or 'ingePresupuestos'
            sub_emp  = _fmt.get('rep_empresa_subtitulo') or 'Sistema de Presupuestos de Obra Pública'
            color_marca_dk = QColor(_fmt.get('rep_color_marca_dk') or _od_r)

            from PySide6.QtGui import QFontMetrics
            # Anchos de las 3 zonas (espejo del `_draw_header` del PDF)
            left_w  = mm(60)
            right_w = mm(58)
            y_top   = mm(2)

            # IZQ: empresa + subtítulo
            painter.setPen(color_marca_dk)
            f = painter.font(); f.setPointSizeF(11); f.setBold(True); painter.setFont(f)
            painter.drawText(QRectF(mm(2), y_top, left_w, mm(5)),
                                Qt.AlignLeft | Qt.AlignVCenter, empresa)
            f.setPointSizeF(7); f.setBold(False); painter.setFont(f)
            painter.setPen(QColor(SLATE_300))
            painter.drawText(QRectF(mm(2), y_top + mm(5), left_w, mm(4)),
                                Qt.AlignLeft | Qt.AlignVCenter, sub_emp)

            # CENTRO: título reporte + nombre proyecto wrapped
            center_x = mm(2) + left_w + mm(4)
            center_w = page_w - mm(2) - left_w - right_w - mm(8)
            f.setPointSizeF(10); f.setBold(True); painter.setFont(f)
            painter.setPen(QColor("#1F2A38"))
            painter.drawText(QRectF(center_x, y_top, center_w, mm(5)),
                                Qt.AlignCenter | Qt.AlignVCenter,
                                'Curva S — Avance Financiero Acumulado')
            # Nombre proyecto centrado abajo del título, hasta 3 líneas
            f.setPointSizeF(7); f.setBold(False); f.setItalic(True); painter.setFont(f)
            painter.setPen(QColor(SLATE_500))
            nom = (proy.get('nombre') or '').strip()
            fm_nom = QFontMetrics(f)
            max_nom_h = mm(11)
            measured = fm_nom.boundingRect(
                QRectF(0, 0, center_w, max_nom_h).toRect(),
                Qt.AlignHCenter | Qt.AlignTop | Qt.TextWordWrap,
                nom,
            )
            nom_h = min(max_nom_h, max(mm(4), measured.height() + 1))
            painter.drawText(
                QRectF(center_x, y_top + mm(5), center_w, nom_h),
                Qt.AlignHCenter | Qt.AlignTop | Qt.TextWordWrap, nom,
            )
            f.setItalic(False); painter.setFont(f)

            # DER: "Costo al: …" + modalidad
            right_x = page_w - mm(2) - right_w
            painter.setPen(QColor(SLATE_700))
            f.setPointSizeF(8); f.setBold(True); painter.setFont(f)
            costo_str = f"Costo al: {_clean_costo_al(proy.get('costo_al'))}"
            painter.drawText(QRectF(right_x, y_top, right_w, mm(5)),
                                Qt.AlignRight | Qt.AlignVCenter, costo_str)
            f.setPointSizeF(7); f.setBold(False); f.setItalic(True); painter.setFont(f)
            painter.setPen(QColor(SLATE_300))
            painter.drawText(QRectF(right_x, y_top + mm(5), right_w, mm(4)),
                                Qt.AlignRight | Qt.AlignVCenter,
                                (proy.get('modalidad') or '').strip())
            f.setItalic(False); painter.setFont(f)

            # Línea separadora slate-100 al final del header
            header_bottom = max(y_top + mm(9), y_top + mm(5) + nom_h + mm(1))
            painter.setPen(QPen(QColor("#CBD5E1"), 0.5))
            painter.drawLine(QPointF(mm(2), header_bottom),
                                QPointF(page_w - mm(2), header_bottom))
            cur_y = header_bottom + mm(4)

            # ── KPI strip — 4 cards horizontales ───────────────────────
            # Estilo espejo del Resumen Ejecutivo: borde superior 2.5pt
            # color de marca en TODAS las cards (no solo la primera),
            # contenido a 3 niveles (LABEL uppercase + VALUE bold + SUB
            # itálica), borde slate-300 alrededor, sin esquinas redondeadas.
            kpi_h = mm(18)   # un poco más alto para la línea SUB
            kpi_count = 4
            gap = mm(2.5)
            kpi_w = (page_w - mm(4) - gap * (kpi_count - 1)) / kpi_count
            n_periods = data.get('n_periods', 0)
            total_gen = data.get('total_general', 0)
            pico_idx  = data.get('pico_idx', 0)
            per_pct_l = data.get('per_pct', [0])
            plazo = proy.get('plazo') or 0
            modalidad = (proy.get('modalidad') or 'Contrata').strip() or 'Contrata'
            kpis = [
                ("COSTO DIRECTO TOTAL",
                    fmt(total_gen, self._cv._moneda, dec),
                    "Acumulado del proyecto"),
                ("PLAZO",
                    f"{plazo} días" if plazo else "—",
                    modalidad),
                ("N° DE PERÍODOS",
                    f"{n_periods}",
                    f"{'semanas' if unidad.lower().startswith('sem') else 'meses'} de duración"),
                ("PERÍODO DE MAYOR AVANCE",
                    f"{unidad} {pico_idx + 1}" if total_gen > 0 else "—",
                    (f"{per_pct_l[pico_idx]:.1f}% del costo"
                     if total_gen > 0 else "")),
            ]
            for i, (lbl, val, sub) in enumerate(kpis):
                kx = mm(2) + i * (kpi_w + gap)
                # Sin marco, sin barra superior — KPIs puramente tipográficas,
                # totalmente limpias (decisión de Marco — minimalista total).
                # LABEL (uppercase, pequeña, slate-500, bold)
                painter.setPen(QColor(SLATE_500))
                f.setPointSizeF(6.5); f.setBold(True); painter.setFont(f)
                painter.drawText(QRectF(kx + mm(3), cur_y + mm(2.5),
                                          kpi_w - mm(6), mm(4)),
                                    Qt.AlignLeft | Qt.AlignVCenter, lbl)
                # VALUE (grande, bold, slate-900 — siempre dark)
                painter.setPen(QColor("#1F2A38"))
                f.setPointSizeF(13); f.setBold(True); painter.setFont(f)
                painter.drawText(QRectF(kx + mm(3), cur_y + mm(7),
                                          kpi_w - mm(6), mm(6.5)),
                                    Qt.AlignLeft | Qt.AlignVCenter, val)
                # SUB (pequeña, itálica, slate-500)
                if sub:
                    painter.setPen(QColor(SLATE_500))
                    f.setPointSizeF(6.5); f.setBold(False)
                    f.setItalic(True); painter.setFont(f)
                    painter.drawText(QRectF(kx + mm(3), cur_y + mm(13.5),
                                              kpi_w - mm(6), mm(4)),
                                        Qt.AlignLeft | Qt.AlignVCenter, sub)
                    f.setItalic(False); painter.setFont(f)
            cur_y += kpi_h + mm(4)

            # ── Chart (rasterizado de la escena, centrado) ─────────────
            # Calcular el target con el mismo aspecto que la escena para
            # que el gráfico quede tight (sin white-space lateral excesivo)
            # y centrado horizontalmente sobre la página.
            band_w = page_w - mm(4)
            band_h_chart = (page_h - cur_y - mm(6)) * 0.50
            scene_rect = self.scene.sceneRect()
            src_aspect = (scene_rect.width() / scene_rect.height()
                          if scene_rect.height() > 0 else 1.6)
            band_aspect = band_w / band_h_chart if band_h_chart > 0 else 1.6
            if band_aspect > src_aspect:
                # franja más ancha que el aspecto: fit por altura y centrar X
                content_h = band_h_chart
                content_w = content_h * src_aspect
                content_x = mm(2) + (band_w - content_w) / 2
                content_y = cur_y
            else:
                # franja más alta que el aspecto: fit por ancho y centrar Y
                content_w = band_w
                content_h = content_w / src_aspect
                content_x = mm(2)
                content_y = cur_y + (band_h_chart - content_h) / 2
            target = QRectF(content_x, content_y, content_w, content_h)
            self.scene.render(painter, target, scene_rect, Qt.KeepAspectRatio)
            cur_y = cur_y + band_h_chart + mm(4)

            # ── Tabla resumen ──────────────────────────────────────────
            tab_y = cur_y
            tab_h = page_h - tab_y - mm(6)
            self._pdf_render_tabla(painter, mm(2), tab_y, page_w - mm(4), tab_h)

            # ── Footer tripartito estilo Resumen Ejecutivo ───────────────
            # IZQ Cliente · CENTRO fecha · DER Página X de N. Línea slate-100
            # de separador arriba. Espejo de `_draw_footer` del PDF principal.
            painter.setPen(QPen(QColor("#CBD5E1"), 0.5))
            painter.drawLine(QPointF(mm(2), page_h - mm(6)),
                                QPointF(page_w - mm(2), page_h - mm(6)))
            painter.setPen(QColor(SLATE_300))
            f.setPointSizeF(7); f.setBold(False); painter.setFont(f)
            cliente_txt = (f"Cliente: {proy['cliente']}"
                              if proy.get('cliente') else '')
            painter.drawText(QRectF(mm(2), page_h - mm(5), page_w - mm(4), mm(5)),
                                Qt.AlignLeft | Qt.AlignVCenter, cliente_txt)
            painter.drawText(QRectF(0, page_h - mm(5), page_w, mm(5)),
                                Qt.AlignCenter, _dt.now().strftime("%d/%m/%Y"))
            # Página: continua si `pie_total` está set (Reporte Completo).
            if pie_total:
                pag_txt = f"Página {1 + (pie_offset or 0)} de {pie_total}"
            else:
                pag_txt = "Página 1 de 1"
            painter.drawText(QRectF(mm(2), page_h - mm(5), page_w - mm(4), mm(5)),
                                Qt.AlignRight | Qt.AlignVCenter, pag_txt)
        finally:
            painter.end()

    def _pdf_render_tabla(self, p, x, y, w, h):
        """Pinta la tabla resumen en el PDF — estilo limpio espejo de
        las tablas de Cronograma Valorizado/Adquisiciones:
          - Header blanco con borde inferior 1.5pt color de marca
          - Gridlines silver-300 0.6pt (horizontal + vertical)
          - Filas alternadas con bg muy sutil
          - Footer con bg silver suave + borde superior 1.5pt slate-700
          - SIN highlight del período pico (decisión de Marco)
        """
        data = self._last_data or {}
        n = data.get('n_periods', 0)
        if n == 0:
            return
        unidad = data.get('unidad', 'Sem')
        sym = self._moneda_simbolo()
        dec = get_decimales_ppto()

        # Color de marca dinámico (sobrio o naranja)
        try:
            from utils.theme import accent_reportes
            _o_t, _, _ = accent_reportes()
        except Exception:
            _o_t = "#F37329"

        # SLATE_900 no está en el scope del módulo — definir localmente
        # (mismo patrón que `_dibujar_cabecera_reporte_gantt`).
        SLATE_900 = "#1F2A38"

        # Columnas: # | Período | Avance | % Per | Acumulado | % Acum
        col_w = [w * c for c in (0.04, 0.18, 0.20, 0.12, 0.20, 0.12)]
        rest = w - sum(col_w)
        col_w[1] += rest
        col_x = []
        cx = x
        for c in col_w:
            col_x.append(cx); cx += c

        # Filas más altas (de mín 8pt → mín 13pt) para más respiro entre
        # períodos. El piso máximo subió también de 16pt → 20pt.
        row_h = h / max(n + 2, 12)
        row_h = max(min(row_h, 20), 13)

        # Gridline muy sutil para horizontales — sin verticales en cuerpo
        # (look "swiss/Bloomberg" minimalista: la alineación de columnas
        # se confía a tabular-nums + alignment, no a líneas).
        ROW_LINE = QColor("#E2E8F0")   # SLATE_100 — muy sutil

        # ── Header ────────────────────────────────────────────────────
        # Sin bg, sin verticales — solo el borde inferior color de marca
        # como divisor con el cuerpo.
        p.save()
        p.setPen(QColor(SLATE_900))
        f = p.font(); f.setPointSize(8); f.setBold(True); p.setFont(f)
        heads = ["#", "Período", f"Avance ({sym})", "% Período",
                  f"Acumulado ({sym})", "% Acumulado"]
        for c, (hx, hw, txt) in enumerate(zip(col_x, col_w, heads)):
            al = Qt.AlignCenter if c in (0, 3, 5) else (
                Qt.AlignRight if c in (2, 4) else Qt.AlignLeft)
            p.drawText(QRectF(hx + 4, y, hw - 8, row_h),
                          int(al) | Qt.AlignVCenter, txt)
        # Borde inferior 1.5pt color de marca
        p.setPen(QPen(QColor(_o_t), 1.5))
        p.drawLine(QPointF(x, y + row_h), QPointF(x + w, y + row_h))
        p.restore()

        # ── Filas ─────────────────────────────────────────────────────
        f = p.font(); f.setPointSize(8); f.setBold(False); p.setFont(f)
        cur_y = y + row_h
        for i in range(n):
            row_vals = [
                (str(i + 1), Qt.AlignCenter),
                (self._periodo_label(unidad, i).replace("\n", " · "), Qt.AlignLeft),
                (fmt(data['total_per'][i], self._cv._moneda, dec), Qt.AlignRight),
                (f"{data['per_pct'][i]:.2f}%", Qt.AlignRight),
                (fmt(data['acum'][i], self._cv._moneda, dec), Qt.AlignRight),
                (f"{data['acum_pct'][i]:.2f}%", Qt.AlignRight),
            ]
            p.setPen(QColor(SLATE_700))
            for c, (txt, al) in enumerate(row_vals):
                p.drawText(QRectF(col_x[c] + 4, cur_y, col_w[c] - 8, row_h),
                              int(al) | Qt.AlignVCenter, str(txt))
            # Horizontal divisor de fila — muy sutil
            p.setPen(QPen(ROW_LINE, 0.4))
            p.drawLine(QPointF(x, cur_y + row_h),
                          QPointF(x + w, cur_y + row_h))
            cur_y += row_h
            if cur_y + row_h > y + h - row_h:
                break

        # ── Footer total ──────────────────────────────────────────────
        # Borde superior 1.5pt slate-700 como divisor con datos, sin bg
        # ni verticales — coherente con el resto.
        p.setPen(QPen(QColor(SLATE_700), 1.5))
        p.drawLine(QPointF(x, cur_y), QPointF(x + w, cur_y))
        p.setPen(QColor(SLATE_900))
        f = p.font(); f.setBold(True); p.setFont(f)
        ftr_vals = [
            ("", Qt.AlignCenter),
            ("TOTAL", Qt.AlignLeft),
            (fmt(data['total_general'], self._cv._moneda, dec), Qt.AlignRight),
            ("100.00%", Qt.AlignRight),
            (fmt(data['total_general'], self._cv._moneda, dec), Qt.AlignRight),
            ("100.00%", Qt.AlignRight),
        ]
        for c, (txt, al) in enumerate(ftr_vals):
            p.drawText(QRectF(col_x[c] + 4, cur_y, col_w[c] - 8, row_h),
                          int(al) | Qt.AlignVCenter, str(txt))

    # ── Word (.docx) y ODT ─────────────────────────────────────────────
    def _scene_png_bytes(self, target_w_px: int = 1400) -> bytes:
        """Rasteriza la escena de la curva a PNG (ancho target_w_px).
        Mantiene el aspect ratio. Devuelve los bytes PNG."""
        from PySide6.QtCore import QBuffer, QIODevice, QSize
        from PySide6.QtGui import QImage
        sr = self.scene.sceneRect()
        if sr.width() <= 0:
            sr = QRectF(0, 0, 800, 400)
        ratio = sr.height() / sr.width() if sr.width() > 0 else 0.55
        w = int(max(600, target_w_px))
        h = int(w * ratio)
        img = QImage(QSize(w, h), QImage.Format_ARGB32)
        img.fill(QColor("white"))
        p = QPainter(img)
        p.setRenderHint(QPainter.Antialiasing, True)
        self.scene.render(p, QRectF(0, 0, w, h), sr)
        p.end()
        buf = QBuffer()
        buf.open(QIODevice.WriteOnly)
        img.save(buf, "PNG")
        return bytes(buf.data())

    def _exportar_word(self):
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os
        if not self._last_data or (self._last_data.get('total_general', 0) <= 0):
            QMessageBox.information(self, "Exportar Word",
                "No hay datos para exportar.")
            return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"curva_s_{self._cv.pid}.docx")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar Curva S a Word", sugerido, "Word (*.docx)"
        )
        if not path:
            return
        try:
            self._build_docx(path)
            self._save_dir(path)
            QMessageBox.information(self, "Exportar Word",
                                       f"Curva S exportada a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")

    def _build_docx(self, path: str):
        """Genera un .docx apaisado — espejo VISUAL del PDF de Curva S.

        Patrón "PDF visible":
        - Título 15pt accent_o, fecha derecha, línea fina accent_o como separador.
        - Nombre del proyecto 11pt bold slate-700 con wrap.
        - Subtítulo (cliente · ubicación · modalidad) 8.5pt slate-500.
        - KPI strip 4 columnas tipográfico puro (LABEL 7pt uppercase slate-500
          bold + VALUE 13pt bold slate-900 + SUB 7pt italic slate-500); SIN
          bordes ni fills (espejo del PDF).
        - Chart embebido centrado.
        - Tabla con header blanco + border-bottom 1.5pt accent, filas sin
          bordes (solo silver-100 thin como divisor horizontal), SIN highlight
          del período pico, TOTAL con border-top 1.5pt slate-700.
        """
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO
        from datetime import datetime as _dt

        try:
            from utils.theme import accent_reportes
            _ohex, _odhex, _ = accent_reportes()
        except Exception:
            _ohex, _odhex = "#F37329", "#C0621A"
        def _rgb(h: str) -> RGBColor:
            h = h.lstrip('#')
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
        _rgb_o, _rgb_od = _rgb(_ohex), _rgb(_odhex)
        _SLATE_900 = RGBColor(0x1F, 0x2A, 0x38)
        _SLATE_700 = RGBColor(0x27, 0x34, 0x45)
        _SLATE_500 = RGBColor(0x48, 0x5A, 0x6C)

        data   = self._last_data
        proy   = self._cv._proy
        sym    = self._moneda_simbolo()
        dec    = get_decimales_ppto()
        unidad = data['unidad']
        n      = data['n_periods']

        doc = Document()
        # Idioma es-PE para que el corrector ortográfico de Word/LibreOffice
        # NO marque las palabras como mal escritas. Y fuente Inter global.
        from core.word_reports import (
            _aplicar_idioma, _aplicar_fuente_global,
            _add_header_marca, _add_footer,
        )
        _aplicar_idioma(doc)
        _aplicar_fuente_global(doc)

        sec = doc.sections[0]
        # Portrait (Marco prefiere vertical — espejo del PDF de Curva S).
        sec.top_margin = Cm(1.2); sec.bottom_margin = Cm(1.2)
        sec.left_margin = Cm(1.2); sec.right_margin = Cm(1.2)

        # Encabezado tripartito empresa | título+proyecto | costo+modalidad
        # (espejo del PDF principal y de Resumen Ejecutivo / Especificaciones).
        _add_header_marca(doc, proy, 'Curva S — Avance Financiero Acumulado')
        # Pie tripartito Cliente | Fecha | Página X de N — mismo espejo.
        _add_footer(doc, proy)

        style = doc.styles['Normal']
        style.font.name = 'Inter'
        style.font.size = Pt(10)
        style.font.color.rgb = _SLATE_700

        def _cell_bg(cell, hex6):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex6.lstrip('#'))
            tcPr.append(shd)

        def _cell_no_borders(cell):
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                bd = OxmlElement(f'w:{side}')
                bd.set(qn('w:val'), 'nil')
                borders.append(bd)
            tcPr.append(borders)

        def _cell_border(cell, *, side='bottom', color='CBD5E1', sz=4):
            """Borde individual; sz en 1/8 pt (sz=12 ≈ 1.5pt)."""
            tcPr = cell._tc.get_or_add_tcPr()
            borders = tcPr.find(qn('w:tcBorders'))
            if borders is None:
                borders = OxmlElement('w:tcBorders')
                tcPr.append(borders)
            for s in ('top', 'left', 'bottom', 'right'):
                tag = qn(f'w:{s}')
                ex = borders.find(tag)
                if ex is None:
                    bd = OxmlElement(f'w:{s}')
                    bd.set(qn('w:val'), 'single' if s == side else 'nil')
                    if s == side:
                        bd.set(qn('w:sz'), str(sz))
                        bd.set(qn('w:color'), color.lstrip('#'))
                    borders.append(bd)
                else:
                    if s == side:
                        ex.set(qn('w:val'), 'single')
                        ex.set(qn('w:sz'), str(sz))
                        ex.set(qn('w:color'), color.lstrip('#'))

        def _para_bottom_border(para, *, sz=4, color='CBD5E1'):
            pPr = para._p.get_or_add_pPr()
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None:
                pBdr = OxmlElement('w:pBdr')
                pPr.append(pBdr)
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(sz))
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), color.lstrip('#'))
            pBdr.append(bottom)

        # ── h2 título del cuerpo (espejo del Resumen Ejecutivo / Specs) ──
        # El encabezado tripartito de la página (empresa | título+proyecto |
        # costo+modalidad) ya está en `_add_header_marca`. El cuerpo solo
        # repite el título como h2 — sin redundar el nombre del proyecto
        # ni la línea cliente/ubicación/modalidad.
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ph.paragraph_format.space_before = Pt(14)
        ph.paragraph_format.space_after = Pt(6)
        rh = ph.add_run("Curva S — Avance Financiero Acumulado")
        rh.font.size = Pt(13); rh.font.bold = True
        rh.font.color.rgb = _rgb_od

        # ── KPI strip tipográfico (4 columnas, SIN bordes ni fills) ───────
        kpis = [
            ("COSTO DIRECTO TOTAL",
                fmt(data['total_general'], self._cv._moneda, dec),
                "Acumulado del proyecto"),
            ("PLAZO",
                f"{proy.get('plazo') or 0} días",
                (proy.get('modalidad') or 'Contrata').strip() or 'Contrata'),
            ("N° DE PERÍODOS",
                f"{n}",
                f"{'semanas' if unidad.lower().startswith('sem') else 'meses'} de duración"),
            ("PERÍODO DE MAYOR AVANCE",
                f"{unidad} {data['pico_idx']+1}",
                f"{data['per_pct'][data['pico_idx']]:.1f}% del costo"),
        ]
        tbl_k = doc.add_table(rows=1, cols=4)
        tbl_k.autofit = False
        col_w = Cm(4.2)   # 4 col × 4.2cm = 16.8cm (cabe en A4 portrait útil 17cm)
        for i, (lbl, val, sub) in enumerate(kpis):
            cell = tbl_k.cell(0, i)
            cell.width = col_w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _cell_no_borders(cell)
            cell.text = ''
            # LABEL
            p_lbl = cell.paragraphs[0]
            p_lbl.paragraph_format.space_after = Pt(2)
            r_lbl = p_lbl.add_run(lbl)
            r_lbl.font.size = Pt(7); r_lbl.font.bold = True
            r_lbl.font.color.rgb = _SLATE_500
            # VALUE
            p_val = cell.add_paragraph()
            p_val.paragraph_format.space_after = Pt(2)
            r_val = p_val.add_run(val)
            r_val.font.size = Pt(13); r_val.font.bold = True
            r_val.font.color.rgb = _SLATE_900
            # SUB
            p_sub = cell.add_paragraph()
            r_sub = p_sub.add_run(sub)
            r_sub.font.size = Pt(7); r_sub.italic = True
            r_sub.font.color.rgb = _SLATE_500
        doc.add_paragraph()

        # ── Chart embebido ────────────────────────────────────────────────
        png_bytes = self._scene_png_bytes(target_w_px=1800)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(8)
        # 6.5" ≈ 16.5cm — cabe holgado en A4 portrait útil 17cm.
        p_img.add_run().add_picture(BytesIO(png_bytes), width=Inches(6.5))

        # ── Tabla resumen ─────────────────────────────────────────────────
        cols = 6
        tbl = doc.add_table(rows=1 + n + 1, cols=cols)
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header — blanco + border-bottom 1.5pt accent (sz=12)
        heads = ["#", "Período", f"Avance ({sym})", "% Período",
                  f"Acumulado ({sym})", "% Acumulado"]
        for c, txt in enumerate(heads):
            cell = tbl.rows[0].cells[c]
            cell.text = ''
            _cell_no_borders(cell)
            _cell_border(cell, side='bottom', color=_ohex, sz=12)
            pp = cell.paragraphs[0]
            r = pp.add_run(txt)
            r.font.bold = True; r.font.size = Pt(8)
            r.font.color.rgb = _SLATE_900
            if c in (0, 3, 5):
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c in (2, 4):
                pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Filas de datos — sin bordes, zebra muy sutil, sin highlight del pico
        for i in range(n):
            row = tbl.rows[i + 1]
            vals = [
                str(i + 1),
                self._periodo_label(unidad, i).replace("\n", " · "),
                fmt(data['total_per'][i], self._cv._moneda, dec),
                f"{data['per_pct'][i]:.2f}%",
                fmt(data['acum'][i], self._cv._moneda, dec),
                f"{data['acum_pct'][i]:.2f}%",
            ]
            zebra = (i % 2 == 1)
            for c, txt in enumerate(vals):
                cell = row.cells[c]
                cell.text = ''
                _cell_no_borders(cell)
                _cell_border(cell, side='bottom', color='E2E8F0', sz=2)
                if zebra:
                    _cell_bg(cell, 'FBFCFD')
                pp = cell.paragraphs[0]
                r = pp.add_run(str(txt))
                r.font.size = Pt(8); r.font.color.rgb = _SLATE_700
                if c in (0, 3, 5):
                    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif c in (2, 4):
                    pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Footer total — sin bg, border-top 1.5pt slate-700
        ftr = tbl.rows[n + 1]
        ftr_vals = ["", "TOTAL",
                      fmt(data['total_general'], self._cv._moneda, dec),
                      "100.00%",
                      fmt(data['total_general'], self._cv._moneda, dec),
                      "100.00%"]
        for c, txt in enumerate(ftr_vals):
            cell = ftr.cells[c]
            cell.text = ''
            _cell_no_borders(cell)
            _cell_border(cell, side='top', color='273445', sz=12)
            pp = cell.paragraphs[0]
            r = pp.add_run(txt)
            r.font.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = _SLATE_900
            if c in (0, 3, 5):
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c in (2, 4):
                pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.save(path)

    def _exportar_odt(self):
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os, subprocess, tempfile
        from core.soffice import find_soffice, mensaje_instalacion
        soffice = find_soffice()
        if not soffice:
            QMessageBox.warning(self, "Exportar ODT", mensaje_instalacion())
            return
        if not self._last_data or (self._last_data.get('total_general', 0) <= 0):
            QMessageBox.information(self, "Exportar ODT", "No hay datos para exportar.")
            return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"curva_s_{self._cv.pid}.odt")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar Curva S a ODT", sugerido, "OpenDocument Text (*.odt)"
        )
        if not path:
            return
        tmp_fd, tmp_docx = tempfile.mkstemp(suffix='.docx', prefix='curvaS_odt_')
        os.close(tmp_fd)
        try:
            self._build_docx(tmp_docx)
            out_dir = os.path.dirname(os.path.abspath(path))
            subprocess.run(
                [soffice, '--headless', '--convert-to', 'odt',
                 '--outdir', out_dir, tmp_docx],
                check=True, capture_output=True, timeout=60,
            )
            generated = os.path.join(
                out_dir, os.path.splitext(os.path.basename(tmp_docx))[0] + '.odt'
            )
            if os.path.exists(generated) and generated != path:
                os.replace(generated, path)
            if not os.path.exists(path):
                raise RuntimeError("LibreOffice no generó el archivo .odt")
            self._save_dir(path)
            QMessageBox.information(self, "Exportar ODT",
                                       f"Curva S exportada a:\n{path}")
        except subprocess.CalledProcessError as e:
            QMessageBox.warning(self, "Exportar ODT",
                f"LibreOffice falló al convertir:\n{e.stderr.decode()[:300]}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")
        finally:
            try: os.unlink(tmp_docx)
            except Exception: pass

    def _exportar_excel(self):
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os
        if not self._last_data or (self._last_data.get('total_general', 0) <= 0):
            QMessageBox.information(self, "Exportar Excel",
                "No hay datos para exportar.")
            return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"curva_s_{self._cv.pid}.xlsx")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar Curva S a Excel", sugerido, "Excel (*.xlsx)"
        )
        if not path:
            return
        try:
            self._build_xlsx(path)
            self._save_dir(path)
            QMessageBox.information(self, "Exportar Excel",
                                       f"Curva S exportada a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")

    def _build_xlsx(self, path: str):
        """Construye .xlsx — espejo VISUAL del PDF de Curva S.

        Mismo patrón "PDF visible" que Valorizado/Adquisiciones:
        - Header tripartito vía `_xlsx_header_pdf_style`.
        - h2 izquierdo "Curva S — Avance Financiero Acumulado" 13pt accent_od.
        - Subtítulo con Costo Directo · Plazo · N° Períodos · Mayor Avance.
        - Sin gridlines, sin bordes en datos; header con `border-bottom` 1.5pt
          accent (igual que el PDF), TOTAL con `border-top` 1.5pt slate-700.
        - Zebra sutil `#FBFCFD`; SIN highlight del pico (espejo del PDF).
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from core.exporter import _xlsx_header_pdf_style
        from utils.theme import accent_reportes

        wb = Workbook(); ws = wb.active; ws.title = "Curva S"
        proy = self._cv._proy
        data = self._last_data or {}
        if not data:
            wb.save(path)
            return
        unidad   = data.get('unidad', 'Sem')
        sym      = self._moneda_simbolo()
        dec      = get_decimales_ppto()

        # Paleta
        o_hex, od_hex, _ = accent_reportes()
        ACC_O    = o_hex.lstrip('#').upper()
        ACC_OD   = od_hex.lstrip('#').upper()
        TXT_DK   = '1F2A38'
        TXT_700  = '273445'
        TXT_500  = '485A6C'
        HDR_BG   = 'F1F5F9'
        ALT_BG   = 'FBFCFD'

        side_acc    = Side(style='medium', color=ACC_O)    # ~1.5pt accent
        side_slate7 = Side(style='medium', color=TXT_700)  # ~1.5pt slate-700
        border_none  = Border()
        border_head  = Border(bottom=side_acc)
        border_total = Border(top=side_slate7)

        align_left   = Alignment(horizontal='left',   vertical='center',
                                  wrap_text=True)
        align_right  = Alignment(horizontal='right',  vertical='center')
        align_center = Alignment(horizontal='center', vertical='center',
                                  wrap_text=True)

        N_COLS = 6  # #, Período, Avance, % Período, Acumulado, % Acumulado

        # Anchos ANTES del header (wrap calc usa width)
        for c, w in enumerate([6, 28, 18, 12, 18, 14], 1):
            ws.column_dimensions[get_column_letter(c)].width = w

        ws.sheet_view.showGridLines = False
        ws.print_options.gridLines = False
        ws.print_options.gridLinesSet = True

        # ── Fila 1-4: header tripartito ──────────────────────────────────
        row = _xlsx_header_pdf_style(ws, proy, 'Curva S', N_COLS)

        # ── h2 izquierdo + subtítulo ─────────────────────────────────────
        label_plural = 'semanas' if unidad.lower().startswith('sem') else 'meses'
        plazo        = int(proy.get('plazo') or 0)
        n_periods    = int(data.get('n_periods', 0) or 0)
        pico_idx     = int(data.get('pico_idx', 0) or 0)
        total_gen    = data.get('total_general', 0) or 0
        per_pct_l    = data.get('per_pct', []) or []
        pico_pct     = per_pct_l[pico_idx] if 0 <= pico_idx < len(per_pct_l) else 0.0

        ws.merge_cells(start_row=row, start_column=1,
                          end_row=row, end_column=N_COLS)
        c_h2 = ws.cell(row=row, column=1,
                          value="Curva S — Avance Financiero Acumulado")
        c_h2.font      = Font(name='Inter', bold=True, size=13, color=ACC_OD)
        c_h2.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[row].height = 20
        row += 1

        from utils.formatting import fmt as _fmt_money
        sub_partes = []
        if total_gen:
            sub_partes.append(f"Costo Directo: {_fmt_money(total_gen, self._cv._moneda, dec)}")
        sub_partes.append(f"Plazo: {plazo} días")
        sub_partes.append(f"Períodos: {n_periods} {label_plural}")
        if total_gen > 0 and n_periods:
            sub_partes.append(f"Mayor avance: {unidad} {pico_idx + 1} ({pico_pct:.1f}%)")
        ws.merge_cells(start_row=row, start_column=1,
                          end_row=row, end_column=N_COLS)
        c_sub = ws.cell(row=row, column=1, value='  ·  '.join(sub_partes))
        c_sub.font      = Font(name='Inter', size=9, color=TXT_500)
        c_sub.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[row].height = 14
        row += 1

        ws.row_dimensions[row].height = 6
        row += 1

        # ── Header de tabla ──────────────────────────────────────────────
        heads = ["#", "Período", f"Avance ({sym})", "% Período",
                  f"Acumulado ({sym})", "% Acumulado"]
        head_fill = PatternFill('solid', fgColor=HDR_BG)
        head_font = Font(name='Inter', bold=True, size=10, color=TXT_DK)
        for c, h in enumerate(heads, 1):
            cell = ws.cell(row=row, column=c, value=h)
            cell.font      = head_font
            cell.fill      = head_fill
            cell.border    = border_head
            cell.alignment = align_center
        head_row = row
        ws.row_dimensions[head_row].height = 24
        row += 1

        # ── Filas de datos ───────────────────────────────────────────────
        data_font = Font(name='Inter', size=10, color=TXT_DK)
        n = int(data.get('n_periods', 0) or 0)
        for i in range(n):
            zebra      = (i % 2 == 1)
            zebra_fill = PatternFill('solid', fgColor=ALT_BG) if zebra else None
            vals = [
                (i + 1,                                            align_center, None),
                (self._periodo_label(unidad, i).replace("\n", " · "), align_left, None),
                (data['total_per'][i],                             align_right,  f'#,##0.{"0" * dec}'),
                (data['per_pct'][i] / 100,                         align_right,  '0.00%'),
                (data['acum'][i],                                  align_right,  f'#,##0.{"0" * dec}'),
                (data['acum_pct'][i] / 100,                        align_right,  '0.00%'),
            ]
            for c, (val, al, nf) in enumerate(vals, 1):
                cell = ws.cell(row=row, column=c, value=val)
                cell.font      = data_font
                cell.alignment = al
                cell.border    = border_none
                if nf:
                    cell.number_format = nf
                if zebra_fill:
                    cell.fill = zebra_fill
            row += 1

        # ── Footer TOTAL ─────────────────────────────────────────────────
        ftr_font = Font(name='Inter', bold=True, size=10, color=TXT_DK)
        ftr_vals = [
            ('',                  align_center, None),
            ('TOTAL',             align_right,  None),
            (data['total_general'], align_right, f'#,##0.{"0" * dec}'),
            (1.0,                 align_right,  '0.00%'),
            (data['total_general'], align_right, f'#,##0.{"0" * dec}'),
            (1.0,                 align_right,  '0.00%'),
        ]
        for c, (val, al, nf) in enumerate(ftr_vals, 1):
            cell = ws.cell(row=row, column=c, value=val)
            cell.font      = ftr_font
            cell.alignment = al
            cell.border    = border_total
            if nf:
                cell.number_format = nf
        ws.row_dimensions[row].height = 20

        # Freeze panes y print titles
        ws.freeze_panes = f'A{head_row + 1}'
        ws.print_title_rows = f'1:{head_row}'

        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.print_options.horizontalCentered = True
        wb.save(path)

    def _exportar_ods(self):
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os, subprocess, tempfile
        from core.soffice import find_soffice, mensaje_instalacion
        soffice = find_soffice()
        if not soffice:
            QMessageBox.warning(self, "Exportar ODS", mensaje_instalacion())
            return
        if not self._last_data or (self._last_data.get('total_general', 0) <= 0):
            QMessageBox.information(self, "Exportar ODS", "No hay datos para exportar.")
            return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"curva_s_{self._cv.pid}.ods")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar Curva S a ODS", sugerido, "OpenDocument Spreadsheet (*.ods)"
        )
        if not path:
            return
        tmp_fd, tmp_xlsx = tempfile.mkstemp(suffix='.xlsx', prefix='curvaS_ods_')
        os.close(tmp_fd)
        try:
            self._build_xlsx(tmp_xlsx)
            out_dir = os.path.dirname(os.path.abspath(path))
            subprocess.run(
                [soffice, '--headless', '--convert-to', 'ods',
                 '--outdir', out_dir, tmp_xlsx],
                check=True, capture_output=True, timeout=60,
            )
            generated = os.path.join(
                out_dir, os.path.splitext(os.path.basename(tmp_xlsx))[0] + '.ods'
            )
            if os.path.exists(generated) and generated != path:
                os.replace(generated, path)
            if not os.path.exists(path):
                raise RuntimeError("LibreOffice no generó el archivo .ods")
            self._save_dir(path)
            QMessageBox.information(self, "Exportar ODS",
                                       f"Curva S exportada a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")
        finally:
            try: os.unlink(tmp_xlsx)
            except Exception: pass


# ════════════════════════════════════════════════════════════════════════════
# 4. ADQUISICIÓN DE INSUMOS
# ════════════════════════════════════════════════════════════════════════════

class _PeriodGroupHeader(QHeaderView):
    """Cabecera horizontal de 2 niveles (estilo Delphin/MS Project) para el
    panel derecho: el PERÍODO (Mes/Sem N + fecha opcional) abarca sus 2
    sub-columnas; debajo van las sub-etiquetas (p.ej. Cantidad | Valoriz.).
    La última sección (Total) abarca ambos niveles.

    Pinta a mano (paintEvent) porque QHeaderView no soporta secciones
    combinadas; usa sectionViewportPosition/sectionSize → respeta el scroll
    horizontal automáticamente."""

    def __init__(self, sub_labels, total_label='Total', parent=None):
        super().__init__(Qt.Horizontal, parent)
        self._periodos = []            # labels de período (pueden tener '\n')
        self._subs = sub_labels        # ('Cant.', 'Valoriz.') / ('Metrado','Valoriz.')
        self._total_label = total_label
        self._h = 36
        self.setSectionsClickable(False)
        self.setHighlightSections(False)

    def set_periodos(self, periodos, height):
        self._periodos = list(periodos)
        self._h = int(height)
        self.updateGeometry()
        vp = self.viewport()
        if vp:
            vp.update()

    def sizeHint(self):
        s = super().sizeHint()
        return QSize(s.width(), self._h)

    def paintEvent(self, ev):
        from PySide6.QtGui import QPainter, QPen
        from PySide6.QtCore import QRect
        painter = QPainter(self.viewport())
        painter.setRenderHint(QPainter.TextAntialiasing, True)
        h = self.height()
        sub_h = 16
        top_h = max(12, h - sub_h)
        bg   = QColor(SLATE_500)
        fg   = QColor('white')
        line = QColor('#3D4D63')
        painter.fillRect(self.viewport().rect(), bg)
        f_top = QFont('Inter'); f_top.setPointSize(8); f_top.setBold(True)
        f_sub = QFont('Inter'); f_sub.setPointSize(7); f_sub.setBold(True)

        for i, plab in enumerate(self._periodos):
            c0 = 2 * i
            x0 = self.sectionViewportPosition(c0)
            w  = self.sectionSize(c0) + self.sectionSize(c0 + 1)
            # Nivel superior: período (abarca el par)
            painter.setPen(QPen(fg)); painter.setFont(f_top)
            painter.drawText(QRect(x0 + 1, 0, w - 2, top_h),
                              Qt.AlignCenter | Qt.TextWordWrap, plab)
            # Separador horizontal entre niveles
            painter.setPen(QPen(line))
            painter.drawLine(x0, top_h, x0 + w, top_h)
            # Nivel inferior: sub-etiquetas
            painter.setFont(f_sub); painter.setPen(QPen(fg))
            for j in range(2):
                xs = self.sectionViewportPosition(c0 + j)
                ws = self.sectionSize(c0 + j)
                painter.drawText(QRect(xs, top_h, ws, sub_h),
                                  Qt.AlignCenter, self._subs[j])
            # Divisor vertical al inicio del período
            painter.setPen(QPen(line))
            painter.drawLine(x0, 0, x0, h)

        # Sección Total (abarca ambos niveles)
        ti = 2 * len(self._periodos)
        xt = self.sectionViewportPosition(ti)
        wt = self.sectionSize(ti)
        painter.setPen(QPen(line)); painter.drawLine(xt, 0, xt, h)
        painter.setPen(QPen(fg)); painter.setFont(f_top)
        painter.drawText(QRect(xt, 0, wt, h), Qt.AlignCenter, self._total_label)
        # Borde inferior
        painter.setPen(QPen(line))
        painter.drawLine(0, h - 1, self.viewport().width(), h - 1)


class InsumosWidget(QWidget):
    """Cronograma de Adquisición de Insumos — espejo visual del Valorizado.
    Columnas izq (frozen): Tipo · Código · Descripción · Und · Precio · % Total.
    Columnas der (scroll): [períodos] · Total. Footer sticky con TOTAL ·
    % PERÍODO · ACUMULADO · % ACUMULADO."""

    def __init__(self, parent: CronogramaView):
        super().__init__(parent)
        self._cv = parent
        self._period_days = 30  # default Mensual (Marco prefiere mes al abrir)
        self._show_fechas = True
        self._last_n_periods = 0
        self._last_total_w = 110
        self._col_selected = None
        self._build_ui()

    def _build_ui(self):
        vl = QVBoxLayout(self)
        vl.setContentsMargins(8, 8, 8, 8)
        vl.setSpacing(6)

        # (KPI strip removido por pedido de Marco — no se veía consistente.)

        # ── Toolbar (período + mostrar fechas + export) ──────────────────
        tb = QFrame()
        tb.setStyleSheet("background:transparent; border:none;")
        tb_hl = QHBoxLayout(tb)
        tb_hl.setContentsMargins(0, 0, 0, 0); tb_hl.setSpacing(8)
        lbl = QLabel("Período:")
        lbl.setStyleSheet(f"color:{SLATE_500}; font-size:11px;")
        tb_hl.addWidget(lbl)
        cmb = QComboBox()
        cmb.addItems(["Semanal", "Mensual"])
        cmb.setCurrentIndex(1)  # default Mensual (espejo de _period_days=30)
        cmb.currentIndexChanged.connect(self._on_periodo)
        cmb.setStyleSheet(
            "QComboBox { min-height:0; padding:2px 8px; font-size:11px;"
            f" border:1px solid {SILVER_300}; border-radius:4px; }}"
        )
        tb_hl.addWidget(cmb)

        from PySide6.QtWidgets import QCheckBox as _QChk
        chk_fechas = _QChk("Mostrar fechas")
        chk_fechas.setChecked(self._show_fechas)
        chk_fechas.setStyleSheet(f"QCheckBox {{ color:{SLATE_500}; font-size:11px; padding:0 8px; }}")
        chk_fechas.toggled.connect(self._on_toggle_fechas)
        tb_hl.addWidget(chk_fechas)
        tb_hl.addStretch()

        from utils.tooltip import set_tooltip
        btn_style = (
            f"QPushButton {{ background:white; border:1px solid {SILVER_300};"
            f" border-radius:4px; font-size:11px; padding:3px 10px;"
            f" min-height:0; color:{SLATE_500}; }}"
            f"QPushButton:hover {{ background:{SILVER_100}; color:{SLATE_700}; }}"
        )
        btn_pdf = QPushButton("📄 PDF"); btn_pdf.setStyleSheet(btn_style)
        btn_pdf.setCursor(Qt.PointingHandCursor)
        btn_pdf.clicked.connect(self._exportar_pdf)
        set_tooltip(btn_pdf, "Exportar adquisiciones a PDF / imagen")
        tb_hl.addWidget(btn_pdf)

        btn_xls = QPushButton("📊 Excel"); btn_xls.setStyleSheet(btn_style)
        btn_xls.setCursor(Qt.PointingHandCursor)
        btn_xls.clicked.connect(self._exportar_excel)
        set_tooltip(btn_xls, "Exportar a Excel (.xlsx)")
        tb_hl.addWidget(btn_xls)

        btn_ods = QPushButton("📑 ODS"); btn_ods.setStyleSheet(btn_style)
        btn_ods.setCursor(Qt.PointingHandCursor)
        btn_ods.clicked.connect(self._exportar_ods)
        set_tooltip(btn_ods, "Exportar a ODS — requiere LibreOffice")
        tb_hl.addWidget(btn_ods)

        vl.addWidget(tb)

        # ── Splitter horizontal (izq frozen + der scroll) ────────────────
        self.split = QSplitter(Qt.Horizontal)
        self.split.setHandleWidth(6)
        self.split.setStyleSheet(
            "QSplitter::handle { background:#CBD5E1; }"
            "QSplitter::handle:horizontal { width:6px; margin:0; }"
            "QSplitter::handle:hover { background:#F37329; }"
        )
        self.split.setChildrenCollapsible(False)

        def _make_panel(is_left):
            fr = QFrame(); v = QVBoxLayout(fr)
            v.setContentsMargins(0, 0, 0, 0); v.setSpacing(0)
            tbl = QTableWidget(0, 0)
            tbl.verticalHeader().setVisible(False)
            tbl.setEditTriggers(QAbstractItemView.NoEditTriggers)
            tbl.setSelectionBehavior(QAbstractItemView.SelectRows)
            tbl.setShowGrid(True); tbl.setAlternatingRowColors(False)
            tbl.setStyleSheet(f"""
                QTableWidget {{ border:1px solid {SILVER_300};
                                border-bottom:none; font-size:11px;
                                gridline-color:#E0E5EC; background:white; }}
                QTableWidget::item {{ padding:3px 8px; }}
                QHeaderView::section {{
                    background:{SLATE_500}; color:white; font-size:10px;
                    font-weight:700; padding:4px 6px; border:none;
                }}
            """)
            if is_left:
                tbl.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                tbl.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            v.addWidget(tbl, stretch=1)
            ftr = QTableWidget(0, 0)
            ftr.verticalHeader().setVisible(False)
            ftr.horizontalHeader().setVisible(False)
            ftr.setEditTriggers(QAbstractItemView.NoEditTriggers)
            ftr.setSelectionMode(QAbstractItemView.NoSelection)
            ftr.setFocusPolicy(Qt.NoFocus)
            ftr.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            ftr.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            ftr.setShowGrid(True)
            ftr.setStyleSheet(f"""
                QTableWidget {{ border:1px solid {SILVER_300};
                                border-top:2px solid {SLATE_500};
                                font-size:11px; background:white;
                                gridline-color:#C5D1E0; }}
                QTableWidget::item {{ padding:4px 8px; }}
            """)
            v.addWidget(ftr)
            return fr, tbl, ftr

        left_fr, self.tbl_l, self.tbl_lf = _make_panel(True)
        right_fr, self.tbl_r, self.tbl_rf = _make_panel(False)

        # Cabecera de 2 niveles: período (Mes/Sem + fecha) abarcando sus 2
        # sub-columnas (Cantidad | Valorización) — estilo Delphin Express.
        self._period_header = _PeriodGroupHeader(('Cantidad', 'Valoriz.'))
        self.tbl_r.setHorizontalHeader(self._period_header)

        self._bg_delegate = _BgFillDelegate(self)
        for tbl_ in (self.tbl_l, self.tbl_r, self.tbl_lf, self.tbl_rf):
            tbl_.setItemDelegate(self._bg_delegate)

        right_pad_qss = "QTableWidget::item { padding:3px 8px; }"
        self.tbl_r.setStyleSheet(self.tbl_r.styleSheet() + right_pad_qss)
        self.tbl_rf.setStyleSheet(self.tbl_rf.styleSheet() + right_pad_qss)

        self.tbl_r.horizontalHeader().setSectionsClickable(True)
        self.tbl_r.horizontalHeader().sectionClicked.connect(
            self._on_period_header_clicked)
        self.tbl_r.cellClicked.connect(self._on_cell_clicked_clear_col)
        self.tbl_l.cellClicked.connect(self._on_cell_clicked_clear_col)

        # ESC global
        from PySide6.QtWidgets import QApplication
        QApplication.instance().installEventFilter(self)

        self.split.addWidget(left_fr); self.split.addWidget(right_fr)
        self.split.setStretchFactor(0, 0); self.split.setStretchFactor(1, 1)
        self.split.setSizes([50 + 80 + 240 + 40 + 90 + 60, 800])
        vl.addWidget(self.split, stretch=1)

        # Sync scrolls
        self.tbl_r.verticalScrollBar().valueChanged.connect(
            self.tbl_l.verticalScrollBar().setValue)
        self.tbl_l.verticalScrollBar().valueChanged.connect(
            self.tbl_r.verticalScrollBar().setValue)
        self.tbl_r.horizontalScrollBar().valueChanged.connect(
            self.tbl_rf.horizontalScrollBar().setValue)

    # ─── Eventos comunes ──────────────────────────────────────────────────

    def eventFilter(self, obj, event):
        from PySide6.QtCore import QEvent
        from PySide6.QtWidgets import QApplication
        if event.type() == QEvent.KeyPress and event.key() == Qt.Key_Escape:
            fw = QApplication.focusWidget()
            if fw is not None and (fw is self or self.isAncestorOf(fw)):
                self._clear_column_selection()
                return True
        return super().eventFilter(obj, event)

    def _clear_column_selection(self):
        for tbl_ in (self.tbl_l, self.tbl_r, self.tbl_rf):
            tbl_.clearSelection()
        self._col_selected = None

    def _on_cell_clicked_clear_col(self, *_):
        if self._col_selected is None:
            return
        self.tbl_rf.clearSelection()
        self._col_selected = None

    def _on_period_header_clicked(self, col: int):
        if col < 0 or col >= self.tbl_r.columnCount():
            return
        from PySide6.QtCore import QItemSelection, QItemSelectionModel
        v_r = self.tbl_r.verticalScrollBar().value()
        v_l = self.tbl_l.verticalScrollBar().value()
        def _sel(tbl, c):
            n = tbl.rowCount()
            if n <= 0:
                return
            mdl = tbl.model()
            sel = QItemSelection(mdl.index(0, c), mdl.index(n - 1, c))
            tbl.selectionModel().select(sel, QItemSelectionModel.ClearAndSelect)
        _sel(self.tbl_r, col); _sel(self.tbl_rf, col)
        self._col_selected = col
        self.tbl_r.verticalScrollBar().setValue(v_r)
        self.tbl_l.verticalScrollBar().setValue(v_l)

    def _on_periodo(self, idx: int):
        self._period_days = 7 if idx == 0 else 30
        self.cargar()

    def _on_toggle_fechas(self, checked: bool):
        self._show_fechas = bool(checked)
        self.cargar()

    def _project_start(self):
        for key in ('fecha_inicio', 'costo_al'):
            fi = (self._cv._proy.get(key) or '').strip()
            if fi:
                try:
                    return datetime.strptime(fi, '%Y-%m-%d')
                except Exception:
                    pass
        return None

    def _periodo_label(self, unidad: str, i: int) -> str:
        base = f"{unidad} {i + 1}"
        if not self._show_fechas:
            return base
        f_ini = self._project_start()
        if f_ini is None:
            return base
        MESES = ["Ene", "Feb", "Mar", "Abr", "May", "Jun",
                  "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
        s_day = i * self._period_days + 1
        e_day = (i + 1) * self._period_days
        d_i = f_ini + timedelta(days=s_day - 1)
        d_f = f_ini + timedelta(days=e_day - 1)
        if d_i.month == d_f.month and d_i.year == d_f.year:
            rng = f"{d_i.day:02d}–{d_f.day:02d} {MESES[d_i.month - 1]}"
        elif d_i.year == d_f.year:
            rng = (f"{d_i.day:02d} {MESES[d_i.month - 1]} – "
                    f"{d_f.day:02d} {MESES[d_f.month - 1]}")
        else:
            rng = (f"{d_i.day:02d}/{d_i.month:02d}/{str(d_i.year)[-2:]} – "
                    f"{d_f.day:02d}/{d_f.month:02d}/{str(d_f.year)[-2:]}")
        return f"{base}\n{rng}"

    def _apply_period_widths(self):
        n = getattr(self, '_last_n_subcols', 0); total_w = self._last_total_w
        if n <= 0:
            return
        vp = max(200, self.tbl_r.viewport().width())
        avail = vp - total_w - 4
        # Mínimo por sub-columna ampliado dinámicamente (calculado en
        # `cargar()`) si TOTAL/ACUMULADO traen montos grandes → no se recortan.
        min_per = getattr(self, '_min_per_w', 78)
        per_w = max(min_per, avail // n)
        for c in range(n):
            self.tbl_r.setColumnWidth(c, per_w)
            self.tbl_rf.setColumnWidth(c, per_w)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self._last_n_periods:
            QTimer.singleShot(0, self._apply_period_widths)

    # ─── Carga de datos ───────────────────────────────────────────────────

    def cargar(self):
        if not hasattr(self._cv, '_partidas'):
            return
        partidas = self._cv._partidas
        tasks = self._cv._tasks
        cmap = self._cv._cron_map
        plazo = self._cv._proy.get('plazo') or 0
        max_ef = max((t['EF'] for t in tasks.values() if t['EF'] > 0),
                      default=plazo)
        n_dias = max(max_ef, plazo, self._period_days)
        n_periods = max(1, (n_dias + self._period_days - 1) // self._period_days)

        unidad = "Sem" if self._period_days == 7 else "Mes"
        moneda = self._cv._moneda

        # Cargar items vía `get_acu_items` para incluir overhead `%MO`/`%MAT`.
        # Distribución proporcional al `parcial_wysiwyg(metrado, PU)` de cada
        # partida (= su contribución al CD) → sum(montos) coincide con CD.
        conn = get_db()
        recursos = {}
        for p in partidas:
            if p['es_titulo']:
                continue
            cd = cmap.get(p['id'], {})
            ini = cd.get('inicio_dia', 1) or 1
            dur = cd.get('duracion', 0) or 0
            t = tasks.get(p['id'], {})
            ini = t.get('ES', ini)
            metrado = p['metrado'] or 0
            pu      = p['precio_unitario'] or 0
            items, _ = get_acu_items(conn, p['id'])
            partida_total = parcial_wysiwyg(metrado, pu) if (metrado and pu) else 0.0
            sum_p = sum((it.get('parcial') or 0) for it in items)
            for it in items:
                if sum_p > 0:
                    ratio = (it.get('parcial') or 0) / sum_p
                    costo = partida_total * ratio
                else:
                    costo = 0.0
                segs_ = cd.get('segmentos', '') or ''
                weekly = distribuir_periodos(
                    segs_, ini, dur, costo, n_periods, self._period_days
                )
                # Cantidad del insumo en esta partida, repartida por el mismo
                # cronograma. Los items de overhead (%MO/%MAT) no tienen
                # cantidad física → quedan en 0.
                und_i = it.get('unidad') or ''
                qty_tot = (it.get('cantidad') or 0) * metrado
                if qty_tot and not und_i.startswith('%'):
                    weekly_qty = distribuir_periodos(
                        segs_, ini, dur, qty_tot, n_periods, self._period_days)
                else:
                    weekly_qty = [0.0] * n_periods
                rid = it['recurso_id']
                if rid not in recursos:
                    recursos[rid] = {
                        'codigo': it.get('codigo')      or '',
                        'desc':   it.get('descripcion') or '',
                        'tipo':   it.get('tipo')        or 'MAT',
                        'unidad': it.get('unidad')      or '',
                        'precio': it.get('precio')      or 0,
                        'weekly': [0.0] * n_periods,
                        'weekly_qty': [0.0] * n_periods,
                    }
                precio_i = it.get('precio') or 0
                if precio_i > recursos[rid]['precio']:
                    recursos[rid]['precio'] = precio_i
                for i, v in enumerate(weekly):
                    recursos[rid]['weekly'][i] += v
                for i, v in enumerate(weekly_qty):
                    recursos[rid]['weekly_qty'][i] += v
        conn.close()

        # MO ordenada por jerarquía (Capataz<Operario<Oficial<Peón), no
        # alfabético — igual que el panel ACU y los reportes de insumos.
        from core.database import _orden_mo
        orden_tipo = {'MO': 0, 'MAT': 1, 'EQ': 2, 'SC': 3}
        sorted_rec = sorted(
            recursos.values(),
            key=lambda r: (orden_tipo.get(r['tipo'], 9),
                           _orden_mo(r['desc']) if r['tipo'] == 'MO' else 0,
                           r['desc']))

        # Total general
        total_general = sum(sum(r['weekly']) for r in sorted_rec)

        # Headers — dos sub-columnas por período: Cantidad y Valorización.
        # Panel izquierdo: Cantidad y Precio = TOTALES del insumo en el proyecto.
        left_headers = ["Tipo", "Código", "Descripción", "Und", "Cantidad", "Precio"]
        # Texto de header del MODELO (lo lee el Excel); en pantalla el header
        # personalizado de 2 niveles dibuja el período combinado.
        period_labels = [self._periodo_label(unidad, i) for i in range(n_periods)]
        right_headers = []
        for plab in period_labels:
            right_headers.append(f"{plab}\nCant.")
            right_headers.append(f"{plab}\nValoriz.")
        right_headers.append("Total")
        N_RIGHT = len(right_headers)
        COL_TOTAL_R = N_RIGHT - 1
        n_sub = 2 * n_periods

        for t_ in (self.tbl_l, self.tbl_r, self.tbl_lf, self.tbl_rf):
            t_.setUpdatesEnabled(False); t_.setRowCount(0); t_.clearSpans()
        self.tbl_l.setColumnCount(len(left_headers))
        self.tbl_l.setHorizontalHeaderLabels(left_headers)
        self.tbl_r.setColumnCount(N_RIGHT)
        self.tbl_r.setHorizontalHeaderLabels(right_headers)
        self.tbl_lf.setColumnCount(len(left_headers))
        self.tbl_rf.setColumnCount(N_RIGHT)

        # Header heights — período (1-2 líneas con fecha) + sub-columna (~16px)
        hdr_h = 50 if self._show_fechas else 34
        self.tbl_r.horizontalHeader().setFixedHeight(hdr_h)
        self.tbl_l.horizontalHeader().setFixedHeight(hdr_h)
        # Cargar la estructura de 2 niveles en el header personalizado.
        if hasattr(self, '_period_header'):
            self._period_header.set_periodos(period_labels, hdr_h)
        self.tbl_r.horizontalHeader().setDefaultAlignment(Qt.AlignCenter)
        self.tbl_l.horizontalHeader().setDefaultAlignment(Qt.AlignCenter)

        # Widths left
        th_l = self.tbl_l.horizontalHeader()
        th_lf = self.tbl_lf.horizontalHeader()
        th_lf.setVisible(False)
        for c, w in enumerate([46, 80, None, 44, 85, 95]):
            if w is None:
                th_l.setSectionResizeMode(c, QHeaderView.Stretch)
                th_lf.setSectionResizeMode(c, QHeaderView.Stretch)
            else:
                th_l.setSectionResizeMode(c, QHeaderView.Interactive)
                self.tbl_l.setColumnWidth(c, w)
                th_lf.setSectionResizeMode(c, QHeaderView.Interactive)
                self.tbl_lf.setColumnWidth(c, w)
        try:
            th_l.sectionResized.disconnect()
        except Exception:
            pass
        th_l.sectionResized.connect(
            lambda i, _o, n: self.tbl_lf.setColumnWidth(i, n))

        # Widths right
        TOTAL_W = 110
        th_r = self.tbl_r.horizontalHeader()
        th_rf = self.tbl_rf.horizontalHeader()
        th_rf.setVisible(False)
        for c in range(N_RIGHT - 1):
            th_r.setSectionResizeMode(c, QHeaderView.Interactive)
            th_rf.setSectionResizeMode(c, QHeaderView.Interactive)
        th_r.setSectionResizeMode(COL_TOTAL_R, QHeaderView.Fixed)
        self.tbl_r.setColumnWidth(COL_TOTAL_R, TOTAL_W)
        th_rf.setSectionResizeMode(COL_TOTAL_R, QHeaderView.Fixed)
        self.tbl_rf.setColumnWidth(COL_TOTAL_R, TOTAL_W)
        th_r.setStretchLastSection(False); th_rf.setStretchLastSection(False)
        th_r.setMinimumSectionSize(78); th_rf.setMinimumSectionSize(78)
        self._last_n_periods = n_periods
        self._last_n_subcols = n_sub
        self._last_total_w = TOTAL_W
        try:
            th_r.sectionResized.disconnect()
        except Exception:
            pass
        th_r.sectionResized.connect(
            lambda i, _o, n: self.tbl_rf.setColumnWidth(i, n))
        QTimer.singleShot(0, self._apply_period_widths)

        alt_bg = QColor("#F1F4F9")
        totales_per = [0.0] * n_periods

        # Etiquetas de grupos por tipo (orden y nombre)
        TIPO_LABEL = {
            'MO':  'MANO DE OBRA',
            'MAT': 'MATERIALES',
            'EQ':  'EQUIPOS Y HERRAMIENTAS',
            'SC':  'SUB-CONTRATOS / SERVICIOS',
        }
        # Título de grupo = banda con tinte suave del tipo + texto del color
        # del tipo (igual que el reporte de Insumos). MO naranja/ámbar, MAT
        # verde, EQ gris-acero, SC morado.
        TIPO_FG = {'MO': '#F39C12', 'MAT': '#27AE60',
                    'EQ': '#607D8B', 'SC': '#7A36B1'}
        TIPO_SOFT = {'MO': '#FEF3DD', 'MAT': '#E3F4EA',
                      'EQ': '#ECEEF1', 'SC': '#F1E7F9'}

        last_tipo = None
        titulo_rows = []   # índices de filas de subtítulo (para sincronizar altura)

        def _add_titulo_grupo(label_txt: str, tipo: str = ''):
            """Inserta una fila de subtítulo que abarca todas las columnas del
            panel izquierdo y el derecho, con banda de tinte suave del tipo y
            el texto en el color del tipo."""
            rr = self.tbl_l.rowCount()
            self.tbl_l.insertRow(rr); self.tbl_r.insertRow(rr)
            self.tbl_l.setRowHeight(rr, 28); self.tbl_r.setRowHeight(rr, 28)
            titulo_rows.append(rr)
            bg_brush = QBrush(QColor(TIPO_SOFT.get(tipo, "#D8E0EC")))

            it = QTableWidgetItem(label_txt)
            it.setBackground(bg_brush)
            it.setForeground(QBrush(QColor(TIPO_FG.get(tipo, SLATE_700))))
            it.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
            f = QFont(); f.setBold(True); f.setPointSize(10); it.setFont(f)
            it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
            self.tbl_l.setItem(rr, 0, it)
            for c in range(1, len(left_headers)):
                it = QTableWidgetItem('')
                it.setBackground(bg_brush)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_l.setItem(rr, c, it)
            self.tbl_l.setSpan(rr, 0, 1, len(left_headers))

            for c in range(N_RIGHT):
                it = QTableWidgetItem('')
                it.setBackground(bg_brush)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_r.setItem(rr, c, it)

        for r_idx, rec in enumerate(sorted_rec):
            # Si cambia el tipo, insertar fila de subtítulo del grupo
            if rec['tipo'] != last_tipo:
                _add_titulo_grupo(TIPO_LABEL.get(rec['tipo'], rec['tipo']), rec['tipo'])
                last_tipo = rec['tipo']

            r = self.tbl_l.rowCount()
            self.tbl_l.insertRow(r); self.tbl_r.insertRow(r)
            row_bg = alt_bg if r % 2 == 1 else None

            cost_total = sum(rec['weekly'])
            qty_total  = sum(rec.get('weekly_qty', []))
            cant_txt = f"{qty_total:,.2f}" if qty_total > 0 else ''
            cells_left = [
                (rec['tipo'],                  Qt.AlignCenter),
                (rec['codigo'],                Qt.AlignLeft  | Qt.AlignVCenter),
                (rec['desc'],                  Qt.AlignLeft  | Qt.AlignVCenter),
                (rec['unidad'],                Qt.AlignCenter),
                (cant_txt,                     Qt.AlignRight | Qt.AlignVCenter),
                (fmt(cost_total, moneda),      Qt.AlignRight | Qt.AlignVCenter),
            ]
            for c, (val, align) in enumerate(cells_left):
                it = QTableWidgetItem(val)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                it.setTextAlignment(align)
                if c == 2 and rec['desc']:
                    it.setToolTip(rec['desc'])
                if row_bg is not None:
                    it.setBackground(QBrush(row_bg))
                self.tbl_l.setItem(r, c, it)

            row_total = 0.0
            wq = rec.get('weekly_qty', [0.0] * n_periods)
            for i in range(n_periods):
                qv = wq[i] if i < len(wq) else 0.0
                vv = rec['weekly'][i] if i < len(rec['weekly']) else 0.0
                # Cantidad (sub-col 2i)
                itq = QTableWidgetItem(f"{qv:,.2f}" if qv > 0 else '')
                itq.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                itq.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                if qv > 0:
                    itq.setForeground(QBrush(QColor(SLATE_500)))
                if row_bg is not None:
                    itq.setBackground(QBrush(row_bg))
                self.tbl_r.setItem(r, 2 * i, itq)
                # Valorización (sub-col 2i+1)
                itv = QTableWidgetItem(fmt(vv, moneda) if vv > 0 else '')
                itv.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                itv.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                if vv > 0:
                    itv.setForeground(QBrush(QColor("#1F2A38")))
                if row_bg is not None:
                    itv.setBackground(QBrush(row_bg))
                self.tbl_r.setItem(r, 2 * i + 1, itv)
                row_total += vv
                totales_per[i] += vv

            it = QTableWidgetItem(fmt(row_total, moneda))
            it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
            it.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            if row_bg is not None:
                it.setBackground(QBrush(row_bg))
            self.tbl_r.setItem(r, COL_TOTAL_R, it)

        # ── Footer (TOTAL · % PERÍODO · ACUMULADO · % ACUMULADO) ─────────
        BG_MONTO = "#E2E8F0"
        BG_PCT   = "#F1F5F9"
        FG_TEXT  = "#1F2A38"
        FG_LBL   = "#273445"

        def _add_summary(label_txt, vals, total_val, *, bg, pct=False):
            bg_brush = QBrush(QColor(bg))
            fg_brush = QBrush(QColor(FG_TEXT))
            lbl_brush = QBrush(QColor(FG_LBL))
            r = self.tbl_lf.rowCount()
            self.tbl_lf.insertRow(r); self.tbl_rf.insertRow(r)
            self.tbl_lf.setRowHeight(r, 28); self.tbl_rf.setRowHeight(r, 28)
            it_lbl = QTableWidgetItem(label_txt)
            it_lbl.setBackground(bg_brush); it_lbl.setForeground(lbl_brush)
            f = QFont(); f.setBold(True); f.setPointSize(10); it_lbl.setFont(f)
            it_lbl.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            it_lbl.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
            self.tbl_lf.setItem(r, 0, it_lbl)
            for c in range(1, len(left_headers)):
                it = QTableWidgetItem('')
                it.setBackground(bg_brush)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_lf.setItem(r, c, it)
            self.tbl_lf.setSpan(r, 0, 1, len(left_headers))
            for i, v in enumerate(vals):
                # Sub-col Cantidad vacía (sumar cantidades de unidades
                # distintas no aplica); el valor va en la sub-col Valorización.
                itq = QTableWidgetItem('')
                itq.setBackground(bg_brush)
                itq.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_rf.setItem(r, 2 * i, itq)
                txt = (f"{v:.1f}%" if v > 0 else '') if pct else (fmt(v, moneda) if v > 0 else '')
                it = QTableWidgetItem(txt)
                it.setBackground(bg_brush); it.setForeground(fg_brush)
                it.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                ff = QFont(); ff.setBold(True); it.setFont(ff)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_rf.setItem(r, 2 * i + 1, it)
            if total_val is not None:
                txt = f"{total_val:.1f}%" if pct else fmt(total_val, moneda)
                it = QTableWidgetItem(txt)
                it.setBackground(bg_brush); it.setForeground(fg_brush)
                it.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                ff = QFont(); ff.setBold(True); it.setFont(ff)
                it.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.tbl_rf.setItem(r, COL_TOTAL_R, it)

        _add_summary("TOTAL", totales_per, total_general, bg=BG_MONTO)
        pct_per = [(v / total_general * 100.0 if total_general else 0.0) for v in totales_per]
        _add_summary("% PERÍODO", pct_per, 100.0 if total_general else 0.0,
                       bg=BG_PCT, pct=True)
        acumulado = []
        s = 0.0
        for v in totales_per:
            s += v; acumulado.append(s)
        _add_summary("ACUMULADO", acumulado, total_general, bg=BG_MONTO)
        pct_acum = [(v / total_general * 100.0 if total_general else 0.0) for v in acumulado]
        _add_summary("% ACUMULADO", pct_acum, 100.0 if total_general else 0.0,
                       bg=BG_PCT, pct=True)

        # Ancho dinámico para que TOTAL/ACUMULADO (montos grandes) no se
        # corten: medir el valor más ancho (= acumulado final = total) y
        # ampliar el mínimo de columna y la columna Total (con tope).
        from PySide6.QtGui import QFontMetrics
        _fb = QFont(); _fb.setBold(True)
        _fm = QFontMetrics(_fb)
        _widest = fmt(total_general, moneda) if total_general else ''
        _need = (_fm.horizontalAdvance(_widest) + 22) if _widest else 78
        self._min_per_w = max(78, min(_need, 170))
        self._last_total_w = max(110, min(_need + 6, 185))
        self.tbl_r.setColumnWidth(COL_TOTAL_R, self._last_total_w)
        self.tbl_rf.setColumnWidth(COL_TOTAL_R, self._last_total_w)

        for t_ in (self.tbl_l, self.tbl_r, self.tbl_lf, self.tbl_rf):
            t_.setUpdatesEnabled(True)
        QTimer.singleShot(0, self._apply_period_widths)
        # Resize por contenido en el panel izquierdo, luego forzar altura
        # consistente para las filas de subtítulo (resizeRowsToContents puede
        # darles altura distinta a causa del setSpan) y replicar al derecho.
        self.tbl_l.resizeRowsToContents()
        titulo_set = set(titulo_rows)
        for r in range(self.tbl_l.rowCount()):
            if r in titulo_set:
                # Subtítulos: altura fija 28px en ambos paneles
                self.tbl_l.setRowHeight(r, 28)
                self.tbl_r.setRowHeight(r, 28)
            else:
                self.tbl_r.setRowHeight(r, self.tbl_l.rowHeight(r))
        h_t = sum(self.tbl_lf.rowHeight(i) for i in range(self.tbl_lf.rowCount())) + 4
        self.tbl_lf.setFixedHeight(h_t); self.tbl_rf.setFixedHeight(h_t)

    # ─── Exportación (reutiliza el diálogo del Valorizado) ────────────────

    def _dir_descargas(self) -> str:
        import os
        from PySide6.QtCore import QSettings
        s = QSettings("ingePresupuestos", "exports")
        g = s.value("last_dir_insumos", "")
        if g and os.path.isdir(g):
            return g
        for c in (os.path.expanduser("~/Descargas"),
                   os.path.expanduser("~/Downloads"),
                   os.path.expanduser("~")):
            if os.path.isdir(c):
                return c
        return os.getcwd()

    def _save_dir(self, path):
        import os
        from PySide6.QtCore import QSettings
        if not path: return
        d = os.path.dirname(path)
        if d and os.path.isdir(d):
            QSettings("ingePresupuestos", "exports").setValue("last_dir_insumos", d)

    def _exportar_pdf(self):
        opts = _DialogExportarReportePdf.preguntar(
            self, "Exportar adquisición de insumos",
            on_preview=self._vista_previa_pdf,
            on_imagen=self._exportar_imagen,
        )
        if opts is None:
            return
        self._render_pdf_a_archivo(opts)

    def _render_pdf_a_archivo(self, opts):
        import os
        try:
            from core.pdf_reports import generar_pdf_archivo
        except Exception as e:
            QMessageBox.warning(self, "Exportar PDF", f"{e}"); return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"adquisiciones_{self._cv.pid}.pdf")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar adquisición de insumos a PDF",
            sugerido, "PDF (*.pdf)")
        if not path:
            return
        try:
            from core import pdf_reports as _pr
            _pr._BUILD_HTML_PAPER['escala'] = ('semana' if self._period_days == 7 else 'mes')
            _pr._BUILD_HTML_PAPER['periodos_por_pagina'] = int(opts.get('periodos_por_pagina', 0) or 0)
            _pr._BUILD_HTML_PAPER['show_fechas'] = bool(opts.get('show_fechas', False))
            generar_pdf_archivo(
                'cronograma_adquisiciones', self._cv.pid, path,
                with_cover=opts.get('with_cover', False),
                paper=opts.get('papel', 'A4'),
                orient=opts.get('orient', 'portrait'),
            )
            self._save_dir(path)
            QMessageBox.information(self, "Exportar PDF",
                                       f"Adquisiciones exportadas a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"No se pudo exportar: {e}")

    def _vista_previa_pdf(self, opts):
        import os, tempfile
        try:
            from core.pdf_reports import generar_pdf_archivo
            from PySide6.QtPdf import QPdfDocument
            from PySide6.QtPdfWidgets import QPdfView
        except Exception as e:
            QMessageBox.warning(self, "Vista previa", f"{e}"); return
        tmp_fd, tmp_pdf = tempfile.mkstemp(suffix='.pdf', prefix='insumos_prev_')
        os.close(tmp_fd)
        try:
            from core import pdf_reports as _pr
            _pr._BUILD_HTML_PAPER['escala'] = ('semana' if self._period_days == 7 else 'mes')
            _pr._BUILD_HTML_PAPER['periodos_por_pagina'] = int(opts.get('periodos_por_pagina', 0) or 0)
            _pr._BUILD_HTML_PAPER['show_fechas'] = bool(opts.get('show_fechas', False))
            generar_pdf_archivo(
                'cronograma_adquisiciones', self._cv.pid, tmp_pdf,
                with_cover=opts.get('with_cover', False),
                paper=opts.get('papel', 'A4'),
                orient=opts.get('orient', 'portrait'),
            )
            if not os.path.exists(tmp_pdf) or os.path.getsize(tmp_pdf) < 500:
                QMessageBox.warning(self, "Vista previa", "PDF vacío.")
                return
            dlg = QDialog(self)
            dlg.setWindowTitle("Vista previa — Adquisición de Insumos")
            dlg.setWindowModality(Qt.ApplicationModal)
            dlg.resize(1100, 800)
            v = QVBoxLayout(dlg); v.setContentsMargins(8, 8, 8, 8); v.setSpacing(6)
            tb = QHBoxLayout()
            bn = (f"QPushButton {{ background:white; color:{SLATE_700};"
                   f" border:1px solid {SILVER_300}; border-radius:4px;"
                   f" padding:6px 14px; }} QPushButton:hover {{ background:#FEF5EB;"
                   " border-color:#F37329; color:#C0621A; }}")
            b1 = QPushButton("🖼  Exportar imagen…"); b1.setStyleSheet(bn)
            b1.clicked.connect(lambda: self._exportar_imagen(opts))
            b2 = QPushButton("📄  Guardar PDF…"); b2.setStyleSheet(bn)
            b2.clicked.connect(lambda: self._render_pdf_a_archivo(opts))
            b3 = QPushButton("Cerrar")
            b3.setStyleSheet(f"QPushButton {{ background:{SLATE_500}; color:white;"
                              f" border:none; border-radius:4px; padding:6px 18px; }}"
                              f"QPushButton:hover {{ background:{SLATE_700}; }}")
            b3.clicked.connect(dlg.accept)
            tb.addWidget(b1); tb.addWidget(b2); tb.addStretch(); tb.addWidget(b3)
            v.addLayout(tb)
            doc = QPdfDocument(dlg); doc.load(tmp_pdf)
            view = QPdfView(dlg); view.setDocument(doc)
            view.setPageMode(QPdfView.PageMode.MultiPage)
            view.setZoomMode(QPdfView.ZoomMode.FitToWidth)
            v.addWidget(view, stretch=1)
            dlg.exec()
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Vista previa", f"{e}")
        finally:
            try: os.unlink(tmp_pdf)
            except Exception: pass

    def _exportar_imagen(self, opts):
        import os, tempfile
        try:
            from core.pdf_reports import generar_pdf_archivo
            from PySide6.QtPdf import QPdfDocument
            from PySide6.QtGui import QImage
        except Exception as e:
            QMessageBox.warning(self, "Exportar imagen", f"{e}"); return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"adquisiciones_{self._cv.pid}.png")
        path, sel = QFileDialog.getSaveFileName(
            self, "Exportar como imagen", sugerido,
            "PNG (*.png);;JPEG (*.jpg *.jpeg)")
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        if 'jpg' in (sel or '').lower() or 'jpeg' in (sel or '').lower():
            if ext not in ('.jpg', '.jpeg'):
                path += '.jpg'
            fmt_ = 'JPG'
        else:
            if ext != '.png':
                path += '.png'
            fmt_ = 'PNG'
        tmp_fd, tmp_pdf = tempfile.mkstemp(suffix='.pdf', prefix='insumos_img_')
        os.close(tmp_fd)
        try:
            from core import pdf_reports as _pr
            _pr._BUILD_HTML_PAPER['escala'] = ('semana' if self._period_days == 7 else 'mes')
            _pr._BUILD_HTML_PAPER['periodos_por_pagina'] = int(opts.get('periodos_por_pagina', 0) or 0)
            _pr._BUILD_HTML_PAPER['show_fechas'] = bool(opts.get('show_fechas', False))
            generar_pdf_archivo(
                'cronograma_adquisiciones', self._cv.pid, tmp_pdf,
                with_cover=opts.get('with_cover', False),
                paper=opts.get('papel', 'A4'),
                orient=opts.get('orient', 'portrait'),
            )
            doc = QPdfDocument(); doc.load(tmp_pdf)
            n = doc.pageCount()
            if n <= 0:
                raise RuntimeError("PDF vacío")
            dpi = 200
            base, real_ext = os.path.splitext(path)
            outs = []
            for i in range(n):
                page_pts = doc.pagePointSize(i)
                w_px = int(page_pts.width()  / 72.0 * dpi)
                h_px = int(page_pts.height() / 72.0 * dpi)
                rendered = doc.render(i, QSize(max(1, w_px), max(1, h_px)))
                canvas = QImage(rendered.size(), QImage.Format_RGB32)
                canvas.fill(QColor("white"))
                pp = QPainter(canvas)
                try: pp.drawImage(0, 0, rendered)
                finally: pp.end()
                out = path if n == 1 else f"{base}_p{i+1}{real_ext}"
                if not canvas.save(out, fmt_, 92 if fmt_ == 'JPG' else -1):
                    raise RuntimeError(f"No se pudo guardar {out}")
                outs.append(out)
            self._save_dir(path)
            msg = ("Imagen exportada:\n" + outs[0]) if n == 1 \
                else f"{n} páginas exportadas:\n{outs[0]}\n…\n{outs[-1]}"
            QMessageBox.information(self, "Exportar imagen", msg)
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Exportar imagen", f"{e}")
        finally:
            try: os.unlink(tmp_pdf)
            except Exception: pass

    def _exportar_excel(self):
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os
        sugerido = os.path.join(self._dir_descargas(),
                                  f"adquisiciones_{self._cv.pid}.xlsx")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar adquisición de insumos a Excel",
            sugerido, "Excel (*.xlsx)")
        if not path: return
        try:
            self._build_xlsx(path)
            self._save_dir(path)
            QMessageBox.information(self, "Exportar Excel",
                                       f"Adquisiciones exportadas a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"{e}")

    def _exportar_ods(self):
        from core.licencia import require_premium
        if not require_premium('export_editable', self):
            return
        import os, subprocess, tempfile
        from core.soffice import find_soffice, mensaje_instalacion
        soffice = find_soffice()
        if not soffice:
            QMessageBox.warning(self, "Exportar ODS", mensaje_instalacion())
            return
        sugerido = os.path.join(self._dir_descargas(),
                                  f"adquisiciones_{self._cv.pid}.ods")
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportar adquisición de insumos a ODS",
            sugerido, "OpenDocument Spreadsheet (*.ods)")
        if not path: return
        tmp_fd, tmp_xlsx = tempfile.mkstemp(suffix='.xlsx', prefix='insumos_ods_')
        os.close(tmp_fd)
        try:
            self._build_xlsx(tmp_xlsx)
            out_dir = os.path.dirname(os.path.abspath(path))
            subprocess.run(
                [soffice, '--headless', '--convert-to', 'ods',
                 '--outdir', out_dir, tmp_xlsx],
                check=True, capture_output=True, timeout=60)
            generated = os.path.join(out_dir,
                os.path.splitext(os.path.basename(tmp_xlsx))[0] + '.ods')
            if os.path.exists(generated) and generated != path:
                os.replace(generated, path)
            if not os.path.exists(path):
                raise RuntimeError("LibreOffice no generó el archivo .ods")
            self._save_dir(path)
            QMessageBox.information(self, "Exportar ODS",
                                       f"Adquisiciones exportadas a:\n{path}")
        except Exception as e:
            import traceback; traceback.print_exc()
            QMessageBox.warning(self, "Error", f"{e}")
        finally:
            try: os.unlink(tmp_xlsx)
            except Exception: pass

    def _build_xlsx(self, path: str):
        """Construye .xlsx — espejo VISUAL del PDF de Adquisición de Insumos.

        Mismo patrón "PDF visible" que Cronograma Valorizado:
        - Header tripartito vía `_xlsx_header_pdf_style`.
        - h2 izquierdo "Cronograma de Adquisición de Insumos — por {label}"
          13pt color `accent_reportes()[1]`.
        - Sin gridlines, sin bordes en datos/títulos/header de columnas;
          solo persiste `border-top` 1.5pt slate-700 en la fila TOTAL.
        - Zebra sutil `#FBFCFD` con contador propio de filas-insumo
          (los títulos de grupo MO/MAT/EQ/SC no rompen la alternancia).
        - print_title_rows + print_title_cols para repetir encabezado y
          cols frozen en cada página al imprimir / convertir a ODS.
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from core.exporter import _xlsx_header_pdf_style
        from utils.theme import accent_reportes

        wb = Workbook(); ws = wb.active; ws.title = "Adquisiciones"
        proy = self._cv._proy

        # Paleta — espejo del PDF
        _, od_hex, _ = accent_reportes()
        ACC_OD   = od_hex.lstrip('#').upper()
        BG_MONTO = 'F1F5F9'                          # TOTAL/ACUMULADO
        BG_PCT   = 'FBFCFD'                          # %PERÍODO/%ACUMULADO
        ALT_BG   = 'FBFCFD'                          # zebra muy sutil
        TXT_DK   = '1F2A38'                          # SLATE_900
        TXT_700  = '273445'                          # SLATE_700
        TXT_500  = '485A6C'                          # SLATE_500

        side_slate7 = Side(style='medium', color=TXT_700)    # ~1.5pt
        border_grid   = Border()
        border_head   = Border()
        border_title  = Border()
        border_total  = Border(top=side_slate7)

        align_left   = Alignment(horizontal='left',   vertical='center',
                                  wrap_text=True)
        align_right  = Alignment(horizontal='right',  vertical='center')
        align_center = Alignment(horizontal='center', vertical='center',
                                  wrap_text=True)

        # SPLIT de Descripción (col C original) en C+D — patrón espejo de
        # Cronograma Valorizado. En cada fila C+D se mergean → user ve 1
        # sola col Descripción.
        n_left_logical = self.tbl_l.columnCount()   # 6 cols de tbl_l
        n_left  = n_left_logical + 1                 # 7 cols físicas
        n_right = self.tbl_r.columnCount()
        n_cols  = n_left + n_right

        # Helper: logical col (0-based, tbl_l) → physical Excel col (1-based)
        def _phys(c_logical: int) -> int:
            if c_logical <= 2:
                return c_logical + 1            # Tipo, Código, Desc1
            return c_logical + 2                # Und, Precio, %Total (shifted)

        def _merge_desc(row_idx: int):
            ws.merge_cells(start_row=row_idx, start_column=3,
                              end_row=row_idx, end_column=4)

        # Anchos de columna ANTES del header (wrap calc usa width).
        # Físicas: A Tipo · B Código · C/D Descripción (split) · E Und ·
        #          F Cantidad · G Precio · H % Total.
        ws.column_dimensions['A'].width = 7    # Tipo
        ws.column_dimensions['B'].width = 12   # Código
        ws.column_dimensions['C'].width = 18   # Descripción parte 1
        ws.column_dimensions['D'].width = 22   # Descripción parte 2
        ws.column_dimensions['E'].width = 7    # Und
        ws.column_dimensions['F'].width = 13   # Cantidad (total)
        ws.column_dimensions['G'].width = 14   # Precio (total)
        for c in range(n_left + 1, n_cols + 1):
            ws.column_dimensions[get_column_letter(c)].width = 14

        ws.sheet_view.showGridLines = False
        ws.print_options.gridLines = False
        ws.print_options.gridLinesSet = True

        # ── Fila 1-3: header tripartito con hueco intencional ─────────────
        # cols_partition=(3, 7, n_cols-1) → Left=A-C (Tipo+Código+Desc1),
        # Center=D-G (Desc2+Und+Precio+%Total) para título+nombre del
        # proyecto, Right=last 2 cols (costo+modalidad). Cols intermedias
        # (períodos) quedan sin contenido en el header.
        if n_cols >= 9:
            row = _xlsx_header_pdf_style(
                ws, proy, 'Cronograma de Adquisición de Insumos', n_cols,
                cols_partition=(3, 7, n_cols - 1),
            )
        else:
            row = _xlsx_header_pdf_style(
                ws, proy, 'Cronograma de Adquisición de Insumos', n_cols,
            )

        # ── h2 izquierdo + subtítulo ─────────────────────────────────────
        label = 'Semana' if self._period_days == 7 else 'Mes'
        label_plural = 'semanas' if self._period_days == 7 else 'meses'
        plazo = int(proy.get('plazo') or 0)
        n_periods = max(0, n_right - 1)

        # Contar insumos (filas que NO son títulos de grupo)
        n_insumos = sum(1 for r in range(self.tbl_l.rowCount())
                        if self.tbl_l.columnSpan(r, 0) <= 1)

        ws.merge_cells(start_row=row, start_column=1,
                          end_row=row, end_column=n_cols)
        c_h2 = ws.cell(row=row, column=1,
                          value=f"Cronograma de Adquisición de Insumos — por {label}")
        c_h2.font      = Font(name='Inter', bold=True, size=13, color=ACC_OD)
        c_h2.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[row].height = 20
        row += 1

        ws.merge_cells(start_row=row, start_column=1,
                          end_row=row, end_column=n_cols)
        c_sub = ws.cell(row=row, column=1,
                          value=f"Plazo: {plazo} días  ·  "
                                f"Períodos: {n_periods} {label_plural}  ·  "
                                f"Insumos: {n_insumos}")
        c_sub.font      = Font(name='Inter', size=9, color=TXT_500)
        c_sub.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[row].height = 14
        row += 1

        ws.row_dimensions[row].height = 6
        row += 1

        # ── Header de la tabla — 2 niveles (período sobre sus 2 sub-cols) ──
        head_fill = PatternFill('solid', fgColor=BG_MONTO)
        head_font = Font(name='Inter', bold=True, size=10, color=TXT_DK)
        n_periods_h = (n_right - 1) // 2
        hdr1 = row
        hdr2 = row + 1
        for col_ in range(1, n_cols + 1):
            ws.cell(hdr1, col_).fill = head_fill
            ws.cell(hdr2, col_).fill = head_fill
        # Cols izquierdas: rótulo merged verticalmente (hdr1:hdr2). El ancho
        # físico de cada col lógica se deriva de _phys (la de Descripción
        # ocupa 2 cols: _phys(c+1)-_phys(c)==2).
        for c in range(n_left_logical):
            pc = _phys(c)
            pc_end = (_phys(c + 1) - 1) if (c + 1) < n_left_logical else n_left
            txt = (self.tbl_l.horizontalHeaderItem(c).text() or '').replace('\n', ' ')
            cell = ws.cell(hdr1, pc, txt)
            cell.font = head_font; cell.alignment = align_center
            ws.merge_cells(start_row=hdr1, start_column=pc, end_row=hdr2, end_column=pc_end)
        # Períodos: nombre (Mes/Sem + fecha) sobre las 2 sub-cols en hdr1,
        # sub-etiquetas (Cantidad | Precio) en hdr2.
        for i in range(n_periods_h):
            mc = n_left + 1 + 2 * i
            vc = mc + 1
            raw_m = (self.tbl_r.horizontalHeaderItem(2 * i).text() or '')
            lines = raw_m.split('\n')
            plabel = '\n'.join(lines[:-1]) if len(lines) > 1 else raw_m
            sub_c  = lines[-1] if lines else 'Cantidad'
            sub_p  = (self.tbl_r.horizontalHeaderItem(2 * i + 1).text() or 'Precio').split('\n')[-1]
            ws.merge_cells(start_row=hdr1, start_column=mc, end_row=hdr1, end_column=vc)
            cp = ws.cell(hdr1, mc, plabel); cp.font = head_font; cp.alignment = align_center
            cc = ws.cell(hdr2, mc, sub_c);  cc.font = head_font; cc.alignment = align_center
            cv = ws.cell(hdr2, vc, sub_p);  cv.font = head_font; cv.alignment = align_center
        # Total: merged vertical
        tc = n_cols
        ws.merge_cells(start_row=hdr1, start_column=tc, end_row=hdr2, end_column=tc)
        ct = ws.cell(hdr1, tc, (self.tbl_r.horizontalHeaderItem(n_right - 1).text() or 'Total'))
        ct.font = head_font; ct.alignment = align_center
        head_row = hdr1
        ws.row_dimensions[hdr1].height = 26 if self._show_fechas else 16
        ws.row_dimensions[hdr2].height = 14
        row += 2

        # ── Filas de datos ───────────────────────────────────────────────
        def _t(it):
            return (it.text() if it is not None else '') or ''

        def _n(s):
            if not s:
                return None
            t = (str(s).replace('S/', '').replace('US$', '').replace('€', '')
                       .strip().replace(',', '').rstrip('%').strip())
            try:
                return float(t)
            except ValueError:
                return s

        def _alto_desc(texto: str) -> float:
            """Alto de fila (pts) para que la Descripción mergeada C+D con
            wrap_text NO se corte — Excel no auto-ajusta el alto en celdas
            MERGEADAS. Estimamos: ancho útil ≈ C(18)+D(22)=40 chars y
            contamos cuántas líneas ocupa el texto (15 pt por línea)."""
            ancho = 40
            lineas = 0
            for seg in (texto or '').split('\n'):
                lineas += max(1, -(-len(seg) // ancho))   # ceil div
            return max(15.0, lineas * 15.0)

        title_font = Font(name='Inter', bold=True, size=10, color=TXT_DK)
        data_font  = Font(name='Inter', size=9, color=TXT_DK)
        # Título de grupo = banda con tinte suave del tipo + texto del color
        # del tipo (igual que el reporte de Insumos).
        _TIPO_COL_BY_LABEL = {
            'MANO DE OBRA': 'F39C12', 'MATERIALES': '27AE60',
            'EQUIPOS Y HERRAMIENTAS': '607D8B',
            'SUB-CONTRATOS / SERVICIOS': '7A36B1',
        }
        _TIPO_SOFT_BY_LABEL = {
            'MANO DE OBRA': 'FEF3DD', 'MATERIALES': 'E3F4EA',
            'EQUIPOS Y HERRAMIENTAS': 'ECEEF1',
            'SUB-CONTRATOS / SERVICIOS': 'F1E7F9',
        }

        # Contador propio de filas-insumo para zebra CONTINUO (los títulos
        # de grupo MO/MAT/EQ/SC no se cuentan, así la alternancia visual
        # se mantiene al cruzar un título).
        insumo_idx = 0
        for r in range(self.tbl_l.rowCount()):
            # Título de grupo: columnSpan(0) cubre todo el panel izquierdo
            is_titulo = self.tbl_l.columnSpan(r, 0) > 1
            if is_titulo:
                label_val = (_t(self.tbl_l.item(r, 0)) or '').upper()
                ws.merge_cells(start_row=row, start_column=1,
                                  end_row=row, end_column=n_cols)
                cell = ws.cell(row=row, column=1, value=label_val)
                cell.font      = Font(name='Inter', bold=True, size=10,
                                       color=_TIPO_COL_BY_LABEL.get(label_val, TXT_DK))
                cell.alignment = align_left
                cell.border    = border_title
                _soft = _TIPO_SOFT_BY_LABEL.get(label_val)
                if _soft:
                    cell.fill = PatternFill('solid', fgColor=_soft)
                ws.row_dimensions[row].height = 22
                row += 1
                continue

            zebra      = (insumo_idx % 2 == 1)
            zebra_fill = (PatternFill('solid', fgColor=ALT_BG) if zebra else None)
            insumo_idx += 1

            # PRE-aplicar fill+border en TODAS las cols del left panel
            # (incluye col 4 = Desc2, MergedCell post-merge)
            for col_ in range(1, n_left + 1):
                cell = ws.cell(row=row, column=col_)
                cell.font = data_font
                cell.border = border_grid
                if zebra_fill:
                    cell.fill = zebra_fill

            for c in range(n_left_logical):
                val  = _t(self.tbl_l.item(r, c))
                cell = ws.cell(row=row, column=_phys(c))
                # Logical: Tipo(0) · Código(1) · Descripción(2) · Und(3) ·
                # Cantidad(4) · Precio(5) — ambos totales del insumo.
                if c in (4, 5):
                    n = _n(val)
                    cell.value = n if isinstance(n, (int, float)) else val
                    if isinstance(n, (int, float)):
                        cell.number_format = '#,##0.00'
                    cell.alignment = align_right
                elif c in (0, 3):
                    cell.value = val
                    cell.alignment = align_center
                else:
                    cell.value = val
                    cell.alignment = align_left
            _merge_desc(row)

            for c in range(n_right):
                val  = _t(self.tbl_r.item(r, c))
                n    = _n(val)
                cell = ws.cell(row=row, column=n_left + c + 1)
                cell.value = n if isinstance(n, (int, float)) else val
                if isinstance(n, (int, float)):
                    cell.number_format = '#,##0.00'
                cell.alignment = align_right
                cell.font   = data_font
                cell.border = border_grid
                if zebra_fill:
                    cell.fill = zebra_fill
            # Alto de fila para que la descripción larga del insumo (C+D
            # mergeada con wrap) no se corte — Excel no lo auto-ajusta.
            ws.row_dimensions[row].height = _alto_desc(_t(self.tbl_l.item(r, 2)))
            row += 1

        # ── Filas de resumen (footer) ─────────────────────────────────────
        ftr_font = Font(name='Inter', bold=True, size=10, color=TXT_DK)
        lbl_font = Font(name='Inter', bold=True, size=10, color=TXT_700)
        for r in range(self.tbl_lf.rowCount()):
            is_pct   = (r in (1, 3))                  # %PERÍODO, %ACUMULADO
            is_total = (r == 0)                        # TOTAL — border-top
            bg       = BG_PCT if is_pct else BG_MONTO
            fill     = PatternFill('solid', fgColor=bg)
            row_bdr  = border_total if is_total else border_grid

            lbl = (_t(self.tbl_lf.item(r, 0)) or '').strip()
            if not lbl:
                for c in range(n_left):
                    t = _t(self.tbl_lf.item(r, c))
                    if t.strip():
                        lbl = t.strip()
                        break
            # PRE-estilar las cols del label ANTES de mergear (setear .fill o
            # .border sobre un MergedCell ya creado lanza error en openpyxl).
            for c in range(n_left):
                cell = ws.cell(row=row, column=c + 1)
                if c == 0:
                    cell.value = lbl
                cell.font      = lbl_font
                cell.fill      = fill
                cell.border    = row_bdr
                cell.alignment = align_right
            ws.merge_cells(start_row=row, start_column=1,
                              end_row=row, end_column=n_left)

            for c in range(n_right):
                val  = _t(self.tbl_rf.item(r, c))
                n    = _n(val)
                cell = ws.cell(row=row, column=n_left + c + 1)
                if isinstance(n, (int, float)):
                    cell.value = n
                    cell.number_format = '0.0"%"' if '%' in val else '#,##0.00'
                else:
                    cell.value = val
                cell.font      = ftr_font
                cell.fill      = fill
                cell.border    = row_bdr
                cell.alignment = align_right
            ws.row_dimensions[row].height = 18
            row += 1

        # ── Líneas divisorias (estilo Delphin / hoja de Marco) ────────────
        # Verticales medias en borde izq., tras el bloque izquierdo (% Total)
        # y borde derecho; finas entre cada mes. Sin horizontales (zebra).
        last_row = row - 1
        MED_CLR  = '94A3B8'      # slate-400
        THIN_CLR = 'D5DCE5'      # slate-200

        def _side(cell, edge, style, color):
            b = cell.border
            kw = {s: getattr(b, s) for s in ('top', 'bottom', 'left', 'right')}
            kw[edge] = Side(style=style, color=color)
            cell.border = Border(**kw)

        med_cols = {1: 'left', n_left: 'right', n_cols: 'right'}
        thin_right = {n_left + 2 * (m + 1) for m in range(n_periods_h)}
        thin_right.discard(n_cols - 1)
        for r in range(hdr1, last_row + 1):
            for col_, edge in med_cols.items():
                _side(ws.cell(r, col_), edge, 'medium', MED_CLR)
            for col_ in thin_right:
                _side(ws.cell(r, col_), 'right', 'thin', THIN_CLR)
        for col_ in range(1, n_cols + 1):
            _side(ws.cell(hdr1, col_), 'top', 'medium', MED_CLR)
            _side(ws.cell(hdr2, col_), 'bottom', 'medium', MED_CLR)
            _side(ws.cell(last_row, col_), 'bottom', 'medium', MED_CLR)

        # Freeze panes después de % Total y bajo el header de 2 filas.
        # Usar coordenada como STRING — pasar un Cell object a freeze_panes
        # falla en openpyxl 3.x cuando la celda destino acaba como
        # MergedCell por colisión con otros merges del workbook.
        ws.freeze_panes = f'{get_column_letter(n_left + 1)}{head_row + 2}'

        # Repetir encabezado (2 filas) y cols frozen en cada página impresa
        ws.print_title_rows = f'1:{head_row + 1}'
        ws.print_title_cols = f'A:{get_column_letter(n_left)}'

        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.print_options.horizontalCentered = True
        # Márgenes compactos espejo de Cronograma Valorizado / Presupuesto.
        from openpyxl.worksheet.page import PageMargins as _PM_adq
        ws.page_margins = _PM_adq(left=0.4, right=0.4, top=0.35, bottom=0.5,
                                     header=0.15, footer=0.2)
        # Pie tripartito: Cliente | Fecha | Página X de N
        from datetime import datetime as _dt_adq
        cliente_adq = (proy.get('cliente') or '').strip()
        fecha_adq   = _dt_adq.now().strftime('%d/%m/%Y')
        prefix_adq  = '&7&K485A6C'   # 7pt + slate-500 (tamaño antes del color)
        ws.oddFooter.left.text   = f"{prefix_adq}Cliente: {cliente_adq}" if cliente_adq else f"{prefix_adq} "
        ws.oddFooter.center.text = f"{prefix_adq}{fecha_adq}"
        ws.oddFooter.right.text  = f"{prefix_adq}Página &P de &N"
        wb.save(path)
